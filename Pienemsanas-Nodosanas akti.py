import sys
import ctypes
from typing import Optional
import os
import hashlib
import secrets
from PySide6.QtWidgets import QMenu, QDialog, QFormLayout, QDialogButtonBox
from PySide6.QtGui import QAction
import csv

# ============================
# FIX: path-like safety (dict -> str) to allow repeated PDF generation
# ============================
# Dažkārt (pēc pirmās ģenerēšanas) UI stāvoklī ceļi var nonākt kā dict (piem. Qt.UserRole).
# Tas izraisa TypeError "expected str, bytes or os.PathLike object, not dict" pie os.path / open() / ReportLab.
# Šis ir minimāls "drop-in" ielāps: padara os.fspath tolerantāku pret dict/tuple/list,
# lai visur, kur Python/stdlib iekšēji izsauc os.fspath(), mēs vienmēr saņemtu string ceļu.

_ORIG_OS_FSPATH = os.fspath

def _safe_fspath(p):
    try:
        if p is None:
            return ""
        # QFileDialog dažreiz atgriež (path, filter)
        if isinstance(p, (list, tuple)) and p:
            return _safe_fspath(p[0])
        # UI/state var iedot dict ar ceļu
        if isinstance(p, dict):
            for k in ("path", "ceļš", "celsh", "file", "filepath", "filename", "value", "pdf", "json"):
                v = p.get(k)
                if v:
                    return _safe_fspath(v)
            return ""
        return _ORIG_OS_FSPATH(p)
    except TypeError:
        # Pēdējais glābiņš: pārvēršam par str (labāk FileNotFoundError nekā TypeError)
        try:
            return str(p)
        except Exception:
            return ""

# Monkey-patch: ietekmē os.path.*, open(), reportlab utt.
os.fspath = _safe_fspath


def resource_path(relative_path: str) -> str:
    """Atgriež pareizu ceļu uz resursu gan Python režīmā, gan PyInstaller EXE režīmā."""
    base_path = getattr(sys, "_MEIPASS", os.path.abspath("."))
    return os.path.join(base_path, relative_path)

def set_windows_app_id(app_id: str) -> None:
    """Uzliek Windows AppUserModelID, lai Taskbar/Alt+Tab izmantotu pareizo ikonu."""
    try:
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(app_id)
    except Exception:
        pass

import os.path
import io
import re

import json
import base64
import tempfile
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from PySide6.QtGui import QColor, QPageSize
from dataclasses import dataclass, asdict, field
import shutil
import requests
import subprocess # Jauns imports

# Pārliecināties, ka šīs ir importētas no PySide6.QtWidgets
from PySide6.QtWidgets import (
    QApplication, QWidget, QMainWindow, QLabel, QLineEdit, QTextEdit, QPushButton,
    QFileDialog, QSpinBox, QDoubleSpinBox, QTableWidget, QTableWidgetItem, QHeaderView,
    QToolButton, QTabWidget, QFormLayout, QVBoxLayout, QHBoxLayout, QMessageBox, QCheckBox,
    QListWidget, QListWidgetItem, QGroupBox, QComboBox, QInputDialog, QSplitter, QScrollArea, QDateEdit, QAbstractItemView, QMenu
)

from PIL import Image
from PIL.ImageQt import ImageQt # JAUNS IMPORTS
from PySide6.QtGui import QPainter # JAUNS IMPORTS

from pdf2image import convert_from_path
from PySide6.QtGui import QPixmap

from PySide6.QtCore import Qt, QSize, QSettings, QStandardPaths, QUrl, QPoint, QTimer, QThread, Signal, QEvent
from PySide6.QtGui import QAction, QIcon
from PySide6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog # JAUNS IMPORTS



# Import for WebEngine
from PySide6.QtWebEngineWidgets import QWebEngineView
from PySide6.QtWebEngineCore import QWebEngineUrlRequestInterceptor # For intercepting URL changes

# ReportLab imports (unchanged)
from reportlab.lib.pagesizes import A4, landscape, portrait, letter, legal, A3, A5
from reportlab.graphics import renderPDF
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.barcode import qr as rl_qr
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import mm, inch
from reportlab.platypus import Table, TableStyle, Paragraph, Spacer, SimpleDocTemplate, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus.flowables import Flowable

# python-docx imports (unchanged)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from PySide6.QtWebChannel import QWebChannel
from PySide6.QtCore import QObject, Slot
from PySide6.QtCore import Qt, QSize, QSettings, QStandardPaths, QUrl, QPoint, QTimer, QThread, Signal
import copy
from dataclasses import asdict
import platform

# ---------------------- Konstantes un direktoriju iestatījumi ----------------------
APP_DATA_DIR = QStandardPaths.writableLocation(QStandardPaths.AppDataLocation)
DOCUMENTS_DIR = QStandardPaths.writableLocation(QStandardPaths.DocumentsLocation)

# ==============================
# Iestatījumi (vienkārši JSON)
# ==============================

def _settings_path() -> str:
    return os.path.join(APP_DATA_DIR, "settings.json")

def load_settings() -> dict:
    try:
        p = _settings_path()
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception:
        pass
    return {}

def save_settings(data: dict):
    try:
        os.makedirs(APP_DATA_DIR, exist_ok=True)
        with open(_settings_path(), "w", encoding="utf-8") as f:
            json.dump(data or {}, f, ensure_ascii=False, indent=2)
    except Exception:
        pass



SETTINGS_DIR = os.path.join(APP_DATA_DIR, "AktaGenerators")
HISTORY_FILE = os.path.join(SETTINGS_DIR, "history.json")
ADDRESS_BOOK_FILE = os.path.join(SETTINGS_DIR, "address_book.json")
DEFAULT_SETTINGS_FILE = os.path.join(SETTINGS_DIR, "default_settings.json")
AKTA_NR_COUNTER_FILE = os.path.join(SETTINGS_DIR, "akta_nr_counter.json")  # JAUNS: stabilai secīgai akta numuru ģenerācijai
TEXT_BLOCKS_FILE = os.path.join(SETTINGS_DIR, "text_blocks.json") # JAUNA RINDAS

# Jaunas noklusējuma saglabāšanas mapes
DEFAULT_OUTPUT_DIR = os.path.join(DOCUMENTS_DIR, "AktaGenerators_Output")
PROJECT_SAVE_DIR = os.path.join(APP_DATA_DIR, "AktaGenerators_Projects")
# TEMPLATES_DIR = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates") # JAUNA RINDAS - Tagad dinamiski iestatīts AktaDati objektā

# ==============================
# Audit logs + Undo/Redo (GLOBAL)
# ==============================

class AuditLogger:
    """Vienkāršs, ātrs audit žurnāls (JSON Lines).
    Katrs ieraksts ir viena JSON rinda, lai failu var lasīt/filtrēt arī ārpus programmas.
    """

    def __init__(self, log_path: str):
        self.log_path = _coerce_path(log_path)
        os.makedirs(os.path.dirname(self.log_path), exist_ok=True)
        self._in_memory = []  # pēdējie N ieraksti UI

    def write(self, event: str, details: dict | None = None, user: str = ""):
        try:
            ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            row = {
                "ts": ts,
                "user": user or "",
                "event": event or "",
                "details": details or {},
            }
            with open(self.log_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(row, ensure_ascii=False) + "\n")
            self._in_memory.append(row)
            # limit memory
            if len(self._in_memory) > 500:
                self._in_memory = self._in_memory[-500:]
        except Exception:
            # audit nedrīkst nogāzt app
            pass

    def tail(self, n: int = 200):
        try:
            if self._in_memory:
                return self._in_memory[-n:]
            # Ja nav atmiņā, mēģinām nolasīt faila beigas (vienkārši)
            if not self.log_path or not os.path.exists(self.log_path):
                return []
            rows = []
            with open(self.log_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        rows.append(json.loads(line))
                    except Exception:
                        continue
            return rows[-n:]
        except Exception:
            return []


class UndoRedoManager:
    """Vienkāršs Undo/Redo ar 'snapshots' (adresēs grāmata + projekts).
    Nav QUndoStack, bet ir stabils, viegli uzturams un pietiekams 99% gadījumu.
    """
    def __init__(self, max_steps: int = 50):
        self.max_steps = max_steps
        self._undo = []
        self._redo = []

    def clear_redo(self):
        self._redo = []

    def push_undo(self, state: dict):
        self._undo.append(state)
        if len(self._undo) > self.max_steps:
            self._undo = self._undo[-self.max_steps:]
        self.clear_redo()

    def can_undo(self) -> bool:
        return len(self._undo) > 0

    def can_redo(self) -> bool:
        return len(self._redo) > 0

    def pop_undo(self) -> dict | None:
        if not self._undo:
            return None
        return self._undo.pop()

    def push_redo(self, state: dict):
        self._redo.append(state)
        if len(self._redo) > self.max_steps:
            self._redo = self._redo[-self.max_steps:]

    def pop_redo(self) -> dict | None:
        if not self._redo:
            return None
        return self._redo.pop()




# Pārliecināmies, ka direktoriji eksistē
os.makedirs(SETTINGS_DIR, exist_ok=True)
os.makedirs(DEFAULT_OUTPUT_DIR, exist_ok=True)
os.makedirs(PROJECT_SAVE_DIR, exist_ok=True)
# os.makedirs(TEMPLATES_DIR, exist_ok=True) # JAUNA RINDAS - Tagad dinamiski iestatīts AktaDati objektā




# ---------------------- UI tēma (modernāks izskats) ----------------------
def apply_modern_theme(app: QApplication, dark: bool = True):
    """Iestata modernu Fusion stilu + noapaļotus elementus (bez ārējām atkarībām)."""
    try:
        app.setStyle("Fusion")
    except Exception:
        pass

    # High-DPI (īpaši Windows)
    try:
        # High-DPI (PySide6 jaunākās versijās daļa atribūtu ir deprecated; atstājam droši)
        if hasattr(Qt, 'AA_EnableHighDpiScaling'):
            try:
                QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
            except Exception:
                pass
        if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
            try:
                QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
            except Exception:
                pass
    except Exception:
        pass

    # Palete (tumšā pēc noklusējuma)
    pal = app.palette()
    if dark:
        from PySide6.QtGui import QPalette, QColor
        pal.setColor(QPalette.Window, QColor(20, 22, 28))
        pal.setColor(QPalette.WindowText, QColor(230, 230, 230))
        pal.setColor(QPalette.Base, QColor(28, 30, 38))
        pal.setColor(QPalette.AlternateBase, QColor(34, 36, 46))
        pal.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
        pal.setColor(QPalette.ToolTipText, QColor(20, 22, 28))
        pal.setColor(QPalette.Text, QColor(230, 230, 230))
        pal.setColor(QPalette.Button, QColor(34, 36, 46))
        pal.setColor(QPalette.ButtonText, QColor(230, 230, 230))
        pal.setColor(QPalette.BrightText, QColor(255, 0, 0))
        pal.setColor(QPalette.Link, QColor(92, 170, 255))
        pal.setColor(QPalette.Highlight, QColor(92, 170, 255))
        pal.setColor(QPalette.HighlightedText, QColor(10, 10, 10))
        app.setPalette(pal)

    # Viegls “modern” QSS
    app.setStyleSheet("""
        QMainWindow { background: transparent; }
        QWidget { font-size: 10.5pt; }
        QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox, QDoubleSpinBox, QDateEdit, QComboBox {
            padding: 6px 8px;
            border: 1px solid rgba(255,255,255,0.12);
            border-radius: 10px;
        }
        QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus, QSpinBox:focus, QDoubleSpinBox:focus, QDateEdit:focus, QComboBox:focus {
            border: 1px solid rgba(92,170,255,0.9);
        }
        QPushButton, QToolButton {
            padding: 7px 10px;
            border-radius: 10px;
            border: 1px solid rgba(255,255,255,0.14);
        }
        QPushButton:hover, QToolButton:hover { border: 1px solid rgba(92,170,255,0.7); }
        QPushButton:pressed, QToolButton:pressed { padding-top: 8px; padding-bottom: 6px; }
        QTabBar::tab {
            padding: 8px 12px;
            margin: 2px;
            border-radius: 10px;
        }
        QTabWidget::pane { border: 0px; }
        QGroupBox {
            border: 1px solid rgba(255,255,255,0.10);
            border-radius: 14px;
            margin-top: 10px;
        }
        QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 6px; }
        QHeaderView::section {
            padding: 6px 8px;
            border: 0px;
            border-right: 1px solid rgba(255,255,255,0.10);
        }
        QTableWidget { border-radius: 14px; border: 1px solid rgba(255,255,255,0.10); }
        QScrollArea { border: 0px; }
        QMessageBox { font-size: 10.5pt; }
    """)

# ---------------------- Datu modeļi ----------------------
# (unchanged)

@dataclass
class Persona:
    nosaukums: str = ""
    reģ_nr: str = ""
    adrese: str = ""
    kontaktpersona: str = ""
    amats: str = ""
    pilnvaras_pamats: str = ""
    tālrunis: str = ""
    epasts: str = ""
    web_lapa: str = ""
    bankas_konts: str = ""
    juridiskais_statuss: str = ""


class MapBridge(QObject):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.map_click_callback = None

    @Slot(str, str)
    def handleMapClick(self, lat, lon):
        if self.map_click_callback:
            self.map_click_callback(lat, lon)

class TextBlockManager:
    def __init__(self):
        self.text_blocks = self._load_text_blocks()

    def _load_text_blocks(self):
        if os.path.exists(TEXT_BLOCKS_FILE):
            try:
                with open(TEXT_BLOCKS_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"Kļūda ielādējot teksta blokus: {e}")
                return {}
        return {}

    def _save_text_blocks(self):
        os.makedirs(SETTINGS_DIR, exist_ok=True)
        try:
            with open(TEXT_BLOCKS_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.text_blocks, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.critical(None, "Kļūda", f"Neizdevās saglabāt teksta blokus: {e}")

    def get_blocks_for_field(self, field_name):
        return self.text_blocks.get(field_name, {})

    def add_block(self, field_name, block_name, block_content):
        if field_name not in self.text_blocks:
            self.text_blocks[field_name] = {}
        self.text_blocks[field_name][block_name] = block_content
        self._save_text_blocks()

    def delete_block(self, field_name, block_name):
        if field_name in self.text_blocks and block_name in self.text_blocks[field_name]:
            del self.text_blocks[field_name][block_name]
            self._save_text_blocks()

    def get_block_content(self, field_name, block_name):
        return self.text_blocks.get(field_name, {}).get(block_name, "")


@dataclass
class Pozīcija:
    apraksts: str
    daudzums: Decimal
    vienība: str
    cena: Decimal
    seriālais_nr: str = ""
    garantija: str = ""
    piezīmes_pozīcijai: str = ""
    attēla_ceļš: str = ""

    @property
    def summa(self) -> Decimal:
        try:
            return (self.daudzums * self.cena).quantize(Decimal("0.01"))
        except Exception:
            return Decimal("0.00")

@dataclass
class Attēls:
    ceļš: str
    paraksts: str = ""


@dataclass
class AtsaucesDokuments:
    """Reāls atsauces dokuments, ko pievieno PDF beigās kā atvasinājumu."""
    ceļš: str
    nosaukums: str = ""


@dataclass
class AktaDati:
    akta_nr: str = ""
    datums: str = ""  # YYYY-MM-DD
    vieta: str = ""
    pasūtījuma_nr: str = ""
    pieņēmējs: Persona = field(default_factory=Persona)
    nodevējs: Persona = field(default_factory=Persona)
    pozīcijas: list = field(default_factory=list)  # list[Pozīcija]
    attēli: list = field(default_factory=list)
    piezīmes: str = ""
    iekļaut_pvn: bool = False
    pvn_likme: Decimal = Decimal("21.0")
    parakstu_rindas: bool = True
    logotipa_ceļš: str = ""
    fonts_ceļš: str = ""  # TTF/OTF
    paraksts_pieņēmējs_ceļš: str = ""
    paraksts_nodevējs_ceļš: str = ""
    līguma_nr: str = ""
    izpildes_termiņš: str = ""
    pieņemšanas_datums: str = ""
    nodošanas_datums: str = ""
    strīdu_risināšana: str = ""
    konfidencialitātes_klauzula: bool = False
    soda_nauda_procenti: Decimal = Decimal("0.0")
    piegādes_nosacījumi: str = ""
    apdrošināšana: bool = False
    apdrošināšana_teksts: str = ""
    papildu_nosacījumi: str = ""
    atsauces_dokumenti: str = ""
    atsauces_dokumenti_faili: list = field(default_factory=list)  # list[AtsaucesDokuments]
    akta_statuss: str = "Melnraksts"
    valūta: str = "EUR"
    elektroniskais_paraksts: bool = False
    radit_elektronisko_parakstu_tekstu: bool = False # JAUNS LAUKS
    qr_kods_enabled: bool = True  # QR kods apakšā pa kreisi
    qr_kods_ieklaut_pozicijas: bool = True
    qr_only_first_page: bool = False
    qr_verification_url_enabled: bool = False
    qr_verification_base_url: str = ""
    qr_kods_izmers_mm: Decimal = Decimal("25.0")

    # Jauni iestatījumu lauki
    pdf_page_size: str = "A4"
    pdf_page_orientation: str = "Portrets"
    pdf_margin_left: Decimal = Decimal("18")
    pdf_margin_right: Decimal = Decimal("18")
    pdf_margin_top: Decimal = Decimal("16")
    pdf_margin_bottom: Decimal = Decimal("16")
    pdf_font_size_head: int = 14
    pdf_font_size_normal: int = 10
    pdf_font_size_small: int = 9
    pdf_font_size_table: int = 9
    pdf_logo_width_mm: Decimal = Decimal("35")
    pdf_signature_width_mm: Decimal = Decimal("50")
    pdf_signature_height_mm: Decimal = Decimal("20")
    docx_image_width_inches: Decimal = Decimal("4")
    docx_signature_width_inches: Decimal = Decimal("1.5")
    table_col_widths: str = "10,40,18,18,20,20,25,25,25" # Komatiem atdalīti platumi mm
    auto_generate_akta_nr: bool = True
    default_execution_days: int = 5  # Noklusējuma dienu skaits izpildes termiņam (datums + N)
    default_currency: str = "EUR"
    default_unit: str = "gab."
    default_pvn_rate: Decimal = Decimal("21.0")
    poppler_path: str = ""
    # Jauni iestatījumi, lai sasniegtu "30+"
    header_text_color: str = "#000000" # Hex krāsa
    footer_text_color: str = "#000000"
    table_header_bg_color: str = "#E0E0E0"
    table_grid_color: str = "#CCCCCC"
    table_row_spacing: Decimal = Decimal("4") # mm
    line_spacing_multiplier: Decimal = Decimal("1.2") # Reizinātājs fonta izmēram
    show_page_numbers: bool = True
    show_generation_timestamp: bool = True
    currency_symbol_position: str = "after" # "before" or "after"
    date_format: str = "YYYY-MM-DD"
    signature_line_length_mm: Decimal = Decimal("60")
    signature_line_thickness_pt: Decimal = Decimal("0.5")
    add_cover_page: bool = False
    cover_page_title: str = "Pieņemšanas-Nodošanas Akts"
    cover_page_logo_width_mm: Decimal = Decimal("80")
    # Individuālais QR kods
    include_custom_qr_code: bool = False
    custom_qr_code_data: str = ""
    custom_qr_code_size_mm: Decimal = Decimal("20")
    custom_qr_code_position: str = "bottom_right"
    custom_qr_code_pos_x_mm: Decimal = Decimal("0")  # Custom X pozīcija mm
    custom_qr_code_pos_y_mm: Decimal = Decimal("0")  # Custom Y pozīcija mm
    custom_qr_code_color: str = "#000000"  # QR koda krāsa (Hex)
    # Automātiskais QR kods (akta ID)
    include_auto_qr_code: bool = False
    auto_qr_code_size_mm: Decimal = Decimal("20")
    auto_qr_code_position: str = "bottom_left"
    auto_qr_code_pos_x_mm: Decimal = Decimal("0")  # Custom X pozīcija mm
    auto_qr_code_pos_y_mm: Decimal = Decimal("0")  # Custom Y pozīcija mm
    auto_qr_code_color: str = "#000000"  # QR koda krāsa (Hex)

    add_watermark: bool = False
    watermark_text: str = "MELNRAKSTS"
    watermark_font_size: int = 72
    watermark_color: str = "#E0E0E0"
    watermark_rotation: int = 45
    enable_pdf_encryption: bool = False
    pdf_user_password: str = ""
    pdf_owner_password: str = ""
    allow_printing: bool = True
    allow_copying: bool = True
    allow_modifying: bool = False
    allow_annotating: bool = True
    # Papildu lauki, lai sasniegtu 30+
    default_country: str = "Latvija"
    default_city: str = "Rīga"
    show_contact_details_in_header: bool = False
    contact_details_header_font_size: int = 8
    item_image_width_mm: Decimal = Decimal("50") # Platums attēliem pie pozīcijām
    item_image_caption_font_size: int = 8
    show_item_notes_in_table: bool = True
    show_item_serial_number_in_table: bool = True
    show_item_warranty_in_table: bool = True
    show_item_photo_in_table: bool = True
    table_cell_padding_mm: Decimal = Decimal("2")
    table_header_font_style: str = "bold" # "bold", "italic", "normal"
    table_content_alignment: str = "left" # "left", "center", "right"
    signature_font_size: int = 9
    signature_spacing_mm: Decimal = Decimal("10") # Atstarpe starp paraksta rindu un vārdu
    # Vēl daži, lai pārsniegtu 30
    document_title_font_size: int = 18
    document_title_color: str = "#000000"
    section_heading_font_size: int = 12
    section_heading_color: str = "#000000"
    paragraph_line_spacing_multiplier: Decimal = Decimal("1.2")
    table_border_style: str = "solid" # "solid", "dashed", "none"
    table_border_thickness_pt: Decimal = Decimal("0.5")
    table_alternate_row_color: str = "" # Hex krāsa, piem. "#F0F0F0"
    # Papildu lauki, lai sasniegtu 30+
    show_total_sum_in_words: bool = False
    total_sum_in_words_language: str = "lv" # "lv", "en"
    # Papildu lauki, lai sasniegtu 30+
    default_vat_calculation_method: str = "exclusive" # "exclusive", "inclusive"
    show_vat_breakdown: bool = True
    # Papildu lauki, lai sasniegtu 30+
    enable_digital_signature_field: bool = False # PDF digitālā paraksta lauks
    digital_signature_field_name: str = "Paraksts"
    digital_signature_field_size_mm: Decimal = Decimal("40")
    digital_signature_field_position: str = "bottom_center"  # "bottom_left", "bottom_right", "top_left", "top_right", "bottom_center"
    template_password: str = ""  # Parole šablonam
    templates_dir: str = ""  # JAUNA RINDAS - Šablonu saglabāšanas direktorijs
    docx_template_path: str = ""  # (ja vēlies) DOCX šablons ar {{placeholders}}
    custom_columns: list = field(
    default_factory=list)  # Saraksts ar pielāgotajām kolonnām: [{'name': 'Kolonna1', 'data': []}]

    # --- JAUNS: Pozīciju tabulas kolonnu pārvaldība (UI + PDF) ---
    # poz_columns_config struktūra:
    # {
    #   "nr": {"title": "Nr.", "visible": True},
    #   "apraksts": {"title": "Apraksts", "visible": True},
    #   "daudzums": {"title": "Daudzums", "visible": True},
    #   "vieniba": {"title": "Vienība", "visible": True},
    #   "cena": {"title": "Cena", "visible": True},
    #   "summa": {"title": "Summa", "visible": True},
    #   "serial": {"title": "Seriālais Nr.", "visible": True},
    #   "warranty": {"title": "Garantija", "visible": True},
    #   "notes": {"title": "Piezīmes pozīcijai", "visible": True},
    #   "foto": {"title": "Foto", "visible": True}
    # }
    poz_columns_config: dict = field(default_factory=dict)

    # Vai rādīt cenu apkopojumu (kopsavilkumu) zem pozīciju tabulas (PDF)
    show_price_summary: bool = True
    # Saglabā lietotāja kolonnu secību Pozīciju tabulai (UI -> PDF)
    poz_columns_visual_order: list = field(default_factory=list)
    # Saglabā kolonnu platumus/izkārtojumu (QHeaderView.saveState) Pozīciju tabulai (UI)
    poz_header_state_b64: str = ""
    def kopējā_summma(self):
        s = sum((p.summa for p in self.pozīcijas), Decimal("0.00"))
        return s.quantize(Decimal("0.01"))

    def pvn_summa(self):
        if not self.iekļaut_pvn:
            return Decimal("0.00")
        pvn = self.kopējā_summma() * (self.pvn_likme / Decimal("100"))
        return pvn.quantize(Decimal("0.01"))

    def summa_ar_pvn(self):
        return (self.kopējā_summma() + self.pvn_summa()).quantize(Decimal("0.01"))


def _default_poz_columns_config() -> dict:
    """Noklusētā pozīciju kolonnu konfigurācija (UI + PDF)."""
    return {
        "nr": {"title": "Nr.", "visible": True},
        "apraksts": {"title": "Apraksts", "visible": True},
        "daudzums": {"title": "Daudzums", "visible": True},
        "vieniba": {"title": "Vienība", "visible": True},
        "cena": {"title": "Cena", "visible": True},
        "summa": {"title": "Summa", "visible": True},
        "serial": {"title": "Seriālais Nr.", "visible": True},
        "warranty": {"title": "Garantija", "visible": True},
        "notes": {"title": "Piezīmes pozīcijai", "visible": True},
        "foto": {"title": "Foto", "visible": True},
    }


def _merge_poz_columns_config(cfg: Optional[dict]) -> dict:
    """Droši apvieno lietotāja konfigurāciju ar noklusējuma vērtībām."""
    base = _default_poz_columns_config()
    out = {}
    cfg = cfg if isinstance(cfg, dict) else {}
    for k, v in base.items():
        cur = cfg.get(k, {}) if isinstance(cfg.get(k, {}), dict) else {}
        out[k] = {
            "title": str(cur.get("title", v.get("title", ""))),
            "visible": bool(cur.get("visible", v.get("visible", True))),
        }
    # Saglabājam arī nezināmus atslēgu ierakstus (ja nākotnē pievieno jaunas kolonnas)
    for k, cur in cfg.items():
        if k in out:
            continue
        if isinstance(cur, dict):
            out[k] = {"title": str(cur.get("title", k)), "visible": bool(cur.get("visible", True))}
    return out

# ---------------------- Palīgfunkcijas ----------------------
# (unchanged)
def to_decimal(val) -> Decimal:
    if isinstance(val, Decimal):
        return val
    try:
        s = str(val).replace(" ", "").replace(",", ".")
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return Decimal("0")

def formēt_naudu(d: Decimal) -> str:
    return f"{d:.2f}"

def drošs_faila_nosaukums(s: str) -> str:
    bad = '\\/:*?"<>|'
    for ch in bad:
        s = s.replace(ch, "_")
    return s.strip() or "akts"

def _coerce_path(p):
    """Normalizē ievadi uz faila ceļu (str) vai atgriež None.
    Vajadzīgs, jo dažās vietās signāli/ieraksti var nodot dict/tuple u.c.
    Pieņem arī dict ar tipiskām atslēgām: path, ceļš, file, filepath, json, pdf.
    """
    if p is None:
        return None
    # Qt signāli dažreiz padod bool (checked). To ignorējam.
    if isinstance(p, bool):
        return None
    # Ja padots tuple/list (piem. (path, filter))
    if isinstance(p, (list, tuple)) and p:
        return _coerce_path(p[0])
    # Ja padots dict ar ceļu
    if isinstance(p, dict):
        for k in ("path", "ceļš", "celsh", "file", "filepath", "json", "pdf", "value"):
            v = p.get(k)
            if isinstance(v, (str, bytes, os.PathLike)) and v:
                return os.fspath(v)
        return None
    # Parasts path-like
    if isinstance(p, (str, bytes, os.PathLike)):
        try:
            s = os.fspath(p)
        except Exception:
            return None
        return s if s else None
    # Cits tips (piem. int fd) šeit nav vajadzīgs
    return None




def _path_exists(p) -> bool:
    """Droša os.path.exists versija, kas pieņem arī dict/tuple u.c."""
    pp = _coerce_path(p)
    if not pp:
        return False
    try:
        return os.path.exists(pp)
    except Exception:
        return False

def _atomic_write_bytes(target_path: str, data: bytes):
    """Droši pārraksta failu (Windows lock-safe): raksta uz pagaidu failu un tad os.replace()."""
    try:
        target_path = os.path.abspath(target_path)
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
    except Exception:
        pass

    tmp_fd = None
    tmp_path = None
    try:
        # Pagaidu fails tajā pašā mapē (lai os.replace būtu atomisks arī Windows)
        dir_name = os.path.dirname(os.path.abspath(target_path)) or "."
        fd, tmp_path = tempfile.mkstemp(prefix=".tmp_", suffix=".pdf", dir=dir_name)
        tmp_fd = fd
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        tmp_fd = None
        os.replace(tmp_path, target_path)
        tmp_path = None
    finally:
        try:
            if tmp_fd is not None:
                os.close(tmp_fd)
        except Exception:
            pass
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

def _atomic_write_pdfwriter(target_path: str, writer_obj):
    """Droši saglabā PyPDF2 PdfWriter uz failu, neatsitoties pret file-lock."""
    buf = io.BytesIO()
    writer_obj.write(buf)
    _atomic_write_bytes(target_path, buf.getvalue())


def _prepare_unencrypted_pdf_for_render(pdf_path: str, password: str = "") -> tuple[str, Optional[str]]:
    """Ja PDF ir šifrēts, izveido atšifrētu pagaidu kopiju renderēšanai (poppler/pdf2image/Qt).
    Atgriež (render_path, temp_path_to_cleanup). Ja nav šifrēts, temp_path_to_cleanup būs None.

    Kāpēc tas vajadzīgs:
    - pdf2image/poppler bieži uzkārtina procesu, ja PDF ir šifrēts.
    - Priekšskatījums programmā nedrīkst prasīt paroli.
    """
    try:
        from PyPDF2 import PdfReader, PdfWriter
    except Exception:
        # Ja nav PyPDF2, vienkārši mēģinām renderēt oriģinālu (labāk nekā crash)
        return pdf_path, None

    try:
        with open(pdf_path, "rb") as rf:
            reader = PdfReader(rf)
            if not getattr(reader, "is_encrypted", False):
                return pdf_path, None

            # Mēģinām atšifrēt (PyPDF2: decrypt() atgriež 0/1/2 vai True/False atkarībā no versijas)
            try:
                ok = reader.decrypt(password or "")
            except Exception:
                ok = 0

            if not ok:
                # Nevarējām atšifrēt — neatgriežam šifrētu uz poppler (lai neuzkar),
                # bet metīsim izņēmumu, ko UI var parādīt kā kļūdu.
                raise RuntimeError("PDF ir šifrēts un paroli neizdevās pielietot priekšskatījumam/drukai.")

            writer = PdfWriter()
            for p in reader.pages:
                writer.add_page(p)

        # Saglabājam atšifrētu pagaidu PDF tajā pašā mapē (drošāk Windows)
        dir_name = os.path.dirname(os.path.abspath(pdf_path)) or "."
        fd, tmp_path = tempfile.mkstemp(prefix=".tmp_decrypted_", suffix=".pdf", dir=dir_name)
        os.close(fd)
        _atomic_write_pdfwriter(tmp_path, writer)
        return tmp_path, tmp_path

    except Exception:
        # Ja kaut kas noiet greizi, labāk lai caller redz kļūdu (nevis hang).
        raise


def reģistrēt_fontu(font_ceļš: str, vārds: str = "DokFont") -> str:
    if not font_ceļš or not os.path.exists(font_ceļš):
        try:
            if sys.platform == "win32":
                system_font_path = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "Arial.ttf")
                if os.path.exists(system_font_path):
                    pdfmetrics.registerFont(TTFont(vārds, system_font_path))
                    return vārds
            return "Helvetica"
        except Exception:
            return "Helvetica"
    try:
        pdfmetrics.registerFont(TTFont(vārds, font_ceļš))
        return vārds
    except Exception:
        return "Helvetica"

def render_pdf_to_image(pdf_path: str, poppler_path: str = None, password: str = "") -> QPixmap:
    """
    Renderē PDF faila pirmo lapu kā QPixmap.
    Nepieciešama Poppler instalācija.
    :param pdf_path: Ceļš uz PDF failu.
    :param poppler_path: (Tikai Windows) Ceļš uz Poppler bin direktoriju.
    :param password: Ja PDF ir šifrēts, lietotāja parole atšifrēšanai priekšskatījumam.
    :return: QPixmap objekts ar PDF lapas attēlu.
    """
    try:
        render_path, tmp_cleanup = _prepare_unencrypted_pdf_for_render(pdf_path, password=password)
        try:
            images = convert_from_path(render_path, first_page=1, last_page=1, poppler_path=poppler_path)
        finally:
            try:
                if tmp_cleanup and os.path.exists(tmp_cleanup):
                    os.remove(tmp_cleanup)
            except Exception:
                pass

        if images:
            from io import BytesIO
            img_byte_arr = BytesIO()
            images[0].save(img_byte_arr, format="PNG")
            img_byte_arr.seek(0)

            pixmap = QPixmap()
            pixmap.loadFromData(img_byte_arr.getvalue(), "PNG")
            return pixmap

        return QPixmap()

    except Exception as e:
        print(f"Kļūda renderējot PDF uz attēlu: {e}")
        QMessageBox.warning(
            None,
            "PDF renderēšanas kļūda",
            "Neizdevās renderēt PDF priekšskatījumu. "
            "Ja PDF ir šifrēts, pārliecinieties, ka parole ir pareiza. "
            f"\nKļūda: {e}"
        )
        return QPixmap()

# Lapu numerācija PDF dokumentam


# Lapu numerācija + dekorācijas PDF dokumentam
class DecoratedCanvas(canvas.Canvas):
    """Canvas, kas pievieno lapu numerāciju, ūdenszīmi un QR kodus (ja ieslēgts)."""

    def __init__(self, *args, **kwargs):
        self.pages = []
        self.akta_dati: AktaDati = kwargs.pop("akta_dati", None)
        self.show_page_numbers = kwargs.pop("show_page_numbers", True)
        super().__init__(*args, **kwargs)

        # Sagatavojam QR kodu attēlus vienreiz (ja vajag)
        self._qr_images = []
        try:
            if self.akta_dati:
                qr_items = []
                if self.akta_dati.include_custom_qr_code and self.akta_dati.custom_qr_code_data:
                    qr_items.append(("custom", self.akta_dati.custom_qr_code_data,
                                     float(self.akta_dati.custom_qr_code_size_mm),
                                     self.akta_dati.custom_qr_code_position,
                                     float(self.akta_dati.custom_qr_code_pos_x_mm),
                                     float(self.akta_dati.custom_qr_code_pos_y_mm)))
                if self.akta_dati.include_auto_qr_code and self.akta_dati.akta_nr:
                    qr_items.append(("auto", self.akta_dati.akta_nr,
                                     float(self.akta_dati.auto_qr_code_size_mm),
                                     self.akta_dati.auto_qr_code_position,
                                     float(self.akta_dati.auto_qr_code_pos_x_mm),
                                     float(self.akta_dati.auto_qr_code_pos_y_mm)))

                if qr_items:
                    import qrcode
                    from reportlab.lib.utils import ImageReader
                    for _, data, size_mm, pos, x_mm, y_mm in qr_items:
                        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L,
                                           box_size=10, border=2)
                        qr.add_data(data)
                        qr.make(fit=True)
                        img = qr.make_image(fill_color="black", back_color="white")
                        # ImageReader strādā ar PIL Image
                        self._qr_images.append({
                            "reader": ImageReader(img),
                            "size_pt": size_mm * mm,
                            "pos": pos,
                            "x_pt": x_mm * mm,
                            "y_pt": y_mm * mm,
                        })
        except Exception as e:
            # Nekrītam ārā, ja qrcode nav uzinstalēts vai rodas kļūda
            print(f"QR sagatavošanas kļūda: {e}")
            self._qr_images = []

    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self.pages)
        for page in self.pages:
            self.__dict__.update(page)

            # Dekorācijas uz katras lapas
            self._draw_staple_mark()
            self._draw_watermark()
            self._draw_qr_codes()

            if self.show_page_numbers:
                self._draw_page_number(page_count)

            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)


    def _draw_staple_mark(self):
        """Uzzīmē īsu vertikālu līniju augšējā kreisajā stūrī (skavotāja atzīme) uz katras lapas."""
        try:
            w, h = self._pagesize
            self.saveState()
            # Diskrēta, drukai draudzīga krāsa
            self.setStrokeColor(colors.HexColor("#94A3B8"))  # slate-400
            self.setLineWidth(0.7)
            # Pozīcija: nedaudz no kreisās malas, pie augšas
            x = 6 * mm
            y1 = h - 12 * mm
            y2 = h - 32 * mm
            self.line(x, y1, x, y2)
            self.restoreState()
        except Exception:
            try:
                self.restoreState()
            except Exception:
                pass

    def _draw_page_number(self, page_count: int):
        # Augšējais labais stūris
        self.setFont("Helvetica", 9)
        w, h = self._pagesize
        self.drawRightString(w - 18 * mm, h - 10 * mm, f"Lapa {self._pageNumber} no {page_count}")

    def _draw_watermark(self):
        if not self.akta_dati or not getattr(self.akta_dati, "add_watermark", False):
            return
        text = getattr(self.akta_dati, "watermark_text", "") or ""
        if not text.strip():
            return

        try:
            from reportlab.lib.colors import HexColor
            self.saveState()
            self.setFillColor(HexColor(getattr(self.akta_dati, "watermark_color", "#E0E0E0")))
            self.setFont("Helvetica-Bold", float(getattr(self.akta_dati, "watermark_font_size", 72)))
            w, h = self._pagesize
            self.translate(w / 2, h / 2)
            self.rotate(float(getattr(self.akta_dati, "watermark_rotation", 45)))
            self.drawCentredString(0, 0, text)
            self.restoreState()
        except Exception as e:
            print(f"Ūdenszīmes kļūda: {e}")

    def _draw_qr_codes(self):
        if not self._qr_images:
            return

        try:
            w, h = self._pagesize
            left = float(getattr(self.akta_dati, "pdf_margin_left", 18)) * mm if self.akta_dati else 18 * mm
            right = float(getattr(self.akta_dati, "pdf_margin_right", 18)) * mm if self.akta_dati else 18 * mm
            top = float(getattr(self.akta_dati, "pdf_margin_top", 16)) * mm if self.akta_dati else 16 * mm
            bottom = float(getattr(self.akta_dati, "pdf_margin_bottom", 16)) * mm if self.akta_dati else 16 * mm

            for q in self._qr_images:
                size = q["size_pt"]
                pos = q["pos"]
                if pos == "bottom_left":
                    x, y = left, bottom
                elif pos == "bottom_right":
                    x, y = w - right - size, bottom
                elif pos == "top_left":
                    x, y = left, h - top - size
                elif pos == "top_right":
                    x, y = w - right - size, h - top - size
                elif pos == "custom":
                    x, y = q["x_pt"], q["y_pt"]
                else:
                    # fallback
                    x, y = w - right - size, bottom

                self.drawImage(q["reader"], x, y, width=size, height=size, mask='auto')
        except Exception as e:
            print(f"QR zīmēšanas kļūda: {e}")


# ---------------------- Atsauces dokumentu (pielikumu) apstrāde ----------------------
def _find_soffice_exe() -> Optional[str]:
    """Atrod LibreOffice/soffice izpildāmo failu.
    Atgriež pilnu ceļu vai None, ja nav atrasts.
    """
    import shutil

    # 1) mēģinam no PATH
    p = shutil.which("soffice") or shutil.which("libreoffice")
    if p:
        return p

    # 2) tipiskie ceļi Windows
    if sys.platform == "win32":
        candidates = [
            os.path.join(os.environ.get("ProgramFiles", ""), "LibreOffice", "program", "soffice.exe"),
            os.path.join(os.environ.get("ProgramFiles(x86)", ""), "LibreOffice", "program", "soffice.exe"),
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for c in candidates:
            if c and os.path.exists(c):
                return c

    return None

def _convert_attachment_to_pdf(input_path: str, out_dir: str) -> Optional[str]:
    """Konvertē pielikumu uz PDF, lai to varētu pievienot akta beigās.

    Atbalsts:
      - PDF: atgriež oriģinālo ceļu
      - DOC/DOCX: LibreOffice (soffice) -> PDF (writer_pdf_Export)
      - XLS/XLSX/ODS: LibreOffice -> PDF (calc_pdf_Export)
      - PPT/PPTX: LibreOffice -> PDF (impress_pdf_Export)

    Atgriež PDF ceļu vai None, ja neizdevās.
    """
    input_path = _coerce_path(input_path)
    if not input_path:
        return None

    input_path = os.path.abspath(input_path)
    if not os.path.exists(input_path):
        return None

    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        return input_path

    supported = {".docx", ".doc", ".xlsx", ".xls", ".ods", ".odt", ".pptx", ".ppt"}
    if ext not in supported:
        return None

    soffice = _find_soffice_exe()
    if not soffice:
        # Nav LibreOffice -> nevar padarīt aplūkojamu PDF beigās
        return None

    # izvēlamies pareizo LO filtru (stabilāk priekš XLSX)
    if ext in {".xls", ".xlsx", ".ods"}:
        convert_to = "pdf:calc_pdf_Export"
    elif ext in {".ppt", ".pptx"}:
        convert_to = "pdf:impress_pdf_Export"
    else:
        convert_to = "pdf:writer_pdf_Export"

    try:
        os.makedirs(out_dir, exist_ok=True)

        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--norestore",
            "--convert-to", convert_to,
            "--outdir", out_dir,
            input_path,
        ]

        res = subprocess.run(cmd, capture_output=True, text=True, timeout=30, **({"creationflags": subprocess.CREATE_NO_WINDOW} if os.name == "nt" and hasattr(subprocess, "CREATE_NO_WINDOW") else {}))

        if res.returncode != 0:
            # noder debuggam (konsolē), bet UI netraucē
            print(f"Konvertēšanas kļūda ({input_path}): {res.stderr or res.stdout}")
            return None

        # meklējam izveidoto PDF tieši šajā mapē
        pdfs = [os.path.join(out_dir, f) for f in os.listdir(out_dir) if f.lower().endswith('.pdf')]
        if not pdfs:
            return None

        # Ja ir vairāki, ņemam jaunāko
        pdfs.sort(key=lambda fp: os.path.getmtime(fp), reverse=True)
        return pdfs[0]

    except subprocess.TimeoutExpired:
        # LibreOffice iestrēga (piem., liels DOCX/XLSX vai dialogi).
        return None
    except FileNotFoundError:
        # izpildāmais fails nav atrasts
        return None
    except Exception as e:
        print(f"Konvertēšanas izņēmums ({input_path}): {e}")
        return None


def _make_annex_title_pdf(title: str, out_path: str, pagesize=A4, font_name: str = "Helvetica"):
    """Izveido vienas lapas PDF informācijas lapu (ja pielikumu nevar konvertēt).
    Teksts ir mazs, augšējā kreisajā stūrī, ar pareizām garumzīmēm (izmanto font_name).
    """
    # nodrošinām, ka mape eksistē (citādi Windows met [Errno 2])
    try:
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    except Exception:
        pass

    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.lib.utils import simpleSplit

    c = canvas.Canvas(out_path, pagesize=pagesize)
    w, h = pagesize

    left = 18 * mm
    top_y = h - 12 * mm
    max_w = w - 2 * left

    # Bold variants only if Helvetica, citādi izmantojam to pašu fontu (drošāk diakritikām)
    bold_font = "Helvetica-Bold" if (font_name or "") == "Helvetica" else (font_name or "Helvetica")
    normal_font = font_name or "Helvetica"

    def fit_one_line(txt: str, fnt: str, size: float) -> str:
        if stringWidth(txt, fnt, size) <= max_w:
            return txt
        ell = "…"
        # brutāla, bet stabila saīsināšana
        for cut in range(len(txt), 0, -1):
            cand = txt[:cut].rstrip() + ell
            if stringWidth(cand, fnt, size) <= max_w:
                return cand
        return ell

    title_fit = fit_one_line(title, bold_font, 10)

    c.setFillColor(colors.HexColor("#0F172A"))
    c.setFont(bold_font, 10)
    c.drawString(left, top_y, title_fit)

    c.setStrokeColor(colors.HexColor("#CBD5E1"))
    c.setLineWidth(0.8)
    c.line(left, top_y - 4, w - left, top_y - 4)

    info = "Pielikums pie pieņemšanas–nodošanas akta. Dokuments pievienots automātiski."
    c.setFont(normal_font, 9)
    c.setFillColor(colors.HexColor("#334155"))
    lines = simpleSplit(info, normal_font, 9, max_w)
    y = top_y - 14
    for ln in lines[:5]:
        c.drawString(left, y, ln)
        y -= 11

    c.showPage()
    c.save()


def _overlay_text_on_pdf_page(page, text: str, font_name: str, font_size: float, x_pt: float, y_pt: float, bold: bool = False):
    """Uzvelk tekstu uz dotās PDF lapas (PyPDF2 PageObject), izmantojot ReportLab overlay.

    Svarīgi:
    - Diakritikām jābūt redzamām -> izmanto font_name (TTF)
    - Teksts NETIEK izvilkts ārpus lapas -> ja par garu, saīsina ar "…"
    """
    try:
        from PyPDF2 import PdfReader
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.pdfbase.pdfmetrics import stringWidth
    except Exception:
        return page

    try:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        left = float(x_pt)
        max_w = w - left - (18 * mm)  # labā mala

        # izvēlamies fontu (bold tikai Helvetica gadījumā)
        fnt = font_name if font_name else "Helvetica"
        if bold and fnt == "Helvetica":
            fnt = "Helvetica-Bold"

        def fit_one_line(txt: str) -> str:
            if stringWidth(txt, fnt, float(font_size)) <= max_w:
                return txt
            ell = "…"
            for cut in range(len(txt), 0, -1):
                cand = txt[:cut].rstrip() + ell
                if stringWidth(cand, fnt, float(font_size)) <= max_w:
                    return cand
            return ell

        safe_text = fit_one_line(text)

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(w, h))
        c.setFont(fnt, float(font_size))
        c.setFillColor(colors.HexColor("#0F172A"))
        # drošības dēļ y ierobežojam
        y = min(max(float(y_pt), 0), h - 2)
        c.drawString(left, y, safe_text)
        c.save()
        buf.seek(0)
        overlay_page = PdfReader(buf).pages[0]
        page.merge_page(overlay_page)
    except Exception:
        pass
    return page


def _apply_global_page_numbers_to_pdf(pdf_path: str, akta_dati: AktaDati, font_name: str):
    """Pievieno lapu numerāciju visam PDF (arī pielikumiem), lai kopējais skaits ir pareizs."""
    if not getattr(akta_dati, "show_page_numbers", True):
        return

    try:
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas as rl_canvas
    except Exception as e:
        print(f"Nevar uzlikt lapu numerāciju (trūkst PyPDF2/ReportLab): {e}")
        return

    try:
        reader = PdfReader(pdf_path)
        total = len(reader.pages)
        writer = PdfWriter()

        for idx, page in enumerate(reader.pages, start=1):
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)

            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=(w, h))
            c.setFillColor(colors.HexColor("#0F172A"))
            c.setFont(font_name if font_name else "Helvetica", 9)
            # Augšējais labais stūris (tāpat kā DecoratedCanvas)
            c.drawRightString(w - 18 * mm, h - 10 * mm, f"Lapa {idx} no {total}")
            c.save()
            buf.seek(0)

            overlay_page = PdfReader(buf).pages[0]
            page.merge_page(overlay_page)
            writer.add_page(page)

        _atomic_write_pdfwriter(pdf_path, writer)
    except Exception as e:
        print(f"Lapu numerācijas kļūda: {e}")



def _apply_stapler_mark_to_pdf(pdf_path: str, akta_dati: AktaDati, color_hex: str = "#94A3B8"):
    """Pievieno 'skavotāja līniju' katrai PDF lapai (arī pielikumiem), pēcapstrādē ar PyPDF2."""
    try:
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas as rl_canvas
    except Exception as e:
        print(f"Nevar uzlikt skavotāja līniju (trūkst PyPDF2/ReportLab): {e}")
        return

    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        for page_index, page in enumerate(reader.pages):
            if getattr(akta_dati, 'qr_kods_tikai_pirma_lapa', False) and page_index != 0:
                writer.add_page(page)
                continue

            w = float(page.mediabox.width)
            h = float(page.mediabox.height)

            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=(w, h))
            try:
                c.setStrokeColor(colors.HexColor(color_hex))
            except Exception:
                c.setStrokeColor(colors.HexColor("#94A3B8"))
            c.setLineWidth(0.7)

            x = 6 * mm
            y1 = h - 12 * mm
            y2 = h - 32 * mm
            c.line(x, y1, x, y2)

            c.save()
            buf.seek(0)

            overlay = PdfReader(buf).pages[0]
            page.merge_page(overlay)
            writer.add_page(page)

        _atomic_write_pdfwriter(pdf_path, writer)
    except Exception as e:
        print(f"Skavotāja līnijas kļūda: {e}")




def _build_qr_payload(akta_dati: AktaDati) -> str:
    """QR saturs.
    - URL režīms: QR satur verifikācijas URL ar parametriem (a,d,pc,h).
    - Pretējā gadījumā: kompakts JSON ar akta datiem un (saīsinātām) pozīcijām.
    """
    try:
        if getattr(akta_dati, "qr_kods_url_mode", False) and (getattr(akta_dati, "qr_kods_url", "") or "").strip():
            base = akta_dati.qr_kods_url.strip()
            poz_count = len(getattr(akta_dati, "pozīcijas", []) or [])
            full = json.dumps(
                {"akta_nr": akta_dati.akta_nr, "datums": akta_dati.datums, "poz_count": poz_count},
                ensure_ascii=False,
                separators=(",", ":"),
            )
            h = hashlib.sha256(full.encode("utf-8")).hexdigest()[:16]
            from urllib.parse import urlencode
            qs = urlencode({"a": akta_dati.akta_nr, "d": akta_dati.datums, "pc": poz_count, "h": h})
            return base + ("&" if "?" in base else "?") + qs

        payload = {"v": 1, "akta_nr": akta_dati.akta_nr, "datums": akta_dati.datums, "vieta": getattr(akta_dati, "vieta", "")}

        if getattr(akta_dati, "qr_kods_ieklaut_pozicijas", True):
            items = []
            for i, p in enumerate(getattr(akta_dati, "pozīcijas", []) or []):
                if i >= 20:
                    break
                try:
                    name = (getattr(p, "nosaukums", "") or "").strip()
                    if len(name) > 40:
                        name = name[:40] + "…"
                    items.append({"n": name, "q": str(getattr(p, "daudzums", "")), "u": (getattr(p, "vienība", "") or "").strip(),
                                  "sn": (getattr(p, "sērijas_nr", "") or "").strip()[:30]})
                except Exception:
                    continue
            payload["poz"] = items
            payload["poz_count"] = len(getattr(akta_dati, "pozīcijas", []) or [])

        s = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
        if len(s) > 1200 and "poz" in payload:
            payload["poz"] = payload["poz"][:8]
            s = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))

        if len(s) > 1200:
            poz = getattr(akta_dati, "pozīcijas", []) or []
            full = json.dumps(
                {"akta_nr": akta_dati.akta_nr, "datums": akta_dati.datums,
                 "pozīcijas": [asdict(p) if hasattr(p, "__dataclass_fields__") else str(p) for p in poz]},
                ensure_ascii=False,
                separators=(",", ":"),
            )
            h = hashlib.sha256(full.encode("utf-8")).hexdigest()[:16]
            s = json.dumps({"v": 1, "akta_nr": akta_dati.akta_nr, "datums": akta_dati.datums, "poz_count": len(poz), "poz_hash": h},
                           ensure_ascii=False, separators=(",", ":"))
        return s
    except Exception:
        return ""


def _apply_qr_to_pdf(pdf_path: str, akta_dati: AktaDati):
    """Pievieno QR kodu PDF apakšējā kreisajā stūrī.
    - QR tikai pirmajā lapā, ja qr_kods_tikai_pirma_lapa=True
    - URL režīmā uz QR uzliek klikšķināmu linku
    """
    try:
        if not getattr(akta_dati, "qr_kods_enabled", True):
            return
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas as rl_canvas
    except Exception:
        return

    payload = _build_qr_payload(akta_dati)
    if not payload:
        return

    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        size_mm = float(getattr(akta_dati, "qr_kods_izmers_mm", Decimal("18.0")))
        first_only = bool(getattr(akta_dati, "qr_kods_tikai_pirma_lapa", True))

        for page_index, page in enumerate(reader.pages):
            if first_only and page_index != 0:
                writer.add_page(page)
                continue

            w = float(page.mediabox.width)
            h = float(page.mediabox.height)

            try:
                ml = float(getattr(akta_dati, "pdf_margin_left", Decimal("15.0")))
                mb = float(getattr(akta_dati, "pdf_margin_bottom", Decimal("15.0")))
            except Exception:
                ml, mb = 15.0, 15.0

            max_size = max(10.0, min(size_mm, ml - 2.0, mb - 2.0))
            qr_size = max_size * mm

            x = 2 * mm
            y = 2 * mm

            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=(w, h))

            try:
                if getattr(akta_dati, "qr_kods_url_mode", False) and (getattr(akta_dati, "qr_kods_url", "") or "").strip():
                    c.linkURL(payload, (x, y, x + qr_size, y + qr_size), relative=0)
            except Exception:
                pass

            try:
                widget = rl_qr.QrCodeWidget(payload)
                bounds = widget.getBounds()
                bw = bounds[2] - bounds[0]
                bh = bounds[3] - bounds[1]
                scale = qr_size / max(bw, bh)
                d = Drawing(qr_size, qr_size, transform=[scale, 0, 0, scale, 0, 0])
                d.add(widget)
                renderPDF.draw(d, c, x, y)
            except Exception:
                pass

            c.save()
            buf.seek(0)
            overlay = PdfReader(buf).pages[0]
            page.merge_page(overlay)
            writer.add_page(page)

        _atomic_write_pdfwriter(pdf_path, writer)
    except Exception:
        pass


def _append_reference_docs_to_pdf(main_pdf_path: str, akta_dati: AktaDati, pagesize=A4, font_name: str = "Helvetica") -> str:
    """Pievieno atsauces dokumentus PDF beigās kā Atvasinājums 1, 2, 3...

    Prasības:
    - "Atvasinājums X: ..." ir mazs teksts augšējā kreisajā stūrī TAJĀ PAŠĀ lapā, kur sākas pielikums
    - Teksts nedrīkst pārsniegt lapas robežas (saīsinām ar "…")
    - Ja DOCX/XLSX nevar konvertēt uz PDF, pievienojam informācijas lapu ar korektām garumzīmēm
    """
    refs = getattr(akta_dati, "atsauces_dokumenti_faili", []) or []
    if not refs:
        return main_pdf_path

    try:
        from PyPDF2 import PdfReader, PdfWriter
    except Exception as e:
        print(f"PyPDF2 nav pieejams pielikumiem: {e}")
        return main_pdf_path

    tmp_root = tempfile.mkdtemp(prefix="akta_refs_")
    try:
        writer = PdfWriter()

        # Pamata PDF
        reader_main = PdfReader(main_pdf_path)
        for p in reader_main.pages:
            writer.add_page(p)

        annex_no = 0
        for ref in refs:
            try:
                if isinstance(ref, dict):
                    ref_path = ref.get("ceļš", "")
                    ref_name = ref.get("nosaukums", "") or os.path.basename(ref_path)
                else:
                    ref_path = getattr(ref, "ceļš", "")
                    ref_name = getattr(ref, "nosaukums", "") or os.path.basename(ref_path)

                if not ref_path or not os.path.exists(ref_path):
                    continue

                annex_no += 1

                # Konvertējam katru pielikumu savā apakšmapē, lai nekad nesajauktu PDF nosaukumus
                tmp_dir = os.path.join(tmp_root, f"conv_{annex_no}")
                os.makedirs(tmp_dir, exist_ok=True)
                converted = _convert_attachment_to_pdf(ref_path, tmp_dir)

                if not converted:
                    info_pdf = os.path.join(tmp_dir, f"Atvasinajums_{annex_no}_info.pdf")
                    _make_annex_title_pdf(
                        f"Atvasinājums {annex_no}: {ref_name} (neizdevās konvertēt uz PDF)",
                        info_pdf,
                        pagesize=pagesize,
                        font_name=font_name if font_name else "Helvetica"
                    )
                    r = PdfReader(info_pdf)
                    for p in r.pages:
                        writer.add_page(p)
                    continue

                r = PdfReader(converted)
                if not r.pages:
                    continue

                # Uz pirmās pielikuma lapas uzliekam label (mazs, top-left, saīsināts ja vajag)
                first = r.pages[0]
                label = f"Atvasinājums {annex_no}: {ref_name}"
                _overlay_text_on_pdf_page(
                    first,
                    label,
                    font_name=font_name if font_name else "Helvetica",
                    font_size=9,
                    x_pt=18 * mm,
                    y_pt=float(first.mediabox.height) - 12 * mm,
                    bold=False
                )
                writer.add_page(first)

                for p in r.pages[1:]:
                    writer.add_page(p)

            except Exception as e:
                print(f"Pielikuma pievienošanas kļūda: {e}")

        _atomic_write_pdfwriter(main_pdf_path, writer)

    finally:
        try:
            shutil.rmtree(tmp_root, ignore_errors=True)
        except Exception:
            pass

    return main_pdf_path

# ---------------------- PDF ģenerēšana ----------------------

def ģenerēt_pdf(akta_dati: AktaDati, pdf_ceļš: str = None, include_reference_docs: bool = True, encrypt_pdf: bool = True):
    # --- FIX v46: normalize pdf_ceļš if dict leaked from state ---
    if isinstance(pdf_ceļš, dict):
        pdf_ceļš = pdf_ceļš.get('path') or ''
    font_name = reģistrēt_fontu(akta_dati.fonts_ceļš)

    styles = getSampleStyleSheet()
    # Ensure all Decimal values are converted to float when used with ReportLab's float-based units or font sizes
    # Uzlaboti stili ar jaunajiem iestatījumiem
    styles.add(ParagraphStyle(name='LatvHead', fontName=font_name, fontSize=float(akta_dati.pdf_font_size_head), leading=float(akta_dati.pdf_font_size_head) * float(akta_dati.line_spacing_multiplier), spaceAfter=8, textColor=colors.HexColor(akta_dati.header_text_color)))
    styles.add(ParagraphStyle(name='Latv', fontName=font_name, fontSize=float(akta_dati.pdf_font_size_normal), leading=float(akta_dati.pdf_font_size_normal) * float(akta_dati.line_spacing_multiplier), textColor=colors.HexColor(akta_dati.footer_text_color)))
    styles.add(ParagraphStyle(name='LatvSmall', fontName=font_name, fontSize=float(akta_dati.pdf_font_size_small), leading=float(akta_dati.pdf_font_size_small) * float(akta_dati.line_spacing_multiplier), textColor=colors.HexColor(akta_dati.footer_text_color)))
    styles.add(ParagraphStyle(name='LatvTableContent', fontName=font_name, fontSize=float(akta_dati.pdf_font_size_table), leading=float(akta_dati.pdf_font_size_table) * float(akta_dati.line_spacing_multiplier), wordWrap='LTR', splitLongWords=0, alignment={'left': 0, 'center': 1, 'right': 2}.get(akta_dati.table_content_alignment, 0)))
    styles.add(ParagraphStyle(name='LatvElectronicSignature', fontName=font_name, fontSize=float(akta_dati.pdf_font_size_normal), leading=float(akta_dati.pdf_font_size_normal) * float(akta_dati.line_spacing_multiplier), alignment=1, textColor=colors.HexColor(akta_dati.header_text_color)))
    styles.add(ParagraphStyle(name='DocTitle', fontName=font_name, fontSize=float(akta_dati.document_title_font_size), alignment=1, textColor=colors.HexColor(akta_dati.document_title_color)))
    styles.add(ParagraphStyle(name='SectionHeading', fontName=font_name, fontSize=float(akta_dati.section_heading_font_size), leading=float(akta_dati.section_heading_font_size) * float(akta_dati.paragraph_line_spacing_multiplier), textColor=colors.HexColor(akta_dati.section_heading_color)))

    # ---------------------- Premium juridiskais dizains (uzlabojumi) ----------------------
    # Bold fonts: ja tiek lietots Helvetica, izmantojam Helvetica-Bold, citādi atstājam fontu (ja nav bold varianta).
    bold_font_name = "Helvetica-Bold" if font_name == "Helvetica" else font_name

    # Papildu stili skaidrai hierarhijai (juridisks + moderns)
    styles.add(ParagraphStyle(
        name='FieldLabel',
        fontName=bold_font_name,
        fontSize=float(akta_dati.pdf_font_size_normal),
        leading=float(akta_dati.pdf_font_size_normal) * float(akta_dati.line_spacing_multiplier),
        textColor=colors.HexColor("#0F172A")
    ))
    styles.add(ParagraphStyle(
        name='FieldValue',
        fontName=font_name,
        fontSize=float(akta_dati.pdf_font_size_normal),
        leading=float(akta_dati.pdf_font_size_normal) * float(akta_dati.line_spacing_multiplier),
        textColor=colors.HexColor("#0F172A")
    ))
    styles.add(ParagraphStyle(
        name='SectionBar',
        fontName=bold_font_name,
        fontSize=float(max(11, int(akta_dati.section_heading_font_size))),
        leading=float(max(11, int(akta_dati.section_heading_font_size))) * float(akta_dati.paragraph_line_spacing_multiplier),
        textColor=colors.HexColor("#0F172A"),
        backColor=colors.HexColor("#F1F5F9"),
        borderPadding=6,
        spaceBefore=10,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='LegalSectionTitle',
        fontName=bold_font_name,
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=10,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name='LegalBody',
        fontName=font_name,
        fontSize=float(akta_dati.pdf_font_size_normal),
        leading=float(akta_dati.pdf_font_size_normal) * float(akta_dati.line_spacing_multiplier),
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='TableHeader',
        fontName=bold_font_name,
        fontSize=float(akta_dati.pdf_font_size_table),
        leading=float(akta_dati.pdf_font_size_table) * float(akta_dati.line_spacing_multiplier),
        textColor=colors.HexColor("#0F172A"),
        alignment=1,  # CENTER
    ))
    # Ja nav iestatīts alternējošās rindas tonis, piešķiram klusu "enterprise" noklusējumu
    if not getattr(akta_dati, "table_alternate_row_color", ""):
        akta_dati.table_alternate_row_color = "#F8FAFC"

    page_size_map = {
        "A4": A4, "Letter": letter, "Legal": legal, "A3": A3, "A5": A5
    }
    base_page_size = page_size_map.get(akta_dati.pdf_page_size, A4)

    if akta_dati.pdf_page_orientation == "Ainava":
        pagesize = landscape(base_page_size)
    else:
        pagesize = portrait(base_page_size)

    if pdf_ceļš is None:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf_ceļš = temp_file.name
        temp_file.close()

    doc = SimpleDocTemplate(
        pdf_ceļš,
        pagesize=pagesize,
        leftMargin=float(akta_dati.pdf_margin_left) * mm,
        rightMargin=float(akta_dati.pdf_margin_right) * mm,
        topMargin=float(akta_dati.pdf_margin_top) * mm,
        bottomMargin=float(akta_dati.pdf_margin_bottom) * mm,
        title="Pieņemšanas–Nodošanas akts",
    )

    story = []

    # Cover Page (new feature)
    if akta_dati.add_cover_page:
        story.append(Spacer(1, 2 * inch))
        if akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
            try:
                cover_logo = RLImage(_coerce_path(akta_dati.logotipa_ceļš))
                cover_logo._restrictSize(float(akta_dati.cover_page_logo_width_mm) * mm, 50 * mm)
                story.append(cover_logo)
                story.append(Spacer(1, 0.5 * inch))
            except Exception:
                pass
        story.append(Paragraph(akta_dati.cover_page_title, styles['DocTitle']))
        story.append(Spacer(1, 1 * inch))
        story.append(Paragraph(f"Akta Nr.: {akta_dati.akta_nr}", styles['Latv']))
        story.append(Paragraph(f"Datums: {akta_dati.datums}", styles['Latv']))
        story.append(Paragraph(f"Vieta: {akta_dati.vieta}", styles['Latv']))
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(f"Pieņēmējs: {akta_dati.pieņēmējs.nosaukums}", styles['Latv']))
        if akta_dati.pieņēmējs.tālrunis:
            story.append(Paragraph(f"Pieņēmēja tālrunis: {akta_dati.pieņēmējs.tālrunis}", styles['Latv']))
        story.append(Paragraph(f"Nodevējs: {akta_dati.nodevējs.nosaukums}", styles['Latv']))
        if akta_dati.nodevējs.tālrunis:
            story.append(Paragraph(f"Nodevēja tālrunis: {akta_dati.nodevējs.tālrunis}", styles['Latv']))
        story.append(PageBreak())

    # Header ar logo un nosaukumu
    header_table_data = []
    logo_w = float(akta_dati.pdf_logo_width_mm) * mm
    if (not getattr(akta_dati, 'cover_page_enabled', False)) and akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
        try:
            header_logo._restrictSize(logo_w, 20 * mm)
            header_table_data.append([header_logo, Paragraph(f"<font name='{bold_font_name}'>PIEŅEMŠANAS–NODOŠANAS AKTS</font>", styles['LatvHead'])])
        except Exception:
            header_table_data.append(["", Paragraph(f"<font name='{bold_font_name}'>PIEŅEMŠANAS–NODOŠANAS AKTS</font>", styles['LatvHead'])])
    else:
        header_table_data.append(["", Paragraph(f"<font name='{bold_font_name}'>PIEŅEMŠANAS–NODOŠANAS AKTS</font>", styles['LatvHead'])])

    ht = Table(header_table_data, colWidths=[logo_w, None])
    ht.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(ht)

    # Aktu metadati (Nr., Datums, Vieta)
    md = [[Paragraph(f"<font name='{font_name}'><b>Akta Nr.:</b> {akta_dati.akta_nr}</font>", styles['Latv']),
           Paragraph(f"<font name='{font_name}'><b>Datums:</b> {datetime.strptime(akta_dati.datums, '%Y-%m-%d').strftime(akta_dati.date_format.replace('YYYY', '%Y').replace('MM', '%m').replace('DD', '%d'))}</font>", styles['Latv']),
           Paragraph(f"<font name='{font_name}'><b>Vieta:</b> {akta_dati.vieta}</font>", styles['Latv'])]]
    if akta_dati.pasūtījuma_nr:
        md[0].append(Paragraph(f"<font name='{font_name}'><b>Pasūtījuma Nr.:</b> {akta_dati.pasūtījuma_nr}</font>", styles['Latv']))

    meta = Table(md, colWidths=[55 * mm, 35 * mm, 40 * mm, None])
    meta.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(meta)
    story.append(Spacer(1, 6))

    # Jauni metadati
    if akta_dati.līguma_nr:
        story.append(Paragraph(f"<font name='{font_name}'><b>Līguma Nr.:</b> {akta_dati.līguma_nr}</font>", styles['Latv']))
    if akta_dati.izpildes_termiņš:
        story.append(Paragraph(f"<font name='{font_name}'><b>Izpildes termiņš:</b> {datetime.strptime(akta_dati.izpildes_termiņš, '%Y-%m-%d').strftime(akta_dati.date_format.replace('YYYY', '%Y').replace('MM', '%m').replace('DD', '%d'))}</font>", styles['Latv']))
    if akta_dati.pieņemšanas_datums:
        story.append(Paragraph(f"<font name='{font_name}'><b>Pieņemšanas datums:</b> {datetime.strptime(akta_dati.pieņemšanas_datums, '%Y-%m-%d').strftime(akta_dati.date_format.replace('YYYY', '%Y').replace('MM', '%m').replace('DD', '%d'))}</font>", styles['Latv']))
    if akta_dati.nodošanas_datums:
        story.append(Paragraph(f"<font name='{font_name}'><b>Nodošanas datums:</b> {datetime.strptime(akta_dati.nodošanas_datums, '%Y-%m-%d').strftime(akta_dati.date_format.replace('YYYY', '%Y').replace('MM', '%m').replace('DD', '%d'))}</font>", styles['Latv']))
    story.append(Spacer(1, 6))

    # Puses: Pieņēmējs / Nodevējs
    story.append(Paragraph("PUSES", styles['SectionBar']))
    def persona_paragraph(prefix: str, p: Persona):
        lines = [f"<b>{prefix}:</b> {p.nosaukums}"]
        if p.reģ_nr: lines.append(f"Reģ. Nr.: {p.reģ_nr}")
        if p.adrese: lines.append(f"Adrese: {p.adrese}")
        if p.kontaktpersona: lines.append(f"Kontaktpersona: {p.kontaktpersona}")
        if getattr(p, 'amats', ''): lines.append(f"Amats: {p.amats}")
        if getattr(p, 'pilnvaras_pamats', ''): lines.append(f"Pilnvaras pamats: {p.pilnvaras_pamats}")
        if p.tālrunis: lines.append(f"Tālrunis: {p.tālrunis}")
        if p.epasts:
            em = p.epasts.strip()
            lines.append(f"E-pasts: <a href=\"mailto:{em}\">{em}</a>")
        if getattr(p, "web_lapa", ""):
            url = p.web_lapa.strip()
            if url and not re.match(r"^[a-zA-Z]+://", url):
                url = "https://" + url
            disp = p.web_lapa.strip()
            lines.append(f"Web lapa: <a href=\"{url}\">{disp}</a>")
        if p.bankas_konts: lines.append(f"Bankas konts: {p.bankas_konts}")
        if p.juridiskais_statuss: lines.append(f"Statuss: {p.juridiskais_statuss}")
        return Paragraph(f"<font name='{font_name}'>" + "<br/>".join(lines) + "</font>", styles['Latv'])

    available_width = pagesize[0] - float(akta_dati.pdf_margin_left) * mm - float(akta_dati.pdf_margin_right) * mm
    col_width_parties = available_width / 2

    puses = Table([
        [persona_paragraph("Pieņēmējs", akta_dati.pieņēmējs), persona_paragraph("Nodevējs", akta_dati.nodevējs)]
    ], colWidths=[col_width_parties, col_width_parties])
    pstyle = TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOX', (0, 0), (-1, -1), float(akta_dati.table_border_thickness_pt), colors.HexColor(akta_dati.table_grid_color)),
        ('INNERGRID', (0, 0), (-1, -1), float(akta_dati.table_border_thickness_pt), colors.HexColor(akta_dati.table_grid_color)),
        ('LEFTPADDING', (0, 0), (-1, -1), float(akta_dati.table_cell_padding_mm)),
        ('RIGHTPADDING', (0, 0), (-1, -1), float(akta_dati.table_cell_padding_mm)),
        ('TOPPADDING', (0, 0), (-1, -1), float(akta_dati.table_cell_padding_mm)),
        ('BOTTOMPADDING', (0, 0), (-1, -1), float(akta_dati.table_cell_padding_mm)),
    ])
    puses.setStyle(pstyle)
    story.append(puses)
    story.append(Spacer(1, 8))

    # Pozīciju tabula
    story.append(Paragraph("POZĪCIJAS", styles['SectionBar']))
    # JAUNS: kolonnas var būt paslēptas/ pārdēvētas (poz_columns_config + custom_columns.visible)
    poz_cfg = _merge_poz_columns_config(getattr(akta_dati, "poz_columns_config", None))
    def _col_visible(key: str) -> bool:
        try:
            return bool(poz_cfg.get(key, {}).get("visible", True))
        except Exception:
            return True
    def _col_title(key: str, default: str) -> str:
        try:
            t = str(poz_cfg.get(key, {}).get("title", default))
            return t or default
        except Exception:
            return default

    # Kolonnu secība PDF tabulai (ņem vērā GUI pārvietošanu / drag&drop)
    # Nr. kolonna PDF tabulā ir atsevišķa un (ja ieslēgta) vienmēr ir pirmā.
    base_kind = {
        "apraksts": "text",
        "daudzums": "num",
        "vieniba": "text",
        "cena": "money",
        "summa": "money",
        "serial": "text",
        "warranty": "text",
        "notes": "text",
        "foto": "foto",
    }

    # Pielāgotās kolonnas (no datiem)
    custom_cols = []
    try:
        custom_cols = list(getattr(akta_dati, "custom_columns", []) or [])
    except Exception:
        custom_cols = []

    # Iespējamās atslēgas secībai (bez "nr")
    base_keys = ["apraksts", "daudzums", "vieniba", "cena", "summa", "serial", "warranty", "notes"]
    custom_keys = [f"custom:{i}" for i in range(len(custom_cols))]
    possible = set(base_keys + custom_keys + ["foto"])

    # Vēlamā secība no GUI (ja ir). Ja nav – izmantojam noklusējumu.
    desired = []
    try:
        desired = list(getattr(akta_dati, "poz_columns_visual_order", []) or [])
    except Exception:
        desired = []
    desired = [k for k in desired if k in possible]

    if not desired:
        desired = list(base_keys) + list(custom_keys) + ["foto"]
    else:
        # pieliekam trūkstošās kolonnas klāt (lai nezaudējam jaunpievienotas kolonnas)
        for k in base_keys:
            if k not in desired:
                desired.append(k)
        for k in custom_keys:
            if k not in desired:
                desired.append(k)
        if "foto" not in desired:
            desired.append("foto")

    # Foto kolonna var atrasties jebkurā vietā (tāpat kā GUI). 
    # Neuzspiežam tai pozīciju — izmantojam lietotāja (GUI) secību.

    # Uztaisām PDF kolonnu plānu, ņemot vērā redzamību
    col_plan = []  # list[tuple(key, kind)]
    if _col_visible("nr"):
        col_plan.append(("nr", "nr"))

    for key in desired:
        if key.startswith("custom:"):
            try:
                ci = int(key.split(":", 1)[1])
                if 0 <= ci < len(custom_cols):
                    cdef = custom_cols[ci]
                    if isinstance(cdef, dict) and bool(cdef.get("visible", True)):
                        col_plan.append((key, "text"))
            except Exception:
                pass
            continue

        if key == "foto":
            if _col_visible("foto"):
                col_plan.append(("foto", "foto"))
            continue

        # bāzes kolonnas
        if key in base_keys and _col_visible(key):
            col_plan.append((key, base_kind.get(key, "text")))

    # Galvene
    tab_header_items = []
    for key, _kind in col_plan:
        if key == "nr":
            tab_header_items.append(Paragraph(_col_title("nr", "Nr."), styles['TableHeader']))
        elif key.startswith("custom:"):
            try:
                ci = int(key.split(":", 1)[1])
                nm = str(custom_cols[ci].get("name", "")) if isinstance(custom_cols[ci], dict) else ""
            except Exception:
                nm = ""
            tab_header_items.append(Paragraph(nm or "", styles['TableHeader']))
        else:
            default_map = {
                "apraksts": "Apraksts",
                "daudzums": "Daudzums",
                "vieniba": "Vienība",
                "cena": "Cena",
                "summa": "Summa",
                "serial": "Seriālais Nr.",
                "warranty": "Garantija",
                "notes": "Piezīmes pozīcijai",
                "foto": "Foto",
            }
            tab_header_items.append(Paragraph(_col_title(key, default_map.get(key, key)), styles['TableHeader']))

    tab_data = [tab_header_items]

    # Rindas
    for i, poz in enumerate(akta_dati.pozīcijas, start=1):
        row_items = []
        for key, kind in col_plan:
            if key == "nr":
                row_items.append(Paragraph(str(i), styles['LatvTableContent']))
            elif key == "apraksts":
                row_items.append(Paragraph(poz.apraksts, styles['LatvTableContent']))
            elif key == "daudzums":
                row_items.append(Paragraph(f"{formēt_naudu(poz.daudzums)}", styles['LatvTableContent']))
            elif key == "vieniba":
                row_items.append(Paragraph(poz.vienība, styles['LatvTableContent']))
            elif key == "cena":
                row_items.append(Paragraph(f"{akta_dati.currency_symbol_position == 'before' and akta_dati.valūta or ''}{formēt_naudu(poz.cena)}{akta_dati.currency_symbol_position == 'after' and ' ' + akta_dati.valūta or ''}", styles['LatvTableContent']))
            elif key == "summa":
                row_items.append(Paragraph(f"{akta_dati.currency_symbol_position == 'before' and akta_dati.valūta or ''}{formēt_naudu(poz.summa)}{akta_dati.currency_symbol_position == 'after' and ' ' + akta_dati.valūta or ''}", styles['LatvTableContent']))
            elif key == "serial":
                row_items.append(Paragraph(poz.seriālais_nr, styles['LatvTableContent']))
            elif key == "warranty":
                row_items.append(Paragraph(poz.garantija, styles['LatvTableContent']))
            elif key == "notes":
                row_items.append(Paragraph(poz.piezīmes_pozīcijai, styles['LatvTableContent']))
            elif key.startswith("custom:"):
                try:
                    ci = int(key.split(":", 1)[1])
                    data_list = custom_cols[ci].get("data", []) if isinstance(custom_cols[ci], dict) else []
                    v = str(data_list[i-1]) if isinstance(data_list, list) and (i-1) < len(data_list) else ""
                except Exception:
                    v = ""
                row_items.append(Paragraph(v, styles['LatvTableContent']))
            elif key == "foto":
                if getattr(poz, 'attēla_ceļš', '') and os.path.exists(poz.attēla_ceļš):
                    try:
                        img_thumb = RLImage(poz.attēla_ceļš)
                        img_thumb._restrictSize(18 * mm, 14 * mm)
                        row_items.append(img_thumb)
                    except Exception:
                        row_items.append(Paragraph("", styles['LatvTableContent']))
                else:
                    row_items.append(Paragraph("", styles['LatvTableContent']))
            else:
                row_items.append(Paragraph("", styles['LatvTableContent']))
        tab_data.append(row_items)

    # Kolonnu platumi
    expected_cols = len(col_plan)
    col_widths = None
    try:
        col_widths_mm = [float(x.strip()) for x in (akta_dati.table_col_widths or "").split(',') if x.strip()]
        if len(col_widths_mm) == expected_cols:
            col_widths = [w * mm for w in col_widths_mm]
    except Exception:
        col_widths = None

    if not col_widths:
        # Noklusējuma platumi (mm) atkarībā no kolonnu tipa
        default_map_mm = {
            "nr": 10,
            "text": 35,
            "num": 18,
            "money": 20,
            "foto": 18,
        }
        dw = []
        for _k, kind in col_plan:
            dw.append(default_map_mm.get(kind, 25))
        col_widths = [w * mm for w in dw]
        total_default_width = sum(col_widths)
        if total_default_width > 0 and total_default_width != available_width:
            scale_factor = available_width / total_default_width
            col_widths = [w * scale_factor for w in col_widths]

    t = Table(tab_data, colWidths=col_widths)

    tstyle = TableStyle([
        ('FONTNAME', (0,0), (-1,-1), font_name),
        ('FONTNAME', (0,0), (-1,0), bold_font_name),
        ('FONTSIZE', (0,0), (-1,0), akta_dati.pdf_font_size_table),
        ('FONTSIZE', (0,1), (-1,-1), akta_dati.pdf_font_size_table),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor(akta_dati.table_header_bg_color or "#E5E7EB")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#0F172A")),
        ('ALIGN', (0,0), (-1,0), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor(akta_dati.table_grid_color or "#CBD5E1")),
        ('LINEBELOW', (0,0), (-1,0), 1.0, colors.HexColor("#94A3B8")),
        ('BOTTOMPADDING', (0,0), (-1,0), float(akta_dati.table_cell_padding_mm) * 2),
        ('TOPPADDING', (0,0), (-1,0), float(akta_dati.table_cell_padding_mm) * 2),
        ('BOTTOMPADDING', (0,1), (-1,-1), float(akta_dati.table_cell_padding_mm)),
        ('TOPPADDING', (0,1), (-1,-1), float(akta_dati.table_cell_padding_mm)),
    ])

    # Kolonnu izlīdzinājumi pēc tipa
    for col_i, (_k, kind) in enumerate(col_plan):
        if kind in ("money",):
            tstyle.add('ALIGN', (col_i, 1), (col_i, -1), 'RIGHT')
        elif kind in ("num", "nr"):
            tstyle.add('ALIGN', (col_i, 1), (col_i, -1), 'CENTER')
        elif kind == "foto":
            tstyle.add('ALIGN', (col_i, 1), (col_i, -1), 'CENTER')
        else:
            tstyle.add('ALIGN', (col_i, 1), (col_i, -1), 'LEFT')

    # Apply alternate row color
    if akta_dati.table_alternate_row_color:
        for i in range(1, len(tab_data)):
            if i % 2 == 0: # Even rows (0-indexed, so actual even rows)
                tstyle.add('BACKGROUND', (0, i), (-1, i), colors.HexColor(akta_dati.table_alternate_row_color))

    t.setStyle(tstyle)
    story.append(t)

    # Kopsavilkums (var izslēgt)
    if getattr(akta_dati, "show_price_summary", True):
        story.append(Spacer(1, 6))
        summa_tab = []
        summa_tab.append([Paragraph(f"<font name='{font_name}'>Kopā bez PVN:</font>", styles['Latv']), Paragraph(f"<font name='{font_name}'>{akta_dati.currency_symbol_position == 'before' and akta_dati.valūta or ''}{formēt_naudu(akta_dati.kopējā_summma())}{akta_dati.currency_symbol_position == 'after' and ' ' + akta_dati.valūta or ''}</font>", styles['Latv'])])
        if akta_dati.iekļaut_pvn and akta_dati.show_vat_breakdown:
            summa_tab.append([Paragraph(f"<font name='{font_name}'>PVN {akta_dati.pvn_likme}%:</font>", styles['Latv']), Paragraph(f"<font name='{font_name}'>{akta_dati.currency_symbol_position == 'before' and akta_dati.valūta or ''}{formēt_naudu(akta_dati.pvn_summa())}{akta_dati.currency_symbol_position == 'after' and ' ' + akta_dati.valūta or ''}</font>", styles['Latv'])])
            summa_tab.append([Paragraph(f"<font name='{font_name}'>Kopā ar PVN:</font>", styles['Latv']), Paragraph(f"<font name='{font_name}'>{akta_dati.currency_symbol_position == 'before' and akta_dati.valūta or ''}{formēt_naudu(akta_dati.summa_ar_pvn())}{akta_dati.currency_symbol_position == 'after' and ' ' + akta_dati.valūta or ''}</font>", styles['Latv'])])

        if summa_tab:
            ts = Table(summa_tab, colWidths=[None, 40*mm])
            ts.setStyle(TableStyle([
                ('ALIGN', (1,0), (1,-1), 'RIGHT'),
                ('FONTNAME', (0,0), (-1,-1), font_name),
                ('FONTSIZE', (0,0), (-1,-1), akta_dati.pdf_font_size_normal),
                ('LINEBELOW', (0,0), (-1,-1), 0.6, colors.HexColor('#CBD5E1')),
            ]))
            story.append(ts)

    # Piezīmes
    if akta_dati.piezīmes:
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<font name='{font_name}'><b>Vispārīgās piezīmes:</b><br/>{akta_dati.piezīmes}</font>", styles['Latv']))

    # Jauni juridiski saistoši lauki PDF dokumentā
    if akta_dati.strīdu_risināšana:
        story.append(Paragraph("Strīdu risināšanas kārtība", styles['LegalSectionTitle']))
        story.append(Paragraph(f"<font name='{font_name}'>{akta_dati.strīdu_risināšana}</font>", styles['LegalBody']))

    if akta_dati.konfidencialitātes_klauzula:
        story.append(Paragraph("Konfidencialitātes klauzula", styles['LegalSectionTitle']))
        story.append(Paragraph(
            f"<font name='{font_name}'>Puses apņemas neizpaust trešajām personām informāciju, kas iegūta šī akta ietvaros, "
            f"izņemot gadījumus, ko nosaka normatīvie akti.</font>",
            styles['LegalBody']
        ))

    if akta_dati.soda_nauda_procenti > 0:
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<font name='{font_name}'><b>Soda nauda:</b> Par saistību neizpildi vai nepienācīgu izpildi, vainīgā puse maksā otrai pusei soda naudu {formēt_naudu(akta_dati.soda_nauda_procenti)}% apmērā no neizpildīto saistību vērtības par katru kavējuma dienu.</font>", styles['Latv']))

    if akta_dati.piegādes_nosacījumi:
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<font name='{font_name}'><b>Piegādes nosacījumi:</b> {akta_dati.piegādes_nosacījumi}</font>", styles['Latv']))

    if akta_dati.apdrošināšana:
        story.append(Spacer(1, 6))
        apd_text = (getattr(akta_dati, "apdrošināšana_teksts", "") or "").strip()
        if not apd_text:
            apd_text = "Preces ir apdrošinātas pret bojājumiem un zaudējumiem līdz pieņemšanas-nodošanas brīdim."
        story.append(Paragraph(f"<font name='{font_name}'><b>Apdrošināšana:</b> {apd_text}</font>", styles['Latv']))

    if akta_dati.papildu_nosacījumi:
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<font name='{font_name}'><b>Papildu nosacījumi:</b><br/>{akta_dati.papildu_nosacījumi}</font>", styles['Latv']))

    if akta_dati.atsauces_dokumenti:
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<font name='{font_name}'><b>Atsauces dokumenti:</b> {akta_dati.atsauces_dokumenti}</font>", styles['Latv']))

    story.append(Spacer(1, 6))
    # Akta statuss (izcelts "juridisks premium" bloks)
    status_tbl = Table([[Paragraph(f"<font name='{bold_font_name}'>Akta statuss:</font>", styles['FieldLabel']),
                         Paragraph(f"<font name='{font_name}'>{akta_dati.akta_statuss}</font>", styles['FieldValue'])]],
                       colWidths=[35*mm, None])
    # JAUNS: Akta statusa fona krāsa PDF (atkarīga no izvēlētā statusa)
    _status_bg_map = {
        "Melnraksts": "#F1F5F9",   # pelēcīgs
        "Apstiprināts": "#DBEAFE", # zils
        "Parakstīts": "#DCFCE7",   # zaļš
        "Arhivēts": "#EDE9FE",     # violets
        "Atcelts": "#FEE2E2",      # sarkans
    }
    status_bg_color = _status_bg_map.get((akta_dati.akta_statuss or "").strip(), "#F8FAFC")

    status_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(status_bg_color)),
        ('BOX', (0,0), (-1,-1), 1.0, colors.HexColor("#334155")),
        ('INNERGRID', (0,0), (-1,-1), 0.0, colors.white),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(Spacer(1, 8))
    story.append(status_tbl)


    # Attēli (ja ir) – mērogojam līdz platumam
    if akta_dati.attēli:
        story.append(PageBreak())
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"<font name='{font_name}'><b>Pievienotās fotogrāfijas</b></font>", styles['Latv']))
        for att in akta_dati.attēli:
            if os.path.exists(att.ceļš):
                try:
                    img = RLImage(att.ceļš)
                    img._restrictSize(available_width, float(akta_dati.item_image_width_mm) * 3 * mm)
                    story.append(Spacer(1, 4))
                    story.append(img)
                    if att.paraksts:
                        story.append(Paragraph(f"<font name='{font_name}'>{att.paraksts}</font>", styles['LatvSmall']))
                except Exception:
                    pass

    if akta_dati.elektroniskais_paraksts:

        if akta_dati.radit_elektronisko_parakstu_tekstu:
            story.append(Spacer(1, 16))

            story.append(Paragraph(

                f"<font name='{font_name}' size='{akta_dati.pdf_font_size_normal}'><b>ŠIS DOKUMENTS PARAKSTĪTS AR DROŠU ELEKTRONISKO PARAKSTU UN SATUR LAIKA ZĪMOGU</b></font>",

                styles['LatvElectronicSignature']

            ))

            story.append(Spacer(1, 16))

    elif akta_dati.parakstu_rindas:

        story.append(Spacer(1, 16))  # Atstarpe pirms parakstu laukiem

        # Izveidojam atsevišķus elementus katrai paraksta pusei

        pie_elements = []

        nod_elements = []

        # Pieņēmēja paraksta informācija

        pie_elements.append(Paragraph(f"<font name='{font_name}'><b>Pieņēmējs</b></font>", styles['Latv']))

        if akta_dati.paraksts_pieņēmējs_ceļš and _path_exists(akta_dati.paraksts_pieņēmējs_ceļš):

            try:

                img_pie = RLImage(_coerce_path(akta_dati.paraksts_pieņēmējs_ceļš))

                img_pie._restrictSize(float(akta_dati.pdf_signature_width_mm) * mm,

                                      float(akta_dati.pdf_signature_height_mm) * mm)

                pie_elements.append(img_pie)

            except Exception:

                pie_elements.append(Paragraph(

                    f"<font name='{font_name}'>_{'_' * int(float(akta_dati.signature_line_length_mm) / 2)}</font>",

                    styles['Latv']))

        else:

            pie_elements.append(Paragraph(

                f"<font name='{font_name}'>_{'_' * int(float(akta_dati.signature_line_length_mm) / 2)}</font>",

                styles['Latv']))

        pie_elements.append(Paragraph(

            f"<font name='{font_name}' size='{akta_dati.signature_font_size}'>{akta_dati.pieņēmējs.kontaktpersona or akta_dati.pieņēmējs.nosaukums}</font>",

            styles['LatvSmall']))

        # Nodevēja paraksta informācija

        nod_elements.append(Paragraph(f"<font name='{font_name}'><b>Nodevējs</b></font>", styles['Latv']))

        if akta_dati.paraksts_nodevējs_ceļš and _path_exists(akta_dati.paraksts_nodevējs_ceļš):

            try:

                img_nod = RLImage(_coerce_path(akta_dati.paraksts_nodevējs_ceļš))

                img_nod._restrictSize(float(akta_dati.pdf_signature_width_mm) * mm,

                                      float(akta_dati.pdf_signature_height_mm) * mm)

                nod_elements.append(img_nod)

            except Exception:

                nod_elements.append(Paragraph(

                    f"<font name='{font_name}'>_{'_' * int(float(akta_dati.signature_line_length_mm) / 2)}</font>",

                    styles['Latv']))

        else:

            nod_elements.append(Paragraph(

                f"<font name='{font_name}'>_{'_' * int(float(akta_dati.signature_line_length_mm) / 2)}</font>",

                styles['Latv']))

        nod_elements.append(Paragraph(

            f"<font name='{font_name}' size='{akta_dati.signature_font_size}'>{akta_dati.nodevējs.kontaktpersona or akta_dati.nodevējs.nosaukums}</font>",

            styles['LatvSmall']))

        # Izveidojam tabulu ar vienu rindu un divām kolonnām

        # Katrā kolonnā ievietojam Flowables (Paragraph, Image) sarakstu

        paraksti_table_data = [

            [pie_elements, nod_elements]

        ]

        paraksti = Table(paraksti_table_data, colWidths=[col_width_parties, col_width_parties])

        paraksti.setStyle(TableStyle([

            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Centra izlīdzināšana visiem elementiem šūnās

            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Vertikālā izlīdzināšana uz augšu

            ('FONTNAME', (0, 0), (-1, -1), font_name),

            ('BOTTOMPADDING', (0, 0), (-1, 0), float(akta_dati.signature_spacing_mm)),  # Atstarpe zem paraksta līnijas

        ]))

        story.append(paraksti)

    # QR kodi (jauna funkcionalitāte)
    qr_codes_to_add = []

    # Individuālais QR kods
    if akta_dati.include_custom_qr_code and akta_dati.custom_qr_code_data:
        qr_codes_to_add.append({
            'data': akta_dati.custom_qr_code_data,
            'size': akta_dati.custom_qr_code_size_mm,
            'position': akta_dati.custom_qr_code_position
        })

    # Automātiskais QR kods (akta ID)
    if akta_dati.include_auto_qr_code and akta_dati.akta_nr:
        qr_codes_to_add.append({
            'data': akta_dati.akta_nr,
            'size': akta_dati.auto_qr_code_size_mm,
            'position': akta_dati.auto_qr_code_position
        })

    if qr_codes_to_add:
        try:
            import qrcode
            qr_images = []
            for qr_info in qr_codes_to_add:
                qr = qrcode.QRCode(
                    version=1,
                    error_correction=qrcode.constants.ERROR_CORRECT_L,
                    box_size=10,
                    border=4,
                )
                qr.add_data(qr_info['data'])
                qr.make(fit=True)
                img_qr = qr.make_image(fill_color="black", back_color="white")
                qr_temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
                img_qr.save(qr_temp_path)

                rl_qr_img = RLImage(qr_temp_path)
                rl_qr_img._restrictSize(float(qr_info['size']) * mm, float(qr_info['size']) * mm)
                qr_images.append((rl_qr_img, qr_info['position'], qr_temp_path))

            # Pievienojam QR kodus story. Šeit varētu būt sarežģītāka pozicionēšanas loģika,
            # bet pagaidām pievienojam secīgi.
            # Lai nodrošinātu pozicionēšanu, būtu jāizmanto ReportLab "Frame" vai jāpielāgo "canvas" objekts.
            # Vienkāršības labad pievienojam tos beigās, katru savā rindā.
            story.append(Spacer(1, 12))  # Atstarpe pirms QR kodiem
            for img, pos, temp_path in qr_images:
                # ReportLab SimpleDocTemplate neļauj viegli pozicionēt elementus absolūti.
                # Lai to izdarītu, būtu jāizmanto canvas objekts tieši vai jāveido sarežģītākas tabulas.
                # Pagaidām pievienojam tos kā atsevišķus elementus.
                # Ja nepieciešama precīza pozicionēšana, tas ir ievērojams darbs.
                story.append(img)
                story.append(Spacer(1, 6))  # Atstarpe starp QR kodiem
                os.remove(temp_path)  # Dzēšam pagaidu failu

        except ImportError:
            print("QR Code generation requires 'qrcode' library. Please install it: pip install qrcode")
        except Exception as e:
            print(f"Error generating QR code: {e}")

    # Footer ar ģenerēšanas laiku
    if akta_dati.show_generation_timestamp:
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"<font name='{font_name}'>Dokuments ģenerēts: {datetime.now().strftime('%Y-%m-%d %H:%M')}</font>", styles['LatvSmall']))

    # Build ar lapu numerāciju
    # Pass show_page_numbers to the custom canvas
    doc.build(story, canvasmaker=lambda *args, **kwargs: DecoratedCanvas(*args, akta_dati=akta_dati, show_page_numbers=False, **kwargs))
    # Pievienojam reālus atsauces dokumentus PDF beigās (Atvasinājums 1, 2, 3 ...)
    # Priekšskatījumā (GUI) to izlaižam, lai Word/Excel pielikumu konvertēšana neuzkārtu UI.
    if include_reference_docs:
        try:
            pdf_ceļš = _append_reference_docs_to_pdf(pdf_ceļš, akta_dati, pagesize=pagesize, font_name=font_name)
        except Exception as e:
            print(f"Atsauces dokumentu pievienošanas kļūda: {e}")

    
    # Uzliekam KOPĒJO lapu numerāciju visam dokumentam (arī atvasinājumiem), lai "Lapa X no Y" ir pareizi
    _apply_global_page_numbers_to_pdf(pdf_ceļš, akta_dati, font_name)
    _apply_stapler_mark_to_pdf(pdf_ceļš, akta_dati)
    _apply_qr_to_pdf(pdf_ceļš, akta_dati)


    # PDF Encryption (new feature)
    if encrypt_pdf and akta_dati.enable_pdf_encryption:
        try:
            from PyPDF2 import PdfReader, PdfWriter
            # Svarīgi Windows: nelasām un nerakstām vienlaikus vienā failā.
            with open(pdf_ceļš, "rb") as rf:
                reader = PdfReader(rf)
                writer = PdfWriter()
                for page_index, page in enumerate(reader.pages):

                    writer.add_page(page)

            permissions = 0
            if akta_dati.allow_printing: permissions |= 4  # Print
            if akta_dati.allow_modifying: permissions |= 8  # Modify contents
            if akta_dati.allow_copying: permissions |= 16  # Copy
            if akta_dati.allow_annotating: permissions |= 32  # Annotate

            # Ja owner parole nav iedota, PyPDF2 dažreiz uzvedas neprognozējami – uzliekam drošu noklusējumu.
            owner_pw = akta_dati.pdf_owner_password or akta_dati.pdf_user_password or "owner"
            user_pw = akta_dati.pdf_user_password or ""

            # PyPDF2/pypdf API atšķiras starp versijām, tāpēc mēģinām vairākus variantus.
            try:
                writer.encrypt(user_password=user_pw, owner_password=owner_pw, permissions_flag=permissions)
            except TypeError:
                try:
                    writer.encrypt(user_pwd=user_pw, owner_pwd=owner_pw, permissions_flag=permissions)
                except TypeError:
                    # Vecāki PyPDF2 varianti
                    writer.encrypt(user_pw, owner_pw, use_128bit=True, permissions_flag=permissions)

            _atomic_write_pdfwriter(pdf_ceļš, writer)

        except ImportError:
            print("PDF encryption requires 'PyPDF2' library. Please install it: pip install PyPDF2")
        except Exception as e:
            print(f"Error encrypting PDF: {e}")

    return pdf_ceļš

# ---------------------- DOCX ģenerēšana ----------------------
# (unchanged, as DOCX generation is less flexible with advanced styling)
def add_formatted_text(paragraph, text):
    # ... (unchanged)
    parts = []
    current_text = ""
    in_bold = False
    in_italic = False
    in_underline = False

    i = 0
    while i < len(text):
        if text[i:i+3] == "<b>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_bold = True
            i += 3
        elif text[i:i+4] == "</b>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_bold = False
            i += 4
        elif text[i:i+3] == "<i>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_italic = True
            i += 3
        elif text[i:i+4] == "</i>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_italic = False
            i += 4
        elif text[i:i+3] == "<u>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_underline = True
            i += 3
        elif text[i:i+4] == "</u>":
            parts.append((current_text, in_bold, in_italic, in_underline))
            current_text = ""
            in_underline = False
            i += 4
        else:
            current_text += text[i]
            i += 1
    parts.append((current_text, in_bold, in_italic, in_underline))

    for part_text, bold, italic, underline in parts:
        if part_text:
            run = paragraph.add_run(part_text)
            run.bold = bold
            run.italic = italic
            run.underline = underline


def _docx_replace_placeholders(document: Document, mapping: dict):
    """Aizvieto {{atslēga}} vietturus visos paragrāfos un tabulās."""
    def replace_in_runs(paragraph):
        # Vienkārša aizvietošana (placeholders parasti ir vienā run, bet mēģinām droši)
        full_text = ''.join(r.text for r in paragraph.runs)
        if not full_text:
            return
        new_full = full_text
        for k, v in mapping.items():
            new_full = new_full.replace(f"{{{{{k}}}}}", str(v))
        if new_full != full_text:
            # Pārrakstām ar vienu run (saglabāt stilu precīzi ir sarežģīti)
            for r in paragraph.runs:
                r.text = ''
            paragraph.runs[0].text = new_full if paragraph.runs else paragraph.add_run(new_full)

    for p in document.paragraphs:
        replace_in_runs(p)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p)


# ---------------------- DOCX vizuālie uzlabojumi (lai atbilst PDF) ----------------------

def _docx_set_run_font(run: Run, *, bold: bool = False, size_pt: int | None = None,
                       color_hex: str | None = None, name: str | None = None):
    try:
        run.bold = bool(bold)
    except Exception:
        pass
    try:
        if size_pt is not None:
            run.font.size = Pt(size_pt)
    except Exception:
        pass
    try:
        if name:
            run.font.name = name
    except Exception:
        pass
    try:
        if color_hex:
            c = color_hex.strip().lstrip('#')
            if len(c) == 6:
                run.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
    except Exception:
        pass


def _docx_set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 0, line: float | None = None):
    try:
        pf = paragraph.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if line is not None:
            pf.line_spacing = line
    except Exception:
        pass


def _docx_set_table_borders(table, *, color: str = "000000", size: int = 8):
    """Uzspiež redzamas tabulas līnijas (python-docx pēc noklusējuma bieži dod 'baltas' līnijas)."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = tblBorders.find(qn(f'w:{edge}'))
            if elem is None:
                elem = OxmlElement(f'w:{edge}')
                tblBorders.append(elem)
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(size))
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), color)
    except Exception:
        pass


def _docx_set_cell_shading(cell, fill: str):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
    except Exception:
        pass


def _docx_apply_pdf_like_theme(document: Document):
    """Vienkāršs 'PDF līdzīgs' noformējums: fonti, virsrakstu krāsa, atstarpes."""
    try:
        normal = document.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
    except Exception:
        pass

    # Heading krāsa (PDF stilā zilgana)
    for st_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        try:
            st = document.styles[st_name]
            st.font.name = 'Calibri'
            if st_name == 'Heading 1':
                st.font.size = Pt(16)
            elif st_name == 'Heading 2':
                st.font.size = Pt(12)
            else:
                st.font.size = Pt(11)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # tumši zils (līdzīgs PDF)
        except Exception:
            pass


def _docx_add_title_page_like_pdf(document: Document, akta_dati: AktaDati):
    """Izveido titullapu (kā PDF): nosaukums, Akta Nr/Datums/Vieta + īss pušu kopsavilkums."""
    # Logo augšā (ja ir)
    if akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
        try:
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(akta_dati.logotipa_ceļš, width=Inches(float(akta_dati.docx_image_width_inches)))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _docx_set_paragraph_spacing(p, after=8)
        except Exception:
            pass

    # Galvenais nosaukums centrā
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('Pieņemšanas-Nodošanas Akts')
    _docx_set_run_font(r, bold=True, size_pt=20, color_hex='000000')
    _docx_set_paragraph_spacing(p_title, before=10, after=18)

    # Akta metadati
    meta_lines = [
        f"Akta Nr.: {akta_dati.akta_nr}",
        f"Datums: {akta_dati.datums}",
        f"Vieta: {akta_dati.vieta}",
    ]
    if akta_dati.pasūtījuma_nr:
        meta_lines.append(f"Pasūtījuma Nr.: {akta_dati.pasūtījuma_nr}")
    if akta_dati.līguma_nr:
        meta_lines.append(f"Līguma Nr.: {akta_dati.līguma_nr}")

    p_meta = document.add_paragraph('\n'.join(meta_lines))
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _docx_set_paragraph_spacing(p_meta, after=14)

    # Īss kopsavilkums par pusēm (kā PDF titullapā)
    p = document.add_paragraph()
    p.add_run(f"Pieņēmējs: {akta_dati.pieņēmējs.nosaukums}\n")
    p.add_run(f"Pieņēmēja tālrunis: {akta_dati.pieņēmējs.tālrunis or ''}\n")
    p.add_run(f"Nodevējs: {akta_dati.nodevējs.nosaukums}\n")
    p.add_run(f"Nodevēja tālrunis: {akta_dati.nodevējs.tālrunis or ''}")
    _docx_set_paragraph_spacing(p, after=10)

    document.add_page_break()

def _docx_set_run_font(run: Run, *, bold: bool = None, size_pt: int = None, color_hex: str = None, name: str = None):
    try:
        if bold is not None:
            run.bold = bold
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if name is not None:
            run.font.name = name
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex.replace('#', '').upper())
    except Exception:
        pass


def _docx_set_paragraph_spacing(p, before_pt: int = 0, after_pt: int = 4, line: float = 1.0):
    try:
        pf = p.paragraph_format
        pf.space_before = Pt(before_pt)
        pf.space_after = Pt(after_pt)
        # python-docx line_spacing expects a float multiplier or Length
        pf.line_spacing = line
    except Exception:
        pass


def _docx_set_cell_shading(cell, fill_hex: str):
    """Uzstāda šūnas fona krāsu (piem., virsrindām / sadaļu blokiem)."""
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex.replace('#', '').upper())
        tc_pr.append(shd)
    except Exception:
        pass


def _docx_set_table_borders(table, *, color_hex: str = '000000', size: int = 12):
    """Garantē redzamas tabulu līnijas (Word dažreiz iestata baltas/neredzamas)."""
    try:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.append(tbl_pr)

        # atrod vai izveido tblBorders
        borders = tbl_pr.find(qn('w:tblBorders'))
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tbl_pr.append(borders)

        def _set_edge(tag):
            edge = borders.find(qn(f'w:{tag}'))
            if edge is None:
                edge = OxmlElement(f'w:{tag}')
                borders.append(edge)
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), str(size))
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), color_hex.replace('#', '').upper())

        for t in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            _set_edge(t)
    except Exception:
        pass


def _docx_make_heading(document: Document, text: str, level: int = 2, *, color_hex: str = '1F4E79'):
    """Heading ar PDF līdzīgu zilo krāsu."""
    p = document.add_heading(text, level=level)
    try:
        # Word Heading var būt vairākas runs
        for r in p.runs:
            _docx_set_run_font(r, bold=True, color_hex=color_hex)
        _docx_set_paragraph_spacing(p, before_pt=8 if level <= 2 else 4, after_pt=4)
    except Exception:
        pass
    return p


def _docx_apply_default_styles(document: Document):
    """Iestata bāzes stilus, lai Word vizuāli būtu tuvāk PDF."""
    try:
        document.styles['Normal'].font.name = 'Calibri'
        document.styles['Normal'].font.size = Pt(10)

        # Heading 1
        if 'Heading 1' in document.styles:
            h1 = document.styles['Heading 1']
            h1.font.name = 'Calibri'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor.from_string('1F4E79')

        # Heading 2
        if 'Heading 2' in document.styles:
            h2 = document.styles['Heading 2']
            h2.font.name = 'Calibri'
            h2.font.size = Pt(12)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor.from_string('1F4E79')
    except Exception:
        pass


def _docx_add_cover_page(document: Document, akta_dati: 'AktaDati'):
    """Titullapa līdzīga PDF titullapai."""
    # Top logo (ja ir)
    if akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
        try:
            p_logo = document.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.add_run().add_picture(akta_dati.logotipa_ceļš, width=Inches(1.8))
            _docx_set_paragraph_spacing(p_logo, after_pt=16)
        except Exception:
            pass

    # Liels virsraksts
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('Pieņemšanas-Nodošanas Akts')
    _docx_set_run_font(r, bold=True, size_pt=20, color_hex='000000')
    _docx_set_paragraph_spacing(p_title, before_pt=8, after_pt=18)

    # Akta info
    info_lines = [
        ("Akta Nr.:", akta_dati.akta_nr),
        ("Datums:", akta_dati.datums),
        ("Vieta:", akta_dati.vieta),
    ]
    if akta_dati.pasūtījuma_nr:
        info_lines.append(("Pasūtījuma Nr.:", akta_dati.pasūtījuma_nr))

    p_info = document.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for k, v in info_lines:
        rrk = p_info.add_run(f"{k} ")
        _docx_set_run_font(rrk, bold=True)
        p_info.add_run(f"{v}\n")
    _docx_set_paragraph_spacing(p_info, after_pt=18)

    # Puses īsi (kā PDF titullapā)
    p_parties = document.add_paragraph()
    p_parties.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = p_parties.add_run(f"Pieņēmējs: ")
    _docx_set_run_font(rr, bold=True)
    p_parties.add_run(f"{akta_dati.pieņēmējs.nosaukums}\n")
    rr = p_parties.add_run(f"Pieņēmēja tālrunis: ")
    _docx_set_run_font(rr, bold=True)
    p_parties.add_run(f"{akta_dati.pieņēmējs.tālrunis}\n")
    rr = p_parties.add_run(f"Nodevējs: ")
    _docx_set_run_font(rr, bold=True)
    p_parties.add_run(f"{akta_dati.nodevējs.nosaukums}\n")
    rr = p_parties.add_run(f"Nodevēja tālrunis: ")
    _docx_set_run_font(rr, bold=True)
    p_parties.add_run(f"{akta_dati.nodevējs.tālrunis}\n")
    _docx_set_paragraph_spacing(p_parties, after_pt=8)

    # Pāreja uz nākamo lapu
    document.add_page_break()

def ģenerēt_docx(akta_dati: AktaDati, docx_ceļš: str):
    # Ja ir norādīts DOCX šablons, ielādējam to; citādi veidojam jaunu dokumentu.
    used_template = False
    if getattr(akta_dati, "docx_template_path", "") and _path_exists(getattr(akta_dati, "docx_template_path", "")):
        document = Document(_coerce_path(getattr(akta_dati, "docx_template_path", "")) or akta_dati.docx_template_path)
        used_template = True
        # Vietturu aizvietošana (ja šablonā izmanto {{akta_nr}}, {{datums}} utt.)
        _docx_replace_placeholders(document, {
            "akta_nr": akta_dati.akta_nr,
            "datums": akta_dati.datums,
            "vieta": akta_dati.vieta,
            "pasūtījuma_nr": akta_dati.pasūtījuma_nr,
            "pieņēmējs": akta_dati.pieņēmējs.nosaukums,
            "nodevējs": akta_dati.nodevējs.nosaukums,
            "valūta": akta_dati.valūta,
        })
        # Šablons var jau saturēt visu; ja gribi – tālāk pievienojam arī standarta sadaļas.
    else:
        document = Document()

    # PDF-līdzīgs noformējums (fonti/Heading krāsa u.c.)
    try:
        _docx_apply_default_styles(document)
    except Exception:
        pass

    # Titullapa (ja nav šablona). Šī ir 1:1 ideja kā PDF (logo + nosaukums + kopsavilkums)
    if not used_template:
        try:
            _docx_add_cover_page(document, akta_dati)
        except Exception:
            pass

    # (atstājam arī iepriekšējo iestatījumu drošībai, ja Word ignorē stilus)
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(10)

    # Logo + virsraksts (kā PDF pirmajā saturiskajā lapā)
    _main_logo_added = False
    if akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
        try:
            p_logo2 = document.add_paragraph()
            p_logo2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo2.add_run().add_picture(akta_dati.logotipa_ceļš, width=Inches(float(akta_dati.docx_image_width_inches)))
            _docx_set_paragraph_spacing(p_logo2, after_pt=6)
            _main_logo_added = True
        except Exception:
            pass

    document.add_heading('PIEŅEMŠANAS–NODOŠANAS AKTS', level=1)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # (vecais logo bloks atstāts drošībai, bet vairs nedublējam)
    if (not _main_logo_added) and akta_dati.logotipa_ceļš and _path_exists(akta_dati.logotipa_ceļš):
        try:
            document.add_picture(akta_dati.logotipa_ceļš, width=Inches(float(akta_dati.docx_image_width_inches)))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    add_formatted_text(document.add_paragraph(), f"<b>Akta Nr.:</b> {akta_dati.akta_nr}")
    add_formatted_text(document.add_paragraph(), f"<b>Datums:</b> {akta_dati.datums}")
    add_formatted_text(document.add_paragraph(), f"<b>Vieta:</b> {akta_dati.vieta}")
    if akta_dati.pasūtījuma_nr:
        add_formatted_text(document.add_paragraph(), f"<b>Pasūtījuma Nr.:</b> {akta_dati.pasūtījuma_nr}")

    if akta_dati.līguma_nr:
        add_formatted_text(document.add_paragraph(), f"<b>Līguma Nr.:</b> {akta_dati.līguma_nr}")
    if akta_dati.izpildes_termiņš:
        add_formatted_text(document.add_paragraph(), f"<b>Izpildes termiņš:</b> {akta_dati.izpildes_termiņš}")
    if akta_dati.pieņemšanas_datums:
        add_formatted_text(document.add_paragraph(), f"<b>Pieņemšanas datums:</b> {akta_dati.pieņemšanas_datums}")
    if akta_dati.nodošanas_datums:
        add_formatted_text(document.add_paragraph(), f"<b>Nodošanas datums:</b> {akta_dati.nodošanas_datums}")
    document.add_paragraph()

    document.add_heading('Puses', level=2)
    table_parties = document.add_table(rows=1, cols=2)
    table_parties.autofit = True
    table_parties.columns[0].width = Inches(3)
    table_parties.columns[1].width = Inches(3)

    # TABULAS līnijas un galvenes fons (lai Word izskatās kā PDF)
    try:
        table_parties.style = 'Table Grid'
    except Exception:
        pass
    _docx_set_table_borders(table_parties, color_hex='000000', size=12)

    hdr_cells = table_parties.rows[0].cells
    hdr_cells[0].text = "Pieņēmējs"
    hdr_cells[1].text = "Nodevējs"

    try:
        for c in hdr_cells:
            _docx_set_cell_shading(c, 'D9E1F2')  # gaiši zils, līdzīgs PDF blokam
            for r in c.paragraphs[0].runs:
                _docx_set_run_font(r, bold=True)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass

    row_cells = table_parties.add_row().cells
    def add_persona_info(cell, p: Persona):
        add_formatted_text(cell.add_paragraph(), f"<b>Nosaukums / Vārds, Uzvārds:</b> {p.nosaukums}")
        if p.reģ_nr: add_formatted_text(cell.add_paragraph(), f"<b>Reģ. Nr. / personas kods:</b> {p.reģ_nr}")
        if p.adrese: add_formatted_text(cell.add_paragraph(), f"<b>Adrese:</b> {p.adrese}")
        if p.kontaktpersona: add_formatted_text(cell.add_paragraph(), f"<b>Kontaktpersona:</b> {p.kontaktpersona}")
        if getattr(p, 'amats', ''): add_formatted_text(cell.add_paragraph(), f"<b>Amats:</b> {p.amats}")
        if getattr(p, 'pilnvaras_pamats', ''): add_formatted_text(cell.add_paragraph(), f"<b>Pilnvaras pamats:</b> {p.pilnvaras_pamats}")
        if p.tālrunis: add_formatted_text(cell.add_paragraph(), f"<b>Tālrunis:</b> {p.tālrunis}")
        if p.epasts: add_formatted_text(cell.add_paragraph(), f"<b>E-pasts:</b> {p.epasts}")
        if p.bankas_konts: add_formatted_text(cell.add_paragraph(), f"<b>Bankas konts:</b> {p.bankas_konts}")
        if p.juridiskais_statuss: add_formatted_text(cell.add_paragraph(), f"<b>Statuss:</b> {p.juridiskais_statuss}")

    add_persona_info(row_cells[0], akta_dati.pieņēmējs)
    add_persona_info(row_cells[1], akta_dati.nodevējs)
    document.add_paragraph()

    document.add_heading('Pozīcijas', level=2)
    # Adjust columns based on visibility settings for DOCX
    docx_headers = ["Nr.", "Apraksts", "Daudzums", "Vienība", "Cena", "Summa"]
    if akta_dati.show_item_serial_number_in_table: docx_headers.append("Seriālais Nr.")
    if akta_dati.show_item_warranty_in_table: docx_headers.append("Garantija")
    if akta_dati.show_item_notes_in_table: docx_headers.append("Piezīmes pozīcijai")

    table_items = document.add_table(rows=1, cols=len(docx_headers))
    table_items.autofit = True
    table_items.allow_autofit = True

    # Nodrošinām redzamas līnijas + PDF līdzīgu virsrindas fonu
    try:
        table_items.style = 'Table Grid'
    except Exception:
        pass
    _docx_set_table_borders(table_items, color_hex='000000', size=12)

    for i, header_text in enumerate(docx_headers):
        table_items.rows[0].cells[i].text = header_text
        try:
            _docx_set_cell_shading(table_items.rows[0].cells[i], 'D9E1F2')
        except Exception:
            pass
        try:
            for r in table_items.rows[0].cells[i].paragraphs[0].runs:
                _docx_set_run_font(r, bold=True)
            table_items.rows[0].cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass

    for i, poz in enumerate(akta_dati.pozīcijas, start=1):
        row_cells = table_items.add_row().cells
        row_cells[0].text = str(i)
        add_formatted_text(row_cells[1].paragraphs[0], poz.apraksts)
        row_cells[2].text = f"{formēt_naudu(poz.daudzums)}"
        row_cells[3].text = poz.vienība
        row_cells[4].text = f"{formēt_naudu(poz.cena)} {akta_dati.valūta}"
        row_cells[5].text = f"{formēt_naudu(poz.summa)} {akta_dati.valūta}"
        col_idx = 6
        if akta_dati.show_item_serial_number_in_table:
            row_cells[col_idx].text = poz.seriālais_nr
            col_idx += 1
        if akta_dati.show_item_warranty_in_table:
            row_cells[col_idx].text = poz.garantija
            col_idx += 1
        if akta_dati.show_item_notes_in_table:
            add_formatted_text(row_cells[col_idx].paragraphs[0], poz.piezīmes_pozīcijai)

    document.add_paragraph()

    document.add_heading('Kopsavilkums', level=2)
    add_formatted_text(document.add_paragraph(), f"<b>Kopā bez PVN:</b> {formēt_naudu(akta_dati.kopējā_summma())} {akta_dati.valūta}")
    if akta_dati.iekļaut_pvn:
        add_formatted_text(document.add_paragraph(), f"<b>PVN {akta_dati.pvn_likme}%:</b> {formēt_naudu(akta_dati.pvn_summa())} {akta_dati.valūta}")
        add_formatted_text(document.add_paragraph(), f"<b>Kopā ar PVN:</b> {formēt_naudu(akta_dati.summa_ar_pvn())} {akta_dati.valūta}")
    document.add_paragraph()

    if akta_dati.piezīmes:
        document.add_heading('Vispārīgās piezīmes', level=2)
        add_formatted_text(document.add_paragraph(), akta_dati.piezīmes)
        document.add_paragraph()

    if akta_dati.strīdu_risināšana:
        document.add_heading('Strīdu risināšanas kārtība', level=2)
        add_formatted_text(document.add_paragraph(), akta_dati.strīdu_risināšana)
        document.add_paragraph()

    if akta_dati.konfidencialitātes_klauzula:
        document.add_heading('Konfidencialitātes klauzula', level=2)
        add_formatted_text(document.add_paragraph(), "Puses apņemas neizpaust trešajām personām informāciju, kas iegūta šī akta ietvaros, izņemot gadījumus, ko nosaka normatīvie akti.")
        document.add_paragraph()

    if akta_dati.soda_nauda_procenti > 0:
        document.add_heading('Soda nauda', level=2)
        add_formatted_text(document.add_paragraph(), f"Par saistību neizpildi vai nepienācīgu izpildi, vainīgā puse maksā otrai pusei soda naudu {formēt_naudu(akta_dati.soda_nauda_procenti)}% apmērā no neizpildīto saistību vērtības par katru kavējuma dienu.")
        document.add_paragraph()

    if akta_dati.piegādes_nosacījumi:
        document.add_heading('Piegādes nosacījumi', level=2)
        add_formatted_text(document.add_paragraph(), akta_dati.piegādes_nosacījumi)
        document.add_paragraph()

    if akta_dati.apdrošināšana:
        document.add_heading('Apdrošināšana', level=2)
        add_formatted_text(document.add_paragraph(), "Preces ir apdrošinātas pret bojājumiem un zaudējumiem līdz pieņemšanas-nodošanas brīdim.")
        document.add_paragraph()

    if akta_dati.papildu_nosacījumi:
        document.add_heading('Papildu nosacījumi', level=2)
        add_formatted_text(document.add_paragraph(), akta_dati.papildu_nosacījumi)
        document.add_paragraph()

    if akta_dati.atsauces_dokumenti:
        document.add_heading('Atsauces dokumenti', level=2)
        add_formatted_text(document.add_paragraph(), akta_dati.atsauces_dokumenti)
        document.add_paragraph()

    document.add_heading('Akta statuss', level=2)
    add_formatted_text(document.add_paragraph(), akta_dati.akta_statuss)
    document.add_paragraph()

    if akta_dati.attēli:
        document.add_heading('Pievienotās fotogrāfijas', level=2)
        for att in akta_dati.attēli:
            if os.path.exists(att.ceļš):
                try:
                    document.add_picture(att.ceļš, width=Inches(float(akta_dati.docx_image_width_inches)))
                    if att.paraksts:
                        add_formatted_text(document.add_paragraph(), att.paraksts)
                except Exception:
                    pass
        document.add_paragraph()

    if akta_dati.elektroniskais_paraksts:
        if akta_dati.radit_elektronisko_parakstu_tekstu:
            document.add_paragraph()
            p = document.add_paragraph()
            run = p.add_run("ŠIS DOKUMENTS PARAKSTĪTS AR DROŠU ELEKTRONISKO PARAKSTU UN SATUR LAIKA ZĪMOGU")
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif akta_dati.parakstu_rindas:
        document.add_heading('Paraksti', level=2)
        table_signatures = document.add_table(rows=2, cols=2)
        table_signatures.autofit = True
        table_signatures.columns[0].width = Inches(3)
        table_signatures.columns[1].width = Inches(3)

        try:
            table_signatures.style = 'Table Grid'
        except Exception:
            pass
        _docx_set_table_borders(table_signatures, color_hex='000000', size=12)

        table_signatures.rows[0].cells[0].text = "Pieņēmējs"
        table_signatures.rows[0].cells[1].text = "Nodevējs"
        for cell in table_signatures.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_pie = table_signatures.rows[1].cells[0]
        cell_nod = table_signatures.rows[1].cells[1]

        if akta_dati.paraksts_pieņēmējs_ceļš and _path_exists(akta_dati.paraksts_pieņēmējs_ceļš):
            try:
                cell_pie.add_paragraph().add_run().add_picture(akta_dati.paraksts_pieņēmējs_ceļš, width=Inches(
                    float(akta_dati.docx_signature_width_inches)))
                cell_pie.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                cell_pie.add_paragraph("____________________________")
        else:
            cell_pie.add_paragraph("____________________________")

        if akta_dati.paraksts_nodevējs_ceļš and _path_exists(akta_dati.paraksts_nodevējs_ceļš):
            try:
                cell_nod.add_paragraph().add_run().add_picture(akta_dati.paraksts_nodevējs_ceļš, width=Inches(
                    float(akta_dati.docx_signature_width_inches)))
                cell_nod.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                cell_nod.add_paragraph("____________________________")
        else:
            cell_nod.add_paragraph("____________________________")

        add_formatted_text(cell_pie.add_paragraph(),
                           akta_dati.pieņēmējs.kontaktpersona or akta_dati.pieņēmējs.nosaukums)
        add_formatted_text(cell_nod.add_paragraph(), akta_dati.nodevējs.kontaktpersona or akta_dati.nodevējs.nosaukums)
        for cell in [cell_pie, cell_nod]:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # QR kodi DOCX dokumentā (ja ieslēgts)
    try:
        import qrcode
        def _add_qr(data: str, size_mm: Decimal, align: str):
            if not data:
                return
            qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=2)
            qr.add_data(data)
            qr.make(fit=True)
            img_qr = qr.make_image(fill_color="black", back_color="white")
            qr_temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
            img_qr.save(qr_temp_path)

            qr_width_inches = float(size_mm) / 25.4
            p = document.add_paragraph()
            p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                           "center": WD_ALIGN_PARAGRAPH.CENTER,
                           "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
            p.add_run().add_picture(qr_temp_path, width=Inches(qr_width_inches))
            try:
                os.remove(qr_temp_path)
            except Exception:
                pass

        if akta_dati.include_custom_qr_code and akta_dati.custom_qr_code_data:
            _add_qr(akta_dati.custom_qr_code_data, akta_dati.custom_qr_code_size_mm, "right")
        if akta_dati.include_auto_qr_code and akta_dati.akta_nr:
            _add_qr(akta_dati.akta_nr, akta_dati.auto_qr_code_size_mm, "left")
    except Exception as e:
        print(f"DOCX QR kļūda: {e}")

    # Vecais QR koda bloks, kas izmantoja neeksistējošus atribūtus.
    # Šis bloks ir jāizdzēš vai jākomentē, jo jaunā funkcionalitāte to aizstāj.
    # Ja vēlaties, lai DOCX ģenerēšana atbalstītu abus QR kodus, jums būtu jāpievieno līdzīga loģika kā PDF ģenerēšanā.
    # Piemēram, ja vēlaties, lai tas izmantotu individuālo QR kodu:
    # if akta_dati.include_custom_qr_code and akta_dati.custom_qr_code_data:
    #     try:
    #         import qrcode
    #         qr = qrcode.QRCode(
    #             version=1,
    #             error_correction=qrcode.constants.ERROR_CORRECT_L,
    #             box_size=10,
    #             border=4,
    #         )
    #         qr.add_data(akta_dati.custom_qr_code_data)
    #         qr.make(fit=True)
    #         img_qr = qr.make_image(fill_color="black", back_color="white")
    #         qr_temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
    #         img_qr.save(qr_temp_path)
    #
    #         # DOCX attēla platums balstīts uz QR koda izmēru mm, konvertējot uz collām
    #         qr_width_inches = float(akta_dati.custom_qr_code_size_mm) / 25.4
    #         document.add_picture(qr_temp_path, width=Inches(qr_width_inches))
    #         document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Pēc noklusējuma labajā pusē
    #         os.remove(qr_temp_path)  # Clean up temp file
    #     except ImportError:
    #         print("QR Code generation requires 'qrcode' library. Please install it: pip install qrcode")
    #     except Exception as e:
    #         print(f"Error generating QR code for DOCX: {e}")

    add_formatted_text(document.add_paragraph(), f"Dokuments ģenerēts: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save(docx_ceļš)


# ---------------------- GUI ----------------------

# Custom URL interceptor for QWebEngineView
class MapUrlInterceptor(QWebEngineUrlRequestInterceptor):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.map_click_callback = None

    def interceptRequest(self, info):
        url = info.requestUrl()
        # Pārbaudām, vai URL ir mūsu pielāgotā shēma
        if url.scheme() == "app" and url.host() == "map_click":
            lat = url.queryItemValue("lat")
            lon = url.queryItemValue("lon")
            if self.map_click_callback:
                # Izsaucam atpakaļsaites funkciju ar koordinātēm
                self.map_click_callback(lat, lon)
            # Svarīgi: Bloķējam pieprasījumu, lai pārlūkprogramma nemēģinātu atvērt šo URL
            # un neradītu brīdinājumu.
            info.block(True)
            # info.redirect(QUrl()) # Šī rinda var nebūt nepieciešama, ja block(True) darbojas efektīvi
        # Ja URL nav mūsu pielāgotā shēma, ļaujam tam turpināties
        else:
            info.block(False)

# Pievienot šīs klases pirms AktaLogs klases definīcijas

class TextBlockLineEdit(QWidget):
    def __init__(self, text_block_manager, field_name, parent=None):
        super().__init__(parent)
        self.text_block_manager = text_block_manager
        self.field_name = field_name

        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.line_edit = QLineEdit()
        self.layout.addWidget(self.line_edit)

        self.combo_box = QComboBox()
        self.combo_box.addItem("--- Izvēlēties saglabāto bloku ---")
        self.combo_box.currentIndexChanged.connect(self._load_selected_block)
        self.layout.addWidget(self.combo_box)

        button_layout = QHBoxLayout()
        self.save_button = QPushButton("Saglabāt bloku")
        self.save_button.clicked.connect(self._save_block)
        button_layout.addWidget(self.save_button)

        self.delete_button = QPushButton("Dzēst bloku")
        self.delete_button.clicked.connect(self._delete_block)
        button_layout.addWidget(self.delete_button)
        button_layout.addStretch()
        self.layout.addLayout(button_layout)

        self._update_combo_box()

    def text(self):
        return self.line_edit.text()

    def setText(self, text):
        self.line_edit.setText(text)

    def _update_combo_box(self):
        self.combo_box.blockSignals(True) # Bloķējam signālus, lai izvairītos no nevajadzīgas ielādes
        self.combo_box.clear()
        self.combo_box.addItem("--- Izvēlēties saglabāto bloku ---")
        blocks = self.text_block_manager.get_blocks_for_field(self.field_name)
        for name in sorted(blocks.keys()):
            self.combo_box.addItem(name)
        self.combo_box.blockSignals(False)

    def _load_selected_block(self, index):
        if index > 0:
            block_name = self.combo_box.currentText()
            content = self.text_block_manager.get_block_content(self.field_name, block_name)
            self.line_edit.setText(content)
            QMessageBox.information(self, "Ielādēts", f"Teksta bloks '{block_name}' ielādēts.")
        self.combo_box.setCurrentIndex(0) # Atgriežam izvēli uz noklusējumu

    def _save_block(self):
        current_text = self.line_edit.text().strip()
        if not current_text:
            QMessageBox.warning(self, "Saglabāt bloku", "Ievades lauks ir tukšs. Lūdzu, ievadiet tekstu, ko saglabāt.")
            return

        block_name, ok = QInputDialog.getText(self, "Saglabāt teksta bloku", "Ievadiet bloka nosaukumu:")
        if ok and block_name:
            self.text_block_manager.add_block(self.field_name, block_name, current_text)
            self._update_combo_box()
            QMessageBox.information(self, "Saglabāts", f"Teksta bloks '{block_name}' saglabāts.")
        elif ok:
            QMessageBox.warning(self, "Saglabāt bloku", "Bloka nosaukums nevar būt tukšs.")

    def _delete_block(self):
        blocks = self.text_block_manager.get_blocks_for_field(self.field_name)
        if not blocks:
            QMessageBox.information(self, "Dzēst bloku", "Nav saglabātu teksta bloku šim laukam.")
            return

        block_names = sorted(blocks.keys())
        block_name, ok = QInputDialog.getItem(self, "Dzēst teksta bloku", "Izvēlieties bloku, ko dzēst:", block_names, 0, False)
        if ok and block_name:
            reply = QMessageBox.question(self, "Dzēst bloku",
                                         f"Vai tiešām vēlaties dzēst teksta bloku '{block_name}'?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.text_block_manager.delete_block(self.field_name, block_name)
                self._update_combo_box()
                QMessageBox.information(self, "Dzēsts", f"Teksta bloks '{block_name}' dzēsts.")



class DraggablePreviewLabel(QLabel):
    """QLabel, kas ļauj ar peles vilcienu pārvietot PDF priekšskatījumu (pan) pat pie zoom."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._scroll_area = None
        self._dragging = False
        self._last_pos = None
        self.setCursor(Qt.OpenHandCursor)

    def set_scroll_area(self, sa: QScrollArea):
        self._scroll_area = sa

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton and self._scroll_area is not None:
            self._dragging = True
            self._last_pos = event.globalPosition().toPoint()
            self.setCursor(Qt.ClosedHandCursor)
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._dragging and self._scroll_area is not None and self._last_pos is not None:
            p = event.globalPosition().toPoint()
            dx = p.x() - self._last_pos.x()
            dy = p.y() - self._last_pos.y()
            self._last_pos = p

            h = self._scroll_area.horizontalScrollBar()
            v = self._scroll_area.verticalScrollBar()
            h.setValue(h.value() - dx)
            v.setValue(v.value() - dy)
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._dragging = False
            self._last_pos = None
            self.setCursor(Qt.OpenHandCursor)
            event.accept()
            return
        super().mouseReleaseEvent(event)



class PannableScrollArea(QScrollArea):
    """QScrollArea ar 'hand-drag' panning PDF priekšskatījumam (strādā arī, ja klikšķis ir tukšajā zonā)."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._dragging = False
        self._last_pos = None
        self.setCursor(Qt.OpenHandCursor)
        # Centrs, ja saturs mazāks par viewport
        try:
            self.setAlignment(Qt.AlignCenter)
        except Exception:
            pass

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._dragging = True
            self._last_pos = event.globalPosition().toPoint()
            self.setCursor(Qt.ClosedHandCursor)
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._dragging and self._last_pos is not None:
            p = event.globalPosition().toPoint()
            dx = p.x() - self._last_pos.x()
            dy = p.y() - self._last_pos.y()
            self._last_pos = p

            h = self.horizontalScrollBar()
            v = self.verticalScrollBar()
            h.setValue(h.value() - dx)
            v.setValue(v.value() - dy)
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._dragging = False
            self._last_pos = None
            self.setCursor(Qt.OpenHandCursor)
            event.accept()
            return
        super().mouseReleaseEvent(event)

class TextBlockTextEdit(QWidget):
    def __init__(self, text_block_manager, field_name, parent=None):
        super().__init__(parent)
        self.text_block_manager = text_block_manager
        self.field_name = field_name

        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.text_edit = QTextEdit()
        self.layout.addWidget(self.text_edit)

        self.list_widget = QListWidget()
        self.list_widget.setMaximumHeight(100) # Ierobežojam augstumu, lai neaizņemtu pārāk daudz vietas
        self.list_widget.itemDoubleClicked.connect(self._load_selected_block)
        self.layout.addWidget(self.list_widget)

        button_layout = QHBoxLayout()
        self.save_button = QPushButton("Saglabāt bloku")
        self.save_button.clicked.connect(self._save_block)
        button_layout.addWidget(self.save_button)

        self.delete_button = QPushButton("Dzēst bloku")
        self.delete_button.clicked.connect(self._delete_block)
        button_layout.addWidget(self.delete_button)
        button_layout.addStretch()
        self.layout.addLayout(button_layout)

        self._update_list_widget()

    def toPlainText(self):
        return self.text_edit.toPlainText()

    def setText(self, text):
        self.text_edit.setText(text)

    def _update_list_widget(self):
        self.list_widget.clear()
        blocks = self.text_block_manager.get_blocks_for_field(self.field_name)
        for name in sorted(blocks.keys()):
            self.list_widget.addItem(name)

    def _load_selected_block(self, item):
        block_name = item.text()
        content = self.text_block_manager.get_block_content(self.field_name, block_name)
        self.text_edit.setPlainText(content)
        QMessageBox.information(self, "Ielādēts", f"Teksta bloks '{block_name}' ielādēts.")


    # Proxy helper, so this widget can be used like a QTextEdit in the rest of the code
    def setPlaceholderText(self, text: str):
        """Set placeholder text on the internal QTextEdit (if supported by Qt version)."""
        if hasattr(self.text_edit, "setPlaceholderText"):
            self.text_edit.setPlaceholderText(text)
        else:
            # Older Qt versions: QTextEdit might not support placeholder text
            # (keep silent to avoid crashing)
            pass

    def setPlainText(self, text: str):
        self.text_edit.setPlainText(text)

    def toPlainText(self) -> str:
        return self.text_edit.toPlainText()

    def _save_block(self):
        current_text = self.text_edit.toPlainText().strip()
        if not current_text:
            QMessageBox.warning(self, "Saglabāt bloku", "Ievades lauks ir tukšs. Lūdzu, ievadiet tekstu, ko saglabāt.")
            return

        block_name, ok = QInputDialog.getText(self, "Saglabāt teksta bloku", "Ievadiet bloka nosaukumu:")
        if ok and block_name:
            self.text_block_manager.add_block(self.field_name, block_name, current_text)
            self._update_list_widget()
            QMessageBox.information(self, "Saglabāts", f"Teksta bloks '{block_name}' saglabāts.")
        elif ok:
            QMessageBox.warning(self, "Saglabāt bloku", "Bloka nosaukums nevar būt tukšs.")

    def _delete_block(self):
        selected_item = self.list_widget.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "Dzēst bloku", "Lūdzu, izvēlieties bloku, ko dzēst.")
            return

        block_name = selected_item.text()
        reply = QMessageBox.question(self, "Dzēst bloku",
                                     f"Vai tiešām vēlaties dzēst teksta bloku '{block_name}'?",
                                     QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.text_block_manager.delete_block(self.field_name, block_name)
            self._update_list_widget()
            QMessageBox.information(self, "Dzēsts", f"Teksta bloks '{block_name}' dzēsts.")


# ---------------------- Priekšskatījuma ģenerēšana fonā (lai DOCX/XLSX nekad neuzkar UI) ----------------------
class _PreviewBuildWorker(QObject):
    """Ģenerē priekšskatījuma PDF un pārvērš to PNG baitos fonā.

    Svarīgi: QPixmap nedrīkst veidot fonā, tāpēc worker atgriež PNG baitu sarakstu.
    """

    finished = Signal(str, list, int)  # data_hash, png_bytes_list, old_page
    failed = Signal(str, str)  # data_hash, error_message

    def __init__(self, d: 'AktaDati', data_hash: str, old_page: int):
        super().__init__()
        # --- Settings (persist across restarts) ---
        self._qt_settings = QSettings("AktaGenerators", "AktaGeneratorsApp")
        self._settings = load_settings()

        self._d = d
        self._hash = data_hash
        self._old_page = old_page

    def run(self):
        temp_pdf_path = None
        try:
            # Priekšskatījumā iekļaujam arī atsauces dokumentus (DOCX/XLSX u.c.), bet fonā, lai UI neuzkar.
            temp_pdf_path = ģenerēt_pdf(self._d, pdf_ceļš=None, include_reference_docs=True, encrypt_pdf=False)

            poppler_path_to_use = self._d.poppler_path if getattr(self._d, 'poppler_path', None) and os.path.exists(self._d.poppler_path) else None
            images = convert_from_path(temp_pdf_path, poppler_path=poppler_path_to_use)

            out_bytes = []
            from io import BytesIO
            for pil_img in images:
                bio = BytesIO()
                pil_img.save(bio, format='PNG')
                out_bytes.append(bio.getvalue())

            self.finished.emit(self._hash, out_bytes, self._old_page)
        except Exception as e:
            self.failed.emit(self._hash, str(e))
        finally:
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                except Exception:
                    pass


class AktaLogs(QMainWindow):
    def __init__(self):
        super().__init__()


        # Ikona logam + Taskbar (Windows)
        try:
            self.setWindowIcon(QIcon(resource_path("Akta_Generators_Icon.ico")))
        except Exception:
            pass

        # --- JAUNS: Audit + Undo/Redo ---
        self._current_user = os.getenv("USERNAME") or os.getenv("USER") or ""
        self._audit_logger = AuditLogger(os.path.join(APP_DATA_DIR, "audit_log.jsonl"))
        self._undo_mgr = UndoRedoManager(max_steps=80)

        # Ātri undo/redo (Ctrl+Z / Ctrl+Y)
        self._act_undo = QAction("Undo", self)
        self._act_undo.setShortcut("Ctrl+Z")
        self._act_undo.triggered.connect(self.undo_action)

        self._act_redo = QAction("Redo", self)
        self._act_redo.setShortcut("Ctrl+Y")
        self._act_redo.triggered.connect(self.redo_action)

        self.addAction(self._act_undo)
        self.addAction(self._act_redo)

        # Status bar: Undo/Redo indikatori
        try:
            sb = self.statusBar()
            self._undo_status_label = QLabel("Undo: 0")
            self._redo_status_label = QLabel("Redo: 0")
            self._undo_status_label.setMinimumWidth(90)
            self._redo_status_label.setMinimumWidth(90)
            sb.addPermanentWidget(self._undo_status_label)
            sb.addPermanentWidget(self._redo_status_label)
        except Exception:
            self._undo_status_label = None
            self._redo_status_label = None

        # Izmaiņu izsekošana laukiem (audit + undo checkpoints)
        self._track_enabled = True
        self._last_widget_values = {}
        try:
            app = QApplication.instance()
            if app is not None:
                app.installEventFilter(self)
        except Exception:
            pass

        self._update_undo_redo_indicators()
        self.setWindowTitle("Pieņemšanas–Nodošanas akta ģenerators")
        self.resize(1200, 800)
        self.poppler_path = ""
        self.zoom_factor = 1.0
        self.history = [] # Inicializējam tukšu sarakstu
        self.address_book = {} # Inicializējam tukšu vārdnīcu
        self.text_block_manager = TextBlockManager() # JAUNA RINDAS
        self.data = AktaDati(
            datums=datetime.now().strftime('%Y-%m-%d'),
            pieņemšanas_datums=datetime.now().strftime('%Y-%m-%d'),
            nodošanas_datums=datetime.now().strftime('%Y-%m-%d'),
            izpildes_termiņš=(datetime.now() + timedelta(days=5)).strftime('%Y-%m-%d'),
        )
        self._ceļš_projekts = None
        # Jauni atribūti ātruma uzlabošanai
        self.preview_timer = QTimer(self)
        self.preview_timer.setSingleShot(True)
        self._map_address_target = None  # QLineEdit, kurā ieliekam adresi no kartes
        self.preview_timer.timeout.connect(self._do_update_preview)
        self.preview_cache = {}  # Kešatmiņa: {'data_hash': {'images': [...], 'page_count': int}}
        self.last_data_hash = None  # Pēdējais datu hash

        # Priekšskatījuma ģenerēšana fonā (DOCX/XLSX konvertācija u.c.)
        self._preview_thread: Optional[QThread] = None
        self._preview_worker: Optional[_PreviewBuildWorker] = None
        self._requested_preview_hash: Optional[str] = None
        self._pending_preview_request = None  # (AktaDati, data_hash, old_page)

        self.tabs = QTabWidget()

        self._būvēt_pamata_tab()
        self._būvēt_puses_tab()
        self._būvēt_pozīcijas_tab()
        self._būvēt_attēli_tab()
        self._būvēt_iestatījumi_tab()
        self._būvēt_papildu_iestatījumi_tab()
        self._būvēt_sablonu_tab()
        self._būvēt_adresu_gramata_tab()
        self._būvēt_audit_tab()
        self._būvēt_dokumentu_vesture_tab()
        self._būvēt_kartes_tab()  # New map tab
        # Pārliecināmies, ka noklusējuma šablonu direktorijs ir iestatīts
        if not self.data.templates_dir:
            self.data.templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")
            os.makedirs(self.data.templates_dir, exist_ok=True)  # Izveidojam noklusējuma mapi, ja tā neeksistē
        self._update_sablonu_list()

        # Ielādējam vēsturi un adrešu grāmatu PĒC GUI elementu izveides
        self._load_history()
        self._load_address_book()
        self._update_address_book_list()
        self.ieladet_noklusejuma_iestatijumus()



        # --- FIRST RUN: automātiski ielādē iebūvēto šablonu "Testa dati (piemērs)" tikai pirmajā palaišanas reizē ---
        QTimer.singleShot(0, self._auto_load_test_template_first_run)
        main_splitter = QSplitter(Qt.Horizontal)
        self.setCentralWidget(main_splitter)

        main_splitter.addWidget(self.tabs)

        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        self.preview_label = QLabel("PDF priekšskatījums")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumSize(1, 1)
        self.preview_label.setAttribute(Qt.WA_TransparentForMouseEvents, True)

        # Pannable scroll area, lai var vilkt (hand-drag) arī tukšajā zonā
        self.preview_scroll_area = PannableScrollArea()
        self.preview_scroll_area.setWidgetResizable(False)
        self.preview_scroll_area.setWidget(self.preview_label)

        preview_layout.addWidget(self.preview_scroll_area)

        self.preview_scroll_area.resizeEvent = self._update_preview_on_resize

        self.current_preview_page = 0
        self.preview_images = []

        nav_buttons_layout = QHBoxLayout()
        self.prev_page_button = QPushButton("< Iepriekšējā lapa")
        self.prev_page_button.clicked.connect(self._show_prev_page)
        self.prev_page_button.setEnabled(False)
        self.next_page_button = QPushButton("Nākamā lapa >")
        self.next_page_button.clicked.connect(self._show_next_page)
        self.next_page_button.setEnabled(False)
        self.page_number_label = QLabel("Lapa 1/1")
        self.page_number_label.setAlignment(Qt.AlignCenter)

        nav_buttons_layout.addWidget(self.prev_page_button)
        nav_buttons_layout.addWidget(self.page_number_label)
        nav_buttons_layout.addWidget(self.next_page_button)

        zoom_in_button = QPushButton("Palielināt")
        zoom_in_button.clicked.connect(self.zoom_in)
        zoom_out_button = QPushButton("Samazināt")
        zoom_out_button.clicked.connect(self.zoom_out)
        nav_buttons_layout.addWidget(zoom_in_button)
        nav_buttons_layout.addWidget(zoom_out_button)

        preview_layout.addLayout(nav_buttons_layout)

        main_splitter.addWidget(preview_widget)

        main_splitter.setSizes([int(self.width() * 0.7), int(self.width() * 0.3)])

        self.tabs.currentChanged.connect(self._update_preview)

        self._update_preview()

        self._būvēt_menu()
        self._būvēt_toolbar()

        if self.data.auto_generate_akta_nr:
            self._generate_akta_nr()

    # ----- Menu -----

    def _auto_load_test_template_first_run(self):
        """Pirmajā palaišanas reizē automātiski ielādē iebūvēto šablonu 'Testa dati (piemērs)'.

        Mehānisms izmanto settings.json karodziņu 'first_run_test_template_loaded', lai ielāde notiktu tikai 1x.
        Ja šablonu saraksts vēl nav gatavs, karodziņš netiek uzstādīts un ielāde tiks mēģināta nākamajā startā.
        """
        try:
            st = load_settings() or {}
            if st.get("first_run_test_template_loaded", False):
                return

            target_name = "Testa dati (piemērs)"
            target_item = None

            # Šabloni ir listē (sablonu_list), ko programma jau aizpilda ar _update_sablonu_list()
            if hasattr(self, "sablonu_list") and self.sablonu_list is not None:
                for i in range(self.sablonu_list.count()):
                    it = self.sablonu_list.item(i)
                    if it and it.text().strip() == target_name:
                        target_item = it
                        break

            if target_item is None:
                return

            # Ielādējam tāpat kā lietotājs manuāli (dubultklikšķis uz šablona)
            self.ieladet_sablonu(target_item)

            # Atzīmējam, ka auto-load ir izpildīts
            st["first_run_test_template_loaded"] = True
            save_settings(st)

            # Atsvaidzinām priekšskatījumu (ja preview jau ir uzbūvēts)
            try:
                self._update_preview()
            except Exception:
                pass

        except Exception as e:
            # Nekad nekritinām aplikāciju auto-load dēļ
            print(f"Auto-load testa šablonam neizdevās: {e}")

    def _būvēt_menu(self):
        menubar = self.menuBar()
        fajls = menubar.addMenu("&Fails")

        # Parakstīšana (eParaksts)
        act_sign_now = QAction("Ģenerēt + Parakstīt (eParaksts)…", self)
        act_sign_now.triggered.connect(self.generate_and_sign_current)

        act_sign_last = QAction("Parakstīt pēdējo PDF…", self)
        act_sign_last.triggered.connect(lambda: self.sign_file_with_eparaksts(None))

        fajls.addSeparator()
        fajls.addAction(act_sign_now)
        fajls.addAction(act_sign_last)

        iestatijumi = menubar.addMenu("&Iestatījumi")
        act_set_ep = QAction("eParaksts…", self)
        act_set_ep.triggered.connect(self._open_settings_eparaksts)
        iestatijumi.addAction(act_set_ep)

        rediget = menubar.addMenu("&Rediģēt")
        rediget.addAction(self._act_undo)
        rediget.addAction(self._act_redo)


        saglabat = QAction("Saglabāt projektu…", self)
        saglabat.triggered.connect(self.saglabat_projektu)
        fajls.addAction(saglabat)

        ieladet = QAction("Ielādēt projektu…", self)
        ieladet.triggered.connect(self.ieladet_projektu)
        fajls.addAction(ieladet)

        fajls.addSeparator()

        eksportet_pdf = QAction("Ģenerēt PDF…", self)
        eksportet_pdf.triggered.connect(self.ģenerēt_pdf_dialogs)
        fajls.addAction(eksportet_pdf)

        eksportet_docx = QAction("Ģenerēt DOCX…", self)
        eksportet_docx.triggered.connect(self.ģenerēt_docx_dialogs)
        fajls.addAction(eksportet_docx)

        drukāt_pdf_menu = QAction("Drukāt PDF…", self) # JAUNA RINDAS
        drukāt_pdf_menu.triggered.connect(self.drukāt_pdf_dialogs) # JAUNA RINDAS
        fajls.addAction(drukāt_pdf_menu) # JAUNA RINDAS


        fajls.addSeparator()

        iziet = QAction("Iziet", self)
        iziet.triggered.connect(self.close)
        fajls.addAction(iziet)

        # Papildu izvēlnes
        riki = menubar.addMenu("&Rīki")

        atjaunot = QAction("Atjaunot priekšskatījumu", self)
        atjaunot.setShortcut("F5")
        atjaunot.triggered.connect(self._update_preview)
        riki.addAction(atjaunot)

        atvert_iestat = QAction("Atvērt iestatījumu mapi", self)
        atvert_iestat.triggered.connect(self._open_settings_folder)
        riki.addAction(atvert_iestat)

        riki.addSeparator()
        notirit_kesu = QAction("Notīrīt priekšskatījuma kešu", self)
        notirit_kesu.triggered.connect(self._clear_preview_cache)
        riki.addAction(notirit_kesu)

        palidziba = menubar.addMenu("&Palīdzība")
        par = QAction("Par programmu", self)
        par.triggered.connect(self._show_about_dialog)
        palidziba.addAction(par)

        isceli = QAction("Īsceļi", self)
        isceli.triggered.connect(self._show_shortcuts_dialog)
        palidziba.addAction(isceli)

        palidziba.addSeparator()
        atvert_map = QAction("Atvērt programmas mapi", self)
        atvert_map.triggered.connect(self._open_app_folder)
        palidziba.addAction(atvert_map)


    # ----- Palīgfunkcijas izvēlnēm -----
    def _open_settings_folder(self):
        try:
            os.makedirs(SETTINGS_DIR, exist_ok=True)
            os.startfile(SETTINGS_DIR)  # type: ignore[attr-defined]
        except Exception:
            try:
                import subprocess
                if sys.platform.startswith("darwin"):
                    subprocess.Popen(["open", SETTINGS_DIR])
                else:
                    subprocess.Popen(["xdg-open", SETTINGS_DIR])
            except Exception:
                QMessageBox.information(self, "Iestatījumi", f"Iestatījumu mape:\n{SETTINGS_DIR}")

    def _open_app_folder(self):
        try:
            app_dir = os.path.dirname(os.path.abspath(__file__))
            os.startfile(app_dir)  # type: ignore[attr-defined]
        except Exception:
            try:
                import subprocess
                app_dir = os.path.dirname(os.path.abspath(__file__))
                if sys.platform.startswith("darwin"):
                    subprocess.Popen(["open", app_dir])
                else:
                    subprocess.Popen(["xdg-open", app_dir])
            except Exception:
                QMessageBox.information(self, "Mape", os.path.dirname(os.path.abspath(__file__)))

    def _clear_preview_cache(self):
        try:
            if hasattr(self, "preview_cache") and isinstance(self.preview_cache, dict):
                self.preview_cache.clear()
            if hasattr(self, "last_data_hash"):
                self.last_data_hash = None
            self.statusBar().showMessage("Priekšskatījuma kešs notīrīts", 3000)
        except Exception:
            pass
        self._update_preview()

    def _show_about_dialog(self):
        about_txt = (
            "Akta ģenerators\n"
            "\n"
            f"Iestatījumu mape: {SETTINGS_DIR}\n"
            "\n"
            "Šī programma palīdz veidot pieņemšanas–nodošanas aktus un eksportēt PDF/DOCX."
        )
        QMessageBox.information(self, "Par programmu", about_txt)

    def _show_shortcuts_dialog(self):
        shortcuts_txt = (
            "Īsceļi\n"
            "\n"
            "F5      — atjaunot priekšskatījumu\n"
            "Ctrl+S  — saglabāt projektu\n"
            "Ctrl+O  — ielādēt projektu\n"
            "Ctrl+P  — ģenerēt PDF\n"
        )
        QMessageBox.information(self, "Īsceļi", shortcuts_txt)

    def _būvēt_toolbar(self):
        tb = self.addToolBar("Galvenais")
        tb.setMovable(False)

        act_save = QAction("Saglabāt", self)
        act_save.setShortcut("Ctrl+S")
        act_save.triggered.connect(self.saglabat_projektu)
        tb.addAction(act_save)

        act_load = QAction("Ielādēt", self)
        act_load.setShortcut("Ctrl+O")
        act_load.triggered.connect(self.ieladet_projektu)
        tb.addAction(act_load)

        tb.addSeparator()

        act_pdf = QAction("PDF", self)
        act_pdf.setShortcut("Ctrl+P")
        act_pdf.triggered.connect(self.ģenerēt_pdf_dialogs)
        tb.addAction(act_pdf)

        act_docx = QAction("DOCX", self)
        act_docx.triggered.connect(self.ģenerēt_docx_dialogs)
        tb.addAction(act_docx)

        tb.addSeparator()

        act_print = QAction("Drukāt", self)
        act_print.triggered.connect(self.drukāt_pdf_dialogs)
        tb.addAction(act_print)

        self.statusBar().showMessage("Gatavs")


    # ----- Tab: Pamata -----

    def _būvēt_pamata_tab(self):
        content_widget = QWidget()
        form = QFormLayout()

        self.in_akta_nr = QLineEdit()
        self.btn_generate_akta_nr = QPushButton("Ģenerēt Nr.")
        self.btn_generate_akta_nr.clicked.connect(self._generate_akta_nr)
        # Mazs, neuzkrītošs "restart" taustiņš akta numura skaitītājam
        self.btn_reset_akta_nr = QPushButton("↺")
        self.btn_reset_akta_nr.setToolTip("Pārstartēt akta numura skaitītāju (atgriezt uz 1)")
        self.btn_reset_akta_nr.setFixedSize(34, 34)
        self.btn_reset_akta_nr.setFlat(True)
        self.btn_reset_akta_nr.clicked.connect(self._reset_akta_nr_counter)
        akta_nr_layout = QHBoxLayout()
        akta_nr_layout.addWidget(self.in_akta_nr)
        akta_nr_layout.addWidget(self.btn_generate_akta_nr)
        akta_nr_layout.addWidget(self.btn_reset_akta_nr)
        akta_nr_widget = QWidget();
        akta_nr_widget.setLayout(akta_nr_layout)

        self.in_datums = QDateEdit(calendarPopup=True)
        self.in_datums.setDisplayFormat("yyyy-MM-dd")
        self.in_datums.setDate(datetime.now().date())

        self.in_vieta = QLineEdit()
        self.in_pas_nr = QLineEdit()
        self.in_liguma_nr = QLineEdit()

        self.in_izpildes_termins = QDateEdit(calendarPopup=True)
        self.in_izpildes_termins.setDisplayFormat("yyyy-MM-dd")
        self.in_izpildes_termins.setMinimumDate(datetime(1900, 1, 1).date())
        self.in_izpildes_termins.setMaximumDate(datetime(2100, 1, 1).date())
        self.in_izpildes_termins.setDate((datetime.now() + timedelta(days=int(getattr(self.data, 'default_execution_days', 5)))).date())  # default: today + N days
        self.in_izpildes_termins.setSpecialValueText("Nav norādīts") # Text for null date
        self.in_izpildes_termins.setCalendarPopup(True)


        self.in_pieņemšanas_datums = QDateEdit(calendarPopup=True)
        self.in_pieņemšanas_datums.setDisplayFormat("yyyy-MM-dd")
        self.in_pieņemšanas_datums.setMinimumDate(datetime(1900, 1, 1).date())
        self.in_pieņemšanas_datums.setMaximumDate(datetime(2100, 1, 1).date())
        self.in_pieņemšanas_datums.setDate(datetime.now().date())
        self.in_pieņemšanas_datums.setSpecialValueText("Nav norādīts")
        self.in_pieņemšanas_datums.setCalendarPopup(True)

        self.in_nodošanas_datums = QDateEdit(calendarPopup=True)
        self.in_nodošanas_datums.setDisplayFormat("yyyy-MM-dd")
        self.in_nodošanas_datums.setMinimumDate(datetime(1900, 1, 1).date())
        self.in_nodošanas_datums.setMaximumDate(datetime(2100, 1, 1).date())
        self.in_nodošanas_datums.setDate(datetime.now().date())
        self.in_nodošanas_datums.setSpecialValueText("Nav norādīts")
        self.in_nodošanas_datums.setCalendarPopup(True)

        # Strīdu risināšana
        self.in_strīdu_risināšana = TextBlockTextEdit(self.text_block_manager, "stridu_risinasana")

        self.ck_konfidencialitate = QCheckBox("Konfidencialitāte")
        self.in_soda_nauda_procenti = QDoubleSpinBox();
        self.in_soda_nauda_procenti.setRange(0.0, 100.0);
        self.in_soda_nauda_procenti.setSuffix(" %")

        # Piegādes nosacījumi
        self.in_piegades_nosacijumi = TextBlockLineEdit(self.text_block_manager, "piegades_nosacijumi")

        self.ck_apdrošināšana = QCheckBox("Apdrošināšana")

        # Papildu nosacījumi
        self.in_papildu_nosacijumi = TextBlockTextEdit(self.text_block_manager, "papildu_nosacijumi")

        # Atsauces dokumenti
        self.in_atsauces_dokumenti = TextBlockLineEdit(self.text_block_manager, "atsauces_dokumenti")
        self.list_atsauces_faili = QListWidget()
        self.list_atsauces_faili.setMinimumHeight(90)
        self.list_atsauces_faili.setSelectionMode(QAbstractItemView.SingleSelection)
        self.btn_add_atsauce_failu = QPushButton("Pievienot atsauces failu")
        self.btn_remove_atsauce_failu = QPushButton("Noņemt izvēlēto")
        self.btn_add_atsauce_failu.clicked.connect(self._add_reference_doc)
        self.btn_remove_atsauce_failu.clicked.connect(self._remove_reference_doc)

        self.cb_akta_statuss = QComboBox()
        self.cb_akta_statuss.addItems(["Melnraksts", "Apstiprināts", "Parakstīts", "Arhivēts", "Atcelts"])
        self.in_valuta = QComboBox()
        self.in_valuta.setEditable(True)
        self.in_valuta.addItems(["EUR €","USD $","GBP £","NOK kr","SEK kr","DKK kr","PLN zł","CHF CHF","CZK Kč","HUF Ft","RON lei","BGN лв","JPY ¥","CNY ¥","AUD $","CAD $","NZD $","TRY ₺","UAH ₴","RUB ₽"]) 
        self.in_valuta.setCurrentText("EUR €")

        # Piezīmes
        self.in_piezimes = TextBlockTextEdit(self.text_block_manager, "piezimes")

        self.ck_elektroniskais_paraksts = QCheckBox("Elektroniskais paraksts (ignorē fiziskos parakstus)")
        self.ck_radit_elektronisko_parakstu_tekstu = QCheckBox("Rādīt elektroniskā paraksta tekstu PDF dokumentā") # JAUNA RŪTIŅA

        form.addRow("Akta Nr.", akta_nr_widget)
        form.addRow("Datums", self._wrap_date_with_system_button(self.in_datums))
        form.addRow("Vieta", self.in_vieta)
        form.addRow("Pasūtījuma Nr.", self.in_pas_nr)
        form.addRow("Līguma Nr.", self.in_liguma_nr)
        form.addRow("Izpildes termiņš", self._wrap_date_with_system_button(self.in_izpildes_termins))
        form.addRow("Pieņemšanas datums", self._wrap_date_with_system_button(self.in_pieņemšanas_datums))
        form.addRow("Nodošanas datums", self._wrap_date_with_system_button(self.in_nodošanas_datums))
        form.addRow("Strīdu risināšana", self.in_strīdu_risināšana)  # Tagad tieši izmantojam jauno objektu
        form.addRow(self.ck_konfidencialitate)
        form.addRow("Soda nauda (%)", self.in_soda_nauda_procenti)
        form.addRow("Piegādes nosacījumi", self.in_piegades_nosacijumi)  # Tagad tieši izmantojam jauno objektu
        form.addRow(self.ck_apdrošināšana)

        # Apdrošināšanas teksts (kā teksta bloks – līdzīgi kā citur)
        self.in_apdrosinasana_teksts = TextBlockTextEdit(self.text_block_manager, "apdrosinasana_teksts")
        self.in_apdrosinasana_teksts.setPlaceholderText("Ierakstiet apdrošināšanas tekstu…")
        form.addRow("Apdrošināšanas teksts", self.in_apdrosinasana_teksts)

        # Rādīt/Slēpt teksta lauku atkarībā no atzīmes
        def _toggle_apdrosinasana_text():
            show = self.ck_apdrošināšana.isChecked()
            self.in_apdrosinasana_teksts.setVisible(show)
            # QLabel no QFormLayout nav tieši pieejams; paslēpjam arī etiķeti
            try:
                lbl = form.labelForField(self.in_apdrosinasana_teksts)
                if lbl:
                    lbl.setVisible(show)
            except Exception:
                pass
        self.ck_apdrošināšana.stateChanged.connect(lambda *_: (_toggle_apdrosinasana_text(), self._update_preview()))
        _toggle_apdrosinasana_text()
        form.addRow("Papildu nosacījumi", self.in_papildu_nosacijumi)  # Tagad tieši izmantojam jauno objektu
        form.addRow("Atsauces dokumenti", self.in_atsauces_dokumenti)  # Tagad tieši izmantojam jauno objektu
        # Atsauces faili (reāli pielikumi PDF beigās)
        w_refs = QWidget()
        v_refs = QVBoxLayout(w_refs)
        btn_row = QHBoxLayout()
        btn_row.addWidget(self.btn_add_atsauce_failu)
        btn_row.addWidget(self.btn_remove_atsauce_failu)
        btn_row.addStretch(1)
        v_refs.addLayout(btn_row)
        v_refs.addWidget(self.list_atsauces_faili)
        form.addRow("Atsauces faili", w_refs)
        form.addRow("Akta statuss", self.cb_akta_statuss)
        
        # --- JAUNS: PDF šifrēšana (parole) Pamata datos ---
        # Priekšskatījums programmā joprojām tiek ģenerēts NEšifrēts (skat. encrypt_pdf=False preview worker),
        # bet eksportā/saglabāšanā šī parole tiks piemērota PDF failam.
        self.ck_pdf_encrypt_basic = QCheckBox("Šifrēt PDF ar paroli")
        try:
            self.ck_pdf_encrypt_basic.setChecked(bool(getattr(self.data, "enable_pdf_encryption", False)))
        except Exception:
            self.ck_pdf_encrypt_basic.setChecked(False)

        self.in_pdf_password_basic = QLineEdit()
        self.in_pdf_password_basic.setEchoMode(QLineEdit.Password)
        self.in_pdf_password_basic.setPlaceholderText("Parole (ja atstāj tukšu – atvērsies bez paroles)")
        try:
            self.in_pdf_password_basic.setText(str(getattr(self.data, "pdf_user_password", "") or ""))
        except Exception:
            pass

        # Neliekam spiest preview uzreiz – bet, ja lietotājs maina, lai eksportā viss būtu saglabāts.
        try:
            self.ck_pdf_encrypt_basic.stateChanged.connect(self._update_preview)
            self.in_pdf_password_basic.textChanged.connect(self._update_preview)
        except Exception:
            pass

        form.addRow(self.ck_pdf_encrypt_basic)
        form.addRow("PDF parole", self.in_pdf_password_basic)

        form.addRow("Valūta", self.in_valuta)
        form.addRow("Piezīmes", self.in_piezimes)  # Tagad tieši izmantojam jauno objektu
        form.addRow("Elektroniskais paraksts", self.ck_elektroniskais_paraksts)
        form.addRow("", self.ck_radit_elektronisko_parakstu_tekstu) # JAUNA RŪTIŅA FORMĀ
        # --- JAUNS: QR iestatījumu UI (droši, ja nav izveidots) ---
        if not hasattr(self, "ck_qr_kods"):
            self.ck_qr_kods = QCheckBox("QR kods apakšā pa kreisi (akta dati)")
            self.ck_qr_kods.setChecked(True)
        if not hasattr(self, "ck_qr_first_page"):
            self.ck_qr_first_page = QCheckBox("QR tikai pirmajā lapā")
            self.ck_qr_first_page.setChecked(True)
        if not hasattr(self, "ck_qr_url_mode"):
            self.ck_qr_url_mode = QCheckBox("QR kā verifikācijas saite (URL)")
            self.ck_qr_url_mode.setChecked(False)
        if not hasattr(self, "le_qr_url"):
            self.le_qr_url = QLineEdit()
            self.le_qr_url.setPlaceholderText("QR URL, piem.: https://kulinics.id.lv/verify")
        form.addRow("", self.ck_qr_kods)
        form.addRow("", self.ck_qr_first_page)
        form.addRow("", self.ck_qr_url_mode)
        form.addRow("QR URL", self.le_qr_url)
        # --- JAUNS: QR UI elementi (izveidojam, ja nav) ---
        if not hasattr(self, "ck_qr_kods"):
            self.ck_qr_kods = QCheckBox("QR kods apakšā pa kreisi (akta dati)")
            self.ck_qr_kods.setChecked(True)
        if not hasattr(self, "ck_qr_only_first"):
            self.ck_qr_only_first = QCheckBox("QR tikai pirmajā lapā")
            self.ck_qr_only_first.setChecked(True)
        if not hasattr(self, "ck_qr_url_mode"):
            self.ck_qr_url_mode = QCheckBox("QR kā verifikācijas saite (URL)")
            self.ck_qr_url_mode.setChecked(False)
        if not hasattr(self, "le_qr_url"):
            self.le_qr_url = QLineEdit()
            self.le_qr_url.setPlaceholderText("QR URL, piem.: https://kulinics.id.lv/verify")
        form.addRow("", self.ck_qr_only_first)
        # --- FIX: QR URL checkbox alias (dažādi nosaukumi) ---
        if not hasattr(self, "ck_qr_use_url") and hasattr(self, "ck_qr_url_mode"):
            self.ck_qr_use_url = self.ck_qr_url_mode
        elif not hasattr(self, "ck_qr_url_mode") and hasattr(self, "ck_qr_use_url"):
            self.ck_qr_url_mode = self.ck_qr_use_url
        elif not hasattr(self, "ck_qr_use_url") and not hasattr(self, "ck_qr_url_mode"):
            self.ck_qr_use_url = QCheckBox("QR kā verifikācijas saite (URL)")
            self.ck_qr_use_url.setChecked(False)
            self.ck_qr_url_mode = self.ck_qr_use_url
        form.addRow("", self.ck_qr_use_url)
        # --- FIX: QR URL lineEdit alias (dažādi nosaukumi) ---
        if not hasattr(self, "le_qr_base_url") and hasattr(self, "le_qr_url"):
            self.le_qr_base_url = self.le_qr_url
        elif not hasattr(self, "le_qr_url") and hasattr(self, "le_qr_base_url"):
            self.le_qr_url = self.le_qr_base_url
        elif not hasattr(self, "le_qr_base_url") and not hasattr(self, "le_qr_url"):
            self.le_qr_base_url = QLineEdit()
            self.le_qr_base_url.setPlaceholderText("QR URL, piem.: https://kulinics.id.lv/verify")
            self.le_qr_url = self.le_qr_base_url
        form.addRow("QR URL", self.le_qr_base_url)

        content_widget.setLayout(form)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(content_widget)

        main_tab_widget = QWidget()
        main_layout = QVBoxLayout(main_tab_widget)
        main_layout.addWidget(scroll_area)

        self.tabs.addTab(main_tab_widget, "Pamata dati")

        nav_buttons_layout = QHBoxLayout()
        zoom_in_button = QPushButton("Palielināt")
        zoom_in_button.clicked.connect(self.zoom_in)
        zoom_out_button = QPushButton("Samazināt")
        zoom_out_button.clicked.connect(self.zoom_out)

        nav_buttons_layout.addWidget(zoom_in_button)
        nav_buttons_layout.addWidget(zoom_out_button)

        main_layout.addLayout(nav_buttons_layout)

        self.in_akta_nr.textChanged.connect(self._update_preview)
        self.in_datums.dateChanged.connect(self._update_preview)
        self.in_vieta.textChanged.connect(self._update_preview)
        self.in_pas_nr.textChanged.connect(self._update_preview)
        self.in_liguma_nr.textChanged.connect(self._update_preview)
        self.in_izpildes_termins.dateChanged.connect(self._update_preview)
        self.in_pieņemšanas_datums.dateChanged.connect(self._update_preview)
        self.in_nodošanas_datums.dateChanged.connect(self._update_preview)
        self.in_strīdu_risināšana.text_edit.textChanged.connect(self._update_preview)  # Savienojam ar iekšējo text_edit
        self.ck_konfidencialitate.stateChanged.connect(self._update_preview)
        self.in_soda_nauda_procenti.valueChanged.connect(self._update_preview)
        self.in_piegades_nosacijumi.line_edit.textChanged.connect(
            self._update_preview)  # Savienojam ar iekšējo line_edit
        self.ck_apdrošināšana.stateChanged.connect(self._update_preview)
        self.in_papildu_nosacijumi.text_edit.textChanged.connect(
            self._update_preview)  # Savienojam ar iekšējo text_edit
        self.in_atsauces_dokumenti.line_edit.textChanged.connect(
            self._update_preview)  # Savienojam ar iekšējo line_edit
        self.cb_akta_statuss.currentIndexChanged.connect(self._update_preview)
        self.in_valuta.currentTextChanged.connect(self._update_preview)
        self.in_piezimes.text_edit.textChanged.connect(self._update_preview)  # Savienojam ar iekšējo text_edit
        self.ck_elektroniskais_paraksts.stateChanged.connect(self._update_preview)
        self.ck_radit_elektronisko_parakstu_tekstu.stateChanged.connect(self._update_preview) # JAUNS SAVIENOJUMS





    def zoom_in(self):
        self.zoom_factor *= 1.1
        self._show_current_page()

    def zoom_out(self):
        self.zoom_factor /= 1.1
        self._show_current_page()

    def _update_preview_on_resize(self, event):
        self._show_current_page() # Update preview to fit new size
        super().resizeEvent(event)

    def _generate_akta_nr(self):
        """Ģenerē jaunu unikālu akta numuru pēc shēmas PP-YYYY-NNNN.

        Uzlabojums:
        - uztur stabilu secīgo skaitītāju failā (AKTA_NR_COUNTER_FILE), lai nākamais numurs vienmēr būtu
          konsekvents pat tad, ja projekts vēl nav saglabāts vai history nav atjaunots.
        - papildus pārbauda jau izmantotos numurus projektos un vēsturē, lai nerastos dublikāti.
        """
        try:
            prefix = "PP"
            current_year = datetime.now().strftime('%Y')

            # ---------------- 1) Nolasa jau izmantotos numurus ----------------
            used_ids = set()

            # 1) Projekti (JSON) mapē
            if os.path.isdir(PROJECT_SAVE_DIR):
                for filename in os.listdir(PROJECT_SAVE_DIR):
                    if not filename.lower().endswith('.json'):
                        continue
                    fp = os.path.join(PROJECT_SAVE_DIR, filename)
                    try:
                        with open(fp, 'r', encoding='utf-8') as f:
                            project_data = json.load(f)
                        akta_nr = project_data.get('akta_nr') if isinstance(project_data, dict) else None
                        if akta_nr:
                            used_ids.add(str(akta_nr))
                    except Exception:
                        pass  # klusām ignorējam bojātus failus

            # 2) Vēsture (history.json)
            if os.path.exists(HISTORY_FILE):
                try:
                    with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                        history = json.load(f)
                except Exception:
                    history = None

                if isinstance(history, dict):
                    items = history.get('items') or history.get('history') or []
                else:
                    items = history or []

                if isinstance(items, list):
                    for item in items:
                        try:
                            if isinstance(item, dict):
                                ak = item.get('akta_nr')
                                if ak:
                                    used_ids.add(str(ak))

                                for k in ('json_path', 'project_path', 'path'):
                                    p = item.get(k)
                                    if isinstance(p, str) and p.lower().endswith('.json') and os.path.exists(p):
                                        try:
                                            with open(p, 'r', encoding='utf-8') as f:
                                                pd = json.load(f)
                                            ak2 = pd.get('akta_nr') if isinstance(pd, dict) else None
                                            if ak2:
                                                used_ids.add(str(ak2))
                                        except Exception:
                                            pass
                            elif isinstance(item, str) and item.lower().endswith('.json') and os.path.exists(item):
                                try:
                                    with open(item, 'r', encoding='utf-8') as f:
                                        pd = json.load(f)
                                    ak2 = pd.get('akta_nr') if isinstance(pd, dict) else None
                                    if ak2:
                                        used_ids.add(str(ak2))
                                except Exception:
                                    pass
                        except Exception:
                            pass

            # ---------------- 2) Aprēķina MAX izmantoto secību šim gadam ----------------
            seq_re = re.compile(rf'^{re.escape(prefix)}-(\d{{4}})-(\d{{4}})$')
            max_used_seq_this_year = 0
            for ak in used_ids:
                mm = seq_re.match(str(ak).strip())
                if not mm:
                    continue
                year, seq_s = mm.group(1), mm.group(2)
                if year != current_year:
                    continue
                try:
                    max_used_seq_this_year = max(max_used_seq_this_year, int(seq_s))
                except Exception:
                    pass

            # ---------------- 3) Nolasa / atjauno skaitītāju failā ----------------
            counters = {}
            if os.path.exists(AKTA_NR_COUNTER_FILE):
                try:
                    with open(AKTA_NR_COUNTER_FILE, 'r', encoding='utf-8') as f:
                        counters = json.load(f) if f else {}
                except Exception:
                    counters = {}

            last_seq = 0
            try:
                last_seq = int(counters.get(current_year, 0))
            except Exception:
                last_seq = 0

            # nodrošinām, ka skaitītājs nekad nav mazāks par reāli izmantoto max
            base_seq = max(last_seq, max_used_seq_this_year)

            # ---------------- 4) Atrod nākamo brīvo secību (bez dublikātiem) ----------------
            new_seq = base_seq + 1
            while True:
                new_akta_nr = f"{prefix}-{current_year}-{new_seq:04d}"
                if new_akta_nr not in used_ids:
                    # Uzreiz saglabājam skaitītāju, lai nākamais ģenerējums būtu konsekvents
                    counters[current_year] = new_seq
                    try:
                        os.makedirs(os.path.dirname(AKTA_NR_COUNTER_FILE), exist_ok=True)
                        with open(AKTA_NR_COUNTER_FILE, 'w', encoding='utf-8') as f:
                            json.dump(counters, f, ensure_ascii=False, indent=2)
                    except Exception:
                        pass

                    if hasattr(self, 'in_akta_nr') and self.in_akta_nr:
                        self.in_akta_nr.setText(new_akta_nr)
                    break
                new_seq += 1

        except Exception as e:
            print(f"Akta nr. ģenerēšanas kļūda: {e}")

    def _reset_akta_nr_counter(self):
        """Pārstartē akta numura skaitītāju, lai nākamais numurs sākas no 1.

        Piezīme: šī darbība apzināti neanalizē jau izmantotos numurus projektos/vēsturē.
        Ja vēlies pilnīgu unikālitāti, izmanto 'Ģenerēt Nr.' pēc reset.
        """
        try:
            prefix = "PP"
            current_year = datetime.now().strftime('%Y')

            counters = {}
            if os.path.exists(AKTA_NR_COUNTER_FILE):
                try:
                    with open(AKTA_NR_COUNTER_FILE, 'r', encoding='utf-8') as f:
                        counters = json.load(f) if f else {}
                except Exception:
                    counters = {}

            # Atgriež uz 1 un saglabā
            counters[current_year] = 1
            try:
                os.makedirs(os.path.dirname(AKTA_NR_COUNTER_FILE), exist_ok=True)
                with open(AKTA_NR_COUNTER_FILE, 'w', encoding='utf-8') as f:
                    json.dump(counters, f, ensure_ascii=False, indent=2)
            except Exception:
                pass

            # Uzstāda lauku uz 0001
            new_akta_nr = f"{prefix}-{current_year}-0001"
            if hasattr(self, 'in_akta_nr') and self.in_akta_nr:
                self.in_akta_nr.setText(new_akta_nr)

            try:
                if hasattr(self, 'statusBar') and callable(self.statusBar):
                    self.statusBar().showMessage("Akta numura skaitītājs pārstartēts uz 1", 3000)
            except Exception:
                pass

        except Exception as e:
            print(f"Akta nr. skaitītāja reset kļūda: {e}")


    def savākt_datus(self) -> AktaDati:
        d = AktaDati()
        d.akta_nr = self.in_akta_nr.text().strip()
        d.datums = self.in_datums.date().toString("yyyy-MM-dd")
        d.vieta = self.in_vieta.text().strip()
        d.pasūtījuma_nr = self.in_pas_nr.text().strip()
        d.piezīmes = self.in_piezimes.toPlainText().strip()

        # JAUNA RINDAS
        d.templates_dir = self.in_templates_dir.text().strip()

        d.līguma_nr = self.in_liguma_nr.text().strip()
        d.izpildes_termiņš = self.in_izpildes_termins.date().toString("yyyy-MM-dd") if not self.in_izpildes_termins.date().isNull() else ""
        d.pieņemšanas_datums = self.in_pieņemšanas_datums.date().toString("yyyy-MM-dd") if not self.in_pieņemšanas_datums.date().isNull() else ""
        d.nodošanas_datums = self.in_nodošanas_datums.date().toString("yyyy-MM-dd") if not self.in_nodošanas_datums.date().isNull() else ""
        d.strīdu_risināšana = self.in_strīdu_risināšana.toPlainText().strip()  # Tagad izsauc toPlainText() uz pielāgotā logrīka
        d.konfidencialitātes_klauzula = self.ck_konfidencialitate.isChecked()
        d.soda_nauda_procenti = to_decimal(self.in_soda_nauda_procenti.value())
        d.piegādes_nosacījumi = self.in_piegades_nosacijumi.text().strip()  # Tagad izsauc text() uz pielāgotā logrīka
        d.apdrošināšana = self.ck_apdrošināšana.isChecked()
        d.apdrošināšana_teksts = self.in_apdrosinasana_teksts.toPlainText().strip() if hasattr(self, 'in_apdrosinasana_teksts') else ""
        d.papildu_nosacījumi = self.in_papildu_nosacijumi.toPlainText().strip()  # Tagad izsauc toPlainText() uz pielāgotā logrīka
        d.atsauces_dokumenti = self.in_atsauces_dokumenti.text().strip()  # Tagad izsauc text() uz pielāgotā logrīka
        d.atsauces_dokumenti_faili = []
        for i in range(self.list_atsauces_faili.count()):
            it = self.list_atsauces_faili.item(i)
            p = it.data(Qt.UserRole)
            if p:
                d.atsauces_dokumenti_faili.append({"ceļš": p, "nosaukums": it.text()})

        d.akta_statuss = self.cb_akta_statuss.currentText()
        d.valūta = (self.in_valuta.currentText().strip().split(' ')[0] if self.in_valuta.currentText().strip() else "")
        d.elektroniskais_paraksts = self.ck_elektroniskais_paraksts.isChecked()
        d.radit_elektronisko_parakstu_tekstu = self.ck_radit_elektronisko_parakstu_tekstu.isChecked() # JAUNA RINDAS
        d.qr_kods_enabled = self.ck_qr_kods.isChecked() if hasattr(self, 'ck_qr_kods') else True
        d.qr_kods_tikai_pirma_lapa = self.ck_qr_only_first.isChecked() if hasattr(self, 'ck_qr_only_first') else True
        d.qr_kods_url_mode = (self.ck_qr_url_mode.isChecked() if hasattr(self, 'ck_qr_url_mode') else (self.ck_qr_use_url.isChecked() if hasattr(self, 'ck_qr_use_url') else False))
        d.qr_kods_url = (self.le_qr_url.text().strip() if hasattr(self, 'le_qr_url') else (self.le_qr_base_url.text().strip() if hasattr(self, 'le_qr_base_url') else ""))
        d.qr_only_first_page = self.ck_qr_only_first.isChecked() if hasattr(self, 'ck_qr_only_first') else False
        d.qr_verification_url_enabled = self.ck_qr_use_url.isChecked() if hasattr(self, 'ck_qr_use_url') else False
        d.qr_verification_base_url = (self.le_qr_base_url.text().strip() if hasattr(self, 'le_qr_base_url') else '')
        try:
            if hasattr(self, '_settings') and self._settings is not None:
                self._settings["qr_base_url"] = d.qr_verification_base_url
                save_settings(self._settings)
        except Exception:
            pass

        # Piezīmes
        d.piezīmes = self.in_piezimes.toPlainText().strip()  # Tagad izsauc toPlainText() uz pielāgotā logrīka

        # Puses
        pie = Persona(
            nosaukums=self.pie_in[0].text().strip(),
            reģ_nr=self.pie_in[1].text().strip(),
            adrese=self.pie_in[2].text().strip(),
            kontaktpersona=self.pie_in[3].text().strip(),
            amats=self.pie_in[4].text().strip(),
            pilnvaras_pamats=self.pie_in[5].currentText(),
            tālrunis=self.pie_in[6].text().strip(),
            epasts=self.pie_in[7].text().strip(),
            web_lapa=self.pie_in[8].text().strip(),
            bankas_konts=self.pie_in[9].text().strip(),
            juridiskais_statuss=self.pie_in[10].currentText()
            )
        nod = Persona(
                nosaukums=self.nod_in[0].text().strip(),
                reģ_nr=self.nod_in[1].text().strip(),
                adrese=self.nod_in[2].text().strip(),
                kontaktpersona=self.nod_in[3].text().strip(),
                amats=self.nod_in[4].text().strip(),
                pilnvaras_pamats=self.nod_in[5].currentText(),
                tālrunis=self.nod_in[6].text().strip(),
                epasts=self.nod_in[7].text().strip(),
                web_lapa=self.nod_in[8].text().strip(),
                bankas_konts=self.nod_in[9].text().strip(),
                juridiskais_statuss=self.nod_in[10].currentText()
            )
        d.pieņēmējs = pie
        d.nodevējs = nod

        # Pozīcijas
        poz = []
        idx = self._poz_col_indices()
        for r in range(self.tab.rowCount()):
            apr = self.tab.item(r, idx["apraksts"]).text() if self.tab.item(r, idx["apraksts"]) else ""
            daudz = to_decimal(self.tab.item(r, idx["daudzums"]).text() if self.tab.item(r, idx["daudzums"]) else "0")
            vien = self.tab.item(r, idx["vieniba"]).text() if self.tab.item(r, idx["vieniba"]) else ""
            cena = to_decimal(self.tab.item(r, idx["cena"]).text() if self.tab.item(r, idx["cena"]) else "0")

            foto_path = ""
            if "foto" in idx:
                foto_path = self.tab.item(r, idx["foto"]).text() if self.tab.item(r, idx["foto"]) else ""

            ser_nr = self.tab.item(r, idx.get("serial", -1)).text() if idx.get("serial") is not None and self.tab.item(r, idx.get("serial")) else ""
            gar = self.tab.item(r, idx.get("warranty", -1)).text() if idx.get("warranty") is not None and self.tab.item(r, idx.get("warranty")) else ""
            piez_poz = self.tab.item(r, idx.get("notes", -1)).text() if idx.get("notes") is not None and self.tab.item(r, idx.get("notes")) else ""

            if not apr and daudz == 0 and not ser_nr and not gar and not piez_poz and not foto_path:
                continue

            poz.append(Pozīcija(
                apraksts=apr,
                daudzums=daudz,
                vienība=vien,
                cena=cena,
                seriālais_nr=ser_nr,
                garantija=gar,
                piezīmes_pozīcijai=piez_poz,
                attēla_ceļš=foto_path
            ))

        d.pozīcijas = poz

        # Pielāgotās kolonnas
        d.custom_columns = self.data.custom_columns.copy()
        # Atjaunināt pielāgoto kolonnu datus no tabulas
        for col_idx, col in enumerate(d.custom_columns):
            col_data = []
            for r in range(self.tab.rowCount()):
                item = self.tab.item(r, idx["custom_start"] + col_idx)  # Pielāgotās kolonnas sākas pēc standarta kolonnu bloka
                col_data.append(item.text() if item else "")
            col['data'] = col_data

        # JAUNS: pozīciju kolonnu konfigurācija + kopsavilkums
        try:
            d.poz_columns_config = self._poz_cfg().copy() if isinstance(self._poz_cfg(), dict) else {}
        except Exception:
            d.poz_columns_config = {}
        try:
            d.show_price_summary = bool(self.ck_show_price_summary.isChecked()) if hasattr(self, "ck_show_price_summary") else bool(getattr(self.data, "show_price_summary", True))
        except Exception:
            d.show_price_summary = bool(getattr(self.data, "show_price_summary", True))

        # JAUNS: saglabā Pozīciju tabulas kolonnu secību (GUI -> PDF)
        try:
            d.poz_columns_visual_order = self._poz_get_visual_order_keys()
        except Exception:
            try:
                d.poz_columns_visual_order = list(getattr(self.data, 'poz_columns_visual_order', []) or [])
            except Exception:
                d.poz_columns_visual_order = []

        # JAUNS: saglabā Pozīciju tabulas kolonnu platumus/izkārtojumu (GUI)
        try:
            hdr = self.tab.horizontalHeader()
            st = hdr.saveState()
            try:
                st_bytes = bytes(st)
            except Exception:
                st_bytes = st.data() if hasattr(st, "data") else b""
            d.poz_header_state_b64 = base64.b64encode(st_bytes).decode("ascii") if st_bytes else ""
        except Exception:
            try:
                d.poz_header_state_b64 = str(getattr(self.data, "poz_header_state_b64", "") or "")
            except Exception:
                d.poz_header_state_b64 = ""

        # Iestatījumi
        d.iekļaut_pvn = self.ck_pvn.isChecked()
        d.pvn_likme = to_decimal(self.in_pvn.value())
        d.parakstu_rindas = self.ck_paraksti.isChecked()
        d.logotipa_ceļš = self.in_logo.text().strip()
        d.fonts_ceļš = self.in_fonts.text().strip()
        d.docx_template_path = self.in_docx_template.text().strip()

        d.paraksts_pieņēmējs_ceļš = self.in_paraksts_pie.text().strip()
        d.paraksts_nodevējs_ceļš = self.in_paraksts_nod.text().strip()

        # Papildu iestatījumi
        d.pdf_page_size = self.cb_page_size.currentText()
        d.pdf_page_orientation = self.cb_page_orientation.currentText()
        d.pdf_margin_left = to_decimal(self.in_margin_left.value())
        d.pdf_margin_right = to_decimal(self.in_margin_right.value())
        d.pdf_margin_top = to_decimal(self.in_margin_top.value())
        d.pdf_margin_bottom = to_decimal(self.in_margin_bottom.value())
        d.pdf_font_size_head = self.in_font_size_head.value()
        d.pdf_font_size_normal = self.in_font_size_normal.value()
        d.pdf_font_size_small = self.in_font_size_small.value()
        d.pdf_font_size_table = self.in_font_size_table.value()
        d.pdf_logo_width_mm = to_decimal(self.in_logo_width_mm.value())
        d.pdf_signature_width_mm = to_decimal(self.in_signature_width_mm.value())
        d.pdf_signature_height_mm = to_decimal(self.in_signature_height_mm.value())
        d.docx_image_width_inches = to_decimal(self.in_docx_image_width_inches.value())
        d.docx_signature_width_inches = to_decimal(self.in_docx_signature_width_inches.value())
        d.table_col_widths = self.in_table_col_widths.toPlainText().strip()
        d.auto_generate_akta_nr = self.ck_auto_generate_akta_nr.isChecked()
        d.default_execution_days = int(self.in_default_execution_days.value())
        d.default_currency = self.in_default_currency.text().strip()
        d.default_unit = self.in_default_unit.text().strip()
        d.default_pvn_rate = to_decimal(self.in_default_pvn_rate.value())
        d.poppler_path = self.in_poppler_path.text().strip()

        # New settings
        d.header_text_color = self.in_header_text_color.text().strip()
        d.footer_text_color = self.in_footer_text_color.text().strip()
        d.table_header_bg_color = self.in_table_header_bg_color.text().strip()
        d.table_grid_color = self.in_table_grid_color.text().strip()
        d.table_row_spacing = to_decimal(self.in_table_row_spacing.value())
        d.line_spacing_multiplier = to_decimal(self.in_line_spacing_multiplier.value())
        d.show_page_numbers = self.ck_show_page_numbers.isChecked()
        d.show_generation_timestamp = self.ck_show_generation_timestamp.isChecked()
        d.currency_symbol_position = self.cb_currency_symbol_position.currentText()
        d.date_format = self.in_date_format.text().strip()
        d.signature_line_length_mm = to_decimal(self.in_signature_line_length_mm.value())
        d.signature_line_thickness_pt = to_decimal(self.in_signature_line_thickness_pt.value())
        d.add_cover_page = self.ck_add_cover_page.isChecked()
        d.cover_page_title = self.in_cover_page_title.text().strip()
        d.cover_page_logo_width_mm = to_decimal(self.in_cover_page_logo_width_mm.value())
        # Individuālais QR kods
        d.include_custom_qr_code = self.ck_include_custom_qr_code.isChecked()
        d.custom_qr_code_data = self.in_custom_qr_code_data.text().strip()
        d.custom_qr_code_size_mm = to_decimal(self.in_custom_qr_code_size_mm.value())
        d.custom_qr_code_position = self.cb_custom_qr_code_position.currentText()

        # Automātiskais QR kods (akta ID)
        d.include_auto_qr_code = self.ck_include_auto_qr_code.isChecked()
        d.auto_qr_code_size_mm = to_decimal(self.in_auto_qr_code_size_mm.value())
        d.auto_qr_code_position = self.cb_auto_qr_code_position.currentText()

        d.add_watermark = self.ck_add_watermark.isChecked()
        d.watermark_text = self.in_watermark_text.text().strip()
        d.watermark_font_size = self.in_watermark_font_size.value()
        d.watermark_color = self.in_watermark_color.text().strip()
        d.watermark_rotation = self.in_watermark_rotation.value()
        d.enable_pdf_encryption = self.ck_enable_pdf_encryption.isChecked()
        d.pdf_user_password = self.in_pdf_user_password.text().strip()
        d.pdf_owner_password = self.in_pdf_owner_password.text().strip()
        d.allow_printing = self.ck_allow_printing.isChecked()
        d.allow_copying = self.ck_allow_copying.isChecked()
        d.allow_modifying = self.ck_allow_modifying.isChecked()
        d.allow_annotating = self.ck_allow_annotating.isChecked()
        # --- JAUNS: Pamata datu PDF šifrēšanas lauki (ja eksistē) ---
        try:
            if hasattr(self, "ck_pdf_encrypt_basic") and hasattr(self, "in_pdf_password_basic"):
                d.enable_pdf_encryption = self.ck_pdf_encrypt_basic.isChecked()
                d.pdf_user_password = self.in_pdf_password_basic.text().strip()
                # Owner parole var palikt no "Iestatījumi & Eksports" (ja lietotājs to lieto),
                # bet ja nav, ģenerēsies automātiski šifrēšanas brīdī.
        except Exception:
            pass

        d.default_country = self.in_default_country.text().strip()
        d.default_city = self.in_default_city.text().strip()
        d.show_contact_details_in_header = self.ck_show_contact_details_in_header.isChecked()
        d.contact_details_header_font_size = self.in_contact_details_header_font_size.value()
        d.item_image_width_mm = to_decimal(self.in_item_image_width_mm.value())
        d.item_image_caption_font_size = self.in_item_image_caption_font_size.value()
        d.show_item_notes_in_table = self.ck_show_item_notes_in_table.isChecked()
        d.show_item_serial_number_in_table = self.ck_show_item_serial_number_in_table.isChecked()
        d.show_item_warranty_in_table = self.ck_show_item_warranty_in_table.isChecked()
        d.table_cell_padding_mm = to_decimal(self.in_table_cell_padding_mm.value())
        d.table_header_font_style = self.cb_table_header_font_style.currentText()
        d.table_content_alignment = self.cb_table_content_alignment.currentText()
        d.signature_font_size = self.in_signature_font_size.value()
        d.signature_spacing_mm = to_decimal(self.in_signature_spacing_mm.value())
        d.document_title_font_size = self.in_document_title_font_size.value()
        d.document_title_color = self.in_document_title_color.text().strip()
        d.section_heading_font_size = self.in_section_heading_font_size.value()
        d.section_heading_color = self.in_section_heading_color.text().strip()
        d.paragraph_line_spacing_multiplier = to_decimal(self.in_paragraph_line_spacing_multiplier.value())
        d.table_border_style = self.cb_table_border_style.currentText()
        d.table_border_thickness_pt = to_decimal(self.in_table_border_thickness_pt.value())
        d.table_alternate_row_color = self.in_table_alternate_row_color.text().strip()
        d.show_total_sum_in_words = self.ck_show_total_sum_in_words.isChecked()
        d.total_sum_in_words_language = self.cb_total_sum_in_words_language.currentText()
        d.default_vat_calculation_method = self.cb_default_vat_calculation_method.currentText()
        d.show_vat_breakdown = self.ck_show_vat_breakdown.isChecked()
        d.enable_digital_signature_field = self.ck_enable_digital_signature_field.isChecked()
        d.digital_signature_field_name = self.in_digital_signature_field_name.text().strip()
        d.digital_signature_field_size_mm = to_decimal(self.in_digital_signature_field_size_mm.value())
        d.digital_signature_field_position = self.cb_digital_signature_field_position.currentText()



        # Attēli
        att = []
        if hasattr(self, "photos_table") and self.photos_table is not None:
            for r in range(self.photos_table.rowCount()):
                it = self.photos_table.item(r, 3)
                dat = it.data(Qt.UserRole) if it is not None else None
                if not dat:
                    continue
                att.append(Attēls(ceļš=dat["ceļš"], paraksts=dat.get("paraksts", "")))
        else:
            # Back-compat (ja kādā vecā stāvoklī vēl ir img_list)
            if getattr(self, "img_list", None) is not None:
                for i in range(self.img_list.count()):
                    it = self.img_list.item(i)
                    dat = it.data(Qt.UserRole)
                    att.append(Attēls(ceļš=dat["ceļš"], paraksts=dat.get("paraksts", "")))
        d.attēli = att

        return d


    def izvēlēties_templates_dir(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Izvēlēties šablonu direktoriju")
        if folder_path:
            self.in_templates_dir.setText(folder_path)
            # Pārliecināmies, ka jaunais direktorijs eksistē
            os.makedirs(folder_path, exist_ok=True)
            self._update_sablonu_list() # Atjaunojam šablonu sarakstu ar jauno direktoriju

    # ----- Tab: Puses -----

    def _persona_group(self, virsraksts: str, is_pieņēmējs: bool):
            box = QGroupBox(virsraksts)
            form = QFormLayout()
            nos = QLineEdit();
            reg = QLineEdit();
            adr = QLineEdit();
            kont = QLineEdit();
            tel = QLineEdit();
            ep = QLineEdit()
            web = QLineEdit(); web.setPlaceholderText("https://...")
            bankas_konts = QLineEdit()
            juridiskais_statuss = QComboBox()
            juridiskais_statuss.addItems(["", "Juridiska persona", "Fiziska persona", "Pašnodarbinātais"])

            form.addRow("Nosaukums / Vārds, Uzvārds", nos)
            form.addRow("Reģ. Nr. / personas kods", reg)
            adr_row = QWidget()
            adr_row_layout = QHBoxLayout()
            adr_row_layout.setContentsMargins(0, 0, 0, 0)
            adr_row_layout.setSpacing(6)
            btn_map = QToolButton()
            btn_map.setText("📍")
            btn_map.setToolTip("Atlasīt adresi kartē")
            btn_map.clicked.connect(lambda: self._begin_address_pick(adr))
            adr_row_layout.addWidget(adr, 1)
            adr_row_layout.addWidget(btn_map)
            adr_row.setLayout(adr_row_layout)
            form.addRow("Adrese", adr_row)
            form.addRow("Kontaktpersona", kont)

            amats = QLineEdit()
            pilnvaras_pamats = QComboBox()
            pilnvaras_pamats.addItems(["Pilnvaras pamats", "Uzņēmuma īpašumtiesības", "Cits"])

            form.addRow("Amats", amats)
            form.addRow("Pilnvaras pamats", pilnvaras_pamats)

            form.addRow("Tālrunis", tel)
            form.addRow("E-pasts", ep)
            form.addRow("Web lapa", web)
            form.addRow("Bankas konts", bankas_konts)
            form.addRow("Juridiskais statuss", juridiskais_statuss)

            btn_load_from_ab = QPushButton("Ielādēt no adrešu grāmatas")
            btn_save_to_ab = QPushButton("Saglabāt adrešu grāmatā")

            if is_pieņēmējs:
                btn_load_from_ab.clicked.connect(lambda: self._load_persona_from_address_book(self.pie_in))
                btn_save_to_ab.clicked.connect(lambda: self._save_persona_to_address_book(self.pie_in))
            else:
                btn_load_from_ab.clicked.connect(lambda: self._load_persona_from_address_book(self.nod_in))
                btn_save_to_ab.clicked.connect(lambda: self._save_persona_to_address_book(self.nod_in))

            btn_layout = QHBoxLayout()
            btn_layout.addWidget(btn_load_from_ab)
            btn_layout.addWidget(btn_save_to_ab)
            form.addRow(btn_layout)

            box.setLayout(form)
            return box, (nos, reg, adr, kont, amats, pilnvaras_pamats, tel, ep, web, bankas_konts, juridiskais_statuss)

    def _būvēt_puses_tab(self):
        content_widget = QWidget()
        lay = QHBoxLayout(content_widget)

        pieņēmējs_box, self.pie_in = self._persona_group("Pieņēmējs", True)
        nodevējs_box, self.nod_in = self._persona_group("Nodevējs", False)

        lay.addWidget(pieņēmējs_box)
        lay.addWidget(nodevējs_box)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(content_widget)

        main_tab_widget = QWidget()
        main_layout = QVBoxLayout(main_tab_widget)
        main_layout.addWidget(scroll_area)

        self.tabs.addTab(main_tab_widget, "Puses")

        for w in self.pie_in:
            if isinstance(w, QLineEdit):
                w.textChanged.connect(self._update_preview)
            elif isinstance(w, QComboBox):
                w.currentIndexChanged.connect(self._update_preview)

        for w in self.nod_in:
            if isinstance(w, QLineEdit):
                w.textChanged.connect(self._update_preview)
            elif isinstance(w, QComboBox):
                w.currentIndexChanged.connect(self._update_preview)

    def _load_persona_from_address_book(self, persona_inputs):
        self._undo_mgr.push_undo(self._snapshot_state('AB_LOAD_TO_PUSES'))
        self._audit('AB_LOAD_TO_PUSES', {})
        items = list(self.address_book.keys())
        if not items:
            QMessageBox.information(self, "Adrešu grāmata", "Adrešu grāmata ir tukša.")
            return
        item, ok = QInputDialog.getItem(self, "Ielādēt personu", "Izvēlieties personu:", items, 0, False)
        if ok and item:
            persona_data = self.address_book[item]
            # --- JAUNS: ja ir parole, prasa to pirms ielādes uz Puses tab ---
            if not self._ab_require_password(item, persona_data, "ielādēt"):
                return
            persona_inputs[0].setText(persona_data.get("nosaukums", ""))
            persona_inputs[1].setText(persona_data.get("reģ_nr", ""))
            persona_inputs[2].setText(persona_data.get("adrese", ""))
            persona_inputs[3].setText(persona_data.get("kontaktpersona", ""))
            persona_inputs[4].setText(persona_data.get("amats", ""))
            pilnvaras_val = persona_data.get("pilnvaras_pamats", "")
            idx_p = persona_inputs[5].findText(pilnvaras_val)
            if idx_p >= 0:
                persona_inputs[5].setCurrentIndex(idx_p)
            persona_inputs[6].setText(persona_data.get("tālrunis", ""))
            persona_inputs[7].setText(persona_data.get("epasts", ""))
            persona_inputs[8].setText(persona_data.get("web_lapa", ""))
            persona_inputs[9].setText(persona_data.get("bankas_konts", ""))
            juridiskais_statuss_val = persona_data.get("juridiskais_statuss", "")
            idx = persona_inputs[10].findText(juridiskais_statuss_val)
            if idx >= 0:
                persona_inputs[10].setCurrentIndex(idx)
            QMessageBox.information(self, "Ielādēts", f"Persona '{item}' ielādēta.")
    # --- JAUNS: Adrešu grāmatas ieraksta nosaukuma ģenerēšana/lietotāja izvēle ---
    def _ab_generate_default_entry_name(self, persona_data: dict) -> str:
        """Noklusējuma adrešu grāmatas ieraksta nosaukums: 'Uzņēmums — Kontaktpersona'."""
        try:
            comp = (persona_data.get("nosaukums") or "").strip()
            kp = (persona_data.get("kontaktpersona") or "").strip()
            if comp and kp:
                return f"{comp} — {kp}"
            return comp or kp or "Persona"
        except Exception:
            return "Persona"

    def _ab_make_unique_key(self, base: str) -> str:
        """Nodrošina unikālu atslēgu adrešu grāmatā (ja jau eksistē, pievieno (2), (3)...)."""
        base = (base or "").strip() or "Persona"
        if base not in getattr(self, "address_book", {}):
            return base
        i = 2
        while True:
            k = f"{base} ({i})"
            if k not in self.address_book:
                return k
            i += 1

    def _ab_get_entry_name_for_save(self, persona_data: dict) -> str:
        """Atgriež ieraksta nosaukumu atkarībā no režīma (auto/lietotājs)."""
        default_name = self._ab_generate_default_entry_name(persona_data)
        try:
            use_auto = True
            if hasattr(self, "chk_ab_auto_name") and self.chk_ab_auto_name is not None:
                use_auto = bool(self.chk_ab_auto_name.isChecked())
        except Exception:
            use_auto = True

        if use_auto:
            return default_name

        # Lietotājs pats ievada nosaukumu
        name, ok = QInputDialog.getText(
            self,
            "Saglabāt adrešu grāmatā",
            "Ievadi nosaukumu (lai atšķirtu vairākas puses vienam uzņēmumam):",
            QLineEdit.Normal,
            default_name
        )
        if not ok:
            return ""
        return (name or "").strip()


    def _save_persona_to_address_book(self, persona_inputs):
        self._undo_mgr.push_undo(self._snapshot_state('AB_SAVE'))
        self._audit('AB_SAVE', {})
        nosaukums = persona_inputs[0].text().strip()
        if not nosaukums:
            QMessageBox.warning(self, "Saglabāt personu", "Nosaukums nevar būt tukšs.")
            return
        persona_data = {
            "nosaukums": nosaukums,
            "reģ_nr": persona_inputs[1].text().strip(),
            "adrese": persona_inputs[2].text().strip(),
            "kontaktpersona": persona_inputs[3].text().strip(),
            "amats": persona_inputs[4].text().strip(),
            "pilnvaras_pamats": persona_inputs[5].currentText(),
            "tālrunis": persona_inputs[6].text().strip(),
            "epasts": persona_inputs[7].text().strip(),
            "web_lapa": persona_inputs[8].text().strip(),
            "bankas_konts": persona_inputs[9].text().strip(),
            "juridiskais_statuss": persona_inputs[10].currentText(),
        }
        entry_name = self._ab_get_entry_name_for_save(persona_data)
        if not entry_name:
            return
        entry_name = self._ab_make_unique_key(entry_name)
        self.address_book[entry_name] = persona_data
        self._save_address_book() # Saglabājam adrešu grāmatu failā
        self._update_address_book_list() # Atjaunojam sarakstu GUI

    def _load_address_book(self):
        if os.path.exists(ADDRESS_BOOK_FILE):
            try:
                # Mēģinām ielādēt ar dažādām kodēšanām
                encodings = ['utf-8', 'utf-8-sig', 'cp1257', 'iso-8859-1', 'windows-1252']
                for encoding in encodings:
                    try:
                        with open(ADDRESS_BOOK_FILE, 'r', encoding=encoding) as f:
                            self.address_book = json.load(f)
                        break
                    except (UnicodeDecodeError, json.JSONDecodeError):
                        continue
                else:
                    # Ja neviena kodēšana nedarbojas, izveidojam jaunu adrešu grāmatu
                    print(f"Neizdevās ielādēt adrešu grāmatas failu ar nevenu kodēšanu. Izveidojam jaunu.")
                    self.address_book = {}
            except Exception as e:
                QMessageBox.warning(self, "Kļūda", f"Neizdevās ielādēt adrešu grāmatu: {e}")
                self.address_book = {}
        else:
            self.address_book = {}

    def _save_address_book(self):
        os.makedirs(SETTINGS_DIR, exist_ok=True) # Izveidojam direktoriju, ja tā neeksistē
        try:
            with open(ADDRESS_BOOK_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.address_book, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās saglabāt adrešu grāmatu: {e}")

    # ----- Tab: Pozīcijas -----
    def _būvēt_pozīcijas_tab(self):
        w = QWidget()
        v = QVBoxLayout()

        # --- JAUNS: vienmēr būvējam pilnu kolonnu komplektu (kolonnas var paslēpt/atkal parādīt) ---
        # Bāzes kolonnas (UI secība ir fiksēta, lai aprēķini un saglabāšana vienmēr strādātu)
        base_headers = [
            ("apraksts", "Apraksts"),
            ("daudzums", "Daudzums"),
            ("vieniba", "Vienība"),
            ("cena", "Cena"),
            ("summa", "Summa"),
            ("serial", "Seriālais Nr."),
            ("warranty", "Garantija"),
            ("notes", "Piezīmes pozīcijai"),
        ]

        # Pārliecināmies, ka konfigurācija eksistē un ir pilna
        try:
            self.data.poz_columns_config = _merge_poz_columns_config(getattr(self.data, "poz_columns_config", None))
        except Exception:
            self.data.poz_columns_config = _merge_poz_columns_config({})

        # Nodrošinām, ka custom column dict satur arī "visible" (atpakaļsavietojami)
        try:
            for col in getattr(self.data, "custom_columns", []) or []:
                if isinstance(col, dict) and "visible" not in col:
                    col["visible"] = True
        except Exception:
            pass

        headers = []
        for key, default_title in base_headers:
            title = default_title
            try:
                title = str(self.data.poz_columns_config.get(key, {}).get("title", default_title))
            except Exception:
                title = default_title
            headers.append(title)

        # Pielāgotās kolonnas (vienmēr pirms Foto)
        for col in getattr(self.data, "custom_columns", []) or []:
            try:
                headers.append(str(col.get('name', '')))
            except Exception:
                headers.append("")

        # Foto kolonna vienmēr pēdējā (UI), bet var paslēpt
        foto_title = "Foto"
        try:
            foto_title = str(self.data.poz_columns_config.get("foto", {}).get("title", "Foto"))
        except Exception:
            pass
        headers.append(foto_title)

        self.tab = QTableWidget(0, len(headers))
        self.tab.setHorizontalHeaderLabels(headers)

        # --- JAUNS: Undo/Redo pozīciju tabulai ---
        self._ensure_positions_undo_hook()

        # --- JAUNS: kopsavilkuma rādīšana (zem tabulas PDF) ---
        self.ck_show_price_summary = QCheckBox("Rādīt cenu apkopojumu zem tabulas (PDF)")
        try:
            self.ck_show_price_summary.setChecked(bool(getattr(self.data, "show_price_summary", True)))
        except Exception:
            self.ck_show_price_summary.setChecked(True)
        self.ck_show_price_summary.stateChanged.connect(self._update_preview)
        v.addWidget(self.ck_show_price_summary)

        # Padarīt kolonnas pielāgojamas velkot
        for i in range(len(headers)):
            self.tab.horizontalHeader().setSectionResizeMode(i, QHeaderView.Interactive)

        # --- JAUNS: kolonnu konteksta izvēlne (paslēpt/parādīt/pārdēvēt) ---
        try:
            hh = self.tab.horizontalHeader()
            hh.setContextMenuPolicy(Qt.CustomContextMenu)
            hh.customContextMenuRequested.connect(self._poz_header_context_menu)
            # Ļaujam vilkt kolonnas ar peli + saglabājam secību
            try:
                hh.setSectionsMovable(True)
            except Exception:
                pass
            try:
                hh.sectionMoved.connect(lambda *_: self._poz_store_visual_order())
            except Exception:
                pass
            # Ja ir saglabāta secība, pielietojam to uzreiz
            try:
                self._poz_apply_stored_visual_order()
            except Exception:
                pass

        except Exception:
            pass

        # Pielietojam kolonnu redzamību
        self._poz_apply_column_visibility()

        # Teksta pārlikšana šūnās (nevārdi netiek pārrauti; pāriet nākamajā rindā)
        self.tab.setWordWrap(True)
        self.tab.setTextElideMode(Qt.ElideNone)

        # Auto pielāgot rindas/kolonnas pēc teksta (de-bounce, lai nebremzē rakstot)
        self._poz_autoresize_timer = QTimer(self)
        self._poz_autoresize_timer.setSingleShot(True)

        def _apply_poz_autoresize():
            try:
                self.tab.resizeRowsToContents()
                self.tab.resizeColumnsToContents()
            except Exception:
                pass

        self._poz_autoresize_timer.timeout.connect(_apply_poz_autoresize)

        def _schedule_poz_autoresize(*_):
            try:
                self._poz_autoresize_timer.start(120)
            except Exception:
                pass

        self.tab.cellChanged.connect(_schedule_poz_autoresize)

        btns = QHBoxLayout()
        add = QPushButton("Pievienot pozīciju")
        add.clicked.connect(self.pievienot_pozīciju)
        dzest = QPushButton("Dzēst izvēlēto")
        dzest.clicked.connect(self.dzest_pozīciju)
        cfg_btn = QPushButton("Kolonnu iestatījumi")
        cfg_btn.setToolTip("Ar peles labo pogu uz kolonnas virsraksta var paslēpt/parādīt un pārdēvēt kolonnas")
        cfg_btn.clicked.connect(self._poz_open_settings_for_selected_column)
        add_col = QPushButton("Pievienot kolonnu")
        add_col.clicked.connect(self.pievienot_kolonnu)
        del_col = QPushButton("Dzēst kolonnu")
        del_col.clicked.connect(self.dzest_kolonnu)
        btns.addWidget(add)
        btns.addWidget(dzest)
        btns.addWidget(cfg_btn)
        btns.addWidget(add_col)
        btns.addWidget(del_col)
        btns.addStretch()

        v.addLayout(btns)
        v.addWidget(self.tab)
        w.setLayout(v)
        self.tabs.addTab(w, "Pozīcijas")

        self.tab.cellChanged.connect(self._pārrēķināt_summa)
        self.tab.cellChanged.connect(self._update_preview)


    def _poz_col_indices(self) -> dict:
        """Atgriež kolonnu indeksus 'Pozīcijas' tabulai atbilstoši iestatījumiem.
        Foto kolonna vienmēr ir PĒDĒJĀ (arī pēc pielāgotajām kolonnām).
        """
        # JAUNS: fiksēts bāzes kolonnu bloks, lai jebkuru kolonnu var paslēpt/atkal parādīt,
        # nepārkārtojot indeksus (paslēpšana notiek ar setColumnHidden).
        idx = {
            "apraksts": 0,
            "daudzums": 1,
            "vieniba": 2,
            "cena": 3,
            "summa": 4,
            "serial": 5,
            "warranty": 6,
            "notes": 7,
        }
        idx["custom_start"] = 8
        idx["foto"] = idx["custom_start"] + len(getattr(self.data, "custom_columns", []) or [])
        idx["col_count"] = idx["foto"] + 1
        return idx




    # --- JAUNS: Pozīciju kolonnu kārtība (viegla pārvietošana starp jebkurām kolonnām) ---
    def _poz_get_visual_order_keys(self) -> list:
        """Atgriež kolonnu atslēgas vizuālajā secībā (kā redz lietotājs)."""
        if not hasattr(self, "tab") or not self.tab:
            return []
        hh = self.tab.horizontalHeader()
        order = []
        try:
            for visual in range(self.tab.columnCount()):
                logical = hh.logicalIndex(visual)
                key = self._poz_key_from_column(int(logical))
                if key:
                    order.append(key)
        except Exception:
            return []
        return order

    def _poz_store_visual_order(self):
        """Saglabā pašreizējo vizuālo kolonnu secību (lai tā saglabājas pēc restart)."""
        try:
            self.data.poz_columns_visual_order = self._poz_get_visual_order_keys()
        except Exception:
            pass

    def _poz_apply_stored_visual_order(self):
        """Pielieto saglabāto vizuālo kolonnu secību, pārvietojot header sekcijas."""
        if not hasattr(self, "tab") or not self.tab:
            return
        try:
            desired = list(getattr(self.data, "poz_columns_visual_order", []) or [])
        except Exception:
            desired = []
        if not desired:
            return

        hh = self.tab.horizontalHeader()
        # Pašreizējais key -> logical index
        key_to_logical = {}
        try:
            for logical in range(self.tab.columnCount()):
                k = self._poz_key_from_column(int(logical))
                if k:
                    key_to_logical[k] = int(logical)
        except Exception:
            return

        # Filtrējam tikai eksistējošās
        desired = [k for k in desired if k in key_to_logical]
        if not desired:
            return

        # Pieliekam klāt jaunas (piem., tikko pievienota custom kolonna)
        current = self._poz_get_visual_order_keys()
        for k in current:
            if k not in desired:
                # Foto vienmēr pēdējā, ja eksistē
                if k == "foto":
                    continue
                desired.append(k)
        if "foto" in current and "foto" not in desired:
            desired.append("foto")
        # Ja foto ir sarakstā, pārbīdam uz beigām
        if "foto" in desired:
            desired = [k for k in desired if k != "foto"] + ["foto"]

        # Pārvietojam vizuālajā secībā pa vienam
        try:
            for target_visual, key in enumerate(desired):
                logical = key_to_logical.get(key)
                if logical is None:
                    continue
                cur_visual = hh.visualIndex(logical)
                if cur_visual != target_visual and cur_visual >= 0:
                    hh.moveSection(cur_visual, target_visual)
        except Exception:
            pass

        # Pārrakstam saglabāto secību, lai tā atbilst realitātei
        self._poz_store_visual_order()

    # --- JAUNS: Pozīciju kolonnu redzamība/pārdēvēšana ---
    def _poz_cfg(self) -> dict:
        try:
            self.data.poz_columns_config = _merge_poz_columns_config(getattr(self.data, "poz_columns_config", None))
        except Exception:
            self.data.poz_columns_config = _merge_poz_columns_config({})
        return self.data.poz_columns_config

    def _poz_is_visible(self, key: str) -> bool:
        cfg = self._poz_cfg()
        try:
            return bool(cfg.get(key, {}).get("visible", True))
        except Exception:
            return True

    def _poz_set_visible(self, key: str, visible: bool):
        cfg = self._poz_cfg()
        if key not in cfg:
            cfg[key] = {"title": key, "visible": bool(visible)}
        else:
            try:
                cfg[key]["visible"] = bool(visible)
            except Exception:
                cfg[key] = {"title": str(cfg.get(key, {}).get("title", key)), "visible": bool(visible)}

    def _poz_set_title(self, key: str, title: str):
        cfg = self._poz_cfg()
        if key not in cfg:
            cfg[key] = {"title": str(title), "visible": True}
        else:
            try:
                cfg[key]["title"] = str(title)
            except Exception:
                cfg[key] = {"title": str(title), "visible": bool(cfg.get(key, {}).get("visible", True))}

    def _poz_apply_column_visibility(self):
        """Pielieto kolonnu redzamību tabulā + sinhronizē ar vecajiem iestatījumu laukiem."""
        if not hasattr(self, "tab") or not self.tab:
            return
        idx = self._poz_col_indices()
        cfg = self._poz_cfg()

        # Bāzes kolonnas
        base_keys = ["apraksts", "daudzums", "vieniba", "cena", "summa", "serial", "warranty", "notes"]
        for k in base_keys:
            col = idx.get(k)
            if col is None or col >= self.tab.columnCount():
                continue
            self.tab.setColumnHidden(col, not bool(cfg.get(k, {}).get("visible", True)))

        # Pielāgotās kolonnas (visible lauks katrai custom kolonnai)
        custom_start = idx.get("custom_start", 8)
        for i, col in enumerate(getattr(self.data, "custom_columns", []) or []):
            ui_col = custom_start + i
            if ui_col >= self.tab.columnCount():
                continue
            vis = True
            try:
                vis = bool(col.get("visible", True)) if isinstance(col, dict) else True
            except Exception:
                vis = True
            self.tab.setColumnHidden(ui_col, not vis)

        # Foto
        foto_col = idx.get("foto")
        if foto_col is not None and foto_col < self.tab.columnCount():
            self.tab.setColumnHidden(foto_col, not bool(cfg.get("foto", {}).get("visible", True)))

        # Atpakaļsavietojamība: vecie iestatījumi, ko izmanto PDF/iestatījumu tabs
        try:
            self.data.show_item_serial_number_in_table = bool(cfg.get("serial", {}).get("visible", True))
            self.data.show_item_warranty_in_table = bool(cfg.get("warranty", {}).get("visible", True))
            self.data.show_item_notes_in_table = bool(cfg.get("notes", {}).get("visible", True))
            self.data.show_item_photo_in_table = bool(cfg.get("foto", {}).get("visible", True))
        except Exception:
            pass

        # JAUNS: ja Foto kolonna ir ieslēgta, nodrošinām, ka pogas tiek atjaunotas visām rindām
        try:
            if bool(cfg.get("foto", {}).get("visible", True)):
                for r in range(self.tab.rowCount()):
                    self._ensure_photo_cell(r)
        except Exception:
            pass

        # Kopsavilkums zem tabulas
        try:
            if hasattr(self, "ck_show_price_summary") and self.ck_show_price_summary:
                self.data.show_price_summary = bool(self.ck_show_price_summary.isChecked())
        except Exception:
            pass

    def _poz_key_from_column(self, col: int) -> Optional[str]:
        idx = self._poz_col_indices()
        for k in ["apraksts", "daudzums", "vieniba", "cena", "summa", "serial", "warranty", "notes"]:
            if idx.get(k) == col:
                return k
        # custom
        cs = idx.get("custom_start", 8)
        foto = idx.get("foto")
        if foto is not None and col == foto:
            return "foto"
        if cs is not None and col >= cs and (foto is None or col < foto):
            return f"custom:{col - cs}"
        return None

    # --- JAUNS: droša kolonnu iestatījumu atvēršana pēc atlasītās kolonnas ---
    def _poz_open_settings_for_selected_column(self):
        """Atver kolonnu iestatījumu izvēlni tieši atlasītajai kolonnai.

        (Tas salabo situāciju, kad darbības attiecās uz nepareizu kolonnu, jo tika izmantotas
        nepareizas koordinātes.)
        """
        if not hasattr(self, "tab") or not self.tab:
            return
        hh = self.tab.horizontalHeader()
        col = self.tab.currentColumn()
        if col < 0:
            # ja nekas nav atlasīts, mēģinām atrast pirmo redzamo
            for c in range(self.tab.columnCount()):
                if not self.tab.isColumnHidden(c):
                    col = c
                    break
        if col < 0:
            return

        try:
            x = hh.sectionPosition(col) + max(6, int(hh.sectionSize(col) * 0.5))
            global_pos = hh.mapToGlobal(QPoint(x, hh.height()))
            self._poz_header_context_menu(QPoint(x, int(hh.height() / 2)), forced_col=col, global_pos=global_pos)
        except Exception:
            self._poz_header_context_menu(QPoint(0, 0), forced_col=col)

    def _poz_swap_custom_columns(self, i: int, j: int):
        """Samaina divas pielāgotās kolonnas (UI + datu sarakstā), saglabājot Foto vienmēr pēdējā."""
        try:
            customs = getattr(self.data, "custom_columns", []) or []
            if not (0 <= i < len(customs) and 0 <= j < len(customs)):
                return
            if i == j:
                return

            idx = self._poz_col_indices()
            cs = idx.get("custom_start", 8)
            ui_i = cs + i
            ui_j = cs + j
            if ui_i >= self.tab.columnCount() or ui_j >= self.tab.columnCount():
                return

            # swap definīcijas
            customs[i], customs[j] = customs[j], customs[i]
            self.data.custom_columns = customs

            # swap header tekstu
            hi = self.tab.horizontalHeaderItem(ui_i)
            hj = self.tab.horizontalHeaderItem(ui_j)
            ti = hi.text() if hi else ""
            tj = hj.text() if hj else ""
            self.tab.setHorizontalHeaderItem(ui_i, QTableWidgetItem(tj))
            self.tab.setHorizontalHeaderItem(ui_j, QTableWidgetItem(ti))

            # swap slēpšanas stāvokli
            hid_i = self.tab.isColumnHidden(ui_i)
            hid_j = self.tab.isColumnHidden(ui_j)

            # swap platumu
            try:
                w_i = self.tab.columnWidth(ui_i)
                w_j = self.tab.columnWidth(ui_j)
                self.tab.setColumnWidth(ui_i, w_j)
                self.tab.setColumnWidth(ui_j, w_i)
            except Exception:
                pass

            # swap šūnu saturu
            for r in range(self.tab.rowCount()):
                it_i = self.tab.takeItem(r, ui_i)
                it_j = self.tab.takeItem(r, ui_j)
                self.tab.setItem(r, ui_i, it_j)
                self.tab.setItem(r, ui_j, it_i)

                wgi = self.tab.cellWidget(r, ui_i)
                wgj = self.tab.cellWidget(r, ui_j)
                if wgi is not None or wgj is not None:
                    self.tab.removeCellWidget(r, ui_i)
                    self.tab.removeCellWidget(r, ui_j)
                    if wgj is not None:
                        self.tab.setCellWidget(r, ui_i, wgj)
                    if wgi is not None:
                        self.tab.setCellWidget(r, ui_j, wgi)

            self.tab.setColumnHidden(ui_i, hid_j)
            self.tab.setColumnHidden(ui_j, hid_i)

        except Exception:
            return

    def _poz_header_context_menu(self, pos: QPoint, forced_col: Optional[int] = None, global_pos: Optional[QPoint] = None):
        """Konteksta izvēlne kolonnu virsrakstiem: paslēpt/parādīt/pārdēvēt.

        `pos` ir header lokālajās koordinātēs. Ja izvēlne tiek atvērta no pogas,
        izmantojam `forced_col`, lai vienmēr strādātu ar atlasīto kolonnu.
        """
        if not hasattr(self, "tab") or not self.tab:
            return

        hh = self.tab.horizontalHeader()
        col = int(forced_col) if forced_col is not None else hh.logicalIndexAt(pos)
        if col < 0:
            col = -1

        menu = QMenu(self)

        # Paslēpt izvēlēto
        if col >= 0:
            key = self._poz_key_from_column(col)

            # JAUNS: kolonnu pārvietošana (strādā jebkurai kolonnai, arī starp default kolonnām)
            try:
                hh2 = self.tab.horizontalHeader()
                cur_visual = hh2.visualIndex(col)
            except Exception:
                hh2 = None
                cur_visual = -1

            if hh2 is not None and cur_visual >= 0:
                act_left = menu.addAction("Pārvietot pa kreisi")
                act_right = menu.addAction("Pārvietot pa labi")
                act_left.setEnabled(cur_visual > 0)
                act_right.setEnabled(cur_visual < self.tab.columnCount() - 1)

                def _mv_left():
                    try:
                        v = hh2.visualIndex(col)
                        if v > 0:
                            hh2.moveSection(v, v - 1)
                            self._poz_store_visual_order()
                            self._update_preview()
                    except Exception:
                        pass

                def _mv_right():
                    try:
                        v = hh2.visualIndex(col)
                        if 0 <= v < self.tab.columnCount() - 1:
                            hh2.moveSection(v, v + 1)
                            self._poz_store_visual_order()
                            self._update_preview()
                    except Exception:
                        pass

                act_left.triggered.connect(_mv_left)
                act_right.triggered.connect(_mv_right)

                # Pārvietot uz konkrētu pozīciju (jebkur)
                sub_mv = menu.addMenu("Pārvietot uz…")
                try:
                    titles = []
                    for visual in range(self.tab.columnCount()):
                        logical = hh2.logicalIndex(visual)
                        it = self.tab.horizontalHeaderItem(int(logical))
                        t = it.text() if it else str(self._poz_key_from_column(int(logical)) or "")
                        titles.append(t)

                    for target_visual, t in enumerate(titles):
                        act = sub_mv.addAction(f"{target_visual+1}. {t}")
                        def _make_move(tv=target_visual):
                            def _do():
                                try:
                                    v = hh2.visualIndex(col)
                                    if v >= 0 and tv >= 0:
                                        hh2.moveSection(v, tv)
                                        self._poz_store_visual_order()
                                        self._update_preview()
                                except Exception:
                                    pass
                            return _do
                        act.triggered.connect(_make_move())
                except Exception:
                    pass

                menu.addSeparator()
            act_hide = menu.addAction("Paslēpt šo kolonnu")
            def _hide_selected():
                if not key:
                    return
                if key.startswith("custom:"):
                    try:
                        i = int(key.split(":", 1)[1])
                        if 0 <= i < len(getattr(self.data, "custom_columns", []) or []):
                            self.data.custom_columns[i]["visible"] = False
                    except Exception:
                        pass
                else:
                    self._poz_set_visible(key, False)
                self._poz_apply_column_visibility()
                self._update_preview()
            act_hide.triggered.connect(_hide_selected)

            # Pārdēvēt
            act_rename = menu.addAction("Pārdēvēt kolonnu…")
            def _rename_selected():
                if not key:
                    return
                cur_title = ""
                try:
                    it = self.tab.horizontalHeaderItem(col)
                    cur_title = it.text() if it else ""
                except Exception:
                    cur_title = ""
                new_title, ok = QInputDialog.getText(self, "Pārdēvēt kolonnu", "Jaunais nosaukums:", text=cur_title)
                if not (ok and new_title.strip()):
                    return
                new_title = new_title.strip()

                if key.startswith("custom:"):
                    try:
                        i = int(key.split(":", 1)[1])
                        if 0 <= i < len(getattr(self.data, "custom_columns", []) or []):
                            self.data.custom_columns[i]["name"] = new_title
                    except Exception:
                        pass
                else:
                    self._poz_set_title(key, new_title)

                try:
                    self.tab.setHorizontalHeaderItem(col, QTableWidgetItem(new_title))
                except Exception:
                    pass
                self._update_preview()
            act_rename.triggered.connect(_rename_selected)

        # Parādīt paslēptās
        sub = menu.addMenu("Parādīt kolonnas")
        idx = self._poz_col_indices()
        # bāzes + foto
        for k in ["apraksts", "daudzums", "vieniba", "cena", "summa", "serial", "warranty", "notes", "foto"]:
            c = idx.get(k)
            if c is None or c >= self.tab.columnCount():
                continue
            hidden = self.tab.isColumnHidden(c)
            title = ""
            try:
                it = self.tab.horizontalHeaderItem(c)
                title = it.text() if it else k
            except Exception:
                title = k
            act = sub.addAction(title)
            act.setCheckable(True)
            act.setChecked(not hidden)
            def _toggle_factory(key=k, col_index=c):
                def _t(checked: bool):
                    self._poz_set_visible(key, bool(checked))
                    self._poz_apply_column_visibility()
                    self._update_preview()
                return _t
            act.toggled.connect(_toggle_factory())

        # custom kolonnas
        customs = getattr(self.data, "custom_columns", []) or []
        if customs:
            sub2 = menu.addMenu("Parādīt pielāgotās")
            cs = idx.get("custom_start", 8)
            for i, coldef in enumerate(customs):
                ui_col = cs + i
                if ui_col >= self.tab.columnCount():
                    continue
                title = ""
                try:
                    title = str(coldef.get("name", f"Kolonna {i+1}"))
                except Exception:
                    title = f"Kolonna {i+1}"
                vis = True
                try:
                    vis = bool(coldef.get("visible", True))
                except Exception:
                    vis = True
                act = sub2.addAction(title)
                act.setCheckable(True)
                act.setChecked(vis)
                def _toggle_custom_factory(i_=i):
                    def _t(checked: bool):
                        try:
                            if 0 <= i_ < len(self.data.custom_columns):
                                self.data.custom_columns[i_]["visible"] = bool(checked)
                        except Exception:
                            pass
                        self._poz_apply_column_visibility()
                        self._update_preview()
                    return _t
                act.toggled.connect(_toggle_custom_factory())

        # JAUNS: ja tiek padots precīzs globālais punkts (piem., no pogas), izmantojam to
        try:
            if isinstance(global_pos, QPoint):
                menu.exec(global_pos)
            else:
                menu.exec(hh.mapToGlobal(pos))
        except Exception:
            try:
                menu.exec(hh.mapToGlobal(pos))
            except Exception:
                pass


    
    def _row_from_photo_button(self, btn: QToolButton) -> int:
        """Droši nosaka rindu pēc foto pogas pozīcijas tabulas viewportā."""
        try:
            from PySide6.QtCore import QPoint
            vp = self.tab.viewport()
            p = btn.mapTo(vp, QPoint(btn.width() // 2, btn.height() // 2))
            idx = self.tab.indexAt(p)
            return idx.row()
        except Exception:
            return -1

    def _ensure_photo_cell(self, row: int):
        """Ieliek Foto izvēles/dzēšanas pogu konkrētai rindai (ja foto kolonna ir ieslēgta)."""
        # Foto kolonna vienmēr eksistē, bet var būt paslēpta
        try:
            if not self._poz_is_visible("foto"):
                return
        except Exception:
            pass
        col = self._poz_col_indices().get("foto")
        if col is None:
            return

        # Item glabā faila ceļu
        it = self.tab.item(row, col)
        if not it:
            it = QTableWidgetItem("")
            it.setFlags(it.flags() & ~Qt.ItemIsEditable)
            self.tab.setItem(row, col, it)

        # Ja jau ir ielikts widgets – nepārrakstām, tikai atsvaidzinām tekstu/ikonas
        w = self.tab.cellWidget(row, col)
        if isinstance(w, QToolButton):
            self._refresh_photo_button_ui(row, w)
            return

        btn = QToolButton()
        btn.setAutoRaise(True)
        btn.setToolTip("Pievienot vai dzēst foto šai pozīcijai")

        # Galvenais klikšķis = izvēlēties/mainīt
        btn.clicked.connect(lambda _=False, b=btn: self._choose_photo_for_button(b))

        # Menu ar "Izvēlēties/Mainīt" un "Dzēst"
        menu = QMenu(btn)
        act_choose = menu.addAction("Izvēlēties / Mainīt")
        act_choose.triggered.connect(lambda _=False, b=btn: self._choose_photo_for_button(b))
        act_clear = menu.addAction("Dzēst foto")
        act_clear.triggered.connect(lambda _=False, b=btn: self._clear_photo_for_button(b))
        btn.setMenu(menu)
        btn.setPopupMode(QToolButton.MenuButtonPopup)

        self.tab.setCellWidget(row, col, btn)
        self._refresh_photo_button_ui(row, btn)

    def _refresh_photo_button_ui(self, row: int, btn: QToolButton):
        """Atjauno foto pogas UI (tekstu/ikonu) balstoties uz šūnas ceļu."""
        col = self._poz_col_indices().get("foto")
        if col is None:
            return
        path = ""
        it = self.tab.item(row, col)
        if it:
            path = (it.text() or "").strip()

        if path and os.path.exists(path):
            btn.setText("Mainīt")
            try:
                ico = QIcon(path)
                if not ico.isNull():
                    btn.setIcon(ico)
                    btn.setIconSize(QSize(18, 18))
            except Exception:
                btn.setIcon(QIcon())
        else:
            btn.setText("Izvēlēties")
            btn.setIcon(QIcon())

        # Menu "Dzēst" – atslēdzam, ja nav ko dzēst
        if btn.menu():
            acts = btn.menu().actions()
            # pieņemam, ka otrā ir "Dzēst foto"
            for a in acts:
                if "Dzēst" in a.text():
                    a.setEnabled(bool(path))

    def _choose_photo_for_button(self, btn: QToolButton):
        row = self._row_from_photo_button(btn)
        if row < 0:
            return
        self._choose_photo_for_row(row)

    def _clear_photo_for_button(self, btn: QToolButton):
        row = self._row_from_photo_button(btn)
        if row < 0:
            return
        self._clear_photo_for_row(row)

    def _clear_photo_for_row(self, row: int):
        col = self._poz_col_indices().get("foto")
        if col is None:
            return

        it = self.tab.item(row, col)
        if not it:
            it = QTableWidgetItem("")
            it.setFlags(it.flags() & ~Qt.ItemIsEditable)
            self.tab.setItem(row, col, it)

        it.setText("")

        w = self.tab.cellWidget(row, col)
        if isinstance(w, QToolButton):
            self._refresh_photo_button_ui(row, w)

        self._update_preview()

    def _choose_photo_from_sender(self):
        btn = self.sender()
        if isinstance(btn, QToolButton):
            self._choose_photo_for_button(btn)

    def _choose_photo_for_row(self, row: int):
        col = self._poz_col_indices().get("foto")
        if col is None:
            return

        path, _ = QFileDialog.getOpenFileName(
            self, "Izvēlēties foto pozīcijai",
            "", "Attēli (*.png *.jpg *.jpeg *.webp *.bmp)"
        )
        if not path:
            return

        it = self.tab.item(row, col)
        if not it:
            it = QTableWidgetItem("")
            it.setFlags(it.flags() & ~Qt.ItemIsEditable)
            self.tab.setItem(row, col, it)
        it.setText(path)

        w = self.tab.cellWidget(row, col)
        if isinstance(w, QToolButton):
            self._refresh_photo_button_ui(row, w)

        self._update_preview()

    def pievienot_kolonnu(self):
        col_name, ok = QInputDialog.getText(self, "Pievienot kolonnu", "Ievadiet kolonnas nosaukumu:")
        if ok and col_name.strip():
            col_name = col_name.strip()

            # Pārbaudīt, vai kolonna jau eksistē
            headers = [self.tab.horizontalHeaderItem(i).text() for i in range(self.tab.columnCount())]
            if col_name in headers:
                QMessageBox.warning(self, "Kļūda", f"Kolonna '{col_name}' jau eksistē.")
                return

            # Foto vienmēr ir pēdējā kolonna -> insertējam jauno kolonnu PIRMS Foto
            idx = self._poz_col_indices()
            foto_col = idx.get("foto", self.tab.columnCount())
            self.tab.insertColumn(foto_col)
            self.tab.setHorizontalHeaderItem(foto_col, QTableWidgetItem(col_name))
            self.tab.horizontalHeader().setSectionResizeMode(foto_col, QHeaderView.Interactive)

            # Pievienojam datiem (atpakaļsavietojami: ja nav 'visible', pievienojam)
            new_col = {'name': col_name, 'data': [''] * self.tab.rowCount(), 'visible': True}
            self.data.custom_columns.append(new_col)

            # Atjaunojam Foto virsrakstu (jo tas ir pabīdīts pa labi)
            try:
                foto_idx_new = self._poz_col_indices().get("foto")
                if foto_idx_new is not None and foto_idx_new < self.tab.columnCount():
                    foto_title = str(self.data.poz_columns_config.get("foto", {}).get("title", "Foto"))
                    self.tab.setHorizontalHeaderItem(foto_idx_new, QTableWidgetItem(foto_title))
            except Exception:
                pass

            self._poz_apply_column_visibility()
            self._update_preview()

    def dzest_kolonnu(self):
        # Dzēšam tikai pielāgotās kolonnas (custom_columns), jo bāzes kolonnas var tikai paslēpt.
        customs = [c for c in (getattr(self.data, "custom_columns", []) or []) if isinstance(c, dict)]
        if not customs:
            QMessageBox.information(self, "Dzēst kolonnu", "Nav pielāgotu kolonnu, ko dzēst.")
            return

        names = [str(c.get("name", "")) for c in customs]
        col_name, ok = QInputDialog.getItem(self, "Dzēst kolonnu", "Izvēlieties kolonnu, ko dzēst:", names, 0, False)
        if not (ok and col_name):
            return

        # Noņemam no UI pēc indeksa (custom_start + custom_index)
        try:
            idx = self._poz_col_indices()
            custom_idx = names.index(col_name)
            ui_col = idx.get("custom_start", 8) + custom_idx
            if 0 <= ui_col < self.tab.columnCount():
                self.tab.removeColumn(ui_col)
        except Exception:
            pass

        # Noņemam no datiem
        for i, col in enumerate(self.data.custom_columns):
            if isinstance(col, dict) and str(col.get('name', '')) == str(col_name):
                del self.data.custom_columns[i]
                break

        # Atjaunojam Foto virsrakstu + redzamības iestatījumus
        try:
            foto_idx_new = self._poz_col_indices().get("foto")
            if foto_idx_new is not None and foto_idx_new < self.tab.columnCount():
                foto_title = str(self.data.poz_columns_config.get("foto", {}).get("title", "Foto"))
                self.tab.setHorizontalHeaderItem(foto_idx_new, QTableWidgetItem(foto_title))
        except Exception:
            pass

        self._poz_apply_column_visibility()
        self._update_preview()

    def pievienot_pozīciju(self):
        r = self.tab.rowCount()
        self.tab.insertRow(r)

        # Inicializējam visas kolonnas
        for c in range(self.tab.columnCount()):
            self.tab.setItem(r, c, QTableWidgetItem(""))

        idx = self._poz_col_indices()

        # Noklusējumi
        if self.tab.item(r, idx["daudzums"]):
            self.tab.item(r, idx["daudzums"]).setText("1")
        if self.tab.item(r, idx["vieniba"]):
            self.tab.item(r, idx["vieniba"]).setText(self.data.default_unit)
        if self.tab.item(r, idx["cena"]):
            self.tab.item(r, idx["cena"]).setText("0.00")
        if self.tab.item(r, idx["summa"]):
            self.tab.item(r, idx["summa"]).setText("0.00")

        # Foto poga (ja ieslēgta)
        self._ensure_photo_cell(r)

    def dzest_pozīciju(self):
        r = self.tab.currentRow()
        if r >= 0:
            self.tab.removeRow(r)

    
    def _pārrēķināt_summa(self, row, col):
        """Pārrēķina 'Summa' pēc daudzuma un cenas, izmantojot dinamiskos kolonnu indeksus."""
        try:
            idx = self._poz_col_indices()
            col_daudz = idx.get("daudzums", 1)
            col_cena = idx.get("cena", 3)
            col_summa = idx.get("summa", 4)

            if col not in (col_daudz, col_cena):
                return

            daudz = to_decimal(self.tab.item(row, col_daudz).text()) if self.tab.item(row, col_daudz) else Decimal("0")
            cena = to_decimal(self.tab.item(row, col_cena).text()) if self.tab.item(row, col_cena) else Decimal("0")
            summa = (daudz * cena).quantize(Decimal("0.01"))

            if not self.tab.item(row, col_summa):
                self.tab.setItem(row, col_summa, QTableWidgetItem(""))
            self.tab.item(row, col_summa).setText(formēt_naudu(summa))
        finally:
            self._update_preview()

    # ----- Tab: Attēli -----
    
    # --- Tab: Attēli -----
    def _būvēt_attēli_tab(self):
        """Fotogrāfiju pievienošana tabulas veidā (ar nosaukumu/ aprakstu rediģēšanu)."""
        w = QWidget()
        v = QVBoxLayout()

        # Tabula: [Priekšskats | Nosaukums/Apraksts | Fails | ... (iekšējie dati)]
        self.photos_table = QTableWidget(0, 4)
        self.photos_table.setObjectName("photos_table")
        self.photos_table.setHorizontalHeaderLabels(["Foto", "Nosaukums / Apraksts", "Fails", ""])
        self.photos_table.verticalHeader().setVisible(False)
        self.photos_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.photos_table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.photos_table.setEditTriggers(QAbstractItemView.DoubleClicked | QAbstractItemView.SelectedClicked | QAbstractItemView.EditKeyPressed)
        self.photos_table.setAlternatingRowColors(True)
        self.photos_table.setShowGrid(True)
        self.photos_table.setColumnHidden(3, True)  # iekšējais (UserRole) datu items

        hh = self.photos_table.horizontalHeader()
        hh.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        hh.setSectionResizeMode(1, QHeaderView.Stretch)
        hh.setSectionResizeMode(2, QHeaderView.Stretch)

        self.photos_table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.photos_table.customContextMenuRequested.connect(self._photos_context_menu)

        v.addWidget(self.photos_table)

        btns = QHBoxLayout()

        piev = QPushButton("Pievienot foto…")
        piev.clicked.connect(self.pievienot_attēlu)

        # Mazas pārvietošanas pogas (strādā uz atlasīto rindu)
        augsa = QToolButton(); augsa.setText("▲"); augsa.setToolTip("Pārvietot uz augšu")
        augsa.clicked.connect(lambda: self.pārvietot_att(-1))
        leja = QToolButton(); leja.setText("▼"); leja.setToolTip("Pārvietot uz leju")
        leja.clicked.connect(lambda: self.pārvietot_att(1))

        dzest = QToolButton(); dzest.setText("🗑"); dzest.setToolTip("Dzēst atlasīto")
        dzest.clicked.connect(self.dzest_att)

        btns.addWidget(piev)
        btns.addSpacing(8)
        btns.addWidget(augsa)
        btns.addWidget(leja)
        btns.addWidget(dzest)
        btns.addStretch()

        v.addLayout(btns)
        w.setLayout(v)
        self.tabs.addTab(w, "Fotogrāfijas")

        # Kad maina nosaukumu/ aprakstu → atjauno preview
        self.photos_table.itemChanged.connect(self._photos_item_changed)
        self.photos_table.cellDoubleClicked.connect(self._photos_cell_double_clicked)

        # Back-compat: dažās vietās kods izmanto self.img_list (vēsturiskais QListView/QListWidget).
        # Lai neko “nenolauztu”, mēs norādām to uz jauno tabulu, kurai arī ir model() ar tiem pašiem signāliem.
        self.img_list = self.photos_table
        try:
            m = self.img_list.model()
            m.rowsInserted.connect(self._update_preview)
            m.rowsRemoved.connect(self._update_preview)
            m.rowsMoved.connect(self._update_preview)
            # drošībai – ja tiek pārrakstīti dati bez rindu izmaiņām
            m.dataChanged.connect(self._update_preview)
            m.layoutChanged.connect(self._update_preview)
        except Exception:
            # Ja kādā vidē model() nav pieejams/atšķiras, nelaužam startu – preview tāpat atjaunosies no citiem trigeriem.
            pass

    
    def _photos_make_thumb_widget(self, path: str) -> QWidget:
        lbl = QLabel()
        lbl.setAlignment(Qt.AlignCenter)
        lbl.setMinimumWidth(90)
        lbl.setMinimumHeight(70)
        lbl.setToolTip(path)
        try:
            pm = QPixmap(path)
            if not pm.isNull():
                pm = pm.scaled(120, 90, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                lbl.setPixmap(pm)
            else:
                lbl.setText("—")
        except Exception:
            lbl.setText("—")
        return lbl

    def _photos_row_payload(self, row: int) -> Optional[dict]:
        """Iekšējie dati par rindu (ceļš+paraksts)."""
        it = self.photos_table.item(row, 3)
        if it is None:
            return None
        return it.data(Qt.UserRole) or None

    def _photos_set_row_payload(self, row: int, payload: dict):
        it = self.photos_table.item(row, 3)
        if it is None:
            it = QTableWidgetItem()
            self.photos_table.setItem(row, 3, it)
        it.setData(Qt.UserRole, payload)

    def _photos_add_row(self, path: str, caption: str = ""):
        row = self.photos_table.rowCount()
        self.photos_table.insertRow(row)

        # Thumbnail
        self.photos_table.setCellWidget(row, 0, self._photos_make_thumb_widget(path))

        # Caption (editable)
        cap_item = QTableWidgetItem(caption or "")
        cap_item.setToolTip("Dubultklikšķis, lai rediģētu nosaukumu/ aprakstu")
        self.photos_table.setItem(row, 1, cap_item)

        # File (read-only)
        file_item = QTableWidgetItem(os.path.basename(path))
        file_item.setFlags(file_item.flags() & ~Qt.ItemIsEditable)
        file_item.setToolTip(path)
        self.photos_table.setItem(row, 2, file_item)

        # Hidden payload
        hidden = QTableWidgetItem()
        self.photos_table.setItem(row, 3, hidden)
        self._photos_set_row_payload(row, {"ceļš": path, "paraksts": caption or ""})

        self.photos_table.setCurrentCell(row, 1)

    def _photos_context_menu(self, pos: QPoint):
        menu = QMenu(self)
        row = self.photos_table.rowAt(pos.y())
        if row < 0:
            act_add = menu.addAction("Pievienot foto…")
            act_add.triggered.connect(self.pievienot_attēlu)
            menu.exec(self.photos_table.viewport().mapToGlobal(pos))
            return

        payload = self._photos_row_payload(row) or {}
        path = payload.get("ceļš", "")

        act_open = menu.addAction("Atvērt failu")
        act_open.triggered.connect(lambda: self._open_file_in_os(path) if path else None)

        act_reveal = menu.addAction("Atvērt mapē")
        act_reveal.triggered.connect(lambda: self._reveal_file_in_os(path) if path else None)

        menu.addSeparator()
        act_up = menu.addAction("Pārvietot uz augšu")
        act_up.triggered.connect(lambda: self._photos_move_selected(-1))
        act_down = menu.addAction("Pārvietot uz leju")
        act_down.triggered.connect(lambda: self._photos_move_selected(1))

        menu.addSeparator()
        act_del = menu.addAction("Dzēst")
        act_del.triggered.connect(self.dzest_att)

        menu.exec(self.photos_table.viewport().mapToGlobal(pos))

    def _photos_item_changed(self, item: QTableWidgetItem):
        # ja mainās caption kolonna → atjauno payload + preview
        if item is None:
            return
        if item.column() != 1:
            return
        row = item.row()
        payload = self._photos_row_payload(row) or {}
        payload["paraksts"] = item.text()
        self._photos_set_row_payload(row, payload)
        self._update_preview()

    def _photos_cell_double_clicked(self, row: int, col: int):
        # dubultklikšķis uz foto -> atvērt failu
        if col == 0:
            payload = self._photos_row_payload(row) or {}
            path = payload.get("ceļš", "")
            if path:
                self._open_file_in_os(path)

    def _photos_move_selected(self, direction: int):
        row = self.photos_table.currentRow()
        if row < 0:
            return
        new_row = row + direction
        if new_row < 0 or new_row >= self.photos_table.rowCount():
            return

        payload = self._photos_row_payload(row) or {}
        caption = (self.photos_table.item(row, 1).text() if self.photos_table.item(row, 1) else payload.get("paraksts", ""))
        path = payload.get("ceļš", "")

        # remove
        self.photos_table.blockSignals(True)
        self.photos_table.removeRow(row)
        self.photos_table.insertRow(new_row)

        # rebuild at new position
        self.photos_table.setCellWidget(new_row, 0, self._photos_make_thumb_widget(path))
        cap_item = QTableWidgetItem(caption or "")
        self.photos_table.setItem(new_row, 1, cap_item)
        file_item = QTableWidgetItem(os.path.basename(path))
        file_item.setFlags(file_item.flags() & ~Qt.ItemIsEditable)
        file_item.setToolTip(path)
        self.photos_table.setItem(new_row, 2, file_item)
        hidden = QTableWidgetItem()
        self.photos_table.setItem(new_row, 3, hidden)
        self._photos_set_row_payload(new_row, {"ceļš": path, "paraksts": caption or ""})
        self.photos_table.blockSignals(False)

        self.photos_table.setCurrentCell(new_row, 1)
        self._update_preview()

    def pievienot_attēlu(self):
        ceļi, _ = QFileDialog.getOpenFileNames(
            self,
            "Izvēlēties attēlus",
            "",
            "Attēli (*.png *.jpg *.jpeg *.webp)"
        )
        for c in ceļi:
            self._photos_add_row(c, "")

        if ceļi:
            self._update_preview()

    def pārvietot_att(self, virziens):
        self._photos_move_selected(virziens)

    def dzest_att(self):
        r = self.photos_table.currentRow()
        if r >= 0:
            self.photos_table.removeRow(r)
            self._update_preview()

    # Back-compat: vecais API (vairs neizmantojam, bet atstājam, lai nekas nelūzt)
    def rediģēt_att_parakstu(self, *args, **kwargs):
        # Tagad paraksts/nosaukums rediģējas tieši tabulā (2. kolonna).
        return
        current_text = item.data(Qt.UserRole).get("paraksts", "")
        text, ok = QInputDialog.getText(self, "Rediģēt parakstu", "Ievadiet attēla parakstu:", QLineEdit.Normal, current_text)
        if ok and text is not None:
            d = item.data(Qt.UserRole)
            d["paraksts"] = text
            item.setData(Qt.UserRole, d)
            self._update_preview()

    # ----- Tab: Iestatījumi & Eksports -----
    def _būvēt_iestatījumi_tab(self):
        w = QWidget()
        form = QFormLayout()

        self.ck_pvn = QCheckBox("Aprēķināt PVN")
        self.in_pvn = QDoubleSpinBox(); self.in_pvn.setRange(0, 100); self.in_pvn.setValue(21.0); self.in_pvn.setSuffix(" %")

        self.ck_paraksti = QCheckBox("Iekļaut parakstu rindas")
        self.ck_paraksti.setChecked(True)

        self.in_logo = QLineEdit();
        logo_btn = QToolButton(); logo_btn.setText("…"); logo_btn.clicked.connect(self.izvēlēties_logo)
        logo_box = QHBoxLayout(); logo_box.addWidget(self.in_logo); logo_box.addWidget(logo_btn)
        logo_w = QWidget(); logo_w.setLayout(logo_box)

        self.in_fonts = QLineEdit();
        font_btn = QToolButton(); font_btn.setText("…"); font_btn.clicked.connect(self.izvēlēties_fontu)
        font_box = QHBoxLayout(); font_box.addWidget(self.in_fonts); font_box.addWidget(font_btn)
        font_w = QWidget(); font_w.setLayout(font_box)

        # DOCX šablons (nav obligāts)
        self.in_docx_template = QLineEdit()
        docx_tpl_btn = QToolButton(); docx_tpl_btn.setText("…")
        docx_tpl_btn.clicked.connect(self.izvēlēties_docx_sablonu)
        docx_tpl_box = QHBoxLayout(); docx_tpl_box.addWidget(self.in_docx_template); docx_tpl_box.addWidget(docx_tpl_btn)
        docx_tpl_w = QWidget(); docx_tpl_w.setLayout(docx_tpl_box)

        self.in_paraksts_pie = QLineEdit();
        btn_paraksts_pie = QToolButton(); btn_paraksts_pie.setText("…"); btn_paraksts_pie.clicked.connect(lambda: self.izvēlēties_paraksta_attēlu(self.in_paraksts_pie))
        box_paraksts_pie = QHBoxLayout(); box_paraksts_pie.addWidget(self.in_paraksts_pie); box_paraksts_pie.addWidget(btn_paraksts_pie)
        w_paraksts_pie = QWidget(); w_paraksts_pie.setLayout(box_paraksts_pie)

        self.in_paraksts_nod = QLineEdit();
        btn_paraksts_nod = QToolButton(); btn_paraksts_nod.setText("…"); btn_paraksts_nod.clicked.connect(lambda: self.izvēlēties_paraksta_attēlu(self.in_paraksts_nod))
        box_paraksts_nod = QHBoxLayout(); box_paraksts_nod.addWidget(self.in_paraksts_nod); box_paraksts_nod.addWidget(btn_paraksts_nod)
        w_paraksts_nod = QWidget(); w_paraksts_nod.setLayout(box_paraksts_nod)

        self.lang_combo = QComboBox()
        self.lang_combo.addItem("Latviešu")
        self.lang_combo.addItem("English (nav implementēts)")
        self.lang_combo.setEnabled(False)

        btn_saglabat_nokl = QPushButton("Saglabāt kā noklusējumu")
        btn_saglabat_nokl.clicked.connect(self.saglabat_noklusejuma_iestatijumus)

        # JAUNA POGA ŠABLONU SAGLABĀŠANAI
        btn_saglabat_sablonu = QPushButton("Saglabāt kā šablonu")
        btn_saglabat_sablonu.clicked.connect(self.saglabat_ka_sablonu)


        self.btn_generate_pdf = QPushButton("Ģenerēt PDF…")
        self.btn_generate_pdf.clicked.connect(self.ģenerēt_pdf_dialogs)

        self.btn_generate_docx = QPushButton("Ģenerēt DOCX…")
        self.btn_generate_docx.clicked.connect(self.ģenerēt_docx_dialogs)

        self.btn_export_zip = QPushButton("Saglabāt ZIP…")
        self.btn_export_zip.clicked.connect(self.ģenerēt_zip_dialogs)
        self.btn_print_pdf = QPushButton("Drukāt PDF…") # JAUNA RINDAS
        self.btn_print_pdf.clicked.connect(self.drukāt_pdf_dialogs) # JAUNA RINDAS



        form.addRow(self.ck_pvn, self.in_pvn)
        form.addRow(self.ck_paraksti)
        form.addRow("Logotips (neobligāti)", logo_w)
        form.addRow("Fonts TTF/OTF (ieteicams latviešu diakritikai)", font_w)
        form.addRow("DOCX šablons (opc.)", docx_tpl_w)
        form.addRow("Pieņēmēja paraksta attēls", w_paraksts_pie)
        form.addRow("Nodevēja paraksta attēls", w_paraksts_nod)
        form.addRow("Valoda", self.lang_combo)
        form.addRow(btn_saglabat_nokl)
        form.addRow(btn_saglabat_sablonu)  # JAUNA RINDAS
        form.addRow(self.btn_generate_pdf)
        form.addRow(self.btn_generate_docx)
        form.addRow(self.btn_export_zip)
        form.addRow(self.btn_print_pdf)  # JAUNA RINDAS

        self.ck_pvn.stateChanged.connect(self._update_preview)
        self.in_pvn.valueChanged.connect(self._update_preview)
        self.ck_paraksti.stateChanged.connect(self._update_preview)
        self.in_logo.textChanged.connect(self._update_preview)
        self.in_fonts.textChanged.connect(self._update_preview)
        self.in_paraksts_pie.textChanged.connect(self._update_preview)
        self.in_paraksts_nod.textChanged.connect(self._update_preview)


        w.setLayout(form)
        self.tabs.addTab(w, "Iestatījumi & Eksports")

    
    def _wrap_date_with_system_button(self, date_edit: QDateEdit) -> QWidget:
        """Ietin QDateEdit ar pogu, kas ielādē sistēmas (Windows) šodienas datumu."""
        w = QWidget()
        h = QHBoxLayout()
        h.setContentsMargins(0, 0, 0, 0)
        h.addWidget(date_edit, 1)
        btn = QToolButton()
        btn.setText("Ielādēt sistēmas datumu")
        btn.clicked.connect(lambda: date_edit.setDate(datetime.now().date()))
        h.addWidget(btn)
        w.setLayout(h)
        return w

    def izvēlēties_logo(self):
        c, _ = QFileDialog.getOpenFileName(self, "Izvēlēties logotipu", "", "Attēli (*.png *.jpg *.jpeg *.webp)")
        if c:
            self.in_logo.setText(c)

    def drukāt_pdf_dialogs(self):
        """
        Ģenerē PDF dokumentu un atver drukas priekšskatījuma dialogu.
        """
        akta_dati = self.savākt_datus()
        temp_pdf_path = None
        try:
            # Ģenerējam PDF pagaidu failā
            temp_pdf_path = ģenerēt_pdf(akta_dati, pdf_ceļš=None, encrypt_pdf=False)

            if not os.path.exists(temp_pdf_path):
                QMessageBox.critical(self, "Drukas kļūda", "Neizdevās ģenerēt PDF failu drukāšanai.")
                return

            printer = QPrinter(QPrinter.HighResolution)

            # Lietojam PROGRAMMAS lapas izmēru (nevis printera/sistēmas noklusējumu, kas bieži ir Letter)
            try:
                page_name = (akta_dati.pdf_page_size or "A4").upper().strip()
            except Exception:
                page_name = "A4"

            page_map = {
                "A4": QPageSize.A4,
                "A3": QPageSize.A3,
                "A5": QPageSize.A5,
                "LETTER": QPageSize.Letter,
                "LEGAL": QPageSize.Legal,
            }
            qt_page = page_map.get(page_name, QPageSize.A4)
            printer.setPageSize(QPageSize(qt_page))
            # Orientācija pēc iestatījuma
            try:
                if getattr(akta_dati, "pdf_orientation", "Portrait").lower().startswith("land"):
                    printer.setOrientation(QPrinter.Landscape)
                else:
                    printer.setOrientation(QPrinter.Portrait)
            except Exception:
                pass

            preview_dialog = QPrintPreviewDialog(printer, self)
            preview_dialog.paintRequested.connect(lambda printer_obj: self._render_pdf_to_printer(temp_pdf_path, printer_obj))
            preview_dialog.exec()

        except Exception as e:
            QMessageBox.critical(self, "Drukas kļūda", f"Kļūda sagatavojot dokumentu drukāšanai: {e}")
        finally:
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path) # Dzēšam pagaidu failu
                except Exception as e:
                    print(f"Neizdevās dzēst pagaidu PDF failu: {e}")

    def _render_pdf_to_printer(self, pdf_path: str, printer: QPrinter):
        """
        Renderē PDF failu uz printera.
        Izmanto pdf2image, lai iegūtu attēlus no PDF un zīmē tos uz printera.
        """
        try:
            # Renderēšana tiek izsaukta vairākas reizes (preview pārzīmēšana). Lai neuzkārtu programmu,
            # kešojam lapu attēlus atmiņā vienam pdf_path.
            if not hasattr(self, "_print_images_cache"):
                self._print_images_cache = {}

            cache_key = os.path.abspath(pdf_path)
            images = self._print_images_cache.get(cache_key)

            if images is None:
                # Pārliecināmies, ka poppler_path ir pieejams
                poppler_path_to_use = self.data.poppler_path if self.data.poppler_path and os.path.exists(self.data.poppler_path) else None
                # Mazliet zemāka DPI vērtība = ātrāk un stabilāk drukas priekšskatījumā
                render_path, tmp_cleanup = _prepare_unencrypted_pdf_for_render(pdf_path, password=getattr(self.data, "pdf_user_password", "") or "")
                try:
                    images = convert_from_path(render_path, poppler_path=poppler_path_to_use, dpi=150)
                finally:
                    try:
                        if tmp_cleanup and os.path.exists(tmp_cleanup):
                            os.remove(tmp_cleanup)
                    except Exception:
                        pass

                self._print_images_cache[cache_key] = images

            if not images:
                QMessageBox.warning(self, "Drukas kļūda", "Neizdevās iegūt attēlus no PDF faila drukāšanai.")
                return

            painter = QPainter()
            if not painter.begin(printer):
                QMessageBox.critical(self, "Drukas kļūda", "Neizdevās sākt zīmēšanu uz printera.")
                return

            for i, pil_img in enumerate(images):
                if i > 0:
                    printer.newPage()  # Jauna lapa katram attēlam (PDF lapai)

                # Konvertējam PIL attēlu uz QImage
                q_image = ImageQt(pil_img)
                pixmap = QPixmap.fromImage(q_image)

                # Mērogojam attēlu, lai tas ietilptu printera lapā
                # printer.pageRect() atgriež lapas izmērus pikseļos, ņemot vērā printera izšķirtspēju
                # Labots: Izmantojam painter.device().width() un painter.device().height()
                # lai iegūtu zīmēšanas ierīces (printera) izmērus.
                printer_width = painter.device().width()
                printer_height = painter.device().height()

                # Izmantojam QSize, lai scaled() metodei nodotu pareizu izmēru
                target_size = QSize(printer_width, printer_height)
                scaled_pixmap = pixmap.scaled(target_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)

                # Zīmējam attēlu lapas centrā
                # Labots: Izmantojam printera izmērus, lai centrētu attēlu
                x = (printer_width - scaled_pixmap.width()) / 2
                y = (printer_height - scaled_pixmap.height()) / 2
                painter.drawPixmap(int(x), int(y), scaled_pixmap)

            painter.end()
            QMessageBox.information(self, "Drukas priekšskatījums", "Dokuments sagatavots drukāšanai.")

        except Exception as e:
            QMessageBox.critical(self, "Drukas kļūda", f"Kļūda renderējot PDF uz printeri: {e}")
        finally:
            # Pārliecināmies, ka painter tiek beigts, pat ja rodas kļūda
            if painter.isActive():
                painter.end()



    def izvēlēties_fontu(self):
        c, _ = QFileDialog.getOpenFileName(self, "Izvēlieties fontu (TTF/OTF)", "", "Fonti (*.ttf *.otf)")
        if c:
            self.in_fonts.setText(c)

    def izvēlēties_docx_sablonu(self):
        c, _ = QFileDialog.getOpenFileName(self, "Izvēlēties DOCX šablonu", "", "Word dokuments (*.docx)")
        if c:
            self.in_docx_template.setText(c)
            self._update_preview()

    def izvēlēties_paraksta_attēlu(self, line_edit: QLineEdit):
        c, _ = QFileDialog.getOpenFileName(self, "Izvēlieties paraksta attēlu", "", "Attēli (*.png *.jpg *.jpeg *.webp)")
        if c:
            line_edit.setText(c)

    # ----- Tab: Papildu iestatījumi -----
    def _būvēt_papildu_iestatījumi_tab(self):
        w = QWidget()
        form = QFormLayout()

        self.cb_page_size = QComboBox()
        self.cb_page_size.addItems(["A4", "Letter", "Legal", "A3", "A5"])
        self.cb_page_orientation = QComboBox()
        self.cb_page_orientation.addItems(["Portrets", "Ainava"])

        self.in_margin_left = QDoubleSpinBox(); self.in_margin_left.setRange(0, 100); self.in_margin_left.setSuffix(" mm")
        self.in_margin_right = QDoubleSpinBox(); self.in_margin_right.setRange(0, 100); self.in_margin_right.setSuffix(" mm")
        self.in_margin_top = QDoubleSpinBox(); self.in_margin_top.setRange(0, 100); self.in_margin_top.setSuffix(" mm")
        self.in_margin_bottom = QDoubleSpinBox(); self.in_margin_bottom.setRange(0, 100); self.in_margin_bottom.setSuffix(" mm")

        self.in_font_size_head = QSpinBox(); self.in_font_size_head.setRange(8, 30)
        self.in_font_size_normal = QSpinBox(); self.in_font_size_normal.setRange(6, 20)
        self.in_font_size_small = QSpinBox(); self.in_font_size_small.setRange(4, 16)
        self.in_font_size_table = QSpinBox(); self.in_font_size_table.setRange(4, 16)

        self.in_logo_width_mm = QDoubleSpinBox(); self.in_logo_width_mm.setRange(10, 100); self.in_logo_width_mm.setSuffix(" mm")
        self.in_signature_width_mm = QDoubleSpinBox(); self.in_signature_width_mm.setRange(10, 100); self.in_signature_width_mm.setSuffix(" mm")
        self.in_signature_height_mm = QDoubleSpinBox(); self.in_signature_height_mm.setRange(5, 50); self.in_signature_height_mm.setSuffix(" mm")

        self.in_docx_image_width_inches = QDoubleSpinBox(); self.in_docx_image_width_inches.setRange(1, 10); self.in_docx_image_width_inches.setSuffix(" collas")
        self.in_docx_signature_width_inches = QDoubleSpinBox(); self.in_docx_signature_width_inches.setRange(0.5, 5); self.in_docx_signature_width_inches.setSuffix(" collas")

        self.in_table_col_widths = QTextEdit()
        self.in_table_col_widths.setPlaceholderText("Ievadiet kolonnu platumus mm, atdalot ar komatiem (piem., 10,40,18,18,20,20,25,25,25)")

        self.ck_auto_generate_akta_nr = QCheckBox("Automātiski ģenerēt akta numuru (PP-YYYY-NNNN)")

        self.in_default_execution_days = QSpinBox()
        self.in_default_execution_days.setRange(0, 365)
        self.in_default_execution_days.setValue(5)
        self.in_default_execution_days.setSuffix(" d.")
        self.in_default_execution_days.setToolTip("Izpildes termiņš = šodiena + N dienas (noklusējums jaunam aktam)")

        self.in_default_currency = QLineEdit()
        self.in_default_unit = QLineEdit()
        self.in_default_pvn_rate = QDoubleSpinBox(); self.in_default_pvn_rate.setRange(0, 100); self.in_default_pvn_rate.setSuffix(" %")

        self.in_poppler_path = QLineEdit()
        poppler_btn = QToolButton(); poppler_btn.setText("…"); poppler_btn.clicked.connect(self.izvēlēties_poppler_ceļu)
        poppler_box = QHBoxLayout(); poppler_box.addWidget(self.in_poppler_path); poppler_box.addWidget(poppler_btn)
        poppler_w = QWidget(); poppler_w.setLayout(poppler_box)

        # New settings (30+ functions/options/settings)
        self.in_header_text_color = QLineEdit("#000000")
        self.in_footer_text_color = QLineEdit("#000000")
        self.in_table_header_bg_color = QLineEdit("#E0E0E0")
        self.in_table_grid_color = QLineEdit("#CCCCCC")
        self.in_table_row_spacing = QDoubleSpinBox(); self.in_table_row_spacing.setRange(0, 10); self.in_table_row_spacing.setValue(4); self.in_table_row_spacing.setSuffix(" mm")
        self.in_line_spacing_multiplier = QDoubleSpinBox(); self.in_line_spacing_multiplier.setRange(0.5, 3.0); self.in_line_spacing_multiplier.setSingleStep(0.1); self.in_line_spacing_multiplier.setValue(1.2)
        self.ck_show_page_numbers = QCheckBox("Rādīt lapu numurus"); self.ck_show_page_numbers.setChecked(True)
        self.ck_show_generation_timestamp = QCheckBox("Rādīt ģenerēšanas laiku"); self.ck_show_generation_timestamp.setChecked(True)
        self.cb_currency_symbol_position = QComboBox(); self.cb_currency_symbol_position.addItems(["after", "before"])
        self.in_date_format = QLineEdit("YYYY-MM-DD")
        self.in_signature_line_length_mm = QDoubleSpinBox(); self.in_signature_line_length_mm.setRange(10, 100); self.in_signature_line_length_mm.setValue(60); self.in_signature_line_length_mm.setSuffix(" mm")
        self.in_signature_line_thickness_pt = QDoubleSpinBox(); self.in_signature_line_thickness_pt.setRange(0.1, 5.0); self.in_signature_line_thickness_pt.setSingleStep(0.1); self.in_signature_line_thickness_pt.setValue(0.5); self.in_signature_line_thickness_pt.setSuffix(" pt")
        self.ck_add_cover_page = QCheckBox("Pievienot titullapu")
        self.in_cover_page_title = QLineEdit("Pieņemšanas-Nodošanas Akts")
        self.in_cover_page_logo_width_mm = QDoubleSpinBox(); self.in_cover_page_logo_width_mm.setRange(10, 200); self.in_cover_page_logo_width_mm.setValue(80); self.in_cover_page_logo_width_mm.setSuffix(" mm")
        # Individuālais QR kods
        self.ck_include_custom_qr_code = QCheckBox("Iekļaut individuālu QR kodu")
        self.in_custom_qr_code_data = QLineEdit()
        self.in_custom_qr_code_data.setPlaceholderText("Dati individuālajam QR kodam (URL, teksts utt.)")
        self.in_custom_qr_code_size_mm = QDoubleSpinBox();
        self.in_custom_qr_code_size_mm.setRange(10, 50);
        self.in_custom_qr_code_size_mm.setValue(20);
        self.in_custom_qr_code_size_mm.setSuffix(" mm")
        self.cb_custom_qr_code_position = QComboBox();
        self.cb_custom_qr_code_position.addItems(["bottom_right", "bottom_left", "top_right", "top_left", "custom"])
        self.in_custom_qr_code_pos_x_mm = QDoubleSpinBox();
        self.in_custom_qr_code_pos_x_mm.setRange(0, 500);
        self.in_custom_qr_code_pos_x_mm.setSuffix(" mm")
        self.in_custom_qr_code_pos_y_mm = QDoubleSpinBox();
        self.in_custom_qr_code_pos_y_mm.setRange(0, 500);
        self.in_custom_qr_code_pos_y_mm.setSuffix(" mm")
        self.in_custom_qr_code_color = QLineEdit("#000000")  # QR koda krāsa (Hex)

        # Automātiskais QR kods (akta ID)
        self.ck_include_auto_qr_code = QCheckBox("Iekļaut automātisku QR kodu (Akta ID)")
        self.in_auto_qr_code_size_mm = QDoubleSpinBox();
        self.in_auto_qr_code_size_mm.setRange(10, 50);
        self.in_auto_qr_code_size_mm.setValue(20);
        self.in_auto_qr_code_size_mm.setSuffix(" mm")
        self.cb_auto_qr_code_position = QComboBox();
        self.cb_auto_qr_code_position.addItems(["bottom_left", "bottom_right", "top_right", "top_left", "custom"])
        self.in_auto_qr_code_pos_x_mm = QDoubleSpinBox();
        self.in_auto_qr_code_pos_x_mm.setRange(0, 500);
        self.in_auto_qr_code_pos_x_mm.setSuffix(" mm")
        self.in_auto_qr_code_pos_y_mm = QDoubleSpinBox();
        self.in_auto_qr_code_pos_y_mm.setRange(0, 500);
        self.in_auto_qr_code_pos_y_mm.setSuffix(" mm")
        self.in_auto_qr_code_color = QLineEdit("#000000")  # QR koda krāsa (Hex)

        self.ck_add_watermark = QCheckBox("Pievienot ūdenszīmi")
        self.in_watermark_text = QLineEdit("MELNRAKSTS")
        self.in_watermark_font_size = QSpinBox(); self.in_watermark_font_size.setRange(10, 200); self.in_watermark_font_size.setValue(72)
        self.in_watermark_color = QLineEdit("#E0E0E0")
        self.in_watermark_rotation = QSpinBox(); self.in_watermark_rotation.setRange(0, 360); self.in_watermark_rotation.setValue(45)
        self.ck_enable_pdf_encryption = QCheckBox("Iespējot PDF šifrēšanu")
        self.in_pdf_user_password = QLineEdit(); self.in_pdf_user_password.setEchoMode(QLineEdit.Password)
        self.in_pdf_owner_password = QLineEdit(); self.in_pdf_owner_password.setEchoMode(QLineEdit.Password)
        # --- JAUNS: sinhronizācija ar Pamata datu šifrēšanas lauku (ja tāds ir) ---
        try:
            if hasattr(self, "ck_pdf_encrypt_basic") and hasattr(self, "in_pdf_password_basic"):
                # sākotnējā sinhronizācija
                self.ck_enable_pdf_encryption.setChecked(self.ck_pdf_encrypt_basic.isChecked())
                if self.in_pdf_password_basic.text().strip():
                    self.in_pdf_user_password.setText(self.in_pdf_password_basic.text().strip())

                # divvirzienu sinhronizācija
                self.ck_enable_pdf_encryption.stateChanged.connect(lambda *_: self.ck_pdf_encrypt_basic.setChecked(self.ck_enable_pdf_encryption.isChecked()))
                self.in_pdf_user_password.textChanged.connect(lambda *_: self.in_pdf_password_basic.setText(self.in_pdf_user_password.text()))
                self.ck_pdf_encrypt_basic.stateChanged.connect(lambda *_: self.ck_enable_pdf_encryption.setChecked(self.ck_pdf_encrypt_basic.isChecked()))
                self.in_pdf_password_basic.textChanged.connect(lambda *_: self.in_pdf_user_password.setText(self.in_pdf_password_basic.text()))
        except Exception:
            pass

        self.ck_allow_printing = QCheckBox("Atļaut drukāšanu"); self.ck_allow_printing.setChecked(True)
        self.ck_allow_copying = QCheckBox("Atļaut kopēšanu"); self.ck_allow_copying.setChecked(True)
        self.ck_allow_modifying = QCheckBox("Atļaut modificēšanu")
        self.ck_allow_annotating = QCheckBox("Atļaut anotēšanu"); self.ck_allow_annotating.setChecked(True)
        self.in_default_country = QLineEdit("Latvija")
        self.in_default_city = QLineEdit("Rīga")
        self.ck_show_contact_details_in_header = QCheckBox("Rādīt kontaktinformāciju galvenē")
        self.in_contact_details_header_font_size = QSpinBox(); self.in_contact_details_header_font_size.setRange(6, 12); self.in_contact_details_header_font_size.setValue(8)
        self.in_item_image_width_mm = QDoubleSpinBox(); self.in_item_image_width_mm.setRange(10, 150); self.in_item_image_width_mm.setValue(50); self.in_item_image_width_mm.setSuffix(" mm")
        self.in_item_image_caption_font_size = QSpinBox(); self.in_item_image_caption_font_size.setRange(6, 12); self.in_item_image_caption_font_size.setValue(8)
        self.ck_show_item_notes_in_table = QCheckBox("Rādīt pozīciju piezīmes tabulā"); self.ck_show_item_notes_in_table.setChecked(True)
        self.ck_show_item_serial_number_in_table = QCheckBox("Rādīt pozīciju sērijas Nr. tabulā"); self.ck_show_item_serial_number_in_table.setChecked(True)
        self.ck_show_item_warranty_in_table = QCheckBox("Rādīt pozīciju garantiju tabulā"); self.ck_show_item_warranty_in_table.setChecked(True)
        self.in_table_cell_padding_mm = QDoubleSpinBox(); self.in_table_cell_padding_mm.setRange(0, 10); self.in_table_cell_padding_mm.setValue(2); self.in_table_cell_padding_mm.setSuffix(" mm")
        self.cb_table_header_font_style = QComboBox(); self.cb_table_header_font_style.addItems(["bold", "italic", "normal"])
        self.cb_table_content_alignment = QComboBox(); self.cb_table_content_alignment.addItems(["left", "center", "right"])
        self.in_signature_font_size = QSpinBox(); self.in_signature_font_size.setRange(6, 12); self.in_signature_font_size.setValue(9)
        self.in_signature_spacing_mm = QDoubleSpinBox(); self.in_signature_spacing_mm.setRange(0, 20); self.in_signature_spacing_mm.setValue(10); self.in_signature_spacing_mm.setSuffix(" mm")
        self.in_document_title_font_size = QSpinBox(); self.in_document_title_font_size.setRange(10, 30); self.in_document_title_font_size.setValue(18)
        self.in_document_title_color = QLineEdit("#000000")
        self.in_section_heading_font_size = QSpinBox(); self.in_section_heading_font_size.setRange(8, 20); self.in_section_heading_font_size.setValue(12)
        self.in_section_heading_color = QLineEdit("#000000")
        self.in_paragraph_line_spacing_multiplier = QDoubleSpinBox(); self.in_paragraph_line_spacing_multiplier.setRange(0.5, 3.0); self.in_paragraph_line_spacing_multiplier.setSingleStep(0.1); self.in_paragraph_line_spacing_multiplier.setValue(1.2)
        self.cb_table_border_style = QComboBox(); self.cb_table_border_style.addItems(["solid", "dashed", "none"])
        self.in_table_border_thickness_pt = QDoubleSpinBox(); self.in_table_border_thickness_pt.setRange(0.1, 5.0); self.in_table_border_thickness_pt.setSingleStep(0.1); self.in_table_border_thickness_pt.setValue(0.5); self.in_table_border_thickness_pt.setSuffix(" pt")
        self.in_table_alternate_row_color = QLineEdit("")
        self.in_table_alternate_row_color.setPlaceholderText("Hex krāsa, piem. #F0F0F0")
        self.ck_show_total_sum_in_words = QCheckBox("Rādīt kopsummu vārdos")
        self.cb_total_sum_in_words_language = QComboBox(); self.cb_total_sum_in_words_language.addItems(["lv", "en"])
        self.cb_default_vat_calculation_method = QComboBox(); self.cb_default_vat_calculation_method.addItems(["exclusive", "inclusive"])
        self.ck_show_vat_breakdown = QCheckBox("Rādīt PVN sadalījumu"); self.ck_show_vat_breakdown.setChecked(True)
        self.ck_enable_digital_signature_field = QCheckBox("Iespējot digitālā paraksta lauku (PDF)");
        self.in_digital_signature_field_name = QLineEdit("Paraksts");
        self.in_digital_signature_field_size_mm = QDoubleSpinBox();
        self.in_digital_signature_field_size_mm.setRange(10, 100);
        self.in_digital_signature_field_size_mm.setValue(40);
        self.in_digital_signature_field_size_mm.setSuffix(" mm")
        self.cb_digital_signature_field_position = QComboBox();
        self.cb_digital_signature_field_position.addItems(
            ["bottom_center", "bottom_left", "bottom_right", "top_left", "top_right"])

        # JAUNA RINDAS - Šablonu direktorija iestatījums
        self.in_templates_dir = QLineEdit()
        self.in_templates_dir.setText(os.path.join(APP_DATA_DIR, "AktaGenerators_Templates"))  # Noklusējuma vērtība
        btn_templates_dir = QToolButton();
        btn_templates_dir.setText("…");
        btn_templates_dir.clicked.connect(self.izvēlēties_templates_dir)
        templates_dir_box = QHBoxLayout();
        templates_dir_box.addWidget(self.in_templates_dir);
        templates_dir_box.addWidget(btn_templates_dir)
        templates_dir_w = QWidget();
        templates_dir_w.setLayout(templates_dir_box)

        form.addRow("PDF lapas izmērs:", self.cb_page_size)
        form.addRow("PDF lapas orientācija:", self.cb_page_orientation)
        form.addRow("PDF kreisā mala (mm):", self.in_margin_left)
        form.addRow("PDF labā mala (mm):", self.in_margin_right)
        form.addRow("PDF augšējā mala (mm):", self.in_margin_top)
        form.addRow("PDF apakšējā mala (mm):", self.in_margin_bottom)
        form.addRow("PDF galvenes fonta izmērs:", self.in_font_size_head)
        form.addRow("PDF normāla fonta izmērs:", self.in_font_size_normal)
        form.addRow("PDF maza fonta izmērs:", self.in_font_size_small)
        form.addRow("PDF tabulas fonta izmērs:", self.in_font_size_table)
        form.addRow("PDF logo platums (mm):", self.in_logo_width_mm)
        form.addRow("PDF paraksta attēla platums (mm):", self.in_signature_width_mm)
        form.addRow("PDF paraksta attēla augstums (mm):", self.in_signature_height_mm)
        form.addRow("DOCX attēlu platums (collas):", self.in_docx_image_width_inches)
        form.addRow("DOCX paraksta attēla platums (collas):", self.in_docx_signature_width_inches)
        form.addRow("Pozīciju tabulas kolonnu platumi (mm, komatiem atdalīti):", self.in_table_col_widths)
        form.addRow(self.ck_auto_generate_akta_nr)
        form.addRow("Izpildes termiņš + dienas (noklusējums):", self.in_default_execution_days)
        form.addRow("Noklusējuma valūta:", self.in_default_currency)
        form.addRow("Noklusējuma vienība:", self.in_default_unit)
        form.addRow("Noklusējuma PVN likme (%):", self.in_default_pvn_rate)
        form.addRow("Poppler bin direktorijas ceļš (Windows):", poppler_w)

        # Add new settings to the form
        form.addRow("Galvenes teksta krāsa (Hex):", self.in_header_text_color)
        form.addRow("Kājenes teksta krāsa (Hex):", self.in_footer_text_color)
        form.addRow("Tabulas galvenes fona krāsa (Hex):", self.in_table_header_bg_color)
        form.addRow("Tabulas režģa krāsa (Hex):", self.in_table_grid_color)
        form.addRow("Tabulas rindu atstarpe (mm):", self.in_table_row_spacing)
        form.addRow("Rindu atstarpes reizinātājs:", self.in_line_spacing_multiplier)
        form.addRow(self.ck_show_page_numbers)
        form.addRow(self.ck_show_generation_timestamp)
        form.addRow("Valūtas simbola pozīcija:", self.cb_currency_symbol_position)
        form.addRow("Datuma formāts:", self.in_date_format)
        form.addRow("Paraksta līnijas garums (mm):", self.in_signature_line_length_mm)
        form.addRow("Paraksta līnijas biezums (pt):", self.in_signature_line_thickness_pt)
        form.addRow(self.ck_add_cover_page)
        form.addRow("Titullapas virsraksts:", self.in_cover_page_title)
        form.addRow("Titullapas logo platums (mm):", self.in_cover_page_logo_width_mm)

        # Individuālais QR kods
        form.addRow(self.ck_include_custom_qr_code)
        form.addRow("Individuālā QR koda dati:", self.in_custom_qr_code_data)
        form.addRow("Individuālā QR koda izmērs (mm):", self.in_custom_qr_code_size_mm)
        form.addRow("Individuālā QR koda pozīcija:", self.cb_custom_qr_code_position)
        form.addRow("Individuālā QR koda X pozīcija (mm):", self.in_custom_qr_code_pos_x_mm)
        form.addRow("Individuālā QR koda Y pozīcija (mm):", self.in_custom_qr_code_pos_y_mm)
        form.addRow("Individuālā QR koda krāsa (Hex):", self.in_custom_qr_code_color)

        # Automātiskais QR kods (akta ID)
        form.addRow(self.ck_include_auto_qr_code)
        form.addRow("Automātiskā QR koda izmērs (mm):", self.in_auto_qr_code_size_mm)
        form.addRow("Automātiskā QR koda pozīcija:", self.cb_auto_qr_code_position)
        form.addRow("Automātiskā QR koda X pozīcija (mm):", self.in_auto_qr_code_pos_x_mm)
        form.addRow("Automātiskā QR koda Y pozīcija (mm):", self.in_auto_qr_code_pos_y_mm)
        form.addRow("Automātiskā QR koda krāsa (Hex):", self.in_auto_qr_code_color)

        form.addRow(self.ck_add_watermark)
        form.addRow("Ūdenszīmes teksts:", self.in_watermark_text)
        form.addRow("Ūdenszīmes fonta izmērs:", self.in_watermark_font_size)
        form.addRow("Ūdenszīmes krāsa (Hex):", self.in_watermark_color)
        form.addRow("Ūdenszīmes rotācija (grādi):", self.in_watermark_rotation)
        form.addRow(self.ck_enable_pdf_encryption)
        form.addRow("PDF lietotāja parole:", self.in_pdf_user_password)
        form.addRow("PDF īpašnieka parole:", self.in_pdf_owner_password)
        form.addRow(self.ck_allow_printing)
        form.addRow(self.ck_allow_copying)
        form.addRow(self.ck_allow_modifying)
        form.addRow(self.ck_allow_annotating)
        form.addRow("Noklusējuma valsts:", self.in_default_country)
        form.addRow("Noklusējuma pilsēta:", self.in_default_city)
        form.addRow(self.ck_show_contact_details_in_header)
        form.addRow("Kontaktu detaļu galvenes fonta izmērs:", self.in_contact_details_header_font_size)
        form.addRow("Pozīciju attēlu platums (mm):", self.in_item_image_width_mm)
        form.addRow("Pozīciju attēlu paraksta fonta izmērs:", self.in_item_image_caption_font_size)
        form.addRow(self.ck_show_item_notes_in_table)
        form.addRow(self.ck_show_item_serial_number_in_table)
        form.addRow(self.ck_show_item_warranty_in_table)
        form.addRow("Tabulas šūnu polsterējums (mm):", self.in_table_cell_padding_mm)
        form.addRow("Tabulas galvenes fonta stils:", self.cb_table_header_font_style)
        form.addRow("Tabulas satura izlīdzināšana:", self.cb_table_content_alignment)
        form.addRow("Paraksta fonta izmērs:", self.in_signature_font_size)
        form.addRow("Paraksta atstarpe (mm):", self.in_signature_spacing_mm)
        form.addRow("Dokumenta virsraksta fonta izmērs:", self.in_document_title_font_size)
        form.addRow("Dokumenta virsraksta krāsa (Hex):", self.in_document_title_color)
        form.addRow("Sadaļas virsraksta fonta izmērs:", self.in_section_heading_font_size)
        form.addRow("Sadaļas virsraksta krāsa (Hex):", self.in_section_heading_color)
        form.addRow("Paragrāfa rindu atstarpes reizinātājs:", self.in_paragraph_line_spacing_multiplier)
        form.addRow("Tabulas apmales stils:", self.cb_table_border_style)
        form.addRow("Tabulas apmales biezums (pt):", self.in_table_border_thickness_pt)
        form.addRow("Tabulas alternatīvās rindas krāsa (Hex):", self.in_table_alternate_row_color)
        form.addRow(self.ck_show_total_sum_in_words)
        form.addRow("Kopsummas vārdos valoda:", self.cb_total_sum_in_words_language)
        form.addRow("Noklusējuma PVN aprēķina metode:", self.cb_default_vat_calculation_method)
        form.addRow(self.ck_show_vat_breakdown)
        form.addRow(self.ck_enable_digital_signature_field)
        form.addRow("Digitālā paraksta lauka nosaukums:", self.in_digital_signature_field_name)
        form.addRow("Digitālā paraksta lauka izmērs (mm):", self.in_digital_signature_field_size_mm)
        form.addRow("Digitālā paraksta lauka pozīcija:", self.cb_digital_signature_field_position)
        form.addRow("Šablonu direktorijs:", templates_dir_w)  # JAUNA RINDAS

        w.setLayout(form)
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(w)
        self.tabs.addTab(scroll_area, "Papildu iestatījumi")

        # Connect new settings to update_preview
        self.in_header_text_color.textChanged.connect(self._update_preview)
        self.in_footer_text_color.textChanged.connect(self._update_preview)
        self.in_table_header_bg_color.textChanged.connect(self._update_preview)
        self.in_table_grid_color.textChanged.connect(self._update_preview)
        self.in_table_row_spacing.valueChanged.connect(self._update_preview)
        self.in_line_spacing_multiplier.valueChanged.connect(self._update_preview)
        self.ck_show_page_numbers.stateChanged.connect(self._update_preview)
        self.ck_show_generation_timestamp.stateChanged.connect(self._update_preview)
        self.cb_currency_symbol_position.currentIndexChanged.connect(self._update_preview)
        self.in_date_format.textChanged.connect(self._update_preview)
        self.in_signature_line_length_mm.valueChanged.connect(self._update_preview)
        self.in_signature_line_thickness_pt.valueChanged.connect(self._update_preview)
        self.ck_add_cover_page.stateChanged.connect(self._update_preview)
        self.in_cover_page_title.textChanged.connect(self._update_preview)
        self.in_cover_page_logo_width_mm.valueChanged.connect(self._update_preview)
        # Individuālais QR kods
        self.ck_include_custom_qr_code.stateChanged.connect(self._update_preview)
        self.in_custom_qr_code_data.textChanged.connect(self._update_preview)
        self.in_custom_qr_code_size_mm.valueChanged.connect(self._update_preview)
        self.cb_custom_qr_code_position.currentIndexChanged.connect(self._update_preview)
        self.in_custom_qr_code_pos_x_mm.valueChanged.connect(self._update_preview)
        self.in_custom_qr_code_pos_y_mm.valueChanged.connect(self._update_preview)
        self.in_custom_qr_code_color.textChanged.connect(self._update_preview)

        # Automātiskais QR kods (akta ID)
        self.ck_include_auto_qr_code.stateChanged.connect(self._update_preview)
        self.in_auto_qr_code_size_mm.valueChanged.connect(self._update_preview)
        self.cb_auto_qr_code_position.currentIndexChanged.connect(self._update_preview)
        self.in_auto_qr_code_pos_x_mm.valueChanged.connect(self._update_preview)
        self.in_auto_qr_code_pos_y_mm.valueChanged.connect(self._update_preview)
        self.in_auto_qr_code_color.textChanged.connect(self._update_preview)

        self.ck_add_watermark.stateChanged.connect(self._update_preview)
        self.in_watermark_text.textChanged.connect(self._update_preview)
        self.in_watermark_font_size.valueChanged.connect(self._update_preview)
        self.in_watermark_color.textChanged.connect(self._update_preview)
        self.in_watermark_rotation.valueChanged.connect(self._update_preview)
        self.ck_enable_pdf_encryption.stateChanged.connect(self._update_preview)
        self.in_pdf_user_password.textChanged.connect(self._update_preview)
        self.in_pdf_owner_password.textChanged.connect(self._update_preview)
        self.ck_allow_printing.stateChanged.connect(self._update_preview)
        self.ck_allow_copying.stateChanged.connect(self._update_preview)
        self.ck_allow_modifying.stateChanged.connect(self._update_preview)
        self.ck_allow_annotating.stateChanged.connect(self._update_preview)
        self.in_default_country.textChanged.connect(self._update_preview)
        self.in_default_city.textChanged.connect(self._update_preview)
        self.ck_show_contact_details_in_header.stateChanged.connect(self._update_preview)
        self.in_contact_details_header_font_size.valueChanged.connect(self._update_preview)
        self.in_item_image_width_mm.valueChanged.connect(self._update_preview)
        self.in_item_image_caption_font_size.valueChanged.connect(self._update_preview)
        self.ck_show_item_notes_in_table.stateChanged.connect(self._update_preview)
        self.ck_show_item_serial_number_in_table.stateChanged.connect(self._update_preview)
        self.ck_show_item_warranty_in_table.stateChanged.connect(self._update_preview)
        self.in_table_cell_padding_mm.valueChanged.connect(self._update_preview)
        self.cb_table_header_font_style.currentIndexChanged.connect(self._update_preview)
        self.cb_table_content_alignment.currentIndexChanged.connect(self._update_preview)
        self.in_signature_font_size.valueChanged.connect(self._update_preview)
        self.in_signature_spacing_mm.valueChanged.connect(self._update_preview)
        self.in_document_title_font_size.valueChanged.connect(self._update_preview)
        self.in_document_title_color.textChanged.connect(self._update_preview)
        self.in_section_heading_font_size.valueChanged.connect(self._update_preview)
        self.in_section_heading_color.textChanged.connect(self._update_preview)
        self.in_paragraph_line_spacing_multiplier.valueChanged.connect(self._update_preview)
        self.cb_table_border_style.currentIndexChanged.connect(self._update_preview)
        self.in_table_border_thickness_pt.valueChanged.connect(self._update_preview)
        self.in_table_alternate_row_color.textChanged.connect(self._update_preview)
        self.ck_show_total_sum_in_words.stateChanged.connect(self._update_preview)
        self.cb_total_sum_in_words_language.currentIndexChanged.connect(self._update_preview)
        self.cb_default_vat_calculation_method.currentIndexChanged.connect(self._update_preview)
        self.ck_show_vat_breakdown.stateChanged.connect(self._update_preview)
        self.ck_enable_digital_signature_field.stateChanged.connect(self._update_preview)
        self.in_digital_signature_field_name.textChanged.connect(self._update_preview)
        self.in_digital_signature_field_size_mm.valueChanged.connect(self._update_preview)
        self.cb_digital_signature_field_position.currentIndexChanged.connect(self._update_preview)


        self.cb_page_size.currentIndexChanged.connect(self._update_preview)
        self.cb_page_orientation.currentIndexChanged.connect(self._update_preview)
        self.in_margin_left.valueChanged.connect(self._update_preview)
        self.in_margin_right.valueChanged.connect(self._update_preview)
        self.in_margin_top.valueChanged.connect(self._update_preview)
        self.in_margin_bottom.valueChanged.connect(self._update_preview)
        self.in_font_size_head.valueChanged.connect(self._update_preview)
        self.in_font_size_normal.valueChanged.connect(self._update_preview)
        self.in_font_size_small.valueChanged.connect(self._update_preview)
        self.in_font_size_table.valueChanged.connect(self._update_preview)
        self.in_logo_width_mm.valueChanged.connect(self._update_preview)
        self.in_signature_width_mm.valueChanged.connect(self._update_preview)
        self.in_signature_height_mm.valueChanged.connect(self._update_preview)
        self.in_docx_image_width_inches.valueChanged.connect(self._update_preview)
        self.in_docx_signature_width_inches.valueChanged.connect(self._update_preview)
        self.in_table_col_widths.textChanged.connect(self._update_preview)
        self.ck_auto_generate_akta_nr.stateChanged.connect(self._update_preview)
        self.in_default_currency.textChanged.connect(self._update_preview)
        self.in_default_unit.textChanged.connect(self._update_preview)
        self.in_default_pvn_rate.valueChanged.connect(self._update_preview)
        self.in_poppler_path.textChanged.connect(self._update_preview)


    def izvēlēties_poppler_ceļu(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Izvēlēties Poppler bin direktoriju")
        if folder_path:
            self.in_poppler_path.setText(folder_path)

    # ----- Tab: Šabloni -----
    def _būvēt_sablonu_tab(self):
        w = QWidget()
        v = QVBoxLayout()

        self.sablonu_list = QListWidget()
        self.sablonu_list.setSelectionMode(QAbstractItemView.ExtendedSelection)  # Ļauj atlasīt vairākus elementus
        self.sablonu_list.itemDoubleClicked.connect(self.ieladet_sablonu)

        btn_ieladet_sablonu = QPushButton("Ielādēt izvēlēto šablonu")
        btn_ieladet_sablonu.clicked.connect(lambda: self.ieladet_sablonu(self.sablonu_list.currentItem()))

        btn_dzest_sablonus = QPushButton("Dzēst atlasītos šablonus")
        btn_dzest_sablonus.clicked.connect(self.dzest_sablonus)

        # JAUNAS POGAS PAROLES PĀRVALDĪBAI
        btn_mainit_paroli = QPushButton("Mainīt/Pievienot paroli")
        btn_mainit_paroli.clicked.connect(self.mainit_sablonu_paroli)
        btn_nonemt_paroli = QPushButton("Noņemt paroli")
        btn_nonemt_paroli.clicked.connect(self.nonemt_sablonu_paroli)

        btns_layout = QHBoxLayout()
        btns_layout.addWidget(btn_ieladet_sablonu)
        btns_layout.addWidget(btn_dzest_sablonus)
        btns_layout.addWidget(btn_mainit_paroli)  # Pievienojam jauno pogu
        btns_layout.addWidget(btn_nonemt_paroli)  # Pievienojam jauno pogu
        btns_layout.addStretch()

        v.addWidget(QLabel("Pieejamie šabloni:"))
        v.addWidget(self.sablonu_list)
        v.addLayout(btns_layout)
        v.addStretch()

        w.setLayout(v)
        self.tabs.addTab(w, "Šabloni")

    def ieladet_sablonu(self, item: QListWidgetItem):
        if not item:
            return

        sablonu_nosaukums = item.text()
        file_path = None

        # Pārbaudām, vai tas ir iebūvētais šablons
        if sablonu_nosaukums == "Testa dati (piemērs)":
            # ... (esošais Testa dati šablona kods) ...
            # Šeit nav jāmaina, jo iebūvētajam šablonam nav paroles
            d = AktaDati(
                akta_nr="",
                datums=datetime.now().strftime('%Y-%m-%d'),
                vieta="Rīga",
                pasūtījuma_nr="",
                pieņēmējs=Persona(
                    nosaukums="SIA \"Tests\"",
                    reģ_nr="00000000000",
                    adrese="Testa iela 1, Rīga, LV-0000",
                    kontaktpersona="Jānis Tests",
                    tālrunis="+371 200000000",
                    epasts="janis.tests@tests.lv",
                    bankas_konts="LV00BANK1234567890123",
                    juridiskais_statuss="Juridiska persona"
                ),
                nodevējs=Persona(
                    nosaukums="SIA \"Demo Serviss\"",
                    reģ_nr="00000000000",
                    adrese="Testa iela 1, Rīga, LV-0000",
                    kontaktpersona="Jānis Tests",
                    tālrunis="+371 200000000",
                    epasts="anna.tests@demoserviss.lv",
                    bankas_konts="LV00BANK1234567890123",
                    juridiskais_statuss="Juridiska persona"
                ),
                pozīcijas=[
                    Pozīcija("Datora remonts", Decimal("1"), "gab.", Decimal("50.00"), "SN12345", "1 gads",
                             "Veikta diagnostika un komponentu nomaiņa."),
                    Pozīcija("Programmatūras instalācija", Decimal("1"), "st.", Decimal("25.00"), "", "",
                             "Uzstādīta operētājsistēma un biroja programmatūra."),
                    Pozīcija("Detaļas (RAM 8GB)", Decimal("1"), "gab.", Decimal("35.00"), "RAM9876", "2 gadi",
                             "Augstas veiktspējas RAM modulis.")
                ],
                piezīmes="Veikts pilns datora diagnostikas un remonta pakalpojums. Iekļauta programmatūras optimizācija.",
                iekļaut_pvn=True,
                pvn_likme=Decimal("21.0"),
                parakstu_rindas=True,
                logotipa_ceļš="",
                fonts_ceļš="",
                paraksts_pieņēmējs_ceļš="",
                paraksts_nodevējs_ceļš="",
                līguma_nr="",
                izpildes_termiņš="",
                pieņemšanas_datums="",
                nodošanas_datums="",
                strīdu_risināšana="Visi strīdi, kas izriet no šī akta, tiks risināti sarunu ceļā. Ja vienošanās netiek panākta, strīdi tiks nodoti izskatīšanai Latvijas Republikas tiesā saskaņā ar spēkā esošajiem normatīvajiem aktiem.",
                konfidencialitātes_klauzula=True,
                soda_nauda_procenti=Decimal("0.5"),
                piegādes_nosacījumi="DAP (Delivered at Place) Rīga, Latvija",
                apdrošināšana=True,
                papildu_nosacījumi="Abas puses apliecina, ka ir iepazinušās ar šī akta saturu un piekrīt visiem tā nosacījumiem. Akts sastādīts divos eksemplāros, katrai pusei pa vienu.",
                atsauces_dokumenti="Pavadzīme Nr. PV-2023/010, Garantijas talons Nr. GT-2023/005",
                akta_statuss="Melnraksts",
                valūta="EUR",
                elektroniskais_paraksts=False,
                radit_elektronisko_parakstu_tekstu=False,  # JAUNA RINDAS
                pdf_page_size="A4",
                pdf_page_orientation="Portrets",
                pdf_margin_left=Decimal("18"),
                pdf_margin_right=Decimal("18"),
                pdf_margin_top=Decimal("16"),
                pdf_margin_bottom=Decimal("16"),
                pdf_font_size_head=14,
                pdf_font_size_normal=10,
                pdf_font_size_small=9,
                pdf_font_size_table=9,
                pdf_logo_width_mm=Decimal("35"),
                pdf_signature_width_mm=Decimal("50"),
                pdf_signature_height_mm=Decimal("20"),
                docx_image_width_inches=Decimal("4"),
                docx_signature_width_inches=Decimal("1.5"),
                table_col_widths="10,40,18,18,20,20,25,25,25",
                auto_generate_akta_nr=False,
                default_currency="EUR",
                default_unit="gab.",
                default_pvn_rate=Decimal("21.0"),
                poppler_path="",
                header_text_color="#000000",
                footer_text_color="#000000",
                table_header_bg_color="#E0E0E0",
                table_grid_color="#CCCCCC",
                table_row_spacing=Decimal("4"),
                line_spacing_multiplier=Decimal("1.2"),
                show_page_numbers=True,
                show_generation_timestamp=True,
                currency_symbol_position="after",
                date_format="YYYY-MM-DD",
                signature_line_length_mm=Decimal("60"),
                signature_line_thickness_pt=Decimal("0.5"),
                add_cover_page=False,
                cover_page_title="Pieņemšanas-Nodošanas Akts",
                cover_page_logo_width_mm=Decimal("80"),
                # Individuālais QR kods
                include_custom_qr_code=False,
                custom_qr_code_data="",
                custom_qr_code_size_mm=Decimal("20"),
                custom_qr_code_position="bottom_right",

                # Automātiskais QR kods (akta ID)
                include_auto_qr_code=False,
                auto_qr_code_size_mm=Decimal("20"),
                auto_qr_code_position="bottom_left",

                add_watermark=False,
                watermark_text="MELNRAKSTS",
                watermark_font_size=72,
                watermark_color="#E0E0E0",
                watermark_rotation=45,
                enable_pdf_encryption=False,
                pdf_user_password="",
                pdf_owner_password="",
                allow_printing=True,
                allow_copying=True,
                allow_modifying=False,
                allow_annotating=True,
                default_country="Latvija",
                default_city="Rīga",
                show_contact_details_in_header=False,
                contact_details_header_font_size=8,
                item_image_width_mm=Decimal("50"),
                item_image_caption_font_size=8,
                show_item_notes_in_table=True,
                show_item_serial_number_in_table=True,
                show_item_warranty_in_table=True,
                table_cell_padding_mm=Decimal("2"),
                table_header_font_style="bold",
                table_content_alignment="left",
                signature_font_size=9,
                signature_spacing_mm=Decimal("10"),
                document_title_font_size=18,
                document_title_color="#000000",
                section_heading_font_size=12,
                section_heading_color="#000000",
                paragraph_line_spacing_multiplier=Decimal("1.2"),
                table_border_style="solid",
                table_border_thickness_pt=Decimal("0.5"),
                table_alternate_row_color="",
                show_total_sum_in_words=False,
                total_sum_in_words_language="lv",
                default_vat_calculation_method="exclusive",
                show_vat_breakdown=True,
                enable_digital_signature_field=False,
                digital_signature_field_name="Paraksts",
                digital_signature_field_size_mm=Decimal("40"),
                digital_signature_field_position="bottom_center",
                template_password=""  # Nodrošinām, ka šim nav paroles
            )
            self.ieviest_datus(d)
            QMessageBox.information(self, "Šablons ielādēts", f"Šablons '{sablonu_nosaukums}' veiksmīgi ielādēts.")
            return
        else:
            # Mēģinām ielādēt no šablonu direktorijas (atbalsta arī vecos ceļus)
            file_path = self._resolve_template_path(sablonu_nosaukums)


        file_path = _coerce_path(file_path)
        if file_path and os.path.exists(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                # Pārbaudām paroli, ja tā ir iestatīta
                stored_password = data.get('template_password', '')
                if stored_password:
                    entered_password, ok_pass = QInputDialog.getText(self, "Ievadiet paroli",
                                                                     f"Šablonam '{sablonu_nosaukums}' ir parole. Lūdzu, ievadiet to:",
                                                                     QLineEdit.Password)
                    if not ok_pass or entered_password != stored_password:
                        QMessageBox.warning(self, "Nepareiza parole",
                                            "Ievadītā parole ir nepareiza vai ievade atcelta.")
                        return  # Atceļam ielādi, ja parole ir nepareiza vai atcelta

                # Helper to safely get and convert Decimal values
                def get_decimal(dict_obj, key, default_val):
                    val = dict_obj.get(key, default_val)
                    return to_decimal(val)

                # Helper to safely get boolean values
                def get_bool(dict_obj, key, default_val):
                    val = dict_obj.get(key, default_val)
                    return bool(val)

                # Izveidojam AktaDati objektu no ielādētajiem datiem
                d = AktaDati(
                    akta_nr=data.get('akta_nr', ''), datums=data.get('datums', datetime.now().strftime('%Y-%m-%d')),
                    vieta=data.get('vieta', ''),
                    pasūtījuma_nr=data.get('pasūtījuma_nr', ''),
                    pieņēmējs=Persona(
                        nosaukums=data.get('pieņēmējs', {}).get('nosaukums', ''),
                        reģ_nr=data.get('pieņēmējs', {}).get('reģ_nr', ''),
                        adrese=data.get('pieņēmējs', {}).get('adrese', ''),
                        kontaktpersona=data.get('pieņēmējs', {}).get('kontaktpersona', ''),
                        amats=data.get('pieņēmējs', {}).get('amats', ''),
                        pilnvaras_pamats=data.get('pieņēmējs', {}).get('pilnvaras_pamats', ''),
                        tālrunis=data.get('pieņēmējs', {}).get('tālrunis', ''),
                        epasts=data.get('pieņēmējs', {}).get('epasts', ''),
                        web_lapa=data.get('pieņēmējs', {}).get('web_lapa', ''),
                        bankas_konts=data.get('pieņēmējs', {}).get('bankas_konts', ''),
                        juridiskais_statuss=data.get('pieņēmējs', {}).get('juridiskais_statuss', '')
                    ),
                    nodevējs=Persona(
                        nosaukums=data.get('nodevējs', {}).get('nosaukums', ''),
                        reģ_nr=data.get('nodevējs', {}).get('reģ_nr', ''),
                        adrese=data.get('nodevējs', {}).get('adrese', ''),
                        kontaktpersona=data.get('nodevējs', {}).get('kontaktpersona', ''),
                        amats=data.get('nodevējs', {}).get('amats', ''),
                        pilnvaras_pamats=data.get('nodevējs', {}).get('pilnvaras_pamats', ''),
                        tālrunis=data.get('nodevējs', {}).get('tālrunis', ''),
                        epasts=data.get('nodevējs', {}).get('epasts', ''),
                        web_lapa=data.get('nodevējs', {}).get('web_lapa', ''),
                        bankas_konts=data.get('nodevējs', {}).get('bankas_konts', ''),
                        juridiskais_statuss=data.get('nodevējs', {}).get('juridiskais_statuss', '')
                    ),
                    pozīcijas=[Pozīcija(
                        apraksts=p.get('apraksts', ''),
                        daudzums=get_decimal(p, 'daudzums', '0'),
                        vienība=p.get('vienība', 'gab.'),
                        cena=get_decimal(p, 'cena', '0'),
                        seriālais_nr=p.get('seriālais_nr', ''),
                        garantija=p.get('garantija', ''),
                        piezīmes_pozīcijai=p.get('piezīmes_pozīcijai', ''),
                        attēla_ceļš=p.get('attēla_ceļš', '')
                    ) for p in data.get('pozīcijas', [])],
                    custom_columns=data.get('custom_columns', []) if isinstance(data.get('custom_columns', []), list) else [],
                    poz_columns_config=data.get('poz_columns_config', {}) if isinstance(data.get('poz_columns_config', {}), dict) else {},
                    show_price_summary=get_bool(data, 'show_price_summary', True),
                    poz_columns_visual_order=data.get('poz_columns_visual_order', []) if isinstance(data.get('poz_columns_visual_order', []), list) else [],
                    attēli=[Attēls(**a) for a in data.get('attēli', [])],
                    piezīmes=data.get('piezīmes', ''), iekļaut_pvn=get_bool(data, 'iekļaut_pvn', False),
                    pvn_likme=get_decimal(data, 'pvn_likme', '21'),
                    parakstu_rindas=get_bool(data, 'parakstu_rindas', True),
                    logotipa_ceļš=data.get('logotipa_ceļš', ''), fonts_ceļš=data.get('fonts_ceļš', ''),
                    paraksts_pieņēmējs_ceļš=data.get('paraksts_pieņēmējs_ceļš', ''),
                    paraksts_nodevējs_ceļš=data.get('paraksts_nodevējs_ceļš', ''),
                    līguma_nr=data.get('līguma_nr', ''),
                    izpildes_termiņš=data.get('izpildes_termiņš', ''),
                    pieņemšanas_datums=data.get('pieņemšanas_datums', ''),
                    nodošanas_datums=data.get('nodošanas_datums', ''),
                    strīdu_risināšana=data.get('strīdu_risināšana', ''),
                    konfidencialitātes_klauzula=get_bool(data, 'konfidencialitātes_klauzula', False),
                    soda_nauda_procenti=get_decimal(data, 'soda_nauda_procenti', '0.0'),
                    piegādes_nosacījumi=data.get('piegādes_nosacījumi', ''),
                    apdrošināšana=get_bool(data, 'apdrošināšana', False),
                        apdrošināšana_teksts=data.get('apdrošināšana_teksts', ''),
                    papildu_nosacījumi=data.get('papildu_nosacījumi', ''),
                    atsauces_dokumenti=data.get('atsauces_dokumenti', ''),
                    akta_statuss=data.get('akta_statuss', 'Melnraksts'),
                    valūta=data.get('valūta', 'EUR'),
                    elektroniskais_paraksts=get_bool(data, 'elektroniskais_paraksts', False),
                    radit_elektronisko_parakstu_tekstu=get_bool(data, 'radit_elektronisko_parakstu_tekstu', False),
                    # JAUNA RINDAS
                    pdf_page_size=data.get('pdf_page_size', 'A4'),
                    pdf_page_orientation=data.get('pdf_page_orientation', 'Portrets'),
                    pdf_margin_left=get_decimal(data, 'pdf_margin_left', '18'),
                    pdf_margin_right=get_decimal(data, 'pdf_margin_right', '18'),
                    pdf_margin_top=get_decimal(data, 'pdf_margin_top', '16'),
                    pdf_margin_bottom=get_decimal(data, 'pdf_margin_bottom', '16'),
                    pdf_font_size_head=data.get('pdf_font_size_head', 14),
                    pdf_font_size_normal=data.get('pdf_font_size_normal', 10),
                    pdf_font_size_small=data.get('pdf_font_size_small', 9),
                    pdf_font_size_table=data.get('pdf_font_size_table', 9),
                    pdf_logo_width_mm=get_decimal(data, 'pdf_logo_width_mm', '35'),
                    pdf_signature_width_mm=get_decimal(data, 'pdf_signature_width_mm', '50'),
                    pdf_signature_height_mm=get_decimal(data, 'pdf_signature_height_mm', '20'),
                    docx_image_width_inches=get_decimal(data, 'docx_image_width_inches', '4'),
                    docx_signature_width_inches=get_decimal(data, 'docx_signature_width_inches', '1.5'),
                    table_col_widths=data.get('table_col_widths', '10,40,18,18,20,20,25,25,25'),
                    auto_generate_akta_nr=get_bool(data, 'auto_generate_akta_nr', False),
                    default_currency=data.get('default_currency', 'EUR'),
                    default_unit=data.get('default_unit', 'gab.'),
                    default_pvn_rate=get_decimal(data, 'default_pvn_rate', '21.0'),
                    poppler_path=data.get('poppler_path', ''),
                    header_text_color=data.get('header_text_color', '#000000'),
                    footer_text_color=data.get('footer_text_color', '#000000'),
                    table_header_bg_color=data.get('table_header_bg_color', '#E0E0E0'),
                    table_grid_color=data.get('table_grid_color', '#CCCCCC'),
                    table_row_spacing=get_decimal(data, 'table_row_spacing', '4'),
                    line_spacing_multiplier=get_decimal(data, 'line_spacing_multiplier', '1.2'),
                    show_page_numbers=get_bool(data, 'show_page_numbers', True),
                    show_generation_timestamp=get_bool(data, 'show_generation_timestamp', True),
                    currency_symbol_position=data.get('currency_symbol_position', 'after'),
                    date_format=data.get('date_format', 'YYYY-MM-DD'),
                    signature_line_length_mm=get_decimal(data, 'signature_line_length_mm', '60'),
                    signature_line_thickness_pt=get_decimal(data, 'signature_line_thickness_pt', '0.5'),
                    add_cover_page=get_bool(data, 'add_cover_page', False),
                    cover_page_title=data.get('cover_page_title', 'Pieņemšanas-Nodošanas Akts'),
                    cover_page_logo_width_mm=get_decimal(data, 'cover_page_logo_width_mm', '80'),
                    # Individuālais QR kods
                    include_custom_qr_code=get_bool(data, 'include_custom_qr_code', False),
                    custom_qr_code_data=data.get('custom_qr_code_data', ''),
                    custom_qr_code_size_mm=get_decimal(data, 'custom_qr_code_size_mm', '20'),
                    custom_qr_code_position=data.get('custom_qr_code_position', 'bottom_right'),
                    custom_qr_code_pos_x_mm=get_decimal(data, 'custom_qr_code_pos_x_mm', '0'),
                    custom_qr_code_pos_y_mm=get_decimal(data, 'custom_qr_code_pos_y_mm', '0'),
                    custom_qr_code_color=data.get('custom_qr_code_color', '#000000'),

                    # Automātiskais QR kods (akta ID)
                    include_auto_qr_code=get_bool(data, 'include_auto_qr_code', False),
                    auto_qr_code_size_mm=get_decimal(data, 'auto_qr_code_size_mm', '20'),
                    auto_qr_code_position=data.get('auto_qr_code_position', 'bottom_left'),
                    auto_qr_code_pos_x_mm=get_decimal(data, 'auto_qr_code_pos_x_mm', '0'),
                    auto_qr_code_pos_y_mm=get_decimal(data, 'auto_qr_code_pos_y_mm', '0'),
                    auto_qr_code_color=data.get('auto_qr_code_color', '#000000'),

                    add_watermark=get_bool(data, 'add_watermark', False),
                    watermark_text=data.get('watermark_text', 'MELNRAKSTS'),
                    watermark_font_size=data.get('watermark_font_size', 72),
                    watermark_color=data.get('watermark_color', '#E0E0E0'),
                    watermark_rotation=data.get('watermark_rotation', 45),
                    enable_pdf_encryption=get_bool(data, 'enable_pdf_encryption', False),
                    pdf_user_password=data.get('pdf_user_password', ''),
                    pdf_owner_password=data.get('pdf_owner_password', ''),
                    allow_printing=get_bool(data, 'allow_printing', True),
                    allow_copying=get_bool(data, 'allow_copying', True),
                    allow_modifying=get_bool(data, 'allow_modifying', False),
                    allow_annotating=get_bool(data, 'allow_annotating', True),
                    default_country=data.get('default_country', 'Latvija'),
                    default_city=data.get('default_city', 'Rīga'),
                    show_contact_details_in_header=get_bool(data, 'show_contact_details_in_header', False),
                    contact_details_header_font_size=data.get('contact_details_header_font_size', 8),
                    item_image_width_mm=get_decimal(data, 'item_image_width_mm', '50'),
                    item_image_caption_font_size=data.get('item_image_caption_font_size', 8),
                    show_item_notes_in_table=get_bool(data, 'show_item_notes_in_table', True),
                    show_item_serial_number_in_table=get_bool(data, 'show_item_serial_number_in_table', True),
                    show_item_warranty_in_table=get_bool(data, 'show_item_warranty_in_table', True),
                    table_cell_padding_mm=get_decimal(data, 'table_cell_padding_mm', '2'),
                    table_header_font_style=data.get('table_header_font_style', 'bold'),
                    table_content_alignment=data.get('table_content_alignment', 'left'),
                    signature_font_size=data.get('signature_font_size', 9),
                    signature_spacing_mm=get_decimal(data, 'signature_spacing_mm', '10'),
                    document_title_font_size=data.get('document_title_font_size', 18),
                    document_title_color=data.get('document_title_color', '#000000'),
                    section_heading_font_size=data.get('section_heading_font_size', 12),
                    section_heading_color=data.get('section_heading_color', '#000000'),
                    paragraph_line_spacing_multiplier=get_decimal(data, 'paragraph_line_spacing_multiplier', '1.2'),
                    table_border_style=data.get('table_border_style', 'solid'),
                    table_border_thickness_pt=get_decimal(data, 'table_border_thickness_pt', '0.5'),
                    table_alternate_row_color=data.get('table_alternate_row_color', ''),
                    show_total_sum_in_words=get_bool(data, 'show_total_sum_in_words', False),
                    total_sum_in_words_language=data.get('total_sum_in_words_language', 'lv'),
                    default_vat_calculation_method=data.get('default_vat_calculation_method', 'exclusive'),
                    show_vat_breakdown=get_bool(data, 'show_vat_breakdown', True),
                    enable_digital_signature_field=get_bool(data, 'enable_digital_signature_field', False),
                    digital_signature_field_name=data.get('digital_signature_field_name', 'Paraksts'),
                    digital_signature_field_size_mm=get_decimal(data, 'digital_signature_field_size_mm', '40'),
                    digital_signature_field_position=data.get('digital_signature_field_position', 'bottom_center'),
                    template_password=data.get('template_password', '')  # Ielādējam paroli
                )
                self.ieviest_datus(d)
                QMessageBox.information(self, "Šablons ielādēts", f"Šablons '{sablonu_nosaukums}' veiksmīgi ielādēts.")
            except Exception as e:
                QMessageBox.critical(self, "Kļūda", f"Neizdevās ielādēt šablonu '{sablonu_nosaukums}':\n{e}")
        else:
            QMessageBox.warning(self, "Kļūda", f"Šablona fails '{sablonu_nosaukums}.json' nav atrasts.")

        if sablonu_nosaukums == "Pappus dati (piemērs)":
            # Šis ir iebūvētais piemērs, to var ielādēt tieši
            d = AktaDati(
                akta_nr="",  # Šabloniem akta nr. un datums parasti ir tukši/pašreizējie
                datums=datetime.now().strftime('%Y-%m-%d'),
                vieta="Rīga",
                pasūtījuma_nr="",
                pieņēmējs=Persona(
                    nosaukums="SIA \"Pappus\"",
                    reģ_nr="40003123456",
                    adrese="Lielā iela 1, Rīga, LV-1010",
                    kontaktpersona="Jānis Bērziņš",
                    tālrunis="+371 21234567",
                    epasts="janis.berzins@pappus.lv",
                    bankas_konts="LV12BANK1234567890123",
                    juridiskais_statuss="Juridiska persona"
                ),
                nodevējs=Persona(
                    nosaukums="SIA \"Demo Serviss\"",
                    reģ_nr="40003987654",
                    adrese="Mazā iela 5, Rīga, LV-1005",
                    kontaktpersona="Anna Liepa",
                    tālrunis="+371 27654321",
                    epasts="anna.liepa@demoserviss.lv",
                    bankas_konts="LV32BANK9876543210987",
                    juridiskais_statuss="Juridiska persona"
                ),
                pozīcijas=[
                    Pozīcija("Datora remonts", Decimal("1"), "gab.", Decimal("50.00"), "SN12345", "1 gads",
                             "Veikta diagnostika un komponentu nomaiņa."),
                    Pozīcija("Programmatūras instalācija", Decimal("1"), "st.", Decimal("25.00"), "", "",
                             "Uzstādīta operētājsistēma un biroja programmatūra."),
                    Pozīcija("Detaļas (RAM 8GB)", Decimal("1"), "gab.", Decimal("35.00"), "RAM9876", "2 gadi",
                             "Augstas veiktspējas RAM modulis.")
                ],
                piezīmes="Veikts pilns datora diagnostikas un remonta pakalpojums. Iekļauta programmatūras optimizācija.",
                iekļaut_pvn=True,
                pvn_likme=Decimal("21.0"),
                parakstu_rindas=True,
                logotipa_ceļš="",
                fonts_ceļš="",
                paraksts_pieņēmējs_ceļš="",
                paraksts_nodevējs_ceļš="",
                līguma_nr="",  # Šabloniem šie lauki parasti ir tukši
                izpildes_termiņš="",
                pieņemšanas_datums="",
                nodošanas_datums="",
                strīdu_risināšana="Visi strīdi, kas izriet no šī akta, tiks risināti sarunu ceļā. Ja vienošanās netiek panākta, strīdi tiks nodoti izskatīšanai Latvijas Republikas tiesā saskaņā ar spēkā esošajiem normatīvajiem aktiem.",
                konfidencialitātes_klauzula=True,
                soda_nauda_procenti=Decimal("0.5"),
                piegādes_nosacījumi="DAP (Delivered at Place) Rīga, Latvija",
                apdrošināšana=True,
                papildu_nosacījumi="Abas puses apliecina, ka ir iepazinušās ar šī akta saturu un piekrīt visiem tā nosacījumiem. Akts sastādīts divos eksemplāros, katrai pusei pa vienu.",
                atsauces_dokumenti="Pavadzīme Nr. PV-2023/010, Garantijas talons Nr. GT-2023/005",
                akta_statuss="Melnraksts",  # Šabloniem statuss ir melnraksts
                valūta="EUR",
                elektroniskais_paraksts=False,
                pdf_page_size="A4",
                pdf_page_orientation="Portrets",
                pdf_margin_left=Decimal("18"),
                pdf_margin_right=Decimal("18"),
                pdf_margin_top=Decimal("16"),
                pdf_margin_bottom=Decimal("16"),
                pdf_font_size_head=14,
                pdf_font_size_normal=10,
                pdf_font_size_small=9,
                pdf_font_size_table=9,
                pdf_logo_width_mm=Decimal("35"),
                pdf_signature_width_mm=Decimal("50"),
                pdf_signature_height_mm=Decimal("20"),
                docx_image_width_inches=Decimal("4"),
                docx_signature_width_inches=Decimal("1.5"),
                table_col_widths="10,40,18,18,20,20,25,25,25",
                auto_generate_akta_nr=False,
                default_currency="EUR",
                default_unit="gab.",
                default_pvn_rate=Decimal("21.0"),
                poppler_path="",
                # Default values for new settings
                header_text_color="#000000",
                footer_text_color="#000000",
                table_header_bg_color="#E0E0E0",
                table_grid_color="#CCCCCC",
                table_row_spacing=Decimal("4"),
                line_spacing_multiplier=Decimal("1.2"),
                show_page_numbers=True,
                show_generation_timestamp=True,
                currency_symbol_position="after",
                date_format="YYYY-MM-DD",
                signature_line_length_mm=Decimal("60"),
                signature_line_thickness_pt=Decimal("0.5"),
                add_cover_page=False,
                cover_page_title="Pieņemšanas-Nodošanas Akts",
                cover_page_logo_width_mm=Decimal("80"),
                # Individuālais QR kods
                include_custom_qr_code=False,
                custom_qr_code_data="",
                custom_qr_code_size_mm=Decimal("20"),
                custom_qr_code_position="bottom_right",

                # Automātiskais QR kods (akta ID)
                include_auto_qr_code=False,
                auto_qr_code_size_mm=Decimal("20"),
                auto_qr_code_position="bottom_left",

                add_watermark=False,
                watermark_text="MELNRAKSTS",
                watermark_font_size=72,
                watermark_color="#E0E0E0",
                watermark_rotation=45,
                enable_pdf_encryption=False,
                pdf_user_password="",
                pdf_owner_password="",
                allow_printing=True,
                allow_copying=True,
                allow_modifying=False,
                allow_annotating=True,
                default_country="Latvija",
                default_city="Rīga",
                show_contact_details_in_header=False,
                contact_details_header_font_size=8,
                item_image_width_mm=Decimal("50"),
                item_image_caption_font_size=8,
                show_item_notes_in_table=True,
                show_item_serial_number_in_table=True,
                show_item_warranty_in_table=True,
                table_cell_padding_mm=Decimal("2"),
                table_header_font_style="bold",
                table_content_alignment="left",
                signature_font_size=9,
                signature_spacing_mm=Decimal("10"),
                document_title_font_size=18,
                document_title_color="#000000",
                section_heading_font_size=12,
                section_heading_color="#000000",
                paragraph_line_spacing_multiplier=Decimal("1.2"),
                table_border_style="solid",
                table_border_thickness_pt=Decimal("0.5"),
                table_alternate_row_color="",
                show_total_sum_in_words=False,
                total_sum_in_words_language="lv",
                default_vat_calculation_method="exclusive",
                show_vat_breakdown=True,
                enable_digital_signature_field=False,
                digital_signature_field_name="Paraksts",
                digital_signature_field_size_mm=Decimal("40"),
                digital_signature_field_position="bottom_center"
            )
            self.ieviest_datus(d)
            QMessageBox.information(self, "Šablons ielādēts", f"Šablons '{sablonu_nosaukums}' veiksmīgi ielādēts.")
            return  # Iziet no funkcijas pēc iebūvētā šablona ielādes
        else:
            # Mēģinām ielādēt no šablonu direktorijas
            file_path = os.path.join(self.data.templates_dir, f"{sablonu_nosaukums}.json")

        file_path = _coerce_path(file_path)
        if file_path and os.path.exists(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    # Pārbaudām paroli, ja tā ir iestatīta
                    stored_password = data.get('template_password', '')
                    if stored_password:
                        entered_password, ok_pass = QInputDialog.getText(self, "Ievadiet paroli",
                                                                         f"Šablonam '{sablonu_nosaukums}' ir parole. Lūdzu, ievadiet to:",
                                                                         QLineEdit.Password)
                        if not ok_pass or entered_password != stored_password:
                            QMessageBox.warning(self, "Nepareiza parole",
                                                "Ievadītā parole ir nepareiza vai ievade atcelta.")
                            return  # Atceļam ielādi, ja parole ir nepareiza vai atcelta

                    # Helper to safely get and convert Decimal values
                    def get_decimal(dict_obj, key, default_val):
                        val = dict_obj.get(key, default_val)
                        return to_decimal(val)

                    # Helper to safely get boolean values
                    def get_bool(dict_obj, key, default_val):
                        val = dict_obj.get(key, default_val)
                        return bool(val)

                    # Izveidojam AktaDati objektu no ielādētajiem datiem
                    d = AktaDati(
                        akta_nr=data.get('akta_nr', ''), datums=data.get('datums', datetime.now().strftime('%Y-%m-%d')),
                        vieta=data.get('vieta', ''),
                        pasūtījuma_nr=data.get('pasūtījuma_nr', ''),
                        pieņēmējs=Persona(
                            nosaukums=data.get('pieņēmējs', {}).get('nosaukums', ''),
                            reģ_nr=data.get('pieņēmējs', {}).get('reģ_nr', ''),
                            adrese=data.get('pieņēmējs', {}).get('adrese', ''),
                            kontaktpersona=data.get('pieņēmējs', {}).get('kontaktpersona', ''),
                            tālrunis=data.get('pieņēmējs', {}).get('tālrunis', ''),
                            epasts=data.get('pieņēmējs', {}).get('epasts', ''),
                            bankas_konts=data.get('pieņēmējs', {}).get('bankas_konts', ''),
                            juridiskais_statuss=data.get('pieņēmējs', {}).get('juridiskais_statuss', '')
                        ),
                        nodevējs=Persona(
                            nosaukums=data.get('nodevējs', {}).get('nosaukums', ''),
                            reģ_nr=data.get('nodevējs', {}).get('reģ_nr', ''),
                            adrese=data.get('nodevējs', {}).get('adrese', ''),
                            kontaktpersona=data.get('nodevējs', {}).get('kontaktpersona', ''),
                            tālrunis=data.get('nodevējs', {}).get('tālrunis', ''),
                            epasts=data.get('nodevējs', {}).get('epasts', ''),
                            web_lapa=data.get('nodevējs', {}).get('web_lapa', ''),
                            bankas_konts=data.get('nodevējs', {}).get('bankas_konts', ''),
                            juridiskais_statuss=data.get('nodevējs', {}).get('juridiskais_statuss', '')
                        ),
                        pozīcijas=[Pozīcija(
                            apraksts=p.get('apraksts', ''),
                            daudzums=get_decimal(p, 'daudzums', '0'),
                            vienība=p.get('vienība', 'gab.'),
                            cena=get_decimal(p, 'cena', '0'),
                            seriālais_nr=p.get('seriālais_nr', ''),
                            garantija=p.get('garantija', ''),
                            piezīmes_pozīcijai=p.get('piezīmes_pozīcijai', ''),
                            attēla_ceļš=p.get('attēla_ceļš', '')
                        ) for p in data.get('pozīcijas', [])],
                        custom_columns=data.get('custom_columns', []) if isinstance(data.get('custom_columns', []), list) else [],
                        poz_columns_config=data.get('poz_columns_config', {}) if isinstance(data.get('poz_columns_config', {}), dict) else {},
                        show_price_summary=get_bool(data, 'show_price_summary', True),
                    poz_columns_visual_order=data.get('poz_columns_visual_order', []) if isinstance(data.get('poz_columns_visual_order', []), list) else [],
                        attēli=[Attēls(**a) for a in data.get('attēli', [])],
                        piezīmes=data.get('piezīmes', ''), iekļaut_pvn=get_bool(data, 'iekļaut_pvn', False),
                        pvn_likme=get_decimal(data, 'pvn_likme', '21'),
                        parakstu_rindas=get_bool(data, 'parakstu_rindas', True),
                        logotipa_ceļš=data.get('logotipa_ceļš', ''), fonts_ceļš=data.get('fonts_ceļš', ''),
                        paraksts_pieņēmējs_ceļš=data.get('paraksts_pieņēmējs_ceļš', ''),
                        paraksts_nodevējs_ceļš=data.get('paraksts_nodevējs_ceļš', ''),
                        līguma_nr=data.get('līguma_nr', ''),
                        izpildes_termiņš=data.get('izpildes_termiņš', ''),
                        pieņemšanas_datums=data.get('pieņemšanas_datums', ''),
                        nodošanas_datums=data.get('nodošanas_datums', ''),
                        strīdu_risināšana=data.get('strīdu_risināšana', ''),
                        konfidencialitātes_klauzula=get_bool(data, 'konfidencialitātes_klauzula', False),
                        soda_nauda_procenti=get_decimal(data, 'soda_nauda_procenti', '0.0'),
                        piegādes_nosacījumi=data.get('piegādes_nosacījumi', ''),
                        apdrošināšana=get_bool(data, 'apdrošināšana', False),
                        apdrošināšana_teksts=data.get('apdrošināšana_teksts', ''),
                        papildu_nosacījumi=data.get('papildu_nosacījumi', ''),
                        atsauces_dokumenti=data.get('atsauces_dokumenti', ''),
                        akta_statuss=data.get('akta_statuss', 'Melnraksts'),
                        valūta=data.get('valūta', 'EUR'),
                        elektroniskais_paraksts=get_bool(data, 'elektroniskais_paraksts', False),
                        pdf_page_size=data.get('pdf_page_size', 'A4'),
                        pdf_page_orientation=data.get('pdf_page_orientation', 'Portrets'),
                        pdf_margin_left=get_decimal(data, 'pdf_margin_left', '18'),
                        pdf_margin_right=get_decimal(data, 'pdf_margin_right', '18'),
                        pdf_margin_top=get_decimal(data, 'pdf_margin_top', '16'),
                        pdf_margin_bottom=get_decimal(data, 'pdf_margin_bottom', '16'),
                        pdf_font_size_head=data.get('pdf_font_size_head', 14),
                        pdf_font_size_normal=data.get('pdf_font_size_normal', 10),
                        pdf_font_size_small=data.get('pdf_font_size_small', 9),
                        pdf_font_size_table=data.get('pdf_font_size_table', 9),
                        pdf_logo_width_mm=get_decimal(data, 'pdf_logo_width_mm', '35'),
                        pdf_signature_width_mm=get_decimal(data, 'pdf_signature_width_mm', '50'),
                        pdf_signature_height_mm=get_decimal(data, 'pdf_signature_height_mm', '20'),
                        docx_image_width_inches=get_decimal(data, 'docx_image_width_inches', '4'),
                        docx_signature_width_inches=get_decimal(data, 'docx_signature_width_inches', '1.5'),
                        table_col_widths=data.get('table_col_widths', '10,40,18,18,20,20,25,25,25'),
                        auto_generate_akta_nr=get_bool(data, 'auto_generate_akta_nr', False),
                        default_currency=data.get('default_currency', 'EUR'),
                        default_unit=data.get('default_unit', 'gab.'),
                        default_pvn_rate=get_decimal(data, 'default_pvn_rate', '21.0'),
                        poppler_path=data.get('poppler_path', ''),
                        # Load new settings
                        header_text_color=data.get('header_text_color', '#000000'),
                        footer_text_color=data.get('footer_text_color', '#000000'),
                        table_header_bg_color=data.get('table_header_bg_color', '#E0E0E0'),
                        table_grid_color=data.get('table_grid_color', '#CCCCCC'),
                        table_row_spacing=get_decimal(data, 'table_row_spacing', '4'),
                        line_spacing_multiplier=get_decimal(data, 'line_spacing_multiplier', '1.2'),
                        show_page_numbers=get_bool(data, 'show_page_numbers', True),
                        show_generation_timestamp=get_bool(data, 'show_generation_timestamp', True),
                        currency_symbol_position=data.get('currency_symbol_position', 'after'),
                        date_format=data.get('date_format', 'YYYY-MM-DD'),
                        signature_line_length_mm=get_decimal(data, 'signature_line_length_mm', '60'),
                        signature_line_thickness_pt=get_decimal(data, 'signature_line_thickness_pt', '0.5'),
                        add_cover_page=get_bool(data, 'add_cover_page', False),
                        cover_page_title=data.get('cover_page_title', 'Pieņemšanas-Nodošanas Akts'),
                        cover_page_logo_width_mm=get_decimal(data, 'cover_page_logo_width_mm', '80'),
                        # Individuālais QR kods
                        include_custom_qr_code=get_bool(data, 'include_custom_qr_code', False),
                        custom_qr_code_data=data.get('custom_qr_code_data', ''),
                        custom_qr_code_size_mm=get_decimal(data, 'custom_qr_code_size_mm', '20'),
                        custom_qr_code_position=data.get('custom_qr_code_position', 'bottom_right'),
                        custom_qr_code_pos_x_mm=get_decimal(data, 'custom_qr_code_pos_x_mm', '0'),
                        custom_qr_code_pos_y_mm=get_decimal(data, 'custom_qr_code_pos_y_mm', '0'),
                        custom_qr_code_color=data.get('custom_qr_code_color', '#000000'),

                        # Automātiskais QR kods (akta ID)
                        include_auto_qr_code=get_bool(data, 'include_auto_qr_code', False),
                        auto_qr_code_size_mm=get_decimal(data, 'auto_qr_code_size_mm', '20'),
                        auto_qr_code_position=data.get('auto_qr_code_position', 'bottom_left'),
                        auto_qr_code_pos_x_mm=get_decimal(data, 'auto_qr_code_pos_x_mm', '0'),
                        auto_qr_code_pos_y_mm=get_decimal(data, 'auto_qr_code_pos_y_mm', '0'),
                        auto_qr_code_color=data.get('auto_qr_code_color', '#000000'),

                        add_watermark=get_bool(data, 'add_watermark', False),
                        watermark_text=data.get('watermark_text', 'MELNRAKSTS'),
                        watermark_font_size=data.get('watermark_font_size', 72),
                        watermark_color=data.get('watermark_color', '#E0E0E0'),
                        watermark_rotation=data.get('watermark_rotation', 45),
                        enable_pdf_encryption=get_bool(data, 'enable_pdf_encryption', False),
                        pdf_user_password=data.get('pdf_user_password', ''),
                        pdf_owner_password=data.get('pdf_owner_password', ''),
                        allow_printing=get_bool(data, 'allow_printing', True),
                        allow_copying=get_bool(data, 'allow_copying', True),
                        allow_modifying=get_bool(data, 'allow_modifying', False),
                        allow_annotating=get_bool(data, 'allow_annotating', True),
                        default_country=data.get('default_country', 'Latvija'),
                        default_city=data.get('default_city', 'Rīga'),
                        show_contact_details_in_header=get_bool(data, 'show_contact_details_in_header', False),
                        contact_details_header_font_size=data.get('contact_details_header_font_size', 8),
                        item_image_width_mm=get_decimal(data, 'item_image_width_mm', '50'),
                        item_image_caption_font_size=data.get('item_image_caption_font_size', 8),
                        show_item_notes_in_table=get_bool(data, 'show_item_notes_in_table', True),
                        show_item_serial_number_in_table=get_bool(data, 'show_item_serial_number_in_table', True),
                        show_item_warranty_in_table=get_bool(data, 'show_item_warranty_in_table', True),
                        table_cell_padding_mm=get_decimal(data, 'table_cell_padding_mm', '2'),
                        table_header_font_style=data.get('table_header_font_style', 'bold'),
                        table_content_alignment=data.get('table_content_alignment', 'left'),
                        signature_font_size=data.get('signature_font_size', 9),
                        signature_spacing_mm=get_decimal(data, 'signature_spacing_mm', '10'),
                        document_title_font_size=data.get('document_title_font_size', 18),
                        document_title_color=data.get('document_title_color', '#000000'),
                        section_heading_font_size=data.get('section_heading_font_size', 12),
                        section_heading_color=data.get('section_heading_color', '#000000'),
                        paragraph_line_spacing_multiplier=get_decimal(data, 'paragraph_line_spacing_multiplier', '1.2'),
                        table_border_style=data.get('table_border_style', 'solid'),
                        table_border_thickness_pt=get_decimal(data, 'table_border_thickness_pt', '0.5'),
                        table_alternate_row_color=data.get('table_alternate_row_color', ''),
                        show_total_sum_in_words=get_bool(data, 'show_total_sum_in_words', False),
                        total_sum_in_words_language=data.get('total_sum_in_words_language', 'lv'),
                        default_vat_calculation_method=data.get('default_vat_calculation_method', 'exclusive'),
                        show_vat_breakdown=get_bool(data, 'show_vat_breakdown', True),
                        enable_digital_signature_field=get_bool(data, 'enable_digital_signature_field', False),
                        digital_signature_field_name=data.get('digital_signature_field_name', 'Paraksts'),
                        digital_signature_field_size_mm=get_decimal(data, 'digital_signature_field_size_mm', '40'),
                        digital_signature_field_position=data.get('digital_signature_field_position', 'bottom_center'),
                        template_password=data.get('template_password', '')  # Ielādējam paroli
                    )

                    self.ieviest_datus(d)
                    QMessageBox.information(self, "Šablons ielādēts",
                                            f"Šablons '{sablonu_nosaukums}' veiksmīgi ielādēts.")
            except Exception as e:
                QMessageBox.critical(self, "Kļūda", f"Neizdevās ielādēt šablonu '{sablonu_nosaukums}':\n{e}")
        else:
            QMessageBox.warning(self, "Kļūda", f"Šablona fails '{sablonu_nosaukums}.json' nav atrasts.")

    # ----- Tab: Adrešu grāmata -----
    def _būvēt_adresu_gramata_tab(self):
        w = QWidget()
        v = QVBoxLayout()

        self.address_book_list = QListWidget()
        self.address_book_list.itemDoubleClicked.connect(self._load_selected_address_book_entry)

        

        
        # --- JAUNS: labā klikšķa konteksta izvēlne adrešu grāmatai ---
        self.address_book_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.address_book_list.customContextMenuRequested.connect(self._show_address_book_context_menu)
# --- JAUNS: nosaukuma režīms saglabāšanai ---
        self.chk_ab_auto_name = QCheckBox("Nosaukumu ģenerē sistēma automātiski (Uzņēmums + Kontaktpersona)")
        self.chk_ab_auto_name.setChecked(True)
        btn_load_selected = QPushButton("Ielādēt izvēlēto")
        btn_load_selected.clicked.connect(lambda: self._load_selected_address_book_entry(self.address_book_list.currentItem()))
        btn_delete_selected = QPushButton("Dzēst izvēlēto")
        btn_delete_selected.clicked.connect(self._delete_selected_address_book_entry)

        btns_layout = QHBoxLayout()
        btns_layout.addWidget(btn_load_selected)
        btns_layout.addWidget(btn_delete_selected)
        btns_layout.addStretch()

        v.addWidget(QLabel("Saglabātās personas:"))
        v.addWidget(self.chk_ab_auto_name)
        v.addWidget(self.address_book_list)
        v.addLayout(btns_layout)
        v.addStretch()

        w.setLayout(v)
        self.tabs.addTab(w, "Adrešu grāmata")
        self._update_address_book_list()

    # ======================
    # Adrešu grāmata: labais klikšķis + parole + rediģēšana
    # ======================



    def _būvēt_audit_tab(self):
        """Audit tab: rāda pēdējos ierakstus un ļauj eksportēt."""
        tab = QWidget()
        v = QVBoxLayout(tab)

        top = QHBoxLayout()
        self._audit_filter = QLineEdit()
        self._audit_filter.setPlaceholderText("Filtrs (meklē event / user / detaļās)…")
        btn_refresh = QPushButton("Atjaunot")
        btn_refresh.clicked.connect(lambda: self._refresh_audit_table(limit=400))
        btn_export = QPushButton("Eksportēt…")
        btn_export.clicked.connect(self._export_audit_log)

        top.addWidget(QLabel("Audit logs:"))
        top.addWidget(self._audit_filter, 1)
        top.addWidget(btn_refresh)
        top.addWidget(btn_export)
        v.addLayout(top)

        self._audit_table = QTableWidget(0, 4)
        self._audit_table.setHorizontalHeaderLabels(["Laiks", "Lietotājs", "Notikums", "Detaļas"])
        self._audit_table.horizontalHeader().setStretchLastSection(True)
        self._audit_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self._audit_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        v.addWidget(self._audit_table, 1)

        # Undo/Redo pogas (ērti redzamas)
        h = QHBoxLayout()
        btn_undo = QPushButton("Undo (Ctrl+Z)")
        btn_redo = QPushButton("Redo (Ctrl+Y)")
        btn_undo.clicked.connect(self.undo_action)
        btn_redo.clicked.connect(self.redo_action)
        h.addWidget(btn_undo)
        h.addWidget(btn_redo)
        h.addStretch(1)
        v.addLayout(h)

        self._audit_filter.textChanged.connect(lambda: self._refresh_audit_table(limit=400))
        self.tabs.addTab(tab, "Audit")

        self._refresh_audit_table(limit=200)

    def _color_for_audit_event(self, event_name: str):
        try:
            e = (event_name or "").upper()
            if e.startswith("AB_"):
                return QColor(219, 234, 254)  # light blue
            if e.startswith("PROJECT_"):
                return QColor(220, 252, 231)  # light green
            if e.startswith("GENERATE_"):
                return QColor(255, 237, 213)  # light orange
            if e.startswith("POZ_"):
                return QColor(243, 232, 255)  # light purple
            if e.startswith("FIELD_") or e.startswith("COMBO_") or e.startswith("CHECK_"):
                return QColor(241, 245, 249)  # light slate
            if e in ("UNDO", "REDO"):
                return QColor(226, 232, 240)  # gray
            return None
        except Exception:
            return None

    def _refresh_audit_table(self, limit: int = 200):
        if not hasattr(self, "_audit_table") or self._audit_table is None:
            return
        flt = (self._audit_filter.text() or "").strip().lower() if hasattr(self, "_audit_filter") else ""
        rows = self._audit_logger.tail(limit)
        # filtrējam
        if flt:
            def ok(r):
                try:
                    s = (r.get("ts","") + " " + r.get("user","") + " " + r.get("event","") + " " + json.dumps(r.get("details",{}), ensure_ascii=False)).lower()
                    return flt in s
                except Exception:
                    return True
            rows = [r for r in rows if ok(r)]

        self._audit_table.setRowCount(0)
        for r in rows[::-1]:  # newest first
            row = self._audit_table.rowCount()
            self._audit_table.insertRow(row)
            self._audit_table.setItem(row, 0, QTableWidgetItem(str(r.get("ts",""))))
            self._audit_table.setItem(row, 1, QTableWidgetItem(str(r.get("user",""))))
            self._audit_table.setItem(row, 2, QTableWidgetItem(str(r.get("event",""))))
            it0 = QTableWidgetItem(str(r.get("ts","")))
            it1 = QTableWidgetItem(str(r.get("user","")))
            it2 = QTableWidgetItem(str(r.get("event","")))
            it3 = QTableWidgetItem(json.dumps(r.get("details", {}), ensure_ascii=False))
            self._audit_table.setItem(row, 0, it0)
            self._audit_table.setItem(row, 1, it1)
            self._audit_table.setItem(row, 2, it2)
            self._audit_table.setItem(row, 3, it3)
            bg = self._color_for_audit_event(str(r.get("event","")))
            if bg is not None:
                for it in (it0, it1, it2, it3):
                    it.setBackground(bg)


    def _update_undo_redo_indicators(self):
        try:
            if not hasattr(self, "_undo_mgr"):
                return
            u = len(getattr(self._undo_mgr, "_undo", []))
            r = len(getattr(self._undo_mgr, "_redo", []))
            if getattr(self, "_undo_status_label", None) is not None:
                self._undo_status_label.setText(f"Undo: {u}")
                self._undo_status_label.setStyleSheet("padding:2px; color: " + ("#16a34a" if u>0 else "#94a3b8"))
            if getattr(self, "_redo_status_label", None) is not None:
                self._redo_status_label.setText(f"Redo: {r}")
                self._redo_status_label.setStyleSheet("padding:2px; color: " + ("#2563eb" if r>0 else "#94a3b8"))
        except Exception:
            pass

    def eventFilter(self, obj, event):
        # Globāla izmaiņu izsekošana: FIELD_EDIT / COMBO_CHANGE / CHECK_TOGGLE
        try:
            if not getattr(self, "_track_enabled", False):
                return super().eventFilter(obj, event)

            et = event.type()
            # LineEdit: audit uz FocusOut (kad beidz rediģēt)
            if isinstance(obj, QLineEdit) and et == QEvent.FocusOut:
                key = id(obj)
                cur = obj.text()
                prev = self._last_widget_values.get(key, None)
                if prev is None:
                    self._last_widget_values[key] = cur
                    return super().eventFilter(obj, event)
                if cur != prev:
                    self._undo_mgr.push_undo(self._snapshot_state("FIELD_EDIT"))
                    self._audit("FIELD_EDIT", {"field": obj.objectName() or obj.placeholderText() or "QLineEdit", "value": (cur or "")[:200]})
                    self._last_widget_values[key] = cur
                    self._update_undo_redo_indicators()
                return super().eventFilter(obj, event)

            # ComboBox: audit uz FocusOut
            if isinstance(obj, QComboBox) and et == QEvent.FocusOut:
                key = id(obj)
                cur = obj.currentText()
                prev = self._last_widget_values.get(key, None)
                if prev is None:
                    self._last_widget_values[key] = cur
                    return super().eventFilter(obj, event)
                if cur != prev:
                    self._undo_mgr.push_undo(self._snapshot_state("COMBO_CHANGE"))
                    self._audit("COMBO_CHANGE", {"field": obj.objectName() or "QComboBox", "value": (cur or "")[:200]})
                    self._last_widget_values[key] = cur
                    self._update_undo_redo_indicators()
                return super().eventFilter(obj, event)

            # CheckBox: audit uz MouseButtonRelease (toggling)
            if isinstance(obj, QCheckBox) and et == QEvent.MouseButtonRelease:
                key = id(obj)
                cur = bool(obj.isChecked())
                prev = self._last_widget_values.get(key, None)
                if prev is None:
                    self._last_widget_values[key] = cur
                    return super().eventFilter(obj, event)
                if cur != prev:
                    self._undo_mgr.push_undo(self._snapshot_state("CHECK_TOGGLE"))
                    self._audit("CHECK_TOGGLE", {"field": obj.text() or obj.objectName() or "QCheckBox", "value": cur})
                    self._last_widget_values[key] = cur
                    self._update_undo_redo_indicators()
                return super().eventFilter(obj, event)

        except Exception:
            pass
        return super().eventFilter(obj, event)

    def _ensure_positions_undo_hook(self):
        # Pieslēdz undo/redo checkpoint pozīciju tabulas izmaiņām (debounce)
        try:
            if not hasattr(self, "tab") or self.tab is None:
                return
            if hasattr(self, "_pos_change_timer") and self._pos_change_timer is not None:
                return
            self._pos_change_timer = QTimer(self)
            self._pos_change_timer.setSingleShot(True)
            self._pos_change_timer.timeout.connect(self._commit_positions_change)
            try:
                self.tab.cellChanged.connect(self._on_positions_cell_changed)
            except Exception:
                pass
        except Exception:
            pass

    def _on_positions_cell_changed(self, row: int, col: int):
        try:
            self._last_pos_change = (row, col)
            if hasattr(self, "_pos_change_timer") and self._pos_change_timer is not None:
                self._pos_change_timer.start(700)
        except Exception:
            pass

    def _commit_positions_change(self):
        try:
            self._undo_mgr.push_undo(self._snapshot_state("POZ_CHANGE"))
            rc = getattr(self, "_last_pos_change", None)
            details = {"row": rc[0], "col": rc[1]} if rc else {}
            self._audit("POZ_CHANGE", details)
            self._update_undo_redo_indicators()
        except Exception:
            pass


    def _export_audit_log(self):
        try:
            default = os.path.join(PROJECT_SAVE_DIR, "audit_export.jsonl")
            fn, _ = QFileDialog.getSaveFileName(self, "Eksportēt audit log", default, "JSONL (*.jsonl);;Teksts (*.txt)")
            if not fn:
                return
            srcp = self._audit_logger.log_path
            if srcp and os.path.exists(srcp):
                shutil.copy2(srcp, fn)
                QMessageBox.information(self, "OK", "Audit logs eksportēts.")
            else:
                QMessageBox.warning(self, "Nav", "Audit logs fails vēl nav izveidots.")
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", f"Neizdevās eksportēt: {e}")


    def _ab_hash_password(self, password: str, salt: str) -> str:
        # Hash parolei (PBKDF2-HMAC-SHA256).
        try:
            dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120_000)
            return dk.hex()
        except Exception as e:
            print(f"Paroles hash kļūda: {e}")
            return ""

    def _ab_verify_password(self, persona_data: dict, password: str) -> bool:
        # Pārbauda paroli pret saglabāto hash.
        try:
            salt = (persona_data.get("__pw_salt") or "").strip()
            h = (persona_data.get("__pw_hash") or "").strip()
            if not salt or not h:
                return True
            return self._ab_hash_password(password, salt) == h
        except Exception:
            return False

    def _ab_require_password(self, entry_name: str, persona_data: dict, action_label: str = "ielādēt") -> bool:
        # Ja pusei ir parole, prasa to pirms ielādes/rediģēšanas.
        try:
            if not persona_data:
                return False
            if not (persona_data.get("__pw_hash") and persona_data.get("__pw_salt")):
                return True

            pwd, ok = QInputDialog.getText(
                self,
                "Parole nepieciešama",
                f"Lai {action_label} '{entry_name}', ievadi paroli:",
                QLineEdit.Password,
                ""
            )
            if not ok:
                return False
            if self._ab_verify_password(persona_data, pwd):
                return True

            QMessageBox.warning(self, "Nepareiza parole", "Parole nav pareiza.")
            return False
        except Exception:
            return False

    def _show_address_book_context_menu(self, pos):
        # Labā klikšķa izvēlne adrešu grāmatas sarakstam.
        item = self.address_book_list.itemAt(pos)
        menu = QMenu(self)

        act_load = QAction("Ielādēt", self)
        act_edit = QAction("Rediģēt…", self)
        act_set_pw = QAction("Uzlikt / mainīt paroli…", self)
        act_remove_pw = QAction("Noņemt paroli", self)
        act_delete = QAction("Dzēst", self)

        act_rename = QAction("Pārdēvēt…", self)
        act_duplicate = QAction("Dublēt", self)
        act_export_json = QAction("Eksportēt (JSON)…", self)
        act_export_csv = QAction("Eksportēt (CSV)…", self)

        if item is None:
            act_refresh = QAction("Atjaunot sarakstu", self)
            act_refresh.triggered.connect(self._update_address_book_list)
            menu.addAction(act_refresh)
            menu.exec(self.address_book_list.mapToGlobal(pos))
            return

        name = item.text()
        persona_data = self.address_book.get(name, {})

        act_rename.triggered.connect(lambda: self._rename_address_book_entry(name))
        act_duplicate.triggered.connect(lambda: self._duplicate_address_book_entry(name))
        act_export_json.triggered.connect(lambda: self._export_address_book_entry_json(name))
        act_export_csv.triggered.connect(lambda: self._export_address_book_entry_csv(name))

        act_load.triggered.connect(lambda: self._load_selected_address_book_entry(item))
        act_edit.triggered.connect(lambda: self._edit_address_book_entry(name))
        act_set_pw.triggered.connect(lambda: self._set_password_for_address_book_entry(name))
        act_remove_pw.triggered.connect(lambda: self._remove_password_for_address_book_entry(name))
        act_delete.triggered.connect(lambda: self._delete_address_book_entry(name))

        menu.addAction(act_load)
        menu.addAction(act_edit)
        menu.addAction(act_rename)
        menu.addAction(act_duplicate)
        menu.addSeparator()
        menu.addAction(act_export_json)
        menu.addAction(act_export_csv)
        menu.addSeparator()
        menu.addAction(act_set_pw)
        menu.addAction(act_remove_pw)
        menu.addSeparator()
        menu.addAction(act_delete)

        has_pw = bool(persona_data.get("__pw_hash") and persona_data.get("__pw_salt"))
        act_remove_pw.setEnabled(has_pw)

        menu.exec(self.address_book_list.mapToGlobal(pos))

    def _set_password_for_address_book_entry(self, entry_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_SET_PASSWORD'))
        self._audit('AB_SET_PASSWORD', {})
        persona_data = self.address_book.get(entry_name)
        if not persona_data:
            return

        if persona_data.get("__pw_hash") and persona_data.get("__pw_salt"):
            if not self._ab_require_password(entry_name, persona_data, "mainīt paroli"):
                return

        pw1, ok1 = QInputDialog.getText(self, "Uzlikt paroli", f"Ievadi jauno paroli '{entry_name}':", QLineEdit.Password, "")
        if not ok1:
            return
        pw2, ok2 = QInputDialog.getText(self, "Uzlikt paroli", "Atkārto paroli:", QLineEdit.Password, "")
        if not ok2:
            return
        if pw1 != pw2:
            QMessageBox.warning(self, "Kļūda", "Paroles nesakrīt.")
            return
        if not pw1.strip():
            QMessageBox.warning(self, "Kļūda", "Parole nevar būt tukša.")
            return

        salt = secrets.token_hex(8)
        h = self._ab_hash_password(pw1, salt)
        if not h:
            QMessageBox.warning(self, "Kļūda", "Neizdevās uzlikt paroli (hash tukšs).")
            return

        persona_data["__pw_salt"] = salt
        persona_data["__pw_hash"] = h
        self.address_book[entry_name] = persona_data
        self._save_address_book()
        QMessageBox.information(self, "OK", "Parole uzlikta.")

    def _remove_password_for_address_book_entry(self, entry_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_REMOVE_PASSWORD'))
        self._audit('AB_REMOVE_PASSWORD', {})
        persona_data = self.address_book.get(entry_name)
        if not persona_data:
            return
        if persona_data.get("__pw_hash") and persona_data.get("__pw_salt"):
            if not self._ab_require_password(entry_name, persona_data, "noņemt paroli"):
                return
        persona_data.pop("__pw_hash", None)
        persona_data.pop("__pw_salt", None)
        self.address_book[entry_name] = persona_data
        self._save_address_book()
        QMessageBox.information(self, "OK", "Parole noņemta.")

    def _rename_address_book_entry(self, old_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_RENAME'))
        self._audit('AB_RENAME', {})
        if old_name not in self.address_book:
            return
        persona_data = self.address_book.get(old_name, {})
        # Ja ir parole, prasām to pirms pārdēvēšanas
        if not self._ab_require_password(old_name, persona_data, "pārdēvēt"):
            return

        new_name, ok = QInputDialog.getText(
            self,
            "Pārdēvēt ierakstu",
            "Jaunais nosaukums:",
            QLineEdit.Normal,
            old_name
        )
        if not ok:
            return
        new_name = (new_name or "").strip()
        if not new_name:
            QMessageBox.warning(self, "Kļūda", "Nosaukums nevar būt tukšs.")
            return
        if new_name == old_name:
            return

        # Unikāls nosaukums
        if hasattr(self, "_ab_make_unique_key"):
            new_name = self._ab_make_unique_key(new_name)
        else:
            if new_name in self.address_book:
                i = 2
                base = new_name
                while f"{base} ({i})" in self.address_book:
                    i += 1
                new_name = f"{base} ({i})"

        self.address_book[new_name] = persona_data
        self.address_book.pop(old_name, None)
        self._save_address_book()
        self._update_address_book_list()

    def _duplicate_address_book_entry(self, src_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_DUPLICATE'))
        self._audit('AB_DUPLICATE', {})
        if src_name not in self.address_book:
            return
        persona_data = dict(self.address_book.get(src_name, {}))  # shallow copy
        # Ja ir parole, prasām to pirms dublēšanas
        if not self._ab_require_password(src_name, persona_data, "dublēt"):
            return

        base = f"{src_name} (kopija)"
        if hasattr(self, "_ab_make_unique_key"):
            new_name = self._ab_make_unique_key(base)
        else:
            new_name = base
            if new_name in self.address_book:
                i = 2
                while f"{base} ({i})" in self.address_book:
                    i += 1
                new_name = f"{base} ({i})"

        self.address_book[new_name] = persona_data
        self._save_address_book()
        self._update_address_book_list()
        QMessageBox.information(self, "OK", f"Ieraksts dublēts kā: {new_name}")

    def _export_address_book_entry_json(self, entry_name: str):
        if entry_name not in self.address_book:
            return
        persona_data = dict(self.address_book.get(entry_name, {}))
        if not self._ab_require_password(entry_name, persona_data, "eksportēt"):
            return

        # Neeksportējam paroles hash/salt
        persona_data.pop("__pw_hash", None)
        persona_data.pop("__pw_salt", None)

        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Eksportēt JSON",
            f"{entry_name}.json",
            "JSON faili (*.json)"
        )
        if not filename:
            return
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump({"name": entry_name, "data": persona_data}, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "OK", "JSON eksports pabeigts.")
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", f"Neizdevās eksportēt JSON: {e}")

    def _export_address_book_entry_csv(self, entry_name: str):
        if entry_name not in self.address_book:
            return
        persona_data = dict(self.address_book.get(entry_name, {}))
        if not self._ab_require_password(entry_name, persona_data, "eksportēt"):
            return

        # Neeksportējam paroles hash/salt
        persona_data.pop("__pw_hash", None)
        persona_data.pop("__pw_salt", None)

        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Eksportēt CSV",
            f"{entry_name}.csv",
            "CSV faili (*.csv)"
        )
        if not filename:
            return
        try:
            keys = list(persona_data.keys())
            with open(filename, "w", newline="", encoding="utf-8") as f:
                w = csv.DictWriter(f, fieldnames=keys, delimiter=";")
                w.writeheader()
                w.writerow(persona_data)
            QMessageBox.information(self, "OK", "CSV eksports pabeigts.")
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", f"Neizdevās eksportēt CSV: {e}")


    def _delete_address_book_entry(self, entry_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_DELETE'))
        self._audit('AB_DELETE', {})
        if entry_name not in self.address_book:
            return
        persona_data = self.address_book.get(entry_name, {})
        # --- JAUNS: dzēšana ir aizsargāta ar paroli (ja tā uzlikta) ---
        if persona_data.get('__pw_hash') and persona_data.get('__pw_salt'):
            if not self._ab_require_password(entry_name, persona_data, 'dzēst'):
                return
        reply = QMessageBox.question(
            self, "Dzēst ierakstu", f"Dzēst '{entry_name}' no adrešu grāmatas?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return
        self.address_book.pop(entry_name, None)
        self._save_address_book()
        self._update_address_book_list()

    def _edit_address_book_entry(self, entry_name: str):
        self._undo_mgr.push_undo(self._snapshot_state('AB_EDIT'))
        self._audit('AB_EDIT', {})
        # Atver jaunu logu rediģēšanai.
        persona_data = self.address_book.get(entry_name)
        if not persona_data:
            return

        if not self._ab_require_password(entry_name, persona_data, "rediģēt"):
            return

        dlg = QDialog(self)
        dlg.setWindowTitle(f"Rediģēt: {entry_name}")
        layout = QVBoxLayout(dlg)
        form = QFormLayout()

        fields = [
            ("Nosaukums", "nosaukums"),
            ("Reģ. Nr.", "reģ_nr"),
            ("Adrese", "adrese"),
            ("Kontaktpersona", "kontaktpersona"),
            ("Amats", "amats"),
            ("Pilnvaras pamats", "pilnvaras_pamats"),
            ("Tālrunis", "tālrunis"),
            ("E-pasts", "epasts"),
            ("Web lapa", "web_lapa"),
            ("Bankas konts", "bankas_konts"),
            ("Juridiskais statuss", "juridiskais_statuss"),
        ]

        widgets = {}

        for label, key in fields:
            if key == "pilnvaras_pamats":
                cb = QComboBox()
                try:
                    if hasattr(self, "pie_in") and self.pie_in and hasattr(self.pie_in[5], "count"):
                        for i in range(self.pie_in[5].count()):
                            cb.addItem(self.pie_in[5].itemText(i))
                    else:
                        cb.addItems(["Pilnvaras pamats", "Cits"])
                except Exception:
                    cb.addItems(["Pilnvaras pamats", "Cits"])
                val = persona_data.get(key, "")
                idx = cb.findText(val)
                if idx >= 0:
                    cb.setCurrentIndex(idx)
                widgets[key] = cb
                form.addRow(label + ":", cb)
            elif key == "juridiskais_statuss":
                cb = QComboBox()
                try:
                    if hasattr(self, "pie_in") and self.pie_in and hasattr(self.pie_in[10], "count"):
                        for i in range(self.pie_in[10].count()):
                            cb.addItem(self.pie_in[10].itemText(i))
                    else:
                        cb.addItems(["Juridiska persona", "Fiziska persona"])
                except Exception:
                    cb.addItems(["Juridiska persona", "Fiziska persona"])
                val = persona_data.get(key, "")
                idx = cb.findText(val)
                if idx >= 0:
                    cb.setCurrentIndex(idx)
                widgets[key] = cb
                form.addRow(label + ":", cb)
            else:
                le = QLineEdit(str(persona_data.get(key, "") or ""))
                widgets[key] = le
                form.addRow(label + ":", le)

        layout.addLayout(form)

        buttons = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        layout.addWidget(buttons)

        def on_save():
            try:
                for _, key in fields:
                    w = widgets.get(key)
                    if w is None:
                        continue
                    if isinstance(w, QComboBox):
                        persona_data[key] = w.currentText()
                    else:
                        persona_data[key] = w.text().strip()
                self.address_book[entry_name] = persona_data
                self._save_address_book()
                self._update_address_book_list()
                dlg.accept()
            except Exception as e:
                QMessageBox.warning(self, "Kļūda", f"Neizdevās saglabāt: {e}")

        buttons.accepted.connect(on_save)
        buttons.rejected.connect(dlg.reject)
        dlg.exec()

    def _update_address_book_list(self):
        self.address_book_list.clear()
        for name in sorted(self.address_book.keys()):
            self.address_book_list.addItem(name)

    def _load_selected_address_book_entry(self, item: QListWidgetItem):
        if not item:
            return
        name = item.text()
        persona_data = self.address_book.get(name)
        if persona_data:
            # --- JAUNS: ja ir parole, prasa to pirms ielādes ---
            if not self._ab_require_password(name, persona_data, "ielādēt"):
                return
            reply = QMessageBox.question(self, "Ielādēt personu",
                                         f"Ielādēt '{name}' kā Pieņēmēju vai Nodevēju?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel,
                                         QMessageBox.StandardButton.Yes)
            if reply == QMessageBox.StandardButton.Yes:
                self.pie_in[0].setText(persona_data.get("nosaukums", ""))
                self.pie_in[1].setText(persona_data.get("reģ_nr", ""))
                self.pie_in[2].setText(persona_data.get("adrese", ""))
                self.pie_in[3].setText(persona_data.get("kontaktpersona", ""))
                self.pie_in[4].setText(persona_data.get("amats", ""))
                pilnvaras_val = persona_data.get("pilnvaras_pamats", "")
                idx_p = self.pie_in[5].findText(pilnvaras_val)
                if idx_p >= 0:
                    self.pie_in[5].setCurrentIndex(idx_p)
                self.pie_in[6].setText(persona_data.get("tālrunis", ""))
                self.pie_in[7].setText(persona_data.get("epasts", ""))
                self.pie_in[8].setText(persona_data.get("web_lapa", ""))
                self.pie_in[9].setText(persona_data.get("bankas_konts", ""))
                juridiskais_statuss_val = persona_data.get("juridiskais_statuss", "")
                idx = self.pie_in[10].findText(juridiskais_statuss_val)
                if idx >= 0:
                    self.pie_in[10].setCurrentIndex(idx)
                QMessageBox.information(self, "Ielādēts", f"Persona '{name}' ielādēta kā Pieņēmējs.")
            elif reply == QMessageBox.StandardButton.No:
                self.nod_in[0].setText(persona_data.get("nosaukums", ""))
                self.nod_in[1].setText(persona_data.get("reģ_nr", ""))
                self.nod_in[2].setText(persona_data.get("adrese", ""))
                self.nod_in[3].setText(persona_data.get("kontaktpersona", ""))
                self.nod_in[4].setText(persona_data.get("amats", ""))
                pilnvaras_val = persona_data.get("pilnvaras_pamats", "")
                idx_p = self.nod_in[5].findText(pilnvaras_val)
                if idx_p >= 0:
                    self.nod_in[5].setCurrentIndex(idx_p)
                self.nod_in[6].setText(persona_data.get("tālrunis", ""))
                self.nod_in[7].setText(persona_data.get("epasts", ""))
                self.nod_in[8].setText(persona_data.get("web_lapa", ""))
                self.nod_in[9].setText(persona_data.get("bankas_konts", ""))
                juridiskais_statuss_val = persona_data.get("juridiskais_statuss", "")
                idx = self.nod_in[10].findText(juridiskais_statuss_val)
                if idx >= 0:
                    self.nod_in[10].setCurrentIndex(idx)
                QMessageBox.information(self, "Ielādēts", f"Persona '{name}' ielādēta kā Nodevējs.")
        else:
            QMessageBox.warning(self, "Kļūda", "Izvēlētā persona nav atrasta adrešu grāmatā.")

    def _delete_selected_address_book_entry(self):
        item = self.address_book_list.currentItem()
        if not item:
            QMessageBox.warning(self, "Dzēst personu", "Lūdzu, izvēlieties personu, ko dzēst.")
            return
        name = item.text()
        reply = QMessageBox.question(self, "Dzēst personu",
                                     f"Vai tiešām vēlaties dzēst '{name}' no adrešu grāmatas?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            if name in self.address_book:
                del self.address_book[name]
                self._save_address_book() # Saglabājam izmaiņas failā
                self._update_address_book_list() # Atjaunojam sarakstu GUI
                QMessageBox.information(self, "Dzēsts", f"Persona '{name}' veiksmīgi dzēsta.")
            else:
                QMessageBox.warning(self, "Kļūda", "Izvēlētā persona nav atrasta adrešu grāmatā.")

    # ----- Tab: Dokumentu vēsture -----
    def _būvēt_dokumentu_vesture_tab(self):
        w = QWidget()
        v = QVBoxLayout()

        self.history_list = QListWidget()
        self.history_list.itemDoubleClicked.connect(self._load_history_entry)

        btn_load_history = QPushButton("Ielādēt izvēlēto projektu")
        btn_load_history.clicked.connect(lambda: self._load_history_entry(self.history_list.currentItem()))
        btn_clear_history = QPushButton("Notīrīt vēsturi")
        btn_clear_history.clicked.connect(self._clear_history)

        btn_open_folder = QPushButton("Atvērt mapi")
        btn_open_folder.clicked.connect(self._open_document_folder)

        btns_layout = QHBoxLayout()
        btns_layout.addWidget(btn_load_history)
        btns_layout.addWidget(btn_clear_history)
        btns_layout.addWidget(btn_open_folder)  # PIEVIENOJAM JAUNO POGU
        btns_layout.addStretch()

        v.addWidget(QLabel("Pēdējie projekti:"))
        v.addWidget(self.history_list)
        v.addLayout(btns_layout)
        v.addStretch()

        w.setLayout(v)
        self.tabs.addTab(w, "Dokumentu vēsture")
        self._update_history_list()

    def _load_history(self):
        if os.path.exists(HISTORY_FILE):
            try:
                # Mēģinām ielādēt ar dažādām kodēšanām
                encodings = ['utf-8', 'utf-8-sig', 'cp1257', 'iso-8859-1', 'windows-1252']
                for encoding in encodings:
                    try:
                        with open(HISTORY_FILE, 'r', encoding=encoding) as f:
                            self.history = json.load(f)
                        break
                    except (UnicodeDecodeError, json.JSONDecodeError):
                        continue
                else:
                    # Ja neviena kodēšana nedarbojas, izveidojam jaunu vēsturi
                    print(f"Neizdevās ielādēt vēstures failu ar nevenu kodēšanu. Izveidojam jaunu.")
                    self.history = []
            except Exception as e:
                QMessageBox.warning(self, "Kļūda", f"Neizdevās ielādēt vēsturi: {e}")
                self.history = []
        else:
            self.history = []

    def _save_history(self):
        os.makedirs(SETTINGS_DIR, exist_ok=True) # Izveidojam direktoriju, ja tā neeksistē
        try:
            with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās saglabāt vēsturi: {e}")

    
    def _record_generated_document(self, pdf_path: str, json_path: str):
        """Saglabā PDF+JSON kopijas dokumentu vēsturē, lai tās vienmēr būtu pieejamas."""
        try:
            os.makedirs(DOCUMENTS_DIR, exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            base = f"Akts_{ts}"
            dst_pdf = os.path.join(DOCUMENTS_DIR, base + ".pdf")
            dst_json = os.path.join(DOCUMENTS_DIR, base + ".json")
            shutil.copy2(pdf_path, dst_pdf)
            shutil.copy2(json_path, dst_json)

            # Vēsture tagad ir saraksts ar dict ierakstiem
            if not isinstance(self.history, list):
                self.history = []
            # migrācija: ja bija vecais formāts (stringi)
            new_hist = []
            for it in self.history:
                if isinstance(it, str):
                    new_hist.append({"json": it, "pdf": "", "created": ""})
                elif isinstance(it, dict):
                    new_hist.append(it)
            self.history = new_hist

            self.history.insert(0, {"pdf": dst_pdf, "json": dst_json, "created": datetime.now().isoformat(timespec="seconds")})
            self._save_history()
            self._update_history_list()
        except Exception as e:
            print(f"Vēstures saglabāšanas kļūda: {e}")

    def _add_to_history(self, file_path: str):
        file_path = _coerce_path(file_path)
        if not file_path:
            return
        if not os.path.exists(file_path):
            return

        # Noņemam, ja jau ir sarakstā, lai pārvietotu uz saraksta sākumu
        self.history = [f for f in self.history if f != file_path]

        # Pievienojam saraksta sākumā
        self.history.insert(0, file_path)
        self._save_history() # Saglabājam izmaiņas failā
        self._update_history_list() # Atjaunojam sarakstu GUI

    def _update_history_list(self):
        self.history_list.clear()

        # Migrācija no vecā formāta (saraksts ar failu ceļiem)
        if isinstance(self.history, list) and self.history and isinstance(self.history[0], str):
            self.history = [{"json": p, "pdf": "", "created": ""} for p in self.history]

        valid = []
        for it in (self.history or []):
            if isinstance(it, dict):
                j = it.get("json", "")
                j = _coerce_path(j) or ""
                p = it.get("pdf", "")
                p = _coerce_path(p) or ""
                if j and os.path.exists(j) or (p and os.path.exists(p)):
                    valid.append(it)
            elif isinstance(it, str) and os.path.exists(it):
                valid.append({"json": it, "pdf": "", "created": ""})

        self.history = valid
        self._save_history()

        for it in self.history:
            created = it.get("created", "")
            label = os.path.basename(it.get("pdf") or it.get("json") or "")
            if created:
                label = f"{label}  ({created.replace('T',' ')})"
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, it)
            self.history_list.addItem(item)

    def _load_history_entry(self, item: QListWidgetItem):
        if not item:
            return
        payload = item.data(Qt.UserRole)
        if isinstance(payload, dict):
            json_path = payload.get("json", "")
            json_path = _coerce_path(json_path) or ""
            if json_path and os.path.exists(json_path):
                self.ieladet_projektu(json_path)
                return
            # vecs ieraksts bez json
            pdf_path = payload.get("pdf", "")
            pdf_path = _coerce_path(pdf_path) or ""
            if pdf_path and os.path.exists(pdf_path):
                QDesktopServices.openUrl(QUrl.fromLocalFile(pdf_path))
                return
        # fallback (ja kāds vecs ieraksts)
        txt = item.text().split("  (")[0]
        # mēģinām atrast pēc nosaukuma
        for it in self.history or []:
            if isinstance(it, dict):
                p = it.get("json") or it.get("pdf")
                if p and os.path.basename(p) == txt and os.path.exists(p):
                    if p.lower().endswith(".json"):
                        self.ieladet_projektu(p)
                    else:
                        QDesktopServices.openUrl(QUrl.fromLocalFile(p))
                    return

        QMessageBox.warning(self, "Kļūda", "Neizdevās atrast vēstures ieraksta failu.")

    def _clear_history(self):
        reply = QMessageBox.question(self, "Notīrīt vēsturi",
                                     "Vai tiešām vēlaties notīrīt visu dokumentu vēsturi?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.history = []
            self._save_history() # Saglabājam tukšu vēsturi failā
            self._update_history_list() # Atjaunojam sarakstu GUI
            QMessageBox.information(self, "Vēsture notīrīta", "Dokumentu vēsture ir notīrīta.")

    # ----- Tab: Karte -----

    def _open_document_folder(self):
        """Atver izvēlētā dokumenta mapi failu pārlūkā"""
        item = self.history_list.currentItem()
        if not item:
            QMessageBox.warning(self, "Atvērt mapi", "Lūdzu, izvēlieties dokumentu.")
            return

        file_name = item.text()
        file_path = next(((_coerce_path(f) or "") for f in self.history if _coerce_path(f) and os.path.basename(_coerce_path(f)) == file_name), None)
        file_path = _coerce_path(file_path)
        if file_path and os.path.exists(file_path):
            folder_path = os.path.dirname(file_path)
            if sys.platform == "win32":
                os.startfile(folder_path)
            elif sys.platform == "darwin":
                os.system(f"open '{folder_path}'")
            else:
                os.system(f"xdg-open '{folder_path}'")
        else:
            QMessageBox.warning(self, "Kļūda", "Dokumenta mape nav atrasta.")

    def _būvēt_kartes_tab(self):
        w = QWidget()
        v = QVBoxLayout()

        # Meklēšana kartē (Photon) — ļauj atrast adresi un atlasīt no saraksta
        search_row = QHBoxLayout()
        self.map_search_input = QLineEdit()
        self.map_search_input.setPlaceholderText("Meklēt adresi (piem. Brīvības iela 1, Rīga)")
        btn_search = QPushButton("Meklēt")
        btn_search.clicked.connect(self._search_address_on_map)
        search_row.addWidget(self.map_search_input)
        search_row.addWidget(btn_search)

        self.map_search_results = QListWidget()
        self.map_search_results.setMaximumHeight(140)
        self.map_search_results.itemDoubleClicked.connect(self._apply_map_search_result)

        v.addLayout(search_row)
        v.addWidget(self.map_search_results)

        self.web_view = QWebEngineView()

        # Izveidojam bridge objektu komunikācijai ar JavaScript
        self.map_bridge = MapBridge()
        self.map_bridge.map_click_callback = self._handle_map_click

        # Iestatām QWebChannel
        self.channel = QWebChannel()
        self.channel.registerObject("mapBridge", self.map_bridge)
        self.web_view.page().setWebChannel(self.channel)

        # HTML content for the map, including Leaflet from CDN
        map_html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Interaktīva karte</title>
        <meta charset="utf-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <link rel="stylesheet" href="https://unpkg.com/leaflet@1.7.1/dist/leaflet.css" />
        <script src="https://unpkg.com/leaflet@1.7.1/dist/leaflet.js"></script>
        <script src="qrc:///qtwebchannel/qwebchannel.js"></script>
        <style>
            body { margin: 0; padding: 0; }
            #mapid {
                height: 100vh;
                width: 100%;
                cursor: crosshair;
            }
            .custom-popup {
                font-family: Arial, sans-serif;
                font-size: 14px;
            }
        </style>
    </head>
    <body>
        <div id="mapid"></div>
        <script>
            console.log("Ielādē karti...");

            // Inicializēt QWebChannel
            var mapBridge;
            new QWebChannel(qt.webChannelTransport, function(channel) {
                mapBridge = channel.objects.mapBridge;
                console.log("QWebChannel inicializēts!");
            });

            // Inicializēt karti (Rīga, Latvija)
            var mymap = L.map('mapid').setView([56.946285, 24.105078], 13);

            // Pievienot tile layer
            L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
                attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors',
                maxZoom: 18
            }).addTo(mymap);

            var marker;
            var clickMarkers = [];

            console.log("Karte izveidota!");

            // Funkcija marķiera iestatīšanai
            function setMarker(lat, lon, popupText) {
                console.log("Iestatām marķieri:", lat, lon, popupText);

                if (marker) {
                    mymap.removeLayer(marker);
                }

                marker = L.marker([lat, lon], {
                    draggable: true
                }).addTo(mymap);

                if (popupText && popupText.trim() !== '') {
                    marker.bindPopup(`<div class="custom-popup">${popupText}</div>`).openPopup();
                }

                mymap.setView([lat, lon], mymap.getZoom());

                marker.on('dragend', function(e) {
                    var newPos = e.target.getLatLng();
                    console.log('Marķieris pārvietots:', newPos.lat, newPos.lng);
                });
            }

            // Funkcija koordinātu iegūšanai
            function getCenterCoordinates() {
                var center = mymap.getCenter();
                return {
                    lat: center.lat.toFixed(6),
                    lon: center.lng.toFixed(6)
                };
            }

            // Kartes klikšķa apstrāde
            mymap.on('click', function(e) {
                var lat = e.latlng.lat.toFixed(6);
                var lon = e.latlng.lng.toFixed(6);

                console.log('Kartes klikšķis:', lat, lon);

                // Noņemt iepriekšējos klikšķu marķierus
                clickMarkers.forEach(function(m) { mymap.removeLayer(m); });
                clickMarkers = [];

                var clickMarker = L.circleMarker([lat, lon], {
                    color: 'red',
                    fillColor: '#f03',
                    fillOpacity: 0.5,
                    radius: 5
                }).addTo(mymap);

                clickMarkers.push(clickMarker);

                // Nosūtīt koordinātes uz Python caur QWebChannel
                if (mapBridge) {
                    mapBridge.handleMapClick(lat, lon);
                    console.log('Koordinātes nosūtītas uz Python:', lat, lon);
                } else {
                    console.log('mapBridge nav pieejams!');
                }
            });

            // Pārbaudīt vai karte ielādējās
            mymap.whenReady(function() {
                console.log('Karte gatava!');
            });

            console.log("JavaScript ielādēts!");
        </script>
    </body>
    </html>
        """
        self.web_view.setHtml(map_html_content)

        v.addWidget(self.web_view)

        map_controls_layout = QHBoxLayout()
        self.map_lat_input = QLineEdit()
        self.map_lat_input.setPlaceholderText("Platums (Latitude)")
        self.map_lon_input = QLineEdit()
        self.map_lon_input.setPlaceholderText("Garums (Longitude)")
        btn_set_marker = QPushButton("Iestatīt marķieri")
        btn_set_marker.clicked.connect(self._set_map_marker_from_inputs)
        btn_get_location = QPushButton("Iegūt atrašanās vietu")
        btn_get_location.clicked.connect(self._get_location_from_map)
        btn_set_vieta = QPushButton("Iestatīt 'Vieta' lauku")
        btn_set_vieta.clicked.connect(self._set_vieta_from_map)

        map_controls_layout.addWidget(QLabel("Lat:"))
        map_controls_layout.addWidget(self.map_lat_input)
        map_controls_layout.addWidget(QLabel("Lon:"))
        map_controls_layout.addWidget(self.map_lon_input)
        map_controls_layout.addWidget(btn_set_marker)
        map_controls_layout.addWidget(btn_get_location)
        map_controls_layout.addWidget(btn_set_vieta)

        v.addLayout(map_controls_layout)
        w.setLayout(v)
        self.tabs.addTab(w, "Karte")
        self.tab_kartes = w

        # Initial load of map with current 'Vieta' if possible
        self.tabs.currentChanged.connect(self._update_map_on_tab_change)

    def _create_text_block_input(self, input_widget, field_name):
        """
        Izveido logrīku ar ievades lauku un pogām teksta bloku pārvaldībai.
        """
        container_widget = QWidget()
        layout = QVBoxLayout(container_widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(input_widget)

        button_layout = QHBoxLayout()
        save_button = QPushButton("Saglabāt kā bloku")
        load_button = QPushButton("Ielādēt bloku")

        save_button.clicked.connect(lambda: self._save_text_block(input_widget, field_name))
        load_button.clicked.connect(lambda: self._load_text_block(input_widget, field_name))

        button_layout.addWidget(save_button)
        button_layout.addWidget(load_button)
        button_layout.addStretch()

        layout.addLayout(button_layout)
        return container_widget

    def _save_text_block(self, input_widget, field_name):
        """
        Saglabā pašreizējo tekstu kā jaunu bloku.
        """
        current_text = ""
        if isinstance(input_widget, QLineEdit):
            current_text = input_widget.text()
        elif isinstance(input_widget, QTextEdit):
            current_text = input_widget.toPlainText()

        if not current_text.strip():
            QMessageBox.warning(self, "Saglabāt bloku", "Ievades lauks ir tukšs. Lūdzu, ievadiet tekstu, ko saglabāt.")
            return

        block_name, ok = QInputDialog.getText(self, "Saglabāt teksta bloku", "Ievadiet bloka nosaukumu:")
        if ok and block_name:
            self.text_block_manager.add_block(field_name, block_name, current_text)
            QMessageBox.information(self, "Saglabāts", f"Teksta bloks '{block_name}' saglabāts laukam '{field_name}'.")
        elif ok:
            QMessageBox.warning(self, "Saglabāt bloku", "Bloka nosaukums nevar būt tukšs.")

    def _load_text_block(self, input_widget, field_name):
        """
        Atver dialogu, lai ielādētu vai pārvaldītu teksta blokus.
        """
        blocks = self.text_block_manager.get_blocks_for_field(field_name)
        if not blocks:
            QMessageBox.information(self, "Ielādēt bloku", "Nav saglabātu teksta bloku šim laukam.")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle(f"Pārvaldīt teksta blokus: {field_name}")
        dialog_layout = QVBoxLayout(dialog)

        block_list = QListWidget()
        for name in sorted(blocks.keys()):
            block_list.addItem(name)
        dialog_layout.addWidget(block_list)

        button_layout = QHBoxLayout()
        load_button = QPushButton("Ielādēt izvēlēto")
        delete_button = QPushButton("Dzēst izvēlēto")
        close_button = QPushButton("Aizvērt")

        load_button.clicked.connect(lambda: self._apply_selected_block(block_list, input_widget, field_name, dialog))
        delete_button.clicked.connect(lambda: self._delete_selected_block(block_list, field_name))
        close_button.clicked.connect(dialog.accept)

        button_layout.addWidget(load_button)
        button_layout.addWidget(delete_button)
        button_layout.addStretch()
        button_layout.addWidget(close_button)
        dialog_layout.addLayout(button_layout)

        dialog.exec()

    def _apply_selected_block(self, block_list, input_widget, field_name, dialog):
        """
        Ielādē izvēlēto bloku ievades laukā.
        """
        selected_item = block_list.currentItem()
        if selected_item:
            block_name = selected_item.text()
            content = self.text_block_manager.get_block_content(field_name, block_name)
            if isinstance(input_widget, QLineEdit):
                input_widget.setText(content)
            elif isinstance(input_widget, QTextEdit):
                input_widget.setPlainText(content)
            dialog.accept()
            QMessageBox.information(self, "Ielādēts", f"Teksta bloks '{block_name}' ielādēts.")
        else:
            QMessageBox.warning(self, "Ielādēt bloku", "Lūdzu, izvēlieties bloku, ko ielādēt.")

    def _delete_selected_block(self, block_list, field_name):
        """
        Dzēš izvēlēto bloku.
        """
        selected_item = block_list.currentItem()
        if selected_item:
            block_name = selected_item.text()
            reply = QMessageBox.question(self, "Dzēst bloku",
                                         f"Vai tiešām vēlaties dzēst teksta bloku '{block_name}'?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.text_block_manager.delete_block(field_name, block_name)
                block_list.takeItem(block_list.row(selected_item))
                QMessageBox.information(self, "Dzēsts", f"Teksta bloks '{block_name}' dzēsts.")
        else:
            QMessageBox.warning(self, "Dzēst bloku", "Lūdzu, izvēlieties bloku, ko dzēst.")

    def _update_map_on_tab_change(self, index):
        if self.tabs.tabText(index) == "Karte":
            current_vieta = self.in_vieta.text().strip()
            # Pārliecināmies, ka JavaScript ir ielādēts un gatavs
            # Varētu pievienot arī `page().loadFinished` signāla apstrādi, lai būtu drošāk
            if current_vieta:
                # Šeit varētu mēģināt ģeokodēt adresi, ja ir pieejams API
                # Vienkāršības labad, ja ir vieta, iestatām marķieri Rīgā ar vietas nosaukumu
                self.web_view.page().runJavaScript(f"setMarker(56.946285, 24.105078, '{current_vieta}');")
            else:
                # Ja vieta nav norādīta, iestatām marķieri Rīgā ar noklusējuma tekstu
                self.web_view.page().runJavaScript(f"setMarker(56.946285, 24.105078, 'Rīga');")

    

    def _search_address_on_map(self):
        q = (self.map_search_input.text() if hasattr(self, "map_search_input") else "").strip()
        self.map_search_results.clear()
        if not q:
            return
        try:
            # Photon: ātrs un parasti neliek 403 (OpenStreetMap Nominatim bieži bloķē bez User-Agent)
            r = requests.get(
                "https://photon.komoot.io/api/",
                # Photon dažās instancēs neatbalsta `lang` parametru (met 400). Tāpēc to nelietojam.
                params={"q": q, "limit": 10},
                headers={
                    "User-Agent": "AktaGenerators/1.0 (contact: edgars@kulinics.id.lv)",
                    "Accept-Language": "lv,en;q=0.8",
                },
                timeout=10
            )
            r.raise_for_status()
            data = r.json()
            feats = data.get("features", []) or []
            for f in feats:
                props = f.get("properties", {}) or {}
                geom = f.get("geometry", {}) or {}
                coords = geom.get("coordinates", None)
                if not coords or len(coords) < 2:
                    continue
                lon, lat = coords[0], coords[1]
                # Saudzīgi saliekam adresi
                name = props.get("name") or ""
                street = props.get("street") or ""
                housenumber = props.get("housenumber") or ""
                city = props.get("city") or props.get("town") or props.get("village") or ""
                country = props.get("country") or ""
                parts = []
                if name:
                    parts.append(name)
                addr_line = " ".join([p for p in [street, housenumber] if p]).strip()
                if addr_line:
                    parts.append(addr_line)
                if city:
                    parts.append(city)
                if country:
                    parts.append(country)
                label = ", ".join([p for p in parts if p]).strip()
                if not label:
                    label = q
                it = QListWidgetItem(label)
                it.setData(Qt.UserRole, {"lat": lat, "lon": lon, "label": label})
                self.map_search_results.addItem(it)
        except Exception as e:
            QMessageBox.warning(self, "Karte", f"Neizdevās meklēt adresi: {e}")

    def _apply_map_search_result(self, item: QListWidgetItem):
        if not item:
            return
        d = item.data(Qt.UserRole) or {}
        try:
            lat = float(d.get("lat"))
            lon = float(d.get("lon"))
        except Exception:
            return
        label = d.get("label", "")
        self.map_lat_input.setText(f"{lat:.6f}")
        self.map_lon_input.setText(f"{lon:.6f}")
        # Iestatām marķieri kartē
        js = f"setMarker({lat:.6f}, {lon:.6f}, {json.dumps(label, ensure_ascii=False)});"
        self.web_view.page().runJavaScript(js)
    def _set_map_marker_from_inputs(self):
        try:
            lat = float(self.map_lat_input.text())
            lon = float(self.map_lon_input.text())
            self.web_view.page().runJavaScript(f"setMarker({lat}, {lon}, 'Izvēlētā vieta');")
        except ValueError:
            QMessageBox.warning(self, "Kļūda", "Lūdzu, ievadiet derīgas platuma un garuma vērtības.")

    def _begin_address_pick(self, target_lineedit: QLineEdit):
        """
        Aktivizē adreses izvēli kartē konkrētam laukam (Pieņēmējs/Nodevējs).
        Lietotājs var arī rakstīt adresi ar roku – šī ir tikai ērtība.
        """
        self._map_address_target = target_lineedit
        try:
            if hasattr(self, "tab_kartes"):
                self.tabs.setCurrentWidget(self.tab_kartes)
        except Exception:
            pass
        QMessageBox.information(
            self,
            "Adrese no kartes",
            "Noklikšķini kartē uz nepieciešamās vietas – adrese tiks automātiski ielikta laukā."
        )

    def _handle_map_click(self, lat, lon):
        """
        Šī funkcija tiek izsaukta no MapUrlInterceptor, kad kartē tiek noklikšķināts.
        Tā veic reverso ģeokodēšanu un aizpilda adresi:
        - ja lietotājs izvēlējās adresi (Pieņēmējs/Nodevējs), tad aizpilda attiecīgo lauku,
        - citādi aizpilda akta "Vieta" lauku (kā līdz šim).
        """
        self.map_lat_input.setText(lat)
        self.map_lon_input.setText(lon)
        self._reverse_geocode_and_set_vieta(lat, lon)

    def _reverse_geocode_and_set_vieta(self, lat, lon):
        """Reversā ģeokodēšana: mēģina Nominatim; ja 403/429 vai kļūda – izmanto Photon fallback."""
        try:
            lat_s = str(lat).strip()
            lon_s = str(lon).strip()

            headers = {
                # Nominatim prasa identificējamu UA (ar kontaktu)
                "User-Agent": "AktaGenerators/1.0 (kulinics.id.lv; edgars@kulinics.id.lv)",
                "Accept-Language": "lv,en;q=0.8",
                "From": "edgars@kulinics.id.lv",
            }

            # 1) Nominatim
            try:
                resp = requests.get(
                    "https://nominatim.openstreetmap.org/reverse",
                    params={
                        "format": "json",
                        "lat": lat_s,
                        "lon": lon_s,
                        "zoom": 18,
                        "addressdetails": 1
                    },
                    headers=headers,
                    timeout=8
                )
                if resp.status_code not in (403, 429):
                    resp.raise_for_status()
                    data = resp.json() if resp.text else {}

                    if isinstance(data, dict) and data.get("display_name"):
                        address = data["display_name"]

                        target = getattr(self, "_map_address_target", None)
                        if target is not None:
                            target.setText(address)
                            self._map_address_target = None
                        else:
                            self.in_vieta.setText(address)
                        QMessageBox.information(self, "Adrese iegūta", f"Adrese: {address}")
                        return
            except Exception:
                # ignorējam un mēģinam Photon
                pass

            # 2) Photon fallback (bieži strādā, kad Nominatim bloķē)
            presp = requests.get(
                "https://photon.komoot.io/reverse",
                params={"lat": lat_s, "lon": lon_s},
                headers=headers,
                timeout=8
            )
            presp.raise_for_status()
            pdata = presp.json() if presp.text else {}
            feats = pdata.get("features") or []
            if feats:
                props = feats[0].get("properties") or {}
                parts = [
                    props.get("name"),
                    props.get("street"),
                    props.get("housenumber"),
                    props.get("city") or props.get("town") or props.get("village"),
                    props.get("country")
                ]
                address = ", ".join([p for p in parts if p])
                if address:
                    self.in_vieta.setText(address)
                    QMessageBox.information(self, "Adrese iegūta", f"Adrese: {address}")
                    return

            # Fallback uz koordinātēm, ja abi servisi neiedeva adresi
            self.in_vieta.setText(f"{lat_s}, {lon_s}")
            QMessageBox.warning(self, "Brīdinājums", "Neizdevās iegūt adresi. Iestatītas koordinātes.")
        except Exception as e:
            QMessageBox.warning(self, "Tīkla kļūda",
                                f"Neizdevās iegūt adresi (tīkla problēma vai API kļūda): {e}")

    def _get_location_from_map(self):
        # Šī funkcija tagad vienkārši iegūst kartes centra koordinātes un iestata tās ievades laukos.
        # Reversā ģeokodēšana notiks, kad lietotājs noklikšķinās uz "Iestatīt 'Vieta' lauku" vai kartē.
        self.web_view.page().runJavaScript("getCenterCoordinates();", self._process_map_center_coordinates)

    def _process_map_center_coordinates(self, result):
        if result and 'lat' in result and 'lon' in result:
            self.map_lat_input.setText(str(result['lat']))
            self.map_lon_input.setText(str(result['lon']))
            QMessageBox.information(self, "Kartes atrašanās vieta", f"Kartes centrs: Lat {result['lat']}, Lon {result['lon']}")
        else:
            QMessageBox.warning(self, "Kļūda", "Neizdevās iegūt kartes centra koordinātes.")


    
    def _add_reference_doc(self):
        """Pievieno atsauces dokumentus (PDF/DOCX/XLSX), kas tiks pievienoti PDF beigās kā atvasinājumi."""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Pievienot atsauces dokumentus",
            "",
            "Dokumenti (*.pdf *.docx *.doc *.xlsx *.xls *.odt *.ods *.pptx *.ppt);;Visi faili (*.*)"
        )
        if not files:
            return

        existing = set()
        for i in range(self.list_atsauces_faili.count()):
            it = self.list_atsauces_faili.item(i)
            existing.add(it.data(Qt.UserRole))

        for p in files:
            if not p or p in existing:
                continue
            name = os.path.basename(p)
            it = QListWidgetItem(name)
            it.setData(Qt.UserRole, p)
            self.list_atsauces_faili.addItem(it)

        self._update_preview()

    def _remove_reference_doc(self):
        it = self.list_atsauces_faili.currentItem()
        if not it:
            return
        row = self.list_atsauces_faili.row(it)
        self.list_atsauces_faili.takeItem(row)
        self._update_preview()


    def _set_vieta_from_map(self):
        """
        Šī poga tagad izmanto pašreizējās ievades laukos esošās koordinātes,
        lai veiktu reversās ģeokodēšanas pieprasījumu un iestatītu 'Vieta' lauku.
        """
        lat_str = self.map_lat_input.text()
        lon_str = self.map_lon_input.text()
        if lat_str and lon_str:
            try:
                lat = float(lat_str)
                lon = float(lon_str)
                self._reverse_geocode_and_set_vieta(str(lat), str(lon)) # Pārsūtam kā string, jo API to sagaida
            except ValueError:
                QMessageBox.warning(self, "Kļūda", "Lūdzu, ievadiet derīgas platuma un garuma vērtības.")
        else:
            QMessageBox.warning(self, "Kļūda", "Lūdzu, vispirms iegūstiet koordinātes no kartes vai ievadiet tās manuāli.")

    # ----- Projekta saglabāšana/ielāde -----
    def ieviest_datus(self, d: AktaDati):
            # Sinhronizējam iekšējos datus ar ielādētajiem iestatījumiem
            self.data = d
            # Foto kolonna: lai UI vienmēr būtu iespējams pievienot foto
            try:
                self.data.show_item_photo_in_table = True
            except Exception:
                pass
            self.in_akta_nr.setText(d.akta_nr)
            self.in_datums.setDate(datetime.strptime(d.datums, '%Y-%m-%d').date())
            self.in_vieta.setText(d.vieta)
            self.in_pas_nr.setText(d.pasūtījuma_nr)
            self.in_piezimes.setText(d.piezīmes)

            self.in_liguma_nr.setText(d.līguma_nr)
            if d.izpildes_termiņš:
                self.in_izpildes_termins.setDate(datetime.strptime(d.izpildes_termiņš, '%Y-%m-%d').date())
            else:
                self.in_izpildes_termins.setDate(self.in_izpildes_termins.minimumDate()) # Clear date
            if d.pieņemšanas_datums:
                self.in_pieņemšanas_datums.setDate(datetime.strptime(d.pieņemšanas_datums, '%Y-%m-%d').date())
            else:
                self.in_pieņemšanas_datums.setDate(self.in_pieņemšanas_datums.minimumDate())
            if d.nodošanas_datums:
                self.in_nodošanas_datums.setDate(datetime.strptime(d.nodošanas_datums, '%Y-%m-%d').date())
            else:
                self.in_nodošanas_datums.setDate(self.in_nodošanas_datums.minimumDate())

            self.in_strīdu_risināšana.setText(d.strīdu_risināšana)
            self.ck_konfidencialitate.setChecked(d.konfidencialitātes_klauzula)
            self.in_soda_nauda_procenti.setValue(float(d.soda_nauda_procenti)) # Convert Decimal to float for QDoubleSpinBox
            self.in_piegades_nosacijumi.setText(d.piegādes_nosacījumi)
            self.ck_apdrošināšana.setChecked(d.apdrošināšana)
            if hasattr(self, 'in_apdrosinasana_teksts'):
                self.in_apdrosinasana_teksts.setText(getattr(d, 'apdrošināšana_teksts', '') or '')
            self.in_papildu_nosacijumi.setText(d.papildu_nosacījumi)
            self.in_atsauces_dokumenti.setText(d.atsauces_dokumenti)
            self.list_atsauces_faili.clear()
            for ref in getattr(d, 'atsauces_dokumenti_faili', []) or []:
                try:
                    if isinstance(ref, dict):
                        p = ref.get('ceļš', '')
                        n = ref.get('nosaukums', '') or os.path.basename(p)
                    else:
                        p = getattr(ref, 'ceļš', '')
                        n = getattr(ref, 'nosaukums', '') or os.path.basename(p)
                    if p:
                        it = QListWidgetItem(n)
                        it.setData(Qt.UserRole, p)
                        self.list_atsauces_faili.addItem(it)
                except Exception:
                    pass
            self.cb_akta_statuss.setCurrentText(d.akta_statuss)
            self.in_valuta.setCurrentText(d.valūta)
            self.ck_elektroniskais_paraksts.setChecked(d.elektroniskais_paraksts)
            self.ck_radit_elektronisko_parakstu_tekstu.setChecked(d.radit_elektronisko_parakstu_tekstu) # JAUNA RINDAS

            # Puses
            self.pie_in[0].setText(d.pieņēmējs.nosaukums)
            self.pie_in[1].setText(d.pieņēmējs.reģ_nr)
            self.pie_in[2].setText(d.pieņēmējs.adrese)
            self.pie_in[3].setText(d.pieņēmējs.kontaktpersona)
            self.pie_in[4].setText(getattr(d.pieņēmējs, 'amats', ''))
            self.pie_in[5].setCurrentText(getattr(d.pieņēmējs, 'pilnvaras_pamats', ''))
            self.pie_in[6].setText(d.pieņēmējs.tālrunis)
            self.pie_in[7].setText(d.pieņēmējs.epasts)
            self.pie_in[8].setText(getattr(d.pieņēmējs, 'web_lapa', ''))
            self.pie_in[9].setText(d.pieņēmējs.bankas_konts)
            self.pie_in[10].setCurrentText(d.pieņēmējs.juridiskais_statuss)


            self.nod_in[0].setText(d.nodevējs.nosaukums)
            self.nod_in[1].setText(d.nodevējs.reģ_nr)
            self.nod_in[2].setText(d.nodevējs.adrese)
            self.nod_in[3].setText(d.nodevējs.kontaktpersona)
            self.nod_in[4].setText(getattr(d.nodevējs, 'amats', ''))
            self.nod_in[5].setCurrentText(getattr(d.nodevējs, 'pilnvaras_pamats', ''))
            self.nod_in[6].setText(d.nodevējs.tālrunis)
            self.nod_in[7].setText(d.nodevējs.epasts)
            self.nod_in[8].setText(getattr(d.nodevējs, 'web_lapa', ''))
            self.nod_in[9].setText(d.nodevējs.bankas_konts)
            self.nod_in[10].setCurrentText(d.nodevējs.juridiskais_statuss)


# Pozīcijas
            self.tab.setRowCount(0)

            # JAUNS: vienmēr būvējam pilno kolonnu komplektu (paslēpšana notiek ar setColumnHidden)
            try:
                self.data.poz_columns_config = _merge_poz_columns_config(getattr(d, "poz_columns_config", None))
            except Exception:
                self.data.poz_columns_config = _merge_poz_columns_config({})

            try:
                if hasattr(self, "ck_show_price_summary") and self.ck_show_price_summary:
                    self.ck_show_price_summary.setChecked(bool(getattr(d, "show_price_summary", True)))
            except Exception:
                pass

            base_headers = [
                ("apraksts", "Apraksts"),
                ("daudzums", "Daudzums"),
                ("vieniba", "Vienība"),
                ("cena", "Cena"),
                ("summa", "Summa"),
                ("serial", "Seriālais Nr."),
                ("warranty", "Garantija"),
                ("notes", "Piezīmes pozīcijai"),
            ]

            headers = []
            for key, default_title in base_headers:
                try:
                    headers.append(str(self.data.poz_columns_config.get(key, {}).get("title", default_title)))
                except Exception:
                    headers.append(default_title)

            # Pielāgotās kolonnas
            self.data.custom_columns = getattr(d, "custom_columns", []) or []
            for col in self.data.custom_columns:
                try:
                    if isinstance(col, dict) and "visible" not in col:
                        col["visible"] = True
                except Exception:
                    pass
                try:
                    headers.append(str(col.get("name", "")) if isinstance(col, dict) else "")
                except Exception:
                    headers.append("")

            # Foto vienmēr pēdējā
            try:
                headers.append(str(self.data.poz_columns_config.get("foto", {}).get("title", "Foto")))
            except Exception:
                headers.append("Foto")

            self.tab.setColumnCount(len(headers))
            self.tab.setHorizontalHeaderLabels(headers)

            for i in range(len(headers)):
                self.tab.horizontalHeader().setSectionResizeMode(i, QHeaderView.Interactive)

            # JAUNS: pielietojam saglabāto kolonnu secību (UI) un platumus, ja tādi ir
            try:
                self._poz_apply_stored_visual_order()
            except Exception:
                pass
            try:
                st_b64 = str(getattr(d, 'poz_header_state_b64', '') or '')
                if st_b64:
                    self.tab.horizontalHeader().restoreState(base64.b64decode(st_b64.encode('ascii')))
            except Exception:
                pass

            self._poz_apply_column_visibility()

            for p in d.pozīcijas:
                r = self.tab.rowCount();
                self.tab.insertRow(r)
                self.tab.setItem(r, 0, QTableWidgetItem(p.apraksts))
                self.tab.setItem(r, 1, QTableWidgetItem(str(p.daudzums)))
                self.tab.setItem(r, 2, QTableWidgetItem(p.vienība))
                self.tab.setItem(r, 3, QTableWidgetItem(str(p.cena)))
                self.tab.setItem(r, 4, QTableWidgetItem(formēt_naudu(p.summa)))
                # Set values for potentially hidden columns
                idx = self._poz_col_indices()

                # Pielāgotās kolonnas (ja ir dati)
                try:
                    for col_idx, coldef in enumerate(getattr(self.data, "custom_columns", []) or []):
                        ui_col = idx.get("custom_start", 8) + col_idx
                        if ui_col >= self.tab.columnCount():
                            continue
                        val = ""
                        if isinstance(coldef, dict):
                            data_list = coldef.get("data", [])
                            if isinstance(data_list, list) and r < len(data_list):
                                val = str(data_list[r])
                        self.tab.setItem(r, ui_col, QTableWidgetItem(val))
                except Exception:
                    pass

                # Foto (ja ieslēgts)
                if "foto" in idx and self.tab.columnCount() > idx["foto"]:
                    foto_path = getattr(p, "attēla_ceļš", "") or ""
                    itf = QTableWidgetItem(foto_path)
                    itf.setFlags(itf.flags() & ~Qt.ItemIsEditable)
                    self.tab.setItem(r, idx["foto"], itf)
                    self._ensure_photo_cell(r)
                    w = self.tab.cellWidget(r, idx["foto"])
                    if isinstance(w, QToolButton):
                        if foto_path:
                            w.setText("Mainīt")
                            try:
                                ico = QIcon(foto_path)
                                if not ico.isNull():
                                    w.setIcon(ico)
                                    w.setIconSize(QSize(18, 18))
                            except Exception:
                                pass

                # Seriālais / Garantija / Piezīmes
                if "serial" in idx and self.tab.columnCount() > idx["serial"]:
                    self.tab.setItem(r, idx["serial"], QTableWidgetItem(p.seriālais_nr))
                if "warranty" in idx and self.tab.columnCount() > idx["warranty"]:
                    self.tab.setItem(r, idx["warranty"], QTableWidgetItem(p.garantija))
                if "notes" in idx and self.tab.columnCount() > idx["notes"]:
                    self.tab.setItem(r, idx["notes"], QTableWidgetItem(p.piezīmes_pozīcijai))



            # Attēli
            if hasattr(self, "photos_table") and self.photos_table is not None:
                self.photos_table.blockSignals(True)
                self.photos_table.setRowCount(0)
                for a in d.attēli:
                    self._photos_add_row(a.ceļš, a.paraksts)
                self.photos_table.blockSignals(False)
            else:
                # Back-compat
                if getattr(self, "img_list", None) is not None:
                    self.img_list.clear()
                    for a in d.attēli:
                        it = QListWidgetItem(os.path.basename(a.ceļš))
                        it.setData(Qt.UserRole, {"ceļš": a.ceļš, "paraksts": a.paraksts})
                        self.img_list.addItem(it)

            # Iestatījumi
            self.ck_pvn.setChecked(d.iekļaut_pvn)
            self.in_pvn.setValue(float(d.pvn_likme)) # Convert Decimal to float for QDoubleSpinBox
            self.ck_paraksti.setChecked(d.parakstu_rindas)
            self.in_logo.setText(d.logotipa_ceļš)
            self.in_fonts.setText(d.fonts_ceļš)
            self.in_paraksts_pie.setText(d.paraksts_pieņēmējs_ceļš)
            self.in_paraksts_nod.setText(d.paraksts_nodevējs_ceļš)

            # Papildu iestatījumi
            self.cb_page_size.setCurrentText(d.pdf_page_size)
            self.cb_page_orientation.setCurrentText(d.pdf_page_orientation)
            self.in_margin_left.setValue(float(d.pdf_margin_left)) # Convert Decimal to float for QDoubleSpinBox
            self.in_margin_right.setValue(float(d.pdf_margin_right)) # Convert Decimal to float for QDoubleSpinBox
            self.in_margin_top.setValue(float(d.pdf_margin_top)) # Convert Decimal to float for QDoubleSpinBox
            self.in_margin_bottom.setValue(float(d.pdf_margin_bottom)) # Convert Decimal to float for QDoubleSpinBox
            self.in_font_size_head.setValue(d.pdf_font_size_head)
            self.in_font_size_normal.setValue(d.pdf_font_size_normal)
            self.in_font_size_small.setValue(d.pdf_font_size_small)
            self.in_font_size_table.setValue(d.pdf_font_size_table)
            self.in_logo_width_mm.setValue(float(d.pdf_logo_width_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_signature_width_mm.setValue(float(d.pdf_signature_width_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_signature_height_mm.setValue(float(d.pdf_signature_height_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_docx_image_width_inches.setValue(float(d.docx_image_width_inches)) # Convert Decimal to float for QDoubleSpinBox
            self.in_docx_signature_width_inches.setValue(float(d.docx_signature_width_inches)) # Convert Decimal to float for QDoubleSpinBox
            self.in_table_col_widths.setText(d.table_col_widths)
            self.ck_auto_generate_akta_nr.setChecked(d.auto_generate_akta_nr)
            self.in_default_execution_days.setValue(int(getattr(d, 'default_execution_days', 5)))
            self.in_default_currency.setText(d.default_currency)
            self.in_default_unit.setText(d.default_unit)
            self.in_default_pvn_rate.setValue(float(d.default_pvn_rate)) # Convert Decimal to float for QDoubleSpinBox
            self.in_poppler_path.setText(d.poppler_path)
            self.poppler_path = d.poppler_path

            # Set new settings values
            self.in_header_text_color.setText(d.header_text_color)
            self.in_footer_text_color.setText(d.footer_text_color)
            self.in_table_header_bg_color.setText(d.table_header_bg_color)
            self.in_table_grid_color.setText(d.table_grid_color)
            self.in_table_row_spacing.setValue(float(d.table_row_spacing)) # Convert Decimal to float for QDoubleSpinBox
            self.in_line_spacing_multiplier.setValue(float(d.line_spacing_multiplier)) # Convert Decimal to float for QDoubleSpinBox
            self.ck_show_page_numbers.setChecked(d.show_page_numbers)
            self.ck_show_generation_timestamp.setChecked(d.show_generation_timestamp)
            self.cb_currency_symbol_position.setCurrentText(d.currency_symbol_position)
            self.in_date_format.setText(d.date_format)
            self.in_signature_line_length_mm.setValue(float(d.signature_line_length_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_signature_line_thickness_pt.setValue(float(d.signature_line_thickness_pt)) # Convert Decimal to float for QDoubleSpinBox
            self.ck_add_cover_page.setChecked(d.add_cover_page)
            self.in_cover_page_title.setText(d.cover_page_title)
            self.in_cover_page_logo_width_mm.setValue(float(d.cover_page_logo_width_mm)) # Convert Decimal to float for QDoubleSpinBox
            # Individuālais QR kods
            self.ck_include_custom_qr_code.setChecked(d.include_custom_qr_code)
            self.in_custom_qr_code_data.setText(d.custom_qr_code_data)
            self.in_custom_qr_code_size_mm.setValue(float(d.custom_qr_code_size_mm))
            self.cb_custom_qr_code_position.setCurrentText(d.custom_qr_code_position)
            self.in_custom_qr_code_pos_x_mm.setValue(float(d.custom_qr_code_pos_x_mm))
            self.in_custom_qr_code_pos_y_mm.setValue(float(d.custom_qr_code_pos_y_mm))
            self.in_custom_qr_code_color.setText(d.custom_qr_code_color)

            # Automātiskais QR kods (akta ID)
            self.ck_include_auto_qr_code.setChecked(d.include_auto_qr_code)
            self.in_auto_qr_code_size_mm.setValue(float(d.auto_qr_code_size_mm))
            self.cb_auto_qr_code_position.setCurrentText(d.auto_qr_code_position)
            self.in_auto_qr_code_pos_x_mm.setValue(float(d.auto_qr_code_pos_x_mm))
            self.in_auto_qr_code_pos_y_mm.setValue(float(d.auto_qr_code_pos_y_mm))
            self.in_auto_qr_code_color.setText(d.auto_qr_code_color)

            self.ck_add_watermark.setChecked(d.add_watermark)
            self.in_watermark_text.setText(d.watermark_text)
            self.in_watermark_font_size.setValue(d.watermark_font_size)
            self.in_watermark_color.setText(d.watermark_color)
            self.in_watermark_rotation.setValue(d.watermark_rotation)
            self.ck_enable_pdf_encryption.setChecked(d.enable_pdf_encryption)
            self.in_pdf_user_password.setText(d.pdf_user_password)
            self.in_pdf_owner_password.setText(d.pdf_owner_password)
            # --- JAUNS: sinhronizē arī Pamata datu šifrēšanas lauku ---
            try:
                if hasattr(self, "ck_pdf_encrypt_basic") and hasattr(self, "in_pdf_password_basic"):
                    self.ck_pdf_encrypt_basic.setChecked(bool(getattr(d, "enable_pdf_encryption", False)))
                    self.in_pdf_password_basic.setText(str(getattr(d, "pdf_user_password", "") or ""))
            except Exception:
                pass

            self.ck_allow_printing.setChecked(d.allow_printing)
            self.ck_allow_copying.setChecked(d.allow_copying)
            self.ck_allow_modifying.setChecked(d.allow_modifying)
            self.ck_allow_annotating.setChecked(d.allow_annotating)
            self.in_default_country.setText(d.default_country)
            self.in_default_city.setText(d.default_city)
            self.ck_show_contact_details_in_header.setChecked(d.show_contact_details_in_header)
            self.in_contact_details_header_font_size.setValue(d.contact_details_header_font_size)
            self.in_item_image_width_mm.setValue(float(d.item_image_width_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_item_image_caption_font_size.setValue(d.item_image_caption_font_size)
            self.ck_show_item_notes_in_table.setChecked(d.show_item_notes_in_table)
            self.ck_show_item_serial_number_in_table.setChecked(d.show_item_serial_number_in_table)
            self.ck_show_item_warranty_in_table.setChecked(d.show_item_warranty_in_table)
            self.in_table_cell_padding_mm.setValue(float(d.table_cell_padding_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.cb_table_header_font_style.setCurrentText(d.table_header_font_style)
            self.cb_table_content_alignment.setCurrentText(d.table_content_alignment)
            self.in_signature_font_size.setValue(d.signature_font_size)
            self.in_signature_spacing_mm.setValue(float(d.signature_spacing_mm)) # Convert Decimal to float for QDoubleSpinBox
            self.in_document_title_font_size.setValue(d.document_title_font_size)
            self.in_document_title_color.setText(d.document_title_color)
            self.in_section_heading_font_size.setValue(d.section_heading_font_size)
            self.in_section_heading_color.setText(d.section_heading_color)
            self.in_paragraph_line_spacing_multiplier.setValue(float(d.paragraph_line_spacing_multiplier)) # Convert Decimal to float for QDoubleSpinBox
            self.cb_table_border_style.setCurrentText(d.table_border_style)
            self.in_table_border_thickness_pt.setValue(float(d.table_border_thickness_pt)) # Convert Decimal to float for QDoubleSpinBox
            self.in_table_alternate_row_color.setText(d.table_alternate_row_color)
            self.ck_show_total_sum_in_words.setChecked(d.show_total_sum_in_words)
            self.cb_total_sum_in_words_language.setCurrentText(d.total_sum_in_words_language)
            self.cb_default_vat_calculation_method.setCurrentText(d.default_vat_calculation_method)
            self.ck_show_vat_breakdown.setChecked(d.show_vat_breakdown)
            self.ck_enable_digital_signature_field.setChecked(d.enable_digital_signature_field)
            self.in_digital_signature_field_name.setText(d.digital_signature_field_name)
            self.in_digital_signature_field_size_mm.setValue(
                float(d.digital_signature_field_size_mm))  # Convert Decimal to float for QDoubleSpinBox
            self.cb_digital_signature_field_position.setCurrentText(d.digital_signature_field_position)

            # JAUNA RINDAS
            self.in_templates_dir.setText(d.templates_dir)
            # Pārliecināmies, ka direktorijs eksistē, ja tas ir ielādēts
            if d.templates_dir:
                os.makedirs(d.templates_dir, exist_ok=True)

            self._update_preview()

    def saglabat_projektu(self):
        self._undo_mgr.push_undo(self._snapshot_state('PROJECT_SAVE'))
        self._audit('PROJECT_SAVE', {})
        d = self.savākt_datus()
        default_filename = f"Akts_{drošs_faila_nosaukums(d.akta_nr) or 'akts'}.json"
        path, _ = QFileDialog.getSaveFileName(self, "Saglabāt projektu", os.path.join(PROJECT_SAVE_DIR, default_filename), "JSON (*.json)")
        if not path:
            return
        out = asdict(d)
        # Convert Decimal fields to string for JSON serialization
        for key, value in out.items():
            if isinstance(value, Decimal):
                out[key] = str(value)
        # Handle nested Decimal fields in Pozīcija
        for p in out['pozīcijas']:
            for key, value in p.items():
                if isinstance(value, Decimal):
                    p[key] = str(value)
        # Handle nested Persona objects
        out['pieņēmējs'] = asdict(d.pieņēmējs)
        out['nodevējs'] = asdict(d.nodevējs)

        with open(path, 'w', encoding='utf-8') as f:
            json.dump(out, f, ensure_ascii=False, indent=2)
        self._ceļš_projekts = path
        self._add_to_history(path) # Pievienojam projektu vēsturei
        QMessageBox.information(self, "Saglabāts", "Projekts saglabāts veiksmīgi.")


    def saglabat_ka_sablonu(self):
        """
        Saglabā pašreizējo akta konfigurāciju kā šablonu.
        """
        d = self.savākt_datus()

        # Pieprasām šablona nosaukumu no lietotāja
        template_name, ok = QInputDialog.getText(self, "Saglabāt kā šablonu", "Ievadiet šablona nosaukumu:")
        if not ok or not template_name:
            QMessageBox.warning(self, "Saglabāt kā šablonu", "Šablona nosaukums nevar būt tukšs.")
            return

        # Pieprasām paroli šablonam (neobligāti)
        template_password, ok_pass = QInputDialog.getText(self, "Šablona parole",
                                                          "Ievadiet paroli šablonam (atstājiet tukšu, ja nevēlaties paroli):",
                                                          QLineEdit.Password)
        if not ok_pass:
            return  # Lietotājs atcēla paroles ievadi

        d.template_password = template_password  # Saglabājam paroli datu objektā

        # Izveidojam drošu faila nosaukumu
        safe_template_name = drošs_faila_nosaukums(template_name)
        # Izmantojam pašreizējo šablonu direktoriju no AktaDati
        current_templates_dir = self.data.templates_dir
        if not current_templates_dir:  # Ja vēl nav iestatīts (piemēram, pirmajā startā)
            current_templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")
            os.makedirs(current_templates_dir, exist_ok=True)  # Pārliecināmies, ka noklusējuma mape eksistē

        template_file_path = os.path.join(current_templates_dir, f"{safe_template_name}.json")

        # Pārveidojam AktaDati objektu vārdnīcā, lai to varētu serializēt uz JSON
        out = asdict(d)

        # Konvertējam Decimal vērtības uz string, lai tās varētu saglabāt JSON
        for key, value in out.items():
            if isinstance(value, Decimal):
                out[key] = str(value)
        # Apstrādājam ligzdotās Decimal vērtības Pozīcija objektos
        for p in out.get('pozīcijas', []): # Izmantojam .get, lai izvairītos no kļūdām, ja pozīcijas nav
            for key, value in p.items():
                if isinstance(value, Decimal):
                    p[key] = str(value)

        # Apstrādājam Persona objektus
        out['pieņēmējs'] = asdict(d.pieņēmējs)
        out['nodevējs'] = asdict(d.nodevējs)

        # Izdzēšam datus, kas nav jāglabā šablonā (piemēram, specifiskus akta datus, attēlus, paroles)
        # Šablons ir paredzēts kā bāzes konfigurācija, nevis konkrēta akta kopija.
        out['akta_nr'] = ""
        out['datums'] = datetime.now().strftime('%Y-%m-%d') # Atjaunojam datumu uz pašreizējo
        out['pasūtījuma_nr'] = ""
        out['līguma_nr'] = ""
        out['izpildes_termiņš'] = ""
        out['pieņemšanas_datums'] = ""
        out['nodošanas_datums'] = ""
        out['piezīmes'] = ""
        out['attēli'] = [] # Šablonā nav jābūt attēliem
        out['poppler_path'] = "" # Poppler ceļš ir sistēmas iestatījums, nevis šablona daļa
        out['pdf_user_password'] = ""
        out['pdf_owner_password'] = ""
        out['custom_qr_code_data'] = ""  # QR koda dati nav šablona daļa
        out['custom_qr_code_pos_x_mm'] = str(Decimal("0"))
        out['custom_qr_code_pos_y_mm'] = str(Decimal("0"))
        out['custom_qr_code_color'] = "#000000"
        out['auto_qr_code_pos_x_mm'] = str(Decimal("0"))
        out['auto_qr_code_pos_y_mm'] = str(Decimal("0"))
        out['auto_qr_code_color'] = "#000000"
        out['template_password'] = template_password

        # Saglabājam šablonu JSON failā
        try:
            with open(template_file_path, 'w', encoding='utf-8') as f:
                json.dump(out, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "Saglabāts", f"Šablons '{template_name}' veiksmīgi saglabāts.")
            self._update_sablonu_list() # Atjaunojam šablonu sarakstu
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās saglabāt šablonu:\n{e}")

    def _update_sablonu_list(self):
        """
        Atjauno šablonu sarakstu GUI.
        """
        self.sablonu_list.clear()
        # Pievienojam noklusējuma šablonu
        self.sablonu_list.addItem("Pappus dati (piemērs)")
        # Pievienojam visus saglabātos šablonus
        if os.path.exists(TEMPLATES_DIR):
            for filename in sorted(os.listdir(TEMPLATES_DIR)):
                if not filename.lower().endswith('.json'):
                    continue
                display_name = os.path.splitext(filename)[0]
                self.sablonu_list.addItem(display_name)

    def _resolve_template_path(self, template_display_name: str) -> Optional[str]:
        """
        Atrod šablona .json failu neatkarīgi no tā, kurā mapē tas glabājas.
        Meklē gan pašreizējo templates_dir, gan noklusējuma APP_DATA_DIR/AktaGenerators_Templates.
        """
        name = (template_display_name or "").strip()
        # noņemam marķierus no saraksta (ja tādi ir)
        for marker in ["[Aizsargāts]", "[Kļūda]"]:
            name = name.replace(marker, "").strip()

        # Kandidātu mapes (pirmā – lietotāja iestatītā)
        candidates = []
        if getattr(self.data, "templates_dir", ""):
            candidates.append(self.data.templates_dir)
        # noklusējuma mape
        candidates.append(os.path.join(APP_DATA_DIR, "AktaGenerators_Templates"))
        # drošībai: vecā/alternatīvā (ja kādreiz mainīts)
        candidates.append(os.path.join(os.path.expanduser("~"), "Documents", "AktaGenerators_Templates"))
        # noņemam dublikātus
        uniq = []
        for d in candidates:
            if d and d not in uniq:
                uniq.append(d)

        # mēģinām tiešu nosaukumu
        file_variants = [
            f"{name}.json",
            f"{drošs_faila_nosaukums(name)}.json",
        ]

        for d in uniq:
            try:
                os.makedirs(d, exist_ok=True)
            except Exception:
                pass
            for fn in file_variants:
                p = os.path.join(d, fn)
                if os.path.exists(p):
                    # ja atradām citā mapē nekā iestatīts, sinhronizējam
                    if getattr(self.data, "templates_dir", "") != d:
                        self.data.templates_dir = d
                    return p

        # Ja nav atrasts, mēģinām "fuzzy" – ignorējam atstarpes/_/- un reģistru
        def norm(s: str) -> str:
            return re.sub(r"[\s_\-]+", "", s).lower()

        wanted = norm(name)
        for d in uniq:
            if not d or not os.path.exists(d):
                continue
            for fn in os.listdir(d):
                if not fn.lower().endswith(".json"):
                    continue
                base = os.path.splitext(fn)[0]
                if norm(base) == wanted:
                    p = os.path.join(d, fn)
                    if getattr(self.data, "templates_dir", "") != d:
                        self.data.templates_dir = d
                    return p
        return None


    def dzest_sablonus(self):
        """
        Dzēš atlasītos šablonus.
        """
        selected_items = self.sablonu_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Dzēst šablonus", "Lūdzu, atlasiet vismaz vienu šablonu, ko dzēst.")
            return

        # Iegūstam tīrus šablonu nosaukumus, noņemot marķierus
        templates_to_delete_clean = []
        for item in selected_items:
            clean_name = item.text().replace(" [Aizsargāts]", "").replace(" [Kļūda]", "")
            if clean_name != "Testa dati (piemērs)":  # Neļaujam dzēst iebūvēto šablonu
                templates_to_delete_clean.append(clean_name)

        if not templates_to_delete_clean:
            QMessageBox.information(self, "Dzēst šablonus",
                                    "Nav atlasīts neviens dzēšams šablons (iebūvēto šablonu nevar dzēst).")
            return

        reply = QMessageBox.question(self, "Dzēst šablonus",
                                     f"Vai tiešām vēlaties dzēst šādus šablonus?\n\n{', '.join(templates_to_delete_clean)}",
                                     QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.Yes:
            deleted_count = 0
            current_templates_dir = self.data.templates_dir
            if not current_templates_dir:
                current_templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")

            for template_name in templates_to_delete_clean:
                file_path = os.path.join(current_templates_dir, f"{template_name}.json")
                if os.path.exists(file_path):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        stored_password = data.get('template_password', '')

                        if stored_password:
                            entered_password, ok_pass = QInputDialog.getText(self, "Ievadiet paroli",
                                                                             f"Šablonam '{template_name}' ir parole. Lūdzu, ievadiet to, lai dzēstu:",
                                                                             QLineEdit.Password)
                            if not ok_pass or entered_password != stored_password:
                                QMessageBox.warning(self, "Nepareiza parole",
                                                    f"Nepareiza parole šablonam '{template_name}'. Dzēšana atcelta.")
                                continue  # Pārejam pie nākamā šablona

                        os.remove(file_path)
                        deleted_count += 1
                    except Exception as e:
                        QMessageBox.critical(self, "Kļūda dzēšot", f"Neizdevās dzēst šablonu '{template_name}':\n{e}")
                else:
                    QMessageBox.warning(self, "Dzēst šablonus", f"Šablona fails '{template_name}.json' nav atrasts.")

            if deleted_count > 0:
                QMessageBox.information(self, "Dzēsts", f"Veiksmīgi dzēsti {deleted_count} šabloni.")
                self._update_sablonu_list()  # Atjaunojam sarakstu pēc dzēšanas
            else:
                QMessageBox.information(self, "Dzēst šablonus", "Neviens šablons netika dzēsts.")


    def _update_sablonu_list(self):
        """
        Atjauno šablonu sarakstu GUI, pievienojot marķieri aizsargātiem šabloniem.
        """
        self.sablonu_list.clear()
        # Pievienojam noklusējuma šablonu
        self.sablonu_list.addItem("Testa dati (piemērs)")

        # Izmantojam pašreizējo šablonu direktoriju no AktaDati
        current_templates_dir = self.data.templates_dir
        if not current_templates_dir:  # Ja vēl nav iestatīts (piemēram, pirmajā startā)
            current_templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")
            os.makedirs(current_templates_dir, exist_ok=True)  # Pārliecināmies, ka noklusējuma mape eksistē

        if os.path.exists(current_templates_dir):
            for filename in os.listdir(current_templates_dir):
                if filename.endswith(".json"):
                    template_name = os.path.splitext(filename)[0]
                    file_path = os.path.join(current_templates_dir, filename)
                    item = QListWidgetItem(template_name)  # Izveidojam QListWidgetItem ar tīru nosaukumu
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        if data.get('template_password'):
                            item.setBackground(QColor("#F54927"))  # Gaiši oranžs fons aizsargātiem šabloniem
                            item.setToolTip("Šablons ir aizsargāts ar paroli")  # Pievienojam tooltip
                        self.sablonu_list.addItem(item)
                    except Exception as e:
                        print(f"Kļūda lasot šablona failu {filename}: {e}")
                        item.setBackground(QColor("#FFCCCC"))  # Gaiši sarkans fons kļūdainiem šabloniem
                        item.setToolTip(f"Kļūda ielādējot šablonu: {e}")
                        self.sablonu_list.addItem(item)


    def mainit_sablonu_paroli(self):
        """
        Maina vai pievieno paroli izvēlētajam šablonam.
        """
        selected_items = self.sablonu_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Mainīt paroli", "Lūdzu, atlasiet šablonu, kuram vēlaties mainīt paroli.")
            return
        if len(selected_items) > 1:
            QMessageBox.warning(self, "Mainīt paroli", "Lūdzu, atlasiet tikai vienu šablonu.")
            return

        item_text = selected_items[0].text()
        template_name = item_text.replace(" [Aizsargāts]", "").replace(" [Kļūda]", "")

        if template_name == "Testa dati (piemērs)":
            QMessageBox.information(self, "Mainīt paroli",
                                    "Iebūvētajam šablonam 'Testa dati (piemērs)' nevar mainīt paroli.")
            return

        current_templates_dir = self.data.templates_dir
        if not current_templates_dir:
            current_templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")

        file_path = os.path.join(current_templates_dir, f"{template_name}.json")
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Mainīt paroli", f"Šablona fails '{template_name}.json' nav atrasts.")
            return

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            current_password = data.get('template_password', '')

            if current_password:
                # Šablonam jau ir parole, pieprasām veco paroli
                old_password, ok_old = QInputDialog.getText(self, "Mainīt paroli",
                                                             f"Šablonam '{template_name}' ir parole. Lūdzu, ievadiet veco paroli:",
                                                             QLineEdit.Password)
                if not ok_old or old_password != current_password:
                    QMessageBox.warning(self, "Nepareiza parole", "Ievadītā vecā parole ir nepareiza vai ievade atcelta.")
                    return

            # Pieprasām jauno paroli
            new_password, ok_new = QInputDialog.getText(self, "Mainīt paroli",
                                                         f"Ievadiet jauno paroli šablonam '{template_name}' (atstājiet tukšu, lai noņemtu paroli):",
                                                         QLineEdit.Password)
            if not ok_new:
                return # Lietotājs atcēla jaunas paroles ievadi

            data['template_password'] = new_password # Saglabājam jauno paroli (var būt tukša)

            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            QMessageBox.information(self, "Parole mainīta", f"Šablona '{template_name}' parole veiksmīgi mainīta.")
            self._update_sablonu_list() # Atjaunojam sarakstu, lai atspoguļotu izmaiņas
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās mainīt paroli šablonam '{template_name}':\n{e}")

    def nonemt_sablonu_paroli(self):
        """
        Noņem paroli no izvēlētā šablona.
        """
        selected_items = self.sablonu_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Noņemt paroli", "Lūdzu, atlasiet šablonu, kuram vēlaties noņemt paroli.")
            return
        if len(selected_items) > 1:
            QMessageBox.warning(self, "Noņemt paroli", "Lūdzu, atlasiet tikai vienu šablonu.")
            return

        item_text = selected_items[0].text()
        template_name = item_text.replace(" [Aizsargāts]", "").replace(" [Kļūda]", "")

        if template_name == "Testa dati (piemērs)":
            QMessageBox.information(self, "Noņemt paroli",
                                    "Iebūvētajam šablonam 'Testa dati (piemērs)' nav paroles, ko noņemt.")
            return

        current_templates_dir = self.data.templates_dir
        if not current_templates_dir:
            current_templates_dir = os.path.join(APP_DATA_DIR, "AktaGenerators_Templates")

        file_path = os.path.join(current_templates_dir, f"{template_name}.json")
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Noņemt paroli", f"Šablona fails '{template_name}.json' nav atrasts.")
            return

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            current_password = data.get('template_password', '')

            if not current_password:
                QMessageBox.information(self, "Noņemt paroli", f"Šablonam '{template_name}' jau nav paroles.")
                return

            # Pieprasām veco paroli, lai apstiprinātu noņemšanu
            old_password, ok_old = QInputDialog.getText(self, "Noņemt paroli",
                                                         f"Šablonam '{template_name}' ir parole. Lūdzu, ievadiet to, lai noņemtu:",
                                                         QLineEdit.Password)
            if not ok_old or old_password != current_password:
                QMessageBox.warning(self, "Nepareiza parole", "Ievadītā parole ir nepareiza vai ievade atcelta.")
                return

            data['template_password'] = "" # Noņemam paroli

            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            QMessageBox.information(self, "Parole noņemta", f"Parole no šablona '{template_name}' veiksmīgi noņemta.")
            self._update_sablonu_list() # Atjaunojam sarakstu, lai atspoguļotu izmaiņas
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās noņemt paroli no šablona '{template_name}':\n{e}")

    def ieladet_projektu(self, path=None):
        self._undo_mgr.push_undo(self._snapshot_state('PROJECT_LOAD'))
        self._audit('PROJECT_LOAD', {})
        path = _coerce_path(path)
        if not path:
            path, _ = QFileDialog.getOpenFileName(self, "Ielādēt projektu", PROJECT_SAVE_DIR, "JSON (*.json)")
        if not path:
            return
        try:
            # Mēģinām ielādēt ar dažādām kodēšanām
            encodings = ['utf-8', 'utf-8-sig', 'cp1257', 'iso-8859-1', 'windows-1252']
            data = None

            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        data = json.load(f)
                    break
                except (UnicodeDecodeError, json.JSONDecodeError):
                    continue

            if data is None:
                raise Exception("Neizdevās ielādēt failu ar nevienu no atbalstītajām kodēšanām")

            # Helper to safely get and convert Decimal values
            # Helper to safely get and convert Decimal values
            def get_decimal(dict_obj, key, default_val):
                val = dict_obj.get(key, default_val)
                return to_decimal(val)

            # Helper to safely get boolean values
            def get_bool(dict_obj, key, default_val):
                val = dict_obj.get(key, default_val)
                return bool(val)

            d = AktaDati(
                    akta_nr=data.get('akta_nr', ''), datums=data.get('datums', ''), vieta=data.get('vieta', ''),
                    pasūtījuma_nr=data.get('pasūtījuma_nr', ''),
                    pieņēmējs=Persona(
                        nosaukums=data.get('pieņēmējs', {}).get('nosaukums', ''),
                        reģ_nr=data.get('pieņēmējs', {}).get('reģ_nr', ''),
                        adrese=data.get('pieņēmējs', {}).get('adrese', ''),
                        kontaktpersona=data.get('pieņēmējs', {}).get('kontaktpersona', ''),
                        amats=data.get('pieņēmējs', {}).get('amats', ''),
                        pilnvaras_pamats=data.get('pieņēmējs', {}).get('pilnvaras_pamats', ''),
                        tālrunis=data.get('pieņēmējs', {}).get('tālrunis', ''),
                        epasts=data.get('pieņēmējs', {}).get('epasts', ''),
                        web_lapa=data.get('pieņēmējs', {}).get('web_lapa', ''),
                        bankas_konts=data.get('pieņēmējs', {}).get('bankas_konts', ''),
                        juridiskais_statuss=data.get('pieņēmējs', {}).get('juridiskais_statuss', '')
                    ),
                    nodevējs=Persona(
                        nosaukums=data.get('nodevējs', {}).get('nosaukums', ''),
                        reģ_nr=data.get('nodevējs', {}).get('reģ_nr', ''),
                        adrese=data.get('nodevējs', {}).get('adrese', ''),
                        kontaktpersona=data.get('nodevējs', {}).get('kontaktpersona', ''),
                        amats=data.get('nodevējs', {}).get('amats', ''),
                        pilnvaras_pamats=data.get('nodevējs', {}).get('pilnvaras_pamats', ''),
                        tālrunis=data.get('nodevējs', {}).get('tālrunis', ''),
                        epasts=data.get('nodevējs', {}).get('epasts', ''),
                        web_lapa=data.get('nodevējs', {}).get('web_lapa', ''),
                        bankas_konts=data.get('nodevējs', {}).get('bankas_konts', ''),
                        juridiskais_statuss=data.get('nodevējs', {}).get('juridiskais_statuss', '')
                    ),
                    pozīcijas=[Pozīcija(
                        apraksts=p.get('apraksts', ''),
                        daudzums=get_decimal(p, 'daudzums', '0'),
                        vienība=p.get('vienība', 'gab.'),
                        cena=get_decimal(p, 'cena', '0'),
                        seriālais_nr=p.get('seriālais_nr', ''),
                        garantija=p.get('garantija', ''),
                        piezīmes_pozīcijai=p.get('piezīmes_pozīcijai', '')
                    ) for p in data.get('pozīcijas', [])],
                    attēli=[Attēls(**a) for a in data.get('attēli', [])],
                    piezīmes=data.get('piezīmes', ''), iekļaut_pvn=get_bool(data, 'iekļaut_pvn', False),
                    pvn_likme=get_decimal(data, 'pvn_likme', '21'),
                    parakstu_rindas=get_bool(data, 'parakstu_rindas', True),
                    logotipa_ceļš=data.get('logotipa_ceļš', ''), fonts_ceļš=data.get('fonts_ceļš', ''),
                    paraksts_pieņēmējs_ceļš=data.get('paraksts_pieņēmējs_ceļš', ''),
                    paraksts_nodevējs_ceļš=data.get('paraksts_nodevējs_ceļš', ''),
                    līguma_nr=data.get('līguma_nr', ''),
                    izpildes_termiņš=data.get('izpildes_termiņš', ''),
                    pieņemšanas_datums=data.get('pieņemšanas_datums', ''),
                    nodošanas_datums=data.get('nodošanas_datums', ''),
                    strīdu_risināšana=data.get('strīdu_risināšana', ''),
                    konfidencialitātes_klauzula=get_bool(data, 'konfidencialitātes_klauzula', False),
                    soda_nauda_procenti=get_decimal(data, 'soda_nauda_procenti', '0.0'),
                    piegādes_nosacījumi=data.get('piegādes_nosacījumi', ''),
                    apdrošināšana=get_bool(data, 'apdrošināšana', False),
                        apdrošināšana_teksts=data.get('apdrošināšana_teksts', ''),
                    papildu_nosacījumi=data.get('papildu_nosacījumi', ''),
                    atsauces_dokumenti=data.get('atsauces_dokumenti', ''),
                    akta_statuss=data.get('akta_statuss', 'Melnraksts'),
                    valūta=data.get('valūta', 'EUR'),
                    elektroniskais_paraksts=get_bool(data, 'elektroniskais_paraksts', False),
                    radit_elektronisko_parakstu_tekstu=get_bool(data, 'radit_elektronisko_parakstu_tekstu', False),
                # JAUNA RINDAS
                    pdf_page_size=data.get('pdf_page_size', 'A4'),
                    pdf_page_orientation=data.get('pdf_page_orientation', 'Portrets'),
                    pdf_margin_left=get_decimal(data, 'pdf_margin_left', '18'),
                    pdf_margin_right=get_decimal(data, 'pdf_margin_right', '18'),
                    pdf_margin_top=get_decimal(data, 'pdf_margin_top', '16'),
                    pdf_margin_bottom=get_decimal(data, 'pdf_margin_bottom', '16'),
                    pdf_font_size_head=data.get('pdf_font_size_head', 14),
                    pdf_font_size_normal=data.get('pdf_font_size_normal', 10),
                    pdf_font_size_small=data.get('pdf_font_size_small', 9),
                    pdf_font_size_table=data.get('pdf_font_size_table', 9),
                    pdf_logo_width_mm=get_decimal(data, 'pdf_logo_width_mm', '35'),
                    pdf_signature_width_mm=get_decimal(data, 'pdf_signature_width_mm', '50'),
                    pdf_signature_height_mm=get_decimal(data, 'pdf_signature_height_mm', '20'),
                    docx_image_width_inches=get_decimal(data, 'docx_image_width_inches', '4'),
                    docx_signature_width_inches=get_decimal(data, 'docx_signature_width_inches', '1.5'),
                    table_col_widths=data.get('table_col_widths', '10,40,18,18,20,20,25,25,25'),
                    auto_generate_akta_nr=get_bool(data, 'auto_generate_akta_nr', False),
                    default_currency=data.get('default_currency', 'EUR'),
                    default_unit=data.get('default_unit', 'gab.'),
                    default_pvn_rate=get_decimal(data, 'default_pvn_rate', '21.0'),
                    poppler_path=data.get('poppler_path', ''),
                    # Load new settings
                    header_text_color=data.get('header_text_color', '#000000'),
                    footer_text_color=data.get('footer_text_color', '#000000'),
                    table_header_bg_color=data.get('table_header_bg_color', '#E0E0E0'),
                    table_grid_color=data.get('table_grid_color', '#CCCCCC'),
                    table_row_spacing=get_decimal(data, 'table_row_spacing', '4'),
                    line_spacing_multiplier=get_decimal(data, 'line_spacing_multiplier', '1.2'),
                    show_page_numbers=get_bool(data, 'show_page_numbers', True),
                    show_generation_timestamp=get_bool(data, 'show_generation_timestamp', True),
                    currency_symbol_position=data.get('currency_symbol_position', 'after'),
                    date_format=data.get('date_format', 'YYYY-MM-DD'),
                    signature_line_length_mm=get_decimal(data, 'signature_line_length_mm', '60'),
                    signature_line_thickness_pt=get_decimal(data, 'signature_line_thickness_pt', '0.5'),
                    add_cover_page=get_bool(data, 'add_cover_page', False),
                    cover_page_title=data.get('cover_page_title', 'Pieņemšanas-Nodošanas Akts'),
                    cover_page_logo_width_mm=get_decimal(data, 'cover_page_logo_width_mm', '80'),
                # Individuālais QR kods
                include_custom_qr_code=get_bool(data, 'include_custom_qr_code', False),
                custom_qr_code_data=data.get('custom_qr_code_data', ''),
                custom_qr_code_size_mm=get_decimal(data, 'custom_qr_code_size_mm', '20'),
                custom_qr_code_position=data.get('custom_qr_code_position', 'bottom_right'),
                custom_qr_code_pos_x_mm=get_decimal(data, 'custom_qr_code_pos_x_mm', '0'),
                custom_qr_code_pos_y_mm=get_decimal(data, 'custom_qr_code_pos_y_mm', '0'),
                custom_qr_code_color=data.get('custom_qr_code_color', '#000000'),

                # Automātiskais QR kods (akta ID)
                include_auto_qr_code=get_bool(data, 'include_auto_qr_code', False),
                auto_qr_code_size_mm=get_decimal(data, 'auto_qr_code_size_mm', '20'),
                auto_qr_code_position=data.get('auto_qr_code_position', 'bottom_left'),
                auto_qr_code_pos_x_mm=get_decimal(data, 'auto_qr_code_pos_x_mm', '0'),
                auto_qr_code_pos_y_mm=get_decimal(data, 'auto_qr_code_pos_y_mm', '0'),
                auto_qr_code_color=data.get('auto_qr_code_color', '#000000'),

                add_watermark=get_bool(data, 'add_watermark', False),
                    watermark_text=data.get('watermark_text', 'MELNRAKSTS'),
                    watermark_font_size=data.get('watermark_font_size', 72),
                    watermark_color=data.get('watermark_color', '#E0E0E0'),
                    watermark_rotation=data.get('watermark_rotation', 45),
                    enable_pdf_encryption=get_bool(data, 'enable_pdf_encryption', False),
                    pdf_user_password=data.get('pdf_user_password', ''),
                    pdf_owner_password=data.get('pdf_owner_password', ''),
                    allow_printing=get_bool(data, 'allow_printing', True),
                    allow_copying=get_bool(data, 'allow_copying', True),
                    allow_modifying=get_bool(data, 'allow_modifying', False),
                    allow_annotating=get_bool(data, 'allow_annotating', True),
                    default_country=data.get('default_country', 'Latvija'),
                    default_city=data.get('default_city', 'Rīga'),
                    show_contact_details_in_header=get_bool(data, 'show_contact_details_in_header', False),
                    contact_details_header_font_size=data.get('contact_details_header_font_size', 8),
                    item_image_width_mm=get_decimal(data, 'item_image_width_mm', '50'),
                    item_image_caption_font_size=data.get('item_image_caption_font_size', 8),
                    show_item_notes_in_table=get_bool(data, 'show_item_notes_in_table', True),
                    show_item_serial_number_in_table=get_bool(data, 'show_item_serial_number_in_table', True),
                    show_item_warranty_in_table=get_bool(data, 'show_item_warranty_in_table', True),
                    table_cell_padding_mm=get_decimal(data, 'table_cell_padding_mm', '2'),
                    table_header_font_style=data.get('table_header_font_style', 'bold'),
                    table_content_alignment=data.get('table_content_alignment', 'left'),
                    signature_font_size=data.get('signature_font_size', 9),
                    signature_spacing_mm=get_decimal(data, 'signature_spacing_mm', '10'),
                    document_title_font_size=data.get('document_title_font_size', 18),
                    document_title_color=data.get('document_title_color', '#000000'),
                    section_heading_font_size=data.get('section_heading_font_size', 12),
                    section_heading_color=data.get('section_heading_color', '#000000'),
                    paragraph_line_spacing_multiplier=get_decimal(data, 'paragraph_line_spacing_multiplier', '1.2'),
                    table_border_style=data.get('table_border_style', 'solid'),
                    table_border_thickness_pt=get_decimal(data, 'table_border_thickness_pt', '0.5'),
                    table_alternate_row_color=data.get('table_alternate_row_color', ''),
                    show_total_sum_in_words=get_bool(data, 'show_total_sum_in_words', False),
                    total_sum_in_words_language=data.get('total_sum_in_words_language', 'lv'),
                    default_vat_calculation_method=data.get('default_vat_calculation_method', 'exclusive'),
                    show_vat_breakdown=get_bool(data, 'show_vat_breakdown', True),
                    enable_digital_signature_field=get_bool(data, 'enable_digital_signature_field', False),
                    digital_signature_field_name=data.get('digital_signature_field_name', 'Paraksts'),
                    digital_signature_field_size_mm=get_decimal(data, 'digital_signature_field_size_mm', '40'),
                    digital_signature_field_position=data.get('digital_signature_field_position', 'bottom_center')
                )
            self.ieviest_datus(d)
            self._ceļš_projekts = path
            self._add_to_history(path) # Pievienojam projektu vēsturei
            QMessageBox.information(self, "Ielādēts", "Projekts ielādēts.")
        except Exception as e:
                QMessageBox.critical(self, "Kļūda", f"Neizdevās ielādēt projektu:\n{e}")

    def saglabat_noklusejuma_iestatijumus(self):
        d = self.savākt_datus()
        os.makedirs(SETTINGS_DIR, exist_ok=True)
        settings_path = DEFAULT_SETTINGS_FILE
        out = asdict(d)
        # Convert Decimal fields to string for JSON serialization
        for key, value in out.items():
            if isinstance(value, Decimal):
                out[key] = str(value)
        # Clear sensitive/dynamic data for default settings
        out['pozīcijas'] = []
        out['attēli'] = []
        # JAUNS: noklusējumam saglabājam pielāgoto kolonnu definīcijas, bet notīram rindu datus
        try:
            if isinstance(out.get('custom_columns'), list):
                for cc in out['custom_columns']:
                    if isinstance(cc, dict) and 'data' in cc:
                        cc['data'] = []
        except Exception:
            pass
        out['pieņēmējs'] = asdict(d.pieņēmējs)
        out['nodevējs'] = asdict(d.nodevējs)
        out['pdf_user_password'] = ""  # Do not save passwords as default
        out['pdf_owner_password'] = ""
        # Individuālais QR kods
        out['custom_qr_code_data'] = ""
        out['include_custom_qr_code'] = False
        out['custom_qr_code_size_mm'] = str(Decimal("20"))
        out['custom_qr_code_position'] = "bottom_right"
        out['custom_qr_code_pos_x_mm'] = str(Decimal("0"))
        out['custom_qr_code_pos_y_mm'] = str(Decimal("0"))
        out['custom_qr_code_color'] = "#000000"

        # Automātiskais QR kods (akta ID)
        out['include_auto_qr_code'] = False
        out['auto_qr_code_size_mm'] = str(Decimal("20"))
        out['auto_qr_code_position'] = "bottom_left"
        out['auto_qr_code_pos_x_mm'] = str(Decimal("0"))
        out['auto_qr_code_pos_y_mm'] = str(Decimal("0"))
        out['auto_qr_code_color'] = "#000000"

        out['templates_dir'] = d.templates_dir  # JAUNA RINDAS - Saglabājam šablonu direktoriju

        try:
            with open(settings_path, 'w', encoding='utf-8') as f:
                json.dump(out, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "Saglabāts", "Pašreizējie iestatījumi saglabāti kā noklusējuma.")
        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās saglabāt noklusējuma iestatījumus:\n{e}")

    def ieladet_noklusejuma_iestatijumus(self):
            os.makedirs(SETTINGS_DIR, exist_ok=True)
            settings_path = DEFAULT_SETTINGS_FILE

            if os.path.exists(settings_path):
                try:
                    with open(settings_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                    # Helper to safely get and convert Decimal values
                    def get_decimal(dict_obj, key, default_val):
                        val = dict_obj.get(key, default_val)
                        return to_decimal(val)

                    # Helper to safely get boolean values
                    def get_bool(dict_obj, key, default_val):
                        val = dict_obj.get(key, default_val)
                        return bool(val)

                    d = AktaDati(
                        akta_nr=data.get('akta_nr', ''), datums=data.get('datums', datetime.now().strftime('%Y-%m-%d')), vieta=data.get('vieta', ''),
                        pasūtījuma_nr=data.get('pasūtījuma_nr', ''),
                        pieņēmējs=Persona(
                            nosaukums=data.get('pieņēmējs', {}).get('nosaukums', ''),
                            reģ_nr=data.get('pieņēmējs', {}).get('reģ_nr', ''),
                            adrese=data.get('pieņēmējs', {}).get('adrese', ''),
                            kontaktpersona=data.get('pieņēmējs', {}).get('kontaktpersona', ''),
                            tālrunis=data.get('pieņēmējs', {}).get('tālrunis', ''),
                            epasts=data.get('pieņēmējs', {}).get('epasts', ''),
                            bankas_konts=data.get('pieņēmējs', {}).get('bankas_konts', ''),
                            juridiskais_statuss=data.get('pieņēmējs', {}).get('juridiskais_statuss', '')
                        ),
                        nodevējs=Persona(
                            nosaukums=data.get('nodevējs', {}).get('nosaukums', ''),
                            reģ_nr=data.get('nodevējs', {}).get('reģ_nr', ''),
                            adrese=data.get('nodevējs', {}).get('adrese', ''),
                            kontaktpersona=data.get('nodevējs', {}).get('kontaktpersona', ''),
                            tālrunis=data.get('nodevējs', {}).get('tālrunis', ''),
                            epasts=data.get('nodevējs', {}).get('epasts', ''),
                            web_lapa=data.get('nodevējs', {}).get('web_lapa', ''),
                            bankas_konts=data.get('nodevējs', {}).get('bankas_konts', ''),
                            juridiskais_statuss=data.get('nodevējs', {}).get('juridiskais_statuss', '')
                        ),
                        pozīcijas=[], # Default settings should not load positions
                        custom_columns=data.get('custom_columns', []) if isinstance(data.get('custom_columns', []), list) else [],
                        poz_columns_config=data.get('poz_columns_config', {}) if isinstance(data.get('poz_columns_config', {}), dict) else {},
                        show_price_summary=get_bool(data, 'show_price_summary', True),
                        poz_columns_visual_order=data.get('poz_columns_visual_order', []) if isinstance(data.get('poz_columns_visual_order', []), list) else [],
                        poz_header_state_b64=data.get('poz_header_state_b64', '') if isinstance(data.get('poz_header_state_b64', ''), str) else '',
                        attēli=[], # Default settings should not load images
                        piezīmes=data.get('piezīmes', ''), iekļaut_pvn=get_bool(data, 'iekļaut_pvn', False),
                        pvn_likme=get_decimal(data, 'pvn_likme', '21'),
                        parakstu_rindas=get_bool(data, 'parakstu_rindas', True),
                        logotipa_ceļš=data.get('logotipa_ceļš', ''), fonts_ceļš=data.get('fonts_ceļš', ''),
                        paraksts_pieņēmējs_ceļš=data.get('paraksts_pieņēmējs_ceļš', ''),
                        paraksts_nodevējs_ceļš=data.get('paraksts_nodevējs_ceļš', ''),
                        līguma_nr=data.get('līguma_nr', ''),
                        izpildes_termiņš=data.get('izpildes_termiņš', ''),
                        pieņemšanas_datums=data.get('pieņemšanas_datums', ''),
                        nodošanas_datums=data.get('nodošanas_datums', ''),
                        strīdu_risināšana=data.get('strīdu_risināšana', ''),
                        konfidencialitātes_klauzula=get_bool(data, 'konfidencialitātes_klauzula', False),
                        soda_nauda_procenti=get_decimal(data, 'soda_nauda_procenti', '0.0'),
                        piegādes_nosacījumi=data.get('piegādes_nosacījumi', ''),
                        apdrošināšana=get_bool(data, 'apdrošināšana', False),
                        apdrošināšana_teksts=data.get('apdrošināšana_teksts', ''),
                        papildu_nosacījumi=data.get('papildu_nosacījumi', ''),
                        atsauces_dokumenti=data.get('atsauces_dokumenti', ''),
                        akta_statuss=data.get('akta_statuss', 'Melnraksts'),
                        valūta=data.get('valūta', 'EUR'),
                        elektroniskais_paraksts=get_bool(data, 'elektroniskais_paraksts', False),
                        radit_elektronisko_parakstu_tekstu=get_bool(data, 'radit_elektronisko_parakstu_tekstu', False),
                        # JAUNA RINDAS
                        pdf_page_size=data.get('pdf_page_size', 'A4'),
                        pdf_page_orientation=data.get('pdf_page_orientation', 'Portrets'),
                        pdf_margin_left=get_decimal(data, 'pdf_margin_left', '18'),
                        pdf_margin_right=get_decimal(data, 'pdf_margin_right', '18'),
                        pdf_margin_top=get_decimal(data, 'pdf_margin_top', '16'),
                        pdf_margin_bottom=get_decimal(data, 'pdf_margin_bottom', '16'),
                        pdf_font_size_head=data.get('pdf_font_size_head', 14),
                        pdf_font_size_normal=data.get('pdf_font_size_normal', 10),
                        pdf_font_size_small=data.get('pdf_font_size_small', 9),
                        pdf_font_size_table=data.get('pdf_font_size_table', 9),
                        pdf_logo_width_mm=get_decimal(data, 'pdf_logo_width_mm', '35'),
                        pdf_signature_width_mm=get_decimal(data, 'pdf_signature_width_mm', '50'),
                        pdf_signature_height_mm=get_decimal(data, 'pdf_signature_height_mm', '20'),
                        docx_image_width_inches=get_decimal(data, 'docx_image_width_inches', '4'),
                        docx_signature_width_inches=get_decimal(data, 'docx_signature_width_inches', '1.5'),
                        table_col_widths=data.get('table_col_widths', '10,40,18,18,20,20,25,25,25'),
                        auto_generate_akta_nr=get_bool(data, 'auto_generate_akta_nr', False),
                        default_currency=data.get('default_currency', 'EUR'),
                        default_unit=data.get('default_unit', 'gab.'),
                        default_pvn_rate=get_decimal(data, 'default_pvn_rate', '21.0'),
                        poppler_path=data.get('poppler_path', ''),
                        # Load new settings
                        header_text_color=data.get('header_text_color', '#000000'),
                        footer_text_color=data.get('footer_text_color', '#000000'),
                        table_header_bg_color=data.get('table_header_bg_color', '#E0E0E0'),
                        table_grid_color=data.get('table_grid_color', '#CCCCCC'),
                        table_row_spacing=get_decimal(data, 'table_row_spacing', '4'),
                        line_spacing_multiplier=get_decimal(data, 'line_spacing_multiplier', '1.2'),
                        show_page_numbers=get_bool(data, 'show_page_numbers', True),
                        show_generation_timestamp=get_bool(data, 'show_generation_timestamp', True),
                        currency_symbol_position=data.get('currency_symbol_position', 'after'),
                        date_format=data.get('date_format', 'YYYY-MM-DD'),
                        signature_line_length_mm=get_decimal(data, 'signature_line_length_mm', '60'),
                        signature_line_thickness_pt=get_decimal(data, 'signature_line_thickness_pt', '0.5'),
                        add_cover_page=get_bool(data, 'add_cover_page', False),
                        cover_page_title=data.get('cover_page_title', 'Pieņemšanas-Nodošanas Akts'),
                        cover_page_logo_width_mm=get_decimal(data, 'cover_page_logo_width_mm', '80'),
                        # Individuālais QR kods
                        include_custom_qr_code=get_bool(data, 'include_custom_qr_code', False),
                        custom_qr_code_data=data.get('custom_qr_code_data', ''),
                        custom_qr_code_size_mm=get_decimal(data, 'custom_qr_code_size_mm', '20'),
                        custom_qr_code_position=data.get('custom_qr_code_position', 'bottom_right'),
                        custom_qr_code_pos_x_mm=get_decimal(data, 'custom_qr_code_pos_x_mm', '0'),
                        custom_qr_code_pos_y_mm=get_decimal(data, 'custom_qr_code_pos_y_mm', '0'),
                        custom_qr_code_color=data.get('custom_qr_code_color', '#000000'),

                        # Automātiskais QR kods (akta ID)
                        include_auto_qr_code=get_bool(data, 'include_auto_qr_code', False),
                        auto_qr_code_size_mm=get_decimal(data, 'auto_qr_code_size_mm', '20'),
                        auto_qr_code_position=data.get('auto_qr_code_position', 'bottom_left'),
                        auto_qr_code_pos_x_mm=get_decimal(data, 'auto_qr_code_pos_x_mm', '0'),
                        auto_qr_code_pos_y_mm=get_decimal(data, 'auto_qr_code_pos_y_mm', '0'),
                        auto_qr_code_color=data.get('auto_qr_code_color', '#000000'),

                        add_watermark=get_bool(data, 'add_watermark', False),
                        watermark_text=data.get('watermark_text', 'MELNRAKSTS'),
                        watermark_font_size=data.get('watermark_font_size', 72),
                        watermark_color=data.get('watermark_color', '#E0E0E0'),
                        watermark_rotation=data.get('watermark_rotation', 45),
                        enable_pdf_encryption=get_bool(data, 'enable_pdf_encryption', False),
                        pdf_user_password=data.get('pdf_user_password', ''),
                        pdf_owner_password=data.get('pdf_owner_password', ''),
                        allow_printing=get_bool(data, 'allow_printing', True),
                        allow_copying=get_bool(data, 'allow_copying', True),
                        allow_modifying=get_bool(data, 'allow_modifying', False),
                        allow_annotating=get_bool(data, 'allow_annotating', True),
                        default_country=data.get('default_country', 'Latvija'),
                        default_city=data.get('default_city', 'Rīga'),
                        show_contact_details_in_header=get_bool(data, 'show_contact_details_in_header', False),
                        contact_details_header_font_size=data.get('contact_details_header_font_size', 8),
                        item_image_width_mm=get_decimal(data, 'item_image_width_mm', '50'),
                        item_image_caption_font_size=data.get('item_image_caption_font_size', 8),
                        show_item_notes_in_table=get_bool(data, 'show_item_notes_in_table', True),
                        show_item_serial_number_in_table=get_bool(data, 'show_item_serial_number_in_table', True),
                        show_item_warranty_in_table=get_bool(data, 'show_item_warranty_in_table', True),
                        table_cell_padding_mm=get_decimal(data, 'table_cell_padding_mm', '2'),
                        table_header_font_style=data.get('table_header_font_style', 'bold'),
                        table_content_alignment=data.get('table_content_alignment', 'left'),
                        signature_font_size=data.get('signature_font_size', 9),
                        signature_spacing_mm=get_decimal(data, 'signature_spacing_mm', '10'),
                        document_title_font_size=data.get('document_title_font_size', 18),
                        document_title_color=data.get('document_title_color', '#000000'),
                        section_heading_font_size=data.get('section_heading_font_size', 12),
                        section_heading_color=data.get('section_heading_color', '#000000'),
                        paragraph_line_spacing_multiplier=get_decimal(data, 'paragraph_line_spacing_multiplier', '1.2'),
                        table_border_style=data.get('table_border_style', 'solid'),
                        table_border_thickness_pt=get_decimal(data, 'table_border_thickness_pt', '0.5'),
                        table_alternate_row_color=data.get('table_alternate_row_color', ''),
                        show_total_sum_in_words=get_bool(data, 'show_total_sum_in_words', False),
                        total_sum_in_words_language=data.get('total_sum_in_words_language', 'lv'),
                        default_vat_calculation_method=data.get('default_vat_calculation_method', 'exclusive'),
                        show_vat_breakdown=get_bool(data, 'show_vat_breakdown', True),
                        enable_digital_signature_field=get_bool(data, 'enable_digital_signature_field', False),
                        digital_signature_field_name=data.get('digital_signature_field_name', 'Paraksts'),
                        digital_signature_field_size_mm=get_decimal(data, 'digital_signature_field_size_mm', '40'),
                        digital_signature_field_position=data.get('digital_signature_field_position', 'bottom_center'),
                        templates_dir=data.get('templates_dir', os.path.join(APP_DATA_DIR, "AktaGenerators_Templates"))
                        # JAUNA RINDAS
                    )
                    self.ieviest_datus(d)
                    self.data = d
                except Exception as e:
                    QMessageBox.warning(self, "Iestatījumu ielādes kļūda",
                                        f"Neizdevās ielādēt noklusējuma iestatījumus: {e}")

    # -----------------------
    # Audit + Undo/Redo API
    # -----------------------
    def _audit(self, event: str, details: dict | None = None):
        try:
            self._audit_logger.write(event, details or {}, user=getattr(self, "_current_user", ""))
            # ja audit tabs ir uzbūvēts, atjaunojam
            if hasattr(self, "_audit_table") and self._audit_table is not None:
                self._refresh_audit_table(limit=200)
        except Exception:
            pass

    def _snapshot_state(self, label: str = "") -> dict:
        """Saglabā stāvokli undo vajadzībām (adrešu grāmata + projekts)."""
        try:
            # adrešu grāmata
            ab = copy.deepcopy(getattr(self, "address_book", {}))

            # projekts (akta dati) – ja savākšana izgāžas, saglabājam tikai AB
            proj = None
            try:
                d = self.savākt_datus()
                proj = asdict(d)
            except Exception:
                proj = None

            # UI: izvēlētais address book entry
            sel = ""
            try:
                it = self.address_book_list.currentItem() if hasattr(self, "address_book_list") else None
                sel = it.text() if it else ""
            except Exception:
                sel = ""

            return {"label": label, "address_book": ab, "project": proj, "ab_selected": sel}
        except Exception:
            return {"label": label, "address_book": copy.deepcopy(getattr(self, "address_book", {}))}

    def _restore_state(self, state: dict):
        """Atjauno stāvokli no undo/redo."""
        if not state:
            return
        # adrešu grāmata
        if "address_book" in state:
            self.address_book = copy.deepcopy(state.get("address_book") or {})
            try:
                self._save_address_book()
            except Exception:
                pass
            try:
                self._update_address_book_list()
            except Exception:
                pass

        # projekts
        proj = state.get("project")
        if proj:
            try:
                # saglabājam/atjaunojam izmantojot esošo loģiku
                d = self._akta_dati_from_dict(proj)
                self.ieviest_datus(d)
            except Exception:
                pass

        # atjaunojam selekciju
        try:
            sel = state.get("ab_selected") or ""
            if sel and hasattr(self, "address_book_list"):
                for i in range(self.address_book_list.count()):
                    if self.address_book_list.item(i).text() == sel:
                        self.address_book_list.setCurrentRow(i)
                        break
        except Exception:
            pass

    def undo_action(self):
        if not self._undo_mgr.can_undo():
            return
        # pašreizējo stāvokli ieliek redo
        try:
            self._undo_mgr.push_redo(self._snapshot_state("redo"))
            st = self._undo_mgr.pop_undo()
            self._restore_state(st)
            self._update_undo_redo_indicators()
            self._audit("UNDO", {"label": st.get("label", "")})
        except Exception:
            pass

    def redo_action(self):
        if not self._undo_mgr.can_redo():
            return
        try:
            self._undo_mgr.push_undo(self._snapshot_state("undo"))
            st = self._undo_mgr.pop_redo()
            self._restore_state(st)
            self._update_undo_redo_indicators()
            self._audit("REDO", {"label": st.get("label", "")})
        except Exception:
            pass

    def _akta_dati_from_dict(self, data: dict) -> AktaDati:
        """Atjauno AktaDati no dict (undo/redo/projekta ielāde)."""
        # Reuse conversion helpers already in file
        def get_decimal(dict_obj, key, default_val):
            val = dict_obj.get(key, default_val)
            return to_decimal(val)

        def get_bool(dict_obj, key, default_val):
            val = dict_obj.get(key, default_val)
            if isinstance(val, bool):
                return val
            if isinstance(val, str):
                return val.lower() in ("true", "1", "yes", "jā", "ja", "y")
            return bool(val)

        d = AktaDati(
            akta_nr=data.get("akta_nr", ""),
            datums=data.get("datums", ""),
            vieta=data.get("vieta", ""),
            pasūtījuma_nr=data.get("pasūtījuma_nr", ""),
            pieņēmējs=Persona(**(data.get("pieņēmējs") or {})),
            nodevējs=Persona(**(data.get("nodevējs") or {})),
            pozīcijas=[],
            attēli=[],
            piezīmes=data.get("piezīmes", ""),
            iekļaut_pvn=get_bool(data, "iekļaut_pvn", True),
            pvn_likme=get_decimal(data, "pvn_likme", "21"),
            piegādes_nosacījumi=data.get("piegādes_nosacījumi", ""),
            papildu_nosacījumi=data.get("papildu_nosacījumi", ""),
            atsauces_dokumenti=data.get("atsauces_dokumenti", ""),
            akta_statuss=data.get("akta_statuss", ""),
            cover_page_enabled=get_bool(data, "cover_page_enabled", True),
            cover_include_logo=get_bool(data, "cover_include_logo", True),
            cover_show_company_name=get_bool(data, "cover_show_company_name", True),
            cover_show_date=get_bool(data, "cover_show_date", True),
            cover_show_place=get_bool(data, "cover_show_place", True),
            cover_show_contacts=get_bool(data, "cover_show_contacts", True),
            cover_show_summary=get_bool(data, "cover_show_summary", True),
            cover_show_signatures=get_bool(data, "cover_show_signatures", True),
            cover_show_attachments=get_bool(data, "cover_show_attachments", True),
            cover_title_text=data.get("cover_title_text", "Pieņemšanas-Nodošanas Akts"),
            cover_subtitle_text=data.get("cover_subtitle_text", ""),
            cover_logo_max_width_mm=get_decimal(data, "cover_logo_max_width_mm", "30"),
            cover_logo_max_height_mm=get_decimal(data, "cover_logo_max_height_mm", "15"),
            digital_signature_enabled=get_bool(data, "digital_signature_enabled", False),
            digital_signature_text=data.get("digital_signature_text", "Dokuments parakstīts ar drošu elektronisko parakstu"),
            digital_signature_show_timestamp=get_bool(data, "digital_signature_show_timestamp", True),
            digital_signature_field_label=data.get("digital_signature_field_label", "Paraksts"),
            digital_signature_field_size_mm=get_decimal(data, "digital_signature_field_size_mm", "40"),
            digital_signature_field_position=data.get("digital_signature_field_position", "bottom_center"),
            qr_kods_enabled=get_bool(data, "qr_kods_enabled", True),
            qr_kods_ieklaut_pozicijas=get_bool(data, "qr_kods_ieklaut_pozicijas", True),
            qr_only_first_page=get_bool(data, "qr_only_first_page", False),
            qr_verification_url_enabled=get_bool(data, "qr_verification_url_enabled", False),
            qr_verification_base_url=data.get("qr_verification_base_url", ""),
            qr_kods_izmers_mm=get_decimal(data, "qr_kods_izmers_mm", "18"),
            qr_kods_tikai_pirma_lapa=get_bool(data, "qr_kods_tikai_pirma_lapa", True),
            qr_kods_url_mode=get_bool(data, "qr_kods_url_mode", False),
            qr_kods_url=data.get("qr_kods_url", ""),
        )

        # pozīcijas
        try:
            for p in (data.get("pozīcijas") or []):
                if isinstance(p, dict):
                    d.pozīcijas.append(Pozīcija(**p))
                else:
                    d.pozīcijas.append(p)
        except Exception:
            pass

        # attēli
        try:
            for a in (data.get("attēli") or []):
                if isinstance(a, dict):
                    d.attēli.append(Attēls(**a))
                else:
                    d.attēli.append(a)
        except Exception:
            pass

        return d



    # -----------------------
    # eParaksts integrācija (praktiska)
    # -----------------------
    def _get_eparaksts_app_path(self) -> str:
        """Atgriež saglabāto eParaksts EXE ceļu.
        Persistējas starp programmas palaišanām (QSettings), ar atpakaļsavietojamību uz settings.json.
        """
        try:
            # 1) QSettings (primārais)
            if not hasattr(self, "_qt_settings") or self._qt_settings is None:
                self._qt_settings = QSettings("AktaGenerators", "AktaGeneratorsApp")
            p = self._qt_settings.value("eparaksts/app_path", "", type=str)
            p = _coerce_path(p) or ""
            if p:
                return p

            # 2) Fallback uz veco settings.json, un migrējam uz QSettings
            if not hasattr(self, "_settings") or self._settings is None:
                self._settings = load_settings()
            p2 = _coerce_path((self._settings or {}).get("eparaksts_app_path", "")) or ""
            if p2:
                try:
                    self._qt_settings.setValue("eparaksts/app_path", p2)
                    self._qt_settings.sync()
                except Exception:
                    pass
            return p2
        except Exception:
            return ""


    def _set_eparaksts_app_path(self, p: str):
        """Saglabā eParaksts EXE ceļu gan QSettings (primāri), gan settings.json (fallback)."""
        try:
            p = _coerce_path(p) or ""
            if not hasattr(self, "_qt_settings") or self._qt_settings is None:
                self._qt_settings = QSettings("AktaGenerators", "AktaGeneratorsApp")
            self._qt_settings.setValue("eparaksts/app_path", p)
            self._qt_settings.sync()

            # atpakaļsavietojamība
            if not hasattr(self, "_settings") or self._settings is None:
                self._settings = load_settings()
            self._settings["eparaksts_app_path"] = p
            save_settings(self._settings)
        except Exception:
            pass


    def _open_settings_eparaksts(self):
        try:
            cur = self._get_eparaksts_app_path()
            msg = ("Norādi eParaksts parakstīšanas programmu (piem., eParakstītājs) EXE failu.\n"
                   "Ja neatstāsi, tiks atvērts PDF ar noklusēto programmu, un parakstīšanu veiksi manuāli.")
            fn, _ = QFileDialog.getOpenFileName(self, "Izvēlēties eParaksts programmu", cur or "", "Programmas (*.exe);;Visi faili (*.*)")
            if not fn:
                return
            self._set_eparaksts_app_path(fn)
            QMessageBox.information(self, "OK", "Saglabāts.")
            self._audit("SETTINGS_EPARAKSTS_APP", {"path": fn})
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", str(e))

    def sign_file_with_eparaksts(self, file_path: str | None = None):
        """Atver failu eParaksts parakstīšanai (ārējā programma)."""
        try:
            p = _coerce_path(file_path or getattr(self, "_last_generated_pdf", ""))
            if not p or not os.path.exists(p):
                QMessageBox.warning(self, "Nav faila", "Nav atrasts PDF, ko parakstīt. Vispirms ģenerē PDF.")
                return

            app = self._get_eparaksts_app_path()

            # Ja lietotājs nav norādījis, mēģinām atrast biežākos ceļus Windows
            if (not app) and platform.system().lower().startswith("win"):
                candidates = [
                    os.path.join(os.environ.get("ProgramFiles", "C:\\Program Files"), "eParakstītājs", "eParakstītājs.exe"),
                    os.path.join(os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)"), "eParakstītājs", "eParakstītājs.exe"),
                ]
                for c in candidates:
                    if os.path.exists(c):
                        app = c
                        break

            if app and os.path.exists(app):
                # sākam sekot parakstītajam failam tajā pašā mapē
                self._start_signed_file_watcher(p)
                # Daudzas e-paraksta programmas pieņem failu kā argumentu (ja nepieņem, vismaz atvērsies).
                subprocess.Popen([app, p], shell=False)
                self._audit("EPARAKSTS_OPEN_APP", {"app": app, "file": p})
            else:
                # Fallback: atver ar noklusēto PDF programmu
                try:
                    if platform.system().lower().startswith("win"):
                        os.startfile(p)  # type: ignore
                    else:
                        subprocess.Popen(["xdg-open", p])
                    self._audit("EPARAKSTS_OPEN_FALLBACK", {"file": p})
                except Exception as e:
                    QMessageBox.warning(self, "Kļūda", f"Neizdevās atvērt failu: {e}")
                    return

                QMessageBox.information(
                    self,
                    "Parakstīšana",
                    "PDF ir atvērts. Paraksti to ar eParaksts/eID rīku (ārējā programmā).\n"
                    "Ja vēlies automātiski atvērt eParakstītāju, iestatos norādi tā EXE ceļu: Fails → Iestatījumi → eParaksts."
                )
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", str(e))



    def _open_folder(self, folder: str):
        try:
            folder = _coerce_path(folder)
            if not folder:
                return
            if platform.system().lower().startswith("win"):
                os.startfile(folder)  # type: ignore
            else:
                subprocess.Popen(["xdg-open", folder])
        except Exception:
            pass

    def _start_signed_file_watcher(self, original_pdf: str):
        """Pēc eParaksts atvēršanas sekojam mapē, vai parādās parakstītais fails (.asice/.edoc/.pdf)."""
        try:
            original_pdf = _coerce_path(original_pdf)
            if not original_pdf:
                return
            folder = os.path.dirname(original_pdf)
            stem = os.path.splitext(os.path.basename(original_pdf))[0]
            self._signed_watch_folder = folder
            self._signed_watch_stem = stem
            self._signed_watch_seen = set(os.listdir(folder)) if os.path.isdir(folder) else set()
            if hasattr(self, "_signed_watch_timer") and self._signed_watch_timer is not None:
                self._signed_watch_timer.stop()
            self._signed_watch_timer = QTimer(self)
            self._signed_watch_timer.setInterval(1200)
            self._signed_watch_timer.timeout.connect(self._poll_signed_file)
            self._signed_watch_timer.start()
        except Exception:
            pass

    def _poll_signed_file(self):
        try:
            folder = getattr(self, "_signed_watch_folder", "")
            stem = getattr(self, "_signed_watch_stem", "")
            if not folder or not os.path.isdir(folder):
                return
            now = set(os.listdir(folder))
            new_files = [f for f in (now - getattr(self, "_signed_watch_seen", set()))]
            self._signed_watch_seen = now

            # Meklējam tipiskos parakstītos failus
            candidates = []
            for f in new_files:
                fl = f.lower()
                if fl.endswith((".asice", ".edoc", ".asics", ".bdoc", ".pdf")):
                    # vēlams ar tādu pašu stem
                    if stem and f.startswith(stem):
                        candidates.append(f)
            if not candidates:
                # arī vecā sarakstā var parādīties ar aizturi — pārbaudam visu mapi
                for f in now:
                    fl = f.lower()
                    if stem and f.startswith(stem) and fl.endswith((".asice", ".edoc", ".asics", ".bdoc")):
                        candidates = [f]
                        break

            if candidates:
                signed = os.path.join(folder, candidates[0])
                self._last_signed_file = signed
                try:
                    if hasattr(self, "_signed_watch_timer") and self._signed_watch_timer is not None:
                        self._signed_watch_timer.stop()
                except Exception:
                    pass
                self._audit("EPARAKSTS_SIGNED_DETECTED", {"file": signed})
                QMessageBox.information(
                    self,
                    "Parakstīts",
                    f"Parakstītais fails atrasts:\n{signed}"
                )
        except Exception:
            pass

    def generate_and_sign_current(self):
        """1) Automātiski ģenerē PDF no šī brīža datiem 2) Atver eParaksts parakstīšanai 3) SeKo parakstītajam failam tajā pašā mapē."""
        try:
            d = self.savākt_datus()

            safe_akta_nr = drošs_faila_nosaukums(d.akta_nr) if d.akta_nr else "akts"
            safe_datums = d.datums.replace("-", "") if d.datums else datetime.now().strftime("%Y%m%d")
            default_name = f"{safe_akta_nr}_{safe_datums}.pdf"

            # izvēlamies saglabāšanas vietu
            default_dir = PROJECT_SAVE_DIR if os.path.isdir(PROJECT_SAVE_DIR) else os.getcwd()
            default_path = os.path.join(default_dir, default_name)

            pdf_path, _ = QFileDialog.getSaveFileName(self, "Saglabāt PDF (pirms parakstīšanas)", default_path, "PDF faili (*.pdf)")
            if not pdf_path:
                return

            # Undo checkpoint + audit
            self._undo_mgr.push_undo(self._snapshot_state("GENERATE_AND_SIGN"))
            self._audit("GENERATE_AND_SIGN_START", {"file": pdf_path})

            # ģenerējam PDF tieši uz izvēlēto vietu
            pdf_ceļš = ģenerēt_pdf(d, pdf_path, include_reference_docs=True, encrypt_pdf=True)
            pdf_ceļš = _coerce_path(pdf_ceļš) or _coerce_path(pdf_path)
            self._last_generated_pdf = pdf_ceļš
            self._audit("GENERATE_AND_SIGN_PDF_DONE", {"file": pdf_ceļš})

            # atveram eParaksts un sākam sekot rezultātam
            self.sign_file_with_eparaksts(pdf_ceļš)
            self._start_signed_file_watcher(pdf_ceļš)

            QMessageBox.information(
                self,
                "Parakstīšana",
                "PDF ir saglabāts un atvērts parakstīšanai.\n"
                "Parakstīto failu saglabā tajā pašā mapē.\n"
                "Programma automātiski mēģinās atrast parakstīto failu."
            )
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", str(e))


    def closeEvent(self, event):
        # Saglabājam vēsturi, adrešu grāmatu un teksta blokus vienmēr
        self._save_history()
        self._save_address_book()
        self.text_block_manager._save_text_blocks()

        # Noklusējuma iestatījumus saglabājam tikai, ja lietotājs to vēlas
        try:
            ans = QMessageBox.question(
                self,
                "Saglabāt noklusējuma iestatījumus?",
                "Vai saglabāt pašreizējos iestatījumus kā noklusējumu nākamajai palaišanai?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                QMessageBox.No
            )
            if ans == QMessageBox.Cancel:
                event.ignore()
                return
            if ans == QMessageBox.Yes:
                self.saglabat_noklusejuma_iestatijumus()
        except Exception:
            # Ja dialoga parādīšana neizdodas, vienkārši neaiztiekam noklusējumu
            pass

        event.accept()

    def _update_preview(self):
        # Atlikt atjaunināšanu par 500 ms, lai izvairītos no biežas PDF ģenerēšanas
        self.preview_timer.start(500)

    def _do_update_preview(self):
        if not hasattr(self, 'preview_label'):
            return

        old_page = getattr(self, 'current_preview_page', 0)
        d = self.savākt_datus()
        # Izveidot hash no datiem, lai pārbaudītu, vai tie mainījušies
        import hashlib
        data_str = json.dumps(asdict(d), sort_keys=True, default=str)
        data_hash = hashlib.md5(data_str.encode('utf-8')).hexdigest()

        # Pārbaudīt kešatmiņu
        if data_hash == self.last_data_hash and self.preview_cache.get(data_hash):
            # Izmantot kešatmiņu
            cached = self.preview_cache[data_hash]
            self.preview_images = cached['images']
            self.current_preview_page = min(self.current_preview_page, len(self.preview_images) - 1)
            self._show_current_page()
            return

        # Smagā daļa (PDF ģenerācija + DOCX/XLSX konvertācija + pdf2image) notiek fonā,
        # lai programma nekad neuzkaras un neaizveras.
        self._requested_preview_hash = data_hash
        self._start_preview_worker(d, data_hash, old_page)


    
    def _is_preview_thread_running(self) -> bool:
        """Droši pārbauda, vai preview QThread vēl skrien (izvairās no 'already deleted' RuntimeError)."""
        if self._preview_thread is None:
            return False
        try:
            return self._preview_thread.isRunning()
        except RuntimeError:
            # C++ objekts jau izdzēsts (deleteLater), bet Python reference palika
            self._preview_thread = None
            self._preview_worker = None
            return False

    def _cleanup_preview_thread_refs(self):
        """Notīra atsauces uz preview thread/worker, kad thread ir beidzies."""
        self._preview_thread = None
        self._preview_worker = None

    def _start_preview_worker(self, d: 'AktaDati', data_hash: str, old_page: int):
        """Startē (vai ieplāno) priekšskatījuma ģenerēšanu fonā."""
        # Ja jau notiek ģenerēšana – nepārtraucam thread vardarbīgi; ieplānojam jaunāko pieprasījumu.
        if self._is_preview_thread_running():
            self._pending_preview_request = (d, data_hash, old_page)
            self.preview_label.setText("Ģenerē priekšskatījumu... (gaida rindā)")
            return

        # Ja vecais thread objekts ir palicis atsaucēs (bet vairs neskrien), droši notīram.
        try:
            if self._preview_thread is not None:
                try:
                    self._preview_thread.quit()
                    self._preview_thread.wait(50)
                except RuntimeError:
                    # jau izdzēsts
                    pass
        except Exception:
            pass

        self.preview_label.setText("Ģenerē priekšskatījumu...")
        QApplication.processEvents()

        thread = QThread(self)
        worker = _PreviewBuildWorker(d, data_hash, old_page)
        worker.moveToThread(thread)

        # Saglabājam atsauces, lai Qt tās negarbāž ārā
        self._preview_thread = thread
        self._preview_worker = worker

        thread.started.connect(worker.run)
        worker.finished.connect(self._on_preview_worker_finished)
        worker.failed.connect(self._on_preview_worker_failed)

        # Dzīves cikls
        worker.finished.connect(thread.quit)
        worker.failed.connect(thread.quit)

        # Kad beidzas, notīram atsauces PIRMS deleteLater (lai nerodas isRunning() uz izdzēsta C++ objekta)
        thread.finished.connect(self._cleanup_preview_thread_refs)
        thread.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)

        thread.start()


    def _on_preview_worker_finished(self, data_hash: str, png_bytes_list: list, old_page: int):
        # Ja pa vidu bija jauns pieprasījums, bet šis nav jaunākais – ignorējam.
        if self._requested_preview_hash is not None and data_hash != self._requested_preview_hash:
            # Ja bija ieplānots jaunāks, palaidīsim to, kad thread beigsies (šeit jau beidzies)
            if self._pending_preview_request:
                d2, h2, op2 = self._pending_preview_request
                self._pending_preview_request = None
                self._requested_preview_hash = h2
                self._start_preview_worker(d2, h2, op2)
            return

        try:
            self.preview_images = []
            for b in png_bytes_list:
                pixmap = QPixmap()
                pixmap.loadFromData(b, 'PNG')
                self.preview_images.append(pixmap)

            # Saglabāt kešatmiņā
            self.preview_cache[data_hash] = {
                'images': self.preview_images.copy(),
                'page_count': len(self.preview_images)
            }
            self.last_data_hash = data_hash

            # Ierobežot kešatmiņas izmēru (saglabāt tikai pēdējos 5)
            if len(self.preview_cache) > 5:
                oldest_key = next(iter(self.preview_cache))
                del self.preview_cache[oldest_key]

            self.current_preview_page = min(max(old_page, 0), len(self.preview_images) - 1) if self.preview_images else 0
            self._show_current_page()

        finally:
            # Palaist pending, ja tāds ir
            if self._pending_preview_request:
                d2, h2, op2 = self._pending_preview_request
                self._pending_preview_request = None
                self._requested_preview_hash = h2
                self._start_preview_worker(d2, h2, op2)


    def _on_preview_worker_failed(self, data_hash: str, error_message: str):
        # Ja tas nav jaunākais pieprasījums, ignorējam.
        if self._requested_preview_hash is not None and data_hash != self._requested_preview_hash:
            if self._pending_preview_request:
                d2, h2, op2 = self._pending_preview_request
                self._pending_preview_request = None
                self._requested_preview_hash = h2
                self._start_preview_worker(d2, h2, op2)
            return

        self.preview_images = []
        self._show_current_page()
        self.preview_label.setText(f"Kļūda priekšskatījumā: {error_message}")
        QMessageBox.critical(
            self,
            "Priekšskatījuma kļūda",
            f"Neizdevās ģenerēt priekšskatījumu (ar atsauces pielikumiem).\n"
            f"Iespējams, trūkst LibreOffice (DOCX/XLSX konvertācijai) vai Poppler (PDF renderēšanai).\n"
            f"Kļūda: {error_message}"
        )

        if self._pending_preview_request:
            d2, h2, op2 = self._pending_preview_request
            self._pending_preview_request = None
            self._requested_preview_hash = h2
            self._start_preview_worker(d2, h2, op2)

    def _show_current_page(self):
        if self.preview_images and 0 <= self.current_preview_page < len(self.preview_images):
            pixmap = self.preview_images[self.current_preview_page]
            label_size = self.preview_scroll_area.viewport().size()
            scaled_pixmap = pixmap.scaled(label_size * self.zoom_factor, Qt.KeepAspectRatio,
                                          Qt.SmoothTransformation)
            self.preview_label.setPixmap(scaled_pixmap)
            # Lai QScrollArea varētu skrollēt (un pan ar peli strādātu), QLabel izmērs jāpielāgo pixmap izmēram
            self.preview_label.resize(scaled_pixmap.size())
            self.page_number_label.setText(f"Lapa {self.current_preview_page + 1}/{len(self.preview_images)}")
            self.prev_page_button.setEnabled(self.current_preview_page > 0)
            self.next_page_button.setEnabled(self.current_preview_page < len(self.preview_images) - 1)
        else:
            self.preview_label.clear()
            self.page_number_label.setText("Lapa 0/0")
            self.prev_page_button.setEnabled(False)
            self.next_page_button.setEnabled(False)

    def _show_prev_page(self):
            if self.current_preview_page > 0:
                self.current_preview_page -= 1
                self._show_current_page()

    def _show_next_page(self):
            if self.current_preview_page < len(self.preview_images) - 1:
                self.current_preview_page += 1
                self._show_current_page()

    
    
    def _savakt_akta_datus(self) -> AktaDati:
        """Savāc datus no UI. Alias vecākiem izsaukumiem (ZIP u.c.)."""
        return self.savākt_datus()

    def _ģenerēt_pdf_failu(self, d: AktaDati, pdf_path: str) -> str:
        """Ģenerē galveno PDF (kā eksportā) un pievieno atsauces pielikumus (Atvasinājumi), ja tādi ir.
        Alias priekš vecākiem izsaukumiem ZIP funkcijā.
        """
        # 1) ģenerējam pamata PDF
        ģenerēt_pdf(d, pdf_path)

        # 2) pieliekam atsauces dokumentus (ja ir)
        try:
            # nododam to pašu fontu, lai latviešu diakritika vienmēr ir korekta
            fn = reģistrēt_fontu(getattr(d, "fonts_ceļš", ""))
            _append_reference_docs_to_pdf(pdf_path, d, font_name=fn)
        except Exception as e:
            print(f"Pielikumu pievienošanas kļūda: {e}")
        return pdf_path

    def ģenerēt_zip_dialogs(self):
        """Saglabā ZIP arhīvu: PDF + visi pielikumi (atsevišķi) + projekta JSON."""
        d0 = self.savākt_datus()
        safe_akta = drošs_faila_nosaukums(d0.akta_nr) if getattr(d0, "akta_nr", "") else "Akts"
        today = datetime.now().strftime("%Y%m%d")
        default_name = f"{safe_akta}_{today}.zip"
        default_path = os.path.join(DEFAULT_OUTPUT_DIR, default_name) if 'DEFAULT_OUTPUT_DIR' in globals() else default_name
        zip_path, _ = QFileDialog.getSaveFileName(self, "Saglabāt ZIP", default_path, "ZIP arhīvs (*.zip)")
        if not zip_path:
            return
        if not zip_path.lower().endswith(".zip"):
            zip_path += ".zip"

        try:
            # Sagatavojam pagaidu failus
            tmp_dir = tempfile.mkdtemp(prefix="akta_zip_")
            pdf_path = os.path.join(tmp_dir, "Akts.pdf")
            json_path = os.path.join(tmp_dir, "projekts.json")

            # Ģenerējam PDF uz pagaidu vietu
            d = self.savākt_datus()
            self._ģenerēt_pdf_failu(d, pdf_path)

            # Saglabājam JSON
            with open(json_path, "w", encoding="utf-8") as f:
                def _json_safe(o):
                    # Pārvērš Decimal un citus JSON-nederīgus tipus par drošu formu
                    from decimal import Decimal
                    if isinstance(o, Decimal):
                        return str(o)
                    if isinstance(o, dict):
                        return {k: _json_safe(v) for k, v in o.items()}
                    if isinstance(o, (list, tuple)):
                        return [_json_safe(v) for v in o]
                    return o

                json.dump(_json_safe(asdict(d)), f, ensure_ascii=False, indent=2)

            import zipfile
            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
                z.write(pdf_path, arcname="Akts.pdf")
                z.write(json_path, arcname="projekts.json")

                # Pielikumi (atsauces dokumenti) kā atsevišķi faili
                used = set()
                for i in range(self.list_atsauces_faili.count()):
                    it = self.list_atsauces_faili.item(i)
                    p = it.data(Qt.UserRole)
                    p = _coerce_path(p)
                    if not p or not os.path.exists(p):
                        continue
                    base = os.path.basename(p)
                    name = base
                    k = 2
                    while name.lower() in used:
                        root, ext = os.path.splitext(base)
                        name = f"{root}_{k}{ext}"
                        k += 1
                    used.add(name.lower())
                    z.write(p, arcname=os.path.join("pielikumi", name))

            # Ierakstām dokumentu vēsturē (kopējam PDF+JSON uz vēstures mapi)
            self._record_generated_document(pdf_path, json_path)

            QMessageBox.information(self, "Gatavs", "ZIP fails saglabāts veiksmīgi.")
        except Exception as e:
            QMessageBox.warning(self, "Kļūda", f"Neizdevās saglabāt ZIP: {e}")

    def ģenerēt_pdf_dialogs(self):
        self._undo_mgr.push_undo(self._snapshot_state('GENERATE_PDF'))
        self._audit('GENERATE_PDF', {})
        d = self.savākt_datus()

        # Izveidojam automātisku faila nosaukumu
        safe_akta_nr = drošs_faila_nosaukums(d.akta_nr) if d.akta_nr else "akts"
        safe_datums = d.datums.replace("-", "") if d.datums else datetime.now().strftime("%Y%m%d")
        base_name = f"Akts_{safe_akta_nr}_{safe_datums}"

        # Izveidojam mapi dokumentam
        doc_folder = os.path.join(DEFAULT_OUTPUT_DIR, base_name)
        counter = 1
        while os.path.exists(doc_folder):
            doc_folder = os.path.join(DEFAULT_OUTPUT_DIR, f"{base_name}_{counter:03d}")
            counter += 1

        os.makedirs(doc_folder, exist_ok=True)

        # Failu ceļi
        pdf_path = os.path.join(doc_folder, f"{os.path.basename(doc_folder)}.pdf")
        json_path = os.path.join(doc_folder, f"{os.path.basename(doc_folder)}.json")

        try:
            # Ģenerējam PDF
            ģenerēt_pdf(d, pdf_path)

            # Saglabājam JSON
            out = asdict(d)
            for key, value in out.items():
                if isinstance(value, Decimal):
                    out[key] = str(value)
            for p in out['pozīcijas']:
                for key, value in p.items():
                    if isinstance(value, Decimal):
                        p[key] = str(value)
            out['pieņēmējs'] = asdict(d.pieņēmējs)
            out['nodevējs'] = asdict(d.nodevējs)

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(out, f, ensure_ascii=False, indent=2)

            self._add_to_history(json_path)
            QMessageBox.information(self, "PDF ģenerēts", f"Dokumenti saglabāti mapē: {doc_folder}")
            # Pēc veiksmīgas ģenerēšanas, atjaunojam akta numuru, ja ieslēgta auto-ģenerēšana
            if self.data.auto_generate_akta_nr:
                self._generate_akta_nr()

        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās ģenerēt PDF:\n{e}")

    def ģenerēt_docx_dialogs(self):
        self._undo_mgr.push_undo(self._snapshot_state('GENERATE_DOCX'))
        self._audit('GENERATE_DOCX', {})
        d = self.savākt_datus()

        # Izveidojam automātisku faila nosaukumu
        safe_akta_nr = drošs_faila_nosaukums(d.akta_nr) if d.akta_nr else "akts"
        safe_datums = d.datums.replace("-", "") if d.datums else datetime.now().strftime("%Y%m%d")
        base_name = f"Akts_{safe_akta_nr}_{safe_datums}"

        # Izveidojam mapi dokumentam
        doc_folder = os.path.join(DEFAULT_OUTPUT_DIR, base_name)
        counter = 1
        while os.path.exists(doc_folder):
            doc_folder = os.path.join(DEFAULT_OUTPUT_DIR, f"{base_name}_{counter:03d}")
            counter += 1

        os.makedirs(doc_folder, exist_ok=True)

        # Failu ceļi
        docx_path = os.path.join(doc_folder, f"{os.path.basename(doc_folder)}.docx")
        json_path = os.path.join(doc_folder, f"{os.path.basename(doc_folder)}.json")

        try:
            # Ģenerējam DOCX
            ģenerēt_docx(d, docx_path)

            # Saglabājam JSON
            out = asdict(d)
            for key, value in out.items():
                if isinstance(value, Decimal):
                    out[key] = str(value)
            for p in out['pozīcijas']:
                for key, value in p.items():
                    if isinstance(value, Decimal):
                        p[key] = str(value)
            out['pieņēmējs'] = asdict(d.pieņēmējs)
            out['nodevējs'] = asdict(d.nodevējs)

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(out, f, ensure_ascii=False, indent=2)

            self._add_to_history(json_path)
            QMessageBox.information(self, "DOCX ģenerēts", f"Dokumenti saglabāti mapē: {doc_folder}")
            # Pēc veiksmīgas ģenerēšanas, atjaunojam akta numuru, ja ieslēgta auto-ģenerēšana
            if self.data.auto_generate_akta_nr:
                self._generate_akta_nr()

        except Exception as e:
            QMessageBox.critical(self, "Kļūda", f"Neizdevās ģenerēt DOCX:\n{e}")

    # =======================
    # HOTFIX: PDF path safety
    # =======================

    # Saglabājam oriģinālo funkciju
    _original_generate_pdf = ģenerēt_pdf

    def ģenerēt_pdf(
            akta_dati: AktaDati,
            pdf_ceļš: str | None = None,
            include_reference_docs: bool = True,
            encrypt_pdf: bool = True
    ) -> str:
        """
        STABILA PDF ģenerēšana.
        VIENMĒR atgriež STRING faila ceļu.
        NEKAD neatgriež dict.
        """

        # 🔒 1. Garantējam, ka pdf_ceļš ir STRING
        if isinstance(pdf_ceļš, dict):
            pdf_ceļš = pdf_ceļš.get("path") or ""
        if pdf_ceļš and not isinstance(pdf_ceļš, (str, os.PathLike)):
            pdf_ceļš = ""

        # 🔒 2. Ja nav dots ceļš – veidojam pagaidu failu
        if not pdf_ceļš:
            fd, pdf_ceļš = tempfile.mkstemp(suffix=".pdf")
            os.close(fd)

        pdf_ceļš = os.fspath(pdf_ceļš)

        # 🔒 3. PDF ģenerēšana (ReportLab)
        font_name = reģistrēt_fontu(akta_dati.fonts_ceļš)

        doc = SimpleDocTemplate(
            pdf_ceļš,
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
            title="Pieņemšanas–Nodošanas akts",
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name="NormalLV",
            fontName=font_name,
            fontSize=10,
            leading=13
        ))

        story = []
        story.append(Paragraph("PIEŅEMŠANAS–NODOŠANAS AKTS", styles["NormalLV"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Akta Nr.: {akta_dati.akta_nr}", styles["NormalLV"]))
        story.append(Paragraph(f"Datums: {akta_dati.datums}", styles["NormalLV"]))
        story.append(Paragraph(f"Vieta: {akta_dati.vieta}", styles["NormalLV"]))

        doc.build(story)

        # 🔒 4. DROŠĪBA: ja fails tiešām eksistē
        if not os.path.exists(pdf_ceļš):
            raise RuntimeError("PDF netika izveidots")

        # 🔒 5. VIENMĒR atgriežam STRING
        return pdf_ceļš


if __name__ == "__main__":
    # Windows: korekta Taskbar ikona + grupēšana
    set_windows_app_id("kulinics.akta_generators")

    app = QApplication(sys.argv)
    try:
        app.setWindowIcon(QIcon(resource_path("Akta_Generators_Icon.ico")))
    except Exception:
        pass

    apply_modern_theme(app, dark=True)
    window = AktaLogs()
    window.show()
    sys.exit(app.exec())

# =======================
# HOTFIX: PDF path safety
# =======================

# Saglabājam oriģinālo funkciju
_original_generate_pdf = ģenerēt_pdf

def ģenerēt_pdf(
    akta_dati: AktaDati,
    pdf_ceļš: str | None = None,
    include_reference_docs: bool = True,
    encrypt_pdf: bool = True
) -> str:
    """
    STABILA PDF ģenerēšana.
    VIENMĒR atgriež STRING faila ceļu.
    NEKAD neatgriež dict.
    """

    # 🔒 1. Garantējam, ka pdf_ceļš ir STRING
    if isinstance(pdf_ceļš, dict):
        pdf_ceļš = pdf_ceļš.get("path") or ""
    if pdf_ceļš and not isinstance(pdf_ceļš, (str, os.PathLike)):
        pdf_ceļš = ""

    # 🔒 2. Ja nav dots ceļš – veidojam pagaidu failu
    if not pdf_ceļš:
        fd, pdf_ceļš = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

    pdf_ceļš = os.fspath(pdf_ceļš)

    # 🔒 3. PDF ģenerēšana (ReportLab)
    font_name = reģistrēt_fontu(akta_dati.fonts_ceļš)

    doc = SimpleDocTemplate(
        pdf_ceļš,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Pieņemšanas–Nodošanas akts",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="NormalLV",
        fontName=font_name,
        fontSize=10,
        leading=13
    ))

    story = []
    story.append(Paragraph("PIEŅEMŠANAS–NODOŠANAS AKTS", styles["NormalLV"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Akta Nr.: {akta_dati.akta_nr}", styles["NormalLV"]))
    story.append(Paragraph(f"Datums: {akta_dati.datums}", styles["NormalLV"]))
    story.append(Paragraph(f"Vieta: {akta_dati.vieta}", styles["NormalLV"]))

    doc.build(story)

    # 🔒 4. DROŠĪBA: ja fails tiešām eksistē
    if not os.path.exists(pdf_ceļš):
        raise RuntimeError("PDF netika izveidots")

    # 🔒 5. VIENMĒR atgriežam STRING
    return pdf_ceļš




def draw_first_page(canvas, doc):
    try:
        draw_header(canvas, doc, show_logo=False)
    except TypeError:
        draw_header(canvas, doc)

def draw_later_pages(canvas, doc):
    try:
        draw_header(canvas, doc, show_logo=False)
    except TypeError:
        pass
